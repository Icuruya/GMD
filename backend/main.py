import re
import io
import json
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Depends
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from typing import Set, List, Dict
from sqlalchemy.orm import Session
from pydantic import BaseModel

# --- Importar la instancia de Celery y la tarea ---
from celery_worker import celery_app, generate_documents_task
from database import SessionLocal, engine
import models

models.Base.metadata.create_all(bind=engine)

# Pydantic model para la entrada de mapeos
class MappingCreate(BaseModel):
    name: str
    template_id: int
    mapping_data: Dict[str, str]
    project_id: int

class ProjectCreate(BaseModel):
    name: str

# --- Regex para encontrar placeholders como [TEXTO] ---
PLACEHOLDER_REGEX = re.compile(r"\[(.*?)\]")

# --- Inicialización de la Aplicación FastAPI ---
app = FastAPI(
    title="Document Generation API",
    description="API para generar documentos a partir de plantillas .docx y datos.",
    version="0.3.0", # Versión actualizada para reflejar la arquitectura de la base de datos
)

# --- Dependencia de la Sesión de la Base de Datos ---
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- Configuración de CORS ---
origins = [
    "http://localhost",
    "http://localhost:3000",
    "http://localhost:9002",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Funciones de Lógica de Documentos (solo la que se necesita aquí) ---

def find_placeholders_in_docx(doc: Document) -> Set[str]:
    """Analiza un documento y extrae todos los placeholders únicos."""
    placeholders = set()
    for p in doc.paragraphs:
        placeholders.update(PLACEHOLDER_REGEX.findall(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    placeholders.update(PLACEHOLDER_REGEX.findall(p.text))
    return placeholders

# --- Endpoints de la API ---

@app.post("/templates", tags=["Templates"])
async def create_template(file: UploadFile = File(...), project_id: int = Form(...), db: Session = Depends(get_db)):
    """Sube una plantilla .docx, la guarda y devuelve sus placeholders."""
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="El archivo debe ser de tipo .docx")
    
    # Guardar el archivo en el directorio de uploads
    file_path = f"uploads/{file.filename}"
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Crear el registro en la base de datos
    db_template = models.Template(name=file.filename, file_path=file_path, project_id=project_id) # owner_id se elimina por ahora
    db.add(db_template)
    db.commit()
    db.refresh(db_template)

    # Analizar los placeholders
    try:
        document = Document(file_path)
        placeholders = find_placeholders_in_docx(document)
        return {"id": db_template.id, "name": db_template.name, "placeholders": sorted(list(placeholders))}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando el archivo: {e}")

@app.get("/templates", tags=["Templates"])
async def get_templates(db: Session = Depends(get_db)):
    """Devuelve una lista de todas las plantillas guardadas."""
    templates = db.query(models.Template).all()
    return templates

@app.get("/projects/{project_id}/templates", tags=["Templates"])
async def get_project_templates(project_id: int, db: Session = Depends(get_db)):
    """Devuelve una lista de plantillas para un proyecto específico."""
    templates = db.query(models.Template).filter(models.Template.project_id == project_id).all()
    return templates

@app.get("/templates/{template_id}", tags=["Templates"])
async def get_template_details(template_id: int, db: Session = Depends(get_db)):
    """Devuelve los detalles de una plantilla específica."""
    db_template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not db_template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")
    return db_template

@app.get("/templates/{template_id}/placeholders", tags=["Templates"])
async def get_template_placeholders(template_id: int, db: Session = Depends(get_db)):
    """Devuelve los placeholders de una plantilla específica."""
    db_template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not db_template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    try:
        document = Document(db_template.file_path)
        placeholders = find_placeholders_in_docx(document)
        return {"placeholders": sorted(list(placeholders))}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando el archivo: {e}")

@app.post("/mappings", tags=["Mappings"])
async def create_mapping(mapping: MappingCreate, db: Session = Depends(get_db)):
    """Guarda una nueva configuración de mapeo para una plantilla."""
    db_project = db.query(models.Project).filter(models.Project.id == mapping.project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")

    db_mapping = models.Mapping(
        name=mapping.name,
        template_id=mapping.template_id,
        mapping_data=mapping.mapping_data,
        project_id=mapping.project_id
    )
    db.add(db_mapping)
    db.commit()
    db.refresh(db_mapping)
    return db_mapping

@app.get("/templates/{template_id}/mappings", tags=["Mappings"])
async def get_template_mappings(template_id: int, db: Session = Depends(get_db)):
    """Devuelve todas las configuraciones de mapeo guardadas para una plantilla específica."""
    mappings = db.query(models.Mapping).filter(models.Mapping.template_id == template_id).all()
    return mappings

@app.get("/projects/{project_id}/mappings", tags=["Mappings"])
async def get_project_mappings(project_id: int, db: Session = Depends(get_db)):
    """Devuelve todas las configuraciones de mapeo guardadas para un proyecto específico."""
    mappings = db.query(models.Mapping).filter(models.Mapping.project_id == project_id).all()
    return mappings

@app.post("/projects", tags=["Projects"])
async def create_project(project: ProjectCreate, db: Session = Depends(get_db)):
    """Crea un nuevo proyecto."""
    # Por ahora, asignaremos el owner_id a 1 (usuario por defecto)
    db_project = models.Project(name=project.name, owner_id=1) 
    db.add(db_project)
    db.commit()
    db.refresh(db_project)
    return db_project

@app.get("/projects", tags=["Projects"])
async def get_projects(db: Session = Depends(get_db)):
    """Devuelve una lista de todos los proyectos."""
    projects = db.query(models.Project).all()
    return projects

@app.get("/projects/{project_id}", tags=["Projects"])
async def get_project_details(project_id: int, db: Session = Depends(get_db)):
    """Devuelve los detalles de un proyecto específico."""
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    return db_project

@app.get("/jobs", tags=["Jobs"])
async def get_all_jobs(db: Session = Depends(get_db)):
    """Devuelve una lista de todos los trabajos de generación."""
    jobs = db.query(models.GenerationJob).all()
    return jobs

@app.get("/projects/{project_id}/jobs", tags=["Jobs"])
async def get_project_jobs(project_id: int, db: Session = Depends(get_db)):
    """Devuelve una lista de trabajos de generación para un proyecto específico."""
    jobs = db.query(models.GenerationJob).filter(models.GenerationJob.project_id == project_id).all()
    return jobs

@app.post("/jobs", tags=["Jobs"])
async def create_generation_job(
    template_id: int = Form(...),
    data_file: UploadFile = File(...),
    mappings_json: str = Form(...),
    num_rows_to_generate: int | None = Form(None),
    project_id: int = Form(...),
    db: Session = Depends(get_db)
):
    """
    Inicia un trabajo de generación de documentos asíncrono.
    Devuelve un ID de trabajo para el seguimiento.
    """
    db_template = db.query(models.Template).filter(models.Template.id == template_id).first()
    if not db_template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")

    try:
        mappings = json.loads(mappings_json)
        data_content = await data_file.read()
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="El formato de mappings_json no es válido.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error leyendo los archivos de entrada: {e}")

    # Lanzar la tarea de Celery en segundo plano
    task = generate_documents_task.delay(
        template_path=db_template.file_path,
        data_content_bytes=data_content,
        data_filename=data_file.filename,
        mappings=mappings,
        num_rows_to_generate=num_rows_to_generate
    )

    # Guardar el trabajo en la base de datos
    db_job = models.GenerationJob(id=task.id, status="PENDING", template_id=template_id, project_id=project_id)
    db.add(db_job)
    db.commit()
    db.refresh(db_job)

    return JSONResponse({"job_id": task.id}, status_code=202)

@app.get("/jobs/{job_id}", tags=["Jobs"])
async def get_job_status(job_id: str, db: Session = Depends(get_db)):
    """
    Consulta el estado de un trabajo de generación.
    """
    db_job = db.query(models.GenerationJob).filter(models.GenerationJob.id == job_id).first()
    if not db_job:
        raise HTTPException(status_code=404, detail="Trabajo no encontrado")

    response = {
        "job_id": db_job.id,
        "status": db_job.status,
        "result_url": f"/jobs/{job_id}/download" if db_job.status == "SUCCESS" else None
    }
    
    return JSONResponse(response)

@app.get("/jobs/{job_id}/download", tags=["Jobs"])
async def download_job_result(job_id: str, db: Session = Depends(get_db)):
    """
    Descarga el archivo ZIP generado por un trabajo completado.
    """
    db_job = db.query(models.GenerationJob).filter(models.GenerationJob.id == job_id).first()

    if not db_job or db_job.status != "SUCCESS":
        raise HTTPException(status_code=404, detail="El trabajo no se ha completado o ha fallado.")

    result_file_path = db_job.result_file_path
    if not result_file_path:
        raise HTTPException(status_code=404, detail="No se encontró el archivo de resultado.")

    # Determinar el tipo de archivo a partir de la extensión
    file_type = "zip" if result_file_path.endswith(".zip") else "docx"
    media_type = "application/zip" if file_type == 'zip' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    filename = f"documentos_{job_id}.zip" if file_type == 'zip' else f"documento_{job_id}.docx"

    return FileResponse(
        path=result_file_path,
        media_type=media_type,
        filename=filename
    )

@app.get("/", tags=["Root"])
async def read_root():
    """Endpoint raíz para verificar que la API está funcionando."""
    return {"message": "Bienvenido a la API de Generación de Documentos (v2 - Asíncrona)."}