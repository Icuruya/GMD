import io
import zipfile
import pandas as pd
from celery import Celery
from docx import Document
from typing import Dict, Any

from database import SessionLocal
import models

# --- Importaciones de la lógica de documentos de main.py ---
# En una aplicación más grande, esto estaría en su propio módulo de utilidades.
import re
PLACEHOLDER_REGEX = re.compile(r"\[(.*?)\]")

def replace_placeholders_in_doc(doc: Document, data: Dict[str, Any]):
    """Reemplaza los placeholders en un documento con los datos proporcionados, preservando el formato."""

    def _replace_in_element(element):
        # Si el elemento es un párrafo, procesarlo directamente
        if hasattr(element, 'runs'):
            full_text = element.text
            original_runs = list(element.runs)

            for key, value in data.items():
                full_text = full_text.replace(f"[{key}]", str(value))

            for run in original_runs:
                run.clear()

            if original_runs:
                new_run = element.add_run(full_text)
                first_run_format = original_runs[0].font
                new_run.font.bold = first_run_format.bold
                new_run.font.italic = first_run_format.italic
                new_run.font.underline = first_run_format.underline
                new_run.font.color.rgb = first_run_format.color.rgb
                new_run.font.size = first_run_format.size
                new_run.font.name = first_run_format.name
            else:
                element.add_run(full_text)
        # Si el elemento es una celda, iterar sobre sus párrafos
        elif hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                _replace_in_element(p)

    # Procesar párrafos
    for p in doc.paragraphs:
        _replace_in_element(p)

    # Procesar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_element(cell)

# --- Configuración de Celery ---
celery_app = Celery(
    "tasks",
    broker="redis://localhost:6379/0",
    backend="redis://localhost:6379/0"
)

celery_app.conf.update(
    task_track_started=True,
)

# --- Tarea Asíncrona de Generación de Documentos ---
@celery_app.task(bind=True)
def generate_documents_task(self, template_path: str, data_content_bytes: bytes, data_filename: str, mappings: dict, num_rows_to_generate: int | None = None):
    """
    Tarea de Celery que genera documentos en segundo plano.
    """
    db = SessionLocal()
    try:
        db_job = db.query(models.GenerationJob).filter(models.GenerationJob.id == self.request.id).first()
        if not db_job:
            return

        db_job.status = 'PROGRESS'
        db.commit()

        # --- 1. Leer el archivo de datos ---
        self.update_state(state='PROGRESS', meta={'status': 'Leyendo archivo de datos...'})
        if data_filename.endswith('.xlsx'):
            df = pd.read_excel(io.BytesIO(data_content_bytes))
        elif data_filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(data_content_bytes))
        else:
            raise ValueError("Formato de archivo de datos no soportado.")

        # --- 2. Generar Documentos y Comprimir en ZIP ---
        self.update_state(state='PROGRESS', meta={'status': 'Generando documentos...'})
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Limitar el DataFrame si num_rows_to_generate está especificado
            if num_rows_to_generate is not None and num_rows_to_generate > 0:
                df_to_process = df.head(num_rows_to_generate)
            else:
                df_to_process = df
            
            total_rows = len(df_to_process)
            for index, row in df_to_process.iterrows():
                doc = Document(template_path)
                
                data_for_row = {}
                for placeholder, column_name in mappings.items():
                    if column_name in df.columns:
                        data_for_row[placeholder] = row[column_name]
                
                replace_placeholders_in_doc(doc, data_for_row)
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                zip_file.writestr(f"documento_{index + 1}.docx", doc_buffer.getvalue())
                # Actualizar el progreso
                self.update_state(state='PROGRESS', meta={'status': f'Procesando fila {index + 1} de {total_rows}'})

        # Si solo se generó un documento, devolverlo directamente
        if total_rows == 1:
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            output_filename = f"documento_{self.request.id}.docx"
            with open(output_filename, "wb") as f:
                f.write(doc_buffer.getvalue())
            db_job.status = 'SUCCESS'
            db_job.result_file_path = output_filename
            db.commit()
            return {'status': 'Completed', 'result': output_filename, 'file_type': 'docx'}
        else:
            zip_buffer.seek(0)
            output_filename = f"generated_docs_{self.request.id}.zip"
            with open(output_filename, "wb") as f:
                f.write(zip_buffer.getvalue())
            db_job.status = 'SUCCESS'
            db_job.result_file_path = output_filename
            db.commit()
            return {'status': 'Completed', 'result': output_filename, 'file_type': 'zip'}

    except Exception as e:
        db_job = db.query(models.GenerationJob).filter(models.GenerationJob.id == self.request.id).first()
        if db_job:
            db_job.status = 'FAILURE'
            db.commit()
        # Capturar el tipo de excepción y el mensaje
        exc_type = type(e).__name__
        exc_message = str(e)
        
        # Actualizar el estado de la tarea con la información completa de la excepción
        self.update_state(
            state='FAILURE',
            meta={
                'status': 'Failed',
                'error': exc_message,
                'exc_type': exc_type
            }
        )
        # Devolver la información de la excepción para que Celery la procese correctamente
        raise e # Re-lanzar la excepción para que Celery la marque como FAILURE
    finally:
        db.close()