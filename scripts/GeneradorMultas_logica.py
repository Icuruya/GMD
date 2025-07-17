# -*- coding: utf-8 -*-
import os
import re
import locale
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import csv
from num2words import num2words
from collections import OrderedDict, defaultdict
import time
import traceback
from datetime import date, datetime
import math
from pathlib import Path
import tempfile
import shutil
import threading
import copy
import unicodedata
from openpyxl.styles import Font, Alignment
import sys # <--- ASEGÚRATE QUE ESTA LÍNEA ESTÉ PRESENTE Y NO COMENTADA
import sqlite3 # Aunque el log principal de escaneo cambia a Excel por PDF,
              # podríamos mantener SQLite para algún log interno de errores de escaneo si es útil,
              # o eliminarlo si ya no es necesario. Por ahora lo mantengo.


# --- CONSTANTES ESPECÍFICAS PARA EL REPORTE DE DESPACHOS ---
NOMBRE_CARPETA_DATOS_DESPACHO = "DATOS DESPACHO"
NOMBRE_ARCHIVO_RELACION_BASE = "RELACION PARA DESPACHOS.xlsx"
NOMBRE_CARPETA_REPORTES_GENERADOS = "REPORTES GENERADOS"
TEXTO_DESPACHO_FIJO = "MALDONADO GALLARDO Y ASOCIADOS, S.C."
COLUMNA_COMENTARIO_NOTIFICADO = "NOTIFICADO" # Valor exacto en la columna COMENTARIO del reporte de oficios
PREFIJO_OFICIO_PDF = "DIDCFMT" # Prefijo constante en los nombres de PDF de oficios

# --- NOMBRES DE COLUMNA PREDEFINIDOS ---
# Para el archivo "Reportes de OFICIOS" (seleccionado de DATOS DESPACHO)
COL_REPORTE_OFICIO_NUMERO = "OFICIO"
COL_REPORTE_OFICIO_FECHA_NOTIF = "FECHA NOTIFICACION"
COL_REPORTE_OFICIO_COMENTARIO = "COMENTARIO"
COL_REPORTE_OFICIO_CONTRIBUYENTE = "NOMBRE CONTRIBUYENTE"

# Para la BD Maestra de Multas (config_multas_actual.ruta_bd_multas)
COL_BD_MAESTRA_OFICIO = "OFICIO"
COL_BD_MAESTRA_PLACA = "PLACAS"
COL_BD_MAESTRA_MONTO = "IMPORTE"

# Configurar locale para español (para nombres de meses y días)
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8') # Alternativa para México
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252') # Para Windows
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale a español. Las fechas podrían aparecer en inglés.")



# --- Openpyxl para manejo de Excel ---
try:
    import openpyxl
    from openpyxl import load_workbook, Workbook
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.utils.exceptions import InvalidFileException # Para manejo de errores
    print("    - Módulo 'openpyxl' importado correctamente.")
except ImportError:
    print("¡ERROR FATAL! El módulo 'openpyxl' no está instalado. Es necesario para la Base de Datos Maestra.")
    print("Por favor, instálalo ejecutando: pip install openpyxl")
    # Considerar salir del script si openpyxl es crucial y no está.
    # sys.exit(1) # Descomentar para salida forzada

# --- OCR Dependencies ---
try:
    import pytesseract
    from PIL import Image
    print("    - Módulos 'pytesseract' y 'PIL' importados para OCR.")
    # --- CONFIGURACIÓN DE TESSERACT ---
    # !!! AJUSTA ESTA RUTA SI TESSERACT NO ESTÁ EN EL PATH !!!
    # Intentar detectar Tesseract en rutas comunes si no está configurado
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        # Añadir más rutas comunes si es necesario
    ]
    pytesseract.pytesseract.tesseract_cmd = '' # Iniciar vacío
    for tess_path in tesseract_paths:
        if Path(tess_path).is_file():
            pytesseract.pytesseract.tesseract_cmd = tess_path
            break
    
    if not pytesseract.pytesseract.tesseract_cmd:
        # Si no se encontró en rutas comunes, intentar obtenerlo del PATH (Tesseract puede manejar esto)
        # o dejar que falle para que el usuario sepa que debe configurarlo.
        print("    - (!) Advertencia: Tesseract OCR no encontrado en rutas comunes. Se intentará usar el PATH del sistema.")
        print("          Si falla, asegúrate de que Tesseract OCR esté instalado y la ruta sea correcta o esté en el PATH.")
        # No se asigna `pytesseract.pytesseract.tesseract_cmd` si no se encontró,
        # para que pytesseract intente encontrarlo por sí mismo. Si no lo hace, lanzará error.

    try:
        tess_version = pytesseract.get_tesseract_version()
        print(f"    - Tesseract OCR detectado (Versión: {tess_version}).")
        if pytesseract.pytesseract.tesseract_cmd:
             print(f"      Ruta configurada/detectada: {pytesseract.pytesseract.tesseract_cmd}")
    except pytesseract.TesseractNotFoundError:
        print("    - (!) ERROR: 'tesseract.exe' no encontrado en rutas comunes ni en el PATH del sistema.")
        print("      Asegúrate de que Tesseract OCR esté instalado y la ruta en el script sea correcta o esté en el PATH.")
    except Exception as e_tess_check:
        print(f"    - (!) Advertencia: Ocurrió un error al verificar la versión de Tesseract: {e_tess_check}")
except ImportError:
    print("¡ERROR FATAL! No se encontraron 'pytesseract' o 'Pillow'. Necesarios para OCR.")
    # Mock objects para permitir que el script se cargue parcialmente sin OCR
    class pytesseract:
        @staticmethod
        def image_to_string(img, lang, config): raise ImportError("pytesseract no instalado.")
        @staticmethod
        def get_tesseract_version(): raise ImportError("pytesseract no instalado.")
        class pytesseract: tesseract_cmd = ""
    class Image:
        @staticmethod
        def frombytes(mode, size, data): raise ImportError("Pillow no instalado.")
        def crop(self, area): raise ImportError("Pillow no instalado.")
        @property
        def size(self): return (0,0)

# --- docx2pdf ---
try:
    from docx2pdf import convert
    print("    - Módulo 'docx2pdf' importado correctamente.")
except ImportError:
    print("¡ADVERTENCIA! El módulo 'docx2pdf' no está instalado o no se pudo importar.")
    def convert(input_path, output_path):
        print(f"ERROR: docx2pdf no disponible. No se pudo convertir {input_path} a PDF.")
        raise ImportError("docx2pdf no está instalado o Microsoft Word no está disponible.")

# EXPLICACIÓN: Se elimina la importación problemática y se define la variable de error
# directamente con el respaldo. Esto elimina la advertencia del editor.
Docx2PdfFileNotFoundError = FileNotFoundError

# --- PyPDF2 ---
try:
    from PyPDF2 import PdfReader as PyPDF2Reader
    from PyPDF2 import PdfWriter as PyPDF2Writer
    from PyPDF2.errors import PdfReadError as PyPDF2ReadError
    print("    - Módulos 'PyPDF2Reader', 'PyPDF2Writer' y 'PyPDF2ReadError' importados.")
except ImportError:
    print("¡ADVERTENCIA! PyPDF2 no encontrado. Necesario para extraer páginas y fusionar.")
    class PyPDF2Reader:
        def __init__(self, stream): raise ImportError("PyPDF2 no instalado.")
        @property
        def pages(self): return []
    class PyPDF2Writer:
        def add_page(self, page): raise ImportError("PyPDF2 no instalado.")
        def write(self, stream): raise ImportError("PyPDF2 no instalado.")
        @property
        def pages(self): return []
    class PyPDF2ReadError(Exception): pass

# --- PyMuPDF (fitz) ---
try:
    import fitz # PyMuPDF
    print("    - Módulo 'fitz' (PyMuPDF) importado correctamente para Escaneo.")
except ImportError:
    print("¡ERROR FATAL! PyMuPDF (fitz) no encontrado. Necesario para el modo Escaneo.")
    class fitz:
        class Document:
            def __init__(self): raise ImportError("PyMuPDF (fitz) no está instalado.")
            def close(self): pass
            def load_page(self, i): raise ImportError("PyPDF2 no instalado.")
            @property
            def page_count(self): return 0
        @staticmethod
        def open(path): raise ImportError("PyMuPDF (fitz) no está instalado.")
        class Matrix:
            def __init__(self, zoom_x, zoom_y): pass

# --- Constantes Globales ---
RUTA_BASE_SCRIPT = Path(__file__).parent





# --- Columnas y Tipos para la Base de Datos Maestra de Multas ---
# (Columnas en el orden deseado para el Excel)
COL_OFICIO_MULTAS = "OFICIO"
COL_NOMBRE_CONTRIBUYENTE_MULTAS = "NOMBRE CONTRIBUYENTE"
COL_DIRECCION_MULTAS = "DIRECCION"
COL_CP_MULTAS = "CP" # Nueva columna para CP directamente en la BD Maestra
COL_MONTO_MULTAS = "MONTO"
COL_ESTADO_MULTAS = "ESTADO"
COL_PRIORIDAD_MULTAS = "PRIORIDAD" # Para ser usado por impresion.py
COL_HOJAS_DOC_MULTAS = "HOJAS POR DOCUMENTO"
COL_FECHA_IMPRESION_MULTAS = "FECHA IMPRESION" # Para ser usado por impresion.py
COL_BASE_DATOS_ESCANEO_MULTAS = "BASE DE DATOS DE ESCANEO" # Nombre del Excel de escaneo
# Columnas adicionales que podrían ser útiles (opcional, añadir si es necesario)
COL_ID_LOTE_MULTAS = "ID Lote" # Para ser usado por impresion.py
COL_IMPRESORA_MULTAS = "Impresora" # Para ser usado por impresion.py
COL_RUTA_PDF_GENERADO_MULTAS = "Ruta PDF Generado" # Ruta al PDF generado por el modo normal

COLUMNAS_BD_MAESTRA_MULTAS = [
    COL_OFICIO_MULTAS, COL_NOMBRE_CONTRIBUYENTE_MULTAS, COL_DIRECCION_MULTAS, COL_CP_MULTAS,
    COL_MONTO_MULTAS, COL_ESTADO_MULTAS, COL_PRIORIDAD_MULTAS, COL_HOJAS_DOC_MULTAS,
    COL_FECHA_IMPRESION_MULTAS, COL_BASE_DATOS_ESCANEO_MULTAS, COL_ID_LOTE_MULTAS,
    COL_IMPRESORA_MULTAS, COL_RUTA_PDF_GENERADO_MULTAS, "MOVIMIENTO"
]

TIPOS_BD_MAESTRA_MULTAS = {
    COL_OFICIO_MULTAS: 'str',
    COL_NOMBRE_CONTRIBUYENTE_MULTAS: 'str',
    COL_DIRECCION_MULTAS: 'str',
    COL_CP_MULTAS: 'str',
    COL_MONTO_MULTAS: 'str', # Almacenar como texto formateado (ej. "$ 1,234.56") o como float y formatear al leer/mostrar
    COL_ESTADO_MULTAS: 'str',
    COL_PRIORIDAD_MULTAS: 'Int64', # Permite NA
    COL_HOJAS_DOC_MULTAS: 'Int64', # Permite NA
    COL_FECHA_IMPRESION_MULTAS: 'str', # O datetime si se prefiere, pero str es más simple para empezar
    COL_BASE_DATOS_ESCANEO_MULTAS: 'str',
    COL_ID_LOTE_MULTAS: 'str',
    COL_IMPRESORA_MULTAS: 'str',
    COL_RUTA_PDF_GENERADO_MULTAS: 'str',
    "MOVIMIENTO": 'str'
}

# --- Estados del Proceso (Sincronizados con Predial y futura impresion_multas.py) ---
ESTADO_PENDIENTE_MULTAS = "PENDIENTE"
ESTADO_GEN_COMPLETO_MULTAS = "GENERADO COMPLETO"
ESTADO_GEN_ULTIMA_MULTAS = "GENERADO ULTIMA HOJA"
ESTADO_IMP_ULTIMA_MULTAS = "ULTIMA PAG IMPRESA"       # Establecido por impresion_multas.py
ESTADO_IMP_COMPLETO_MULTAS = "IMPRESION COMPLETADA"  # Establecido por impresion_multas.py
ESTADO_ERROR_GENERACION_MULTAS = "ERROR GENERACION"
ESTADO_PDF_NO_ENCONTRADO_MULTAS = "PDF NO ENCONTRADO" # Si el PDF generado no se encuentra al momento de imprimir
ESTADO_GEN_RESTO_MULTAS = "GENERADO RESTO (SIN ULTIMA)"


# --- Modos de Generación (Internos del Script) ---
MODO_GEN_COMPLETO = "MODO_COMPLETO"
MODO_GEN_ULTIMA_PAG = "MODO_ULTIMA_PAG"
MODO_GEN_RESTO_DOC = "MODO_RESTO_DOC" # Para generar completo basado en ESTADO_IMP_ULTIMA_MULTAS



# --- Configuración para Escaneo ---
SOURCE_PDF_FOLDER_SCAN_MULTAS = RUTA_BASE_SCRIPT / "PDFs_A_Escanear_Multas" # Carpeta de entrada para PDFs a escanear
# La carpeta de salida para los escaneos es SUB_DOCUMENTOS_ESCANEADOS_MULTAS

# Columnas para los Excel individuales de escaneo (BD_BlinderX.xlsx)
COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL = ["Oficio", "NombreArchivo", "Paginas", "Direccion", "CP", "Monto", "Error"]
# "Error" para indicar si el oficio no se encontró en la BD Maestra.

# --- Patrón para buscar el número de OFICIO (Usado en Escaneo) ---
OFICIO_REGEX_SCAN_MULTAS = re.compile(r"\bDIDCFMT(\d+)\b", re.IGNORECASE) # Mismo que el original

# --- Constantes de Formato (Generación - similar al original) ---
FONT_NAME_DEFAULT_MULTAS = "Gabarito" # O la fuente que uses
FONT_SIZE_DEFAULT_MULTAS = Pt(7)
FONT_NAME_TABLE_MULTAS = "Gabarito"
FONT_SIZE_TABLE_MULTAS = Pt(3.5)

# --- Diccionario Global de Errores (para registrar fallos durante la ejecución) ---
failed_operations_log = defaultdict(list) # { "categoria_error": ["detalle1", "detalle2"] }

# --- Lock para acceso concurrente a la BD Maestra (si se usaran hilos en el futuro) ---
excel_lock_multas = threading.Lock()


# --- Funciones de Utilidad ---

def limpiar_texto(texto_original):
    """
    Limpia un texto para ser usado en nombres de archivo o identificadores.
    Quita acentos, reemplaza caracteres no alfanuméricos (excepto guiones y espacios)
    por guion bajo, y luego reemplaza espacios por guion bajo.
    """
    if texto_original is None or pd.isna(texto_original):
        return "VACIO"
    
    texto_str = str(texto_original)
    
    # Quitar acentos
    try:
        nfkd_form = unicodedata.normalize('NFKD', texto_str)
        texto_sin_acentos = "".join([c for c in nfkd_form if not unicodedata.combining(c)])
    except NameError: # Fallback si unicodedata no está disponible
        texto_sin_acentos = texto_str
        print("    - (!) Advertencia: Módulo 'unicodedata' no disponible. La limpieza de acentos puede no ser completa.")

    # Reemplazar caracteres no alfanuméricos (excepto guiones) por guion bajo
    texto_limpio_paso1 = re.sub(r'[^\w\s\-]', '_', texto_sin_acentos, flags=re.UNICODE)
    # Reemplazar múltiples espacios o guiones bajos por uno solo y quitar los de los extremos
    texto_limpio_paso2 = re.sub(r'[\s_]+', '_', texto_limpio_paso1).strip('_')
    
    return texto_limpio_paso2 if texto_limpio_paso2 else "VACIO"

def obtener_siguiente_letra_lote(letra_actual):
    """
    Obtiene la siguiente letra del alfabeto para el lote.
    Ej: 'A' -> 'B', 'Z' -> 'AA', 'AZ' -> 'BA'.
    """
    if not letra_actual or not letra_actual.isalpha():
        return 'A'
    
    letra_actual = letra_actual.upper()
    if letra_actual == 'Z' * len(letra_actual):
        return 'A' * (len(letra_actual) + 1)
    
    letras = list(letra_actual)
    i = len(letras) - 1
    while i >= 0:
        if letras[i] == 'Z':
            letras[i] = 'A'
            i -= 1
        else:
            letras[i] = chr(ord(letras[i]) + 1)
            break
    return "".join(letras)

def obtener_ultimo_id_de_lote_especifico(df_bd_maestra, letra_lote):
    """
    Encuentra el número más alto usado en un lote específico (ej. 'A').
    Devuelve el último número encontrado, o 0 si el lote está vacío.
    """
    if df_bd_maestra.empty or "ID" not in df_bd_maestra.columns:
        return 0

    # Filtra el DataFrame para obtener solo los registros del lote deseado
    df_lote = df_bd_maestra[df_bd_maestra['ID'].str.startswith(f"{letra_lote}-", na=False)].copy()

    if df_lote.empty:
        return 0 # No hay registros en este lote, empezamos desde cero

    # Extrae la parte numérica, convierte a número y encuentra el máximo
    df_lote['numero_id'] = pd.to_numeric(df_lote['ID'].str.split('-').str[1], errors='coerce')

    ultimo_numero = df_lote['numero_id'].max()

    return 0 if pd.isna(ultimo_numero) else int(ultimo_numero)

def formatear_fecha_reporte(fecha_obj):
    """Formatea la fecha como 'DÍA DD de MES AÑO' en mayúsculas. EJ: MIERCOLES 04 de JUNIO 2025"""
    if not isinstance(fecha_obj, datetime):
        return "FECHA INVÁLIDA"
    
    # Intentar con el locale configurado
    try:
        # Primero, intentar obtener la fecha formateada con el locale que se intentó configurar
        fecha_formateada = fecha_obj.strftime("%A %d de %B de %Y").upper()
        # Comprobar si hay caracteres mal codificados comunes (ej. Ã© por é)
        # Esta es una heurística simple; puede no cubrir todos los casos.
        if 'Ã©' in fecha_formateada or 'Ã³' in fecha_formateada or 'Ã¡' in fecha_formateada or 'Ã±' in fecha_formateada or 'Ã' in fecha_formateada :
            # Si se detecta posible mala codificación, intentar reconstruir con nombres fijos
            # Esto asume que el problema es con los nombres de día/mes del locale.
            raise ValueError("Posible mala codificación detectada en strftime.")
        return fecha_formateada
    except Exception as e_strftime: # Incluye ValueError de la comprobación anterior
        # print(f"  Advertencia: strftime con locale falló o detectó mala codificación ({e_strftime}). Usando fallback manual.")
        # Fallback manual para evitar problemas de codificación de locale en algunos sistemas
        dias_es = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
        meses_es = ["ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO", "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]
        
        try:
            dia_semana = dias_es[fecha_obj.weekday()]
            dia_mes = fecha_obj.day
            mes_anio = meses_es[fecha_obj.month - 1]
            anio = fecha_obj.year
            return f"{dia_semana} {dia_mes:02d} DE {mes_anio} {anio}" # Formato DD DE MES YYYY
        except Exception as e_manual:
            print(f"  Error crítico formateando fecha manualmente: {e_manual}")
            return "ERROR FORMATO FECHA MANUAL"

# Añade esta nueva función en GeneradorMultas_logica.py, cerca de normalizar_oficio_multas
def convertir_a_formato_con_barras(oficio_str_normalizado):
    """
    Convierte un oficio en formato normalizado (ej. DIDCFMT12345)
    al formato con barras (ej. DI/DCF/MT/12345).
    Si el formato de entrada no es el esperado, devuelve el string original.
    """
    if isinstance(oficio_str_normalizado, str) and oficio_str_normalizado.startswith("DIDCFMT") and oficio_str_normalizado[7:].isdigit():
        return f"DI/DCF/MT/{oficio_str_normalizado[7:]}"
    # Si ya tiene barras o es un formato no reconocido, devolverlo tal cual.
    # Esto es para evitar doble conversión si accidentalmente se le pasa un formato con barras.
    if isinstance(oficio_str_normalizado, str) and re.match(r'DI\s*/\s*DCF\s*/\s*MT\s*/\s*\d+', oficio_str_normalizado):
         return oficio_str_normalizado # Ya está en formato con barras
    return oficio_str_normalizado # Devolver original si no coincide con DIDCFMT...

def extraer_fecha_oficio_de_pdf(ruta_pdf_obj):
    """
    Extrae la fecha de un texto como 'Monterrey, Nuevo León, DD de MMMM de AAAA.'
    Devuelve 'DD de MES_MAYUSCULAS AAAA' (ej. '30 de MAYO 2025') o None si no se encuentra.
    Intenta ser más robusto examinando bloques de texto.
    """
    if not ruta_pdf_obj or not ruta_pdf_obj.exists():
        print(f"    - Advertencia (extracción fecha): PDF no encontrado en {ruta_pdf_obj.name if ruta_pdf_obj else 'Ruta Nula'}")
        return None
    
    # Patrón mejorado para más flexibilidad con espacios y comas opcionales al final
    # (?:...) es un grupo no capturante
    # \s* maneja cero o más espacios
    # [,\.]? maneja una coma o punto opcional
    patron_fecha = re.compile(
        r"(?:Monterrey,\s*Nuevo\s*León,?\s*)?"  # "Monterrey, Nuevo León," opcional, con coma opcional al final
        r"(\d{1,2}\s+de\s+[a-zA-ZáéíóúÁÉÍÓÚñÑ]+\s+de\s+\d{4})" # "DD de MES de AAAA" (MES con acentos y ñ)
        r"[\.,]?",  # Punto o coma opcional al final de la fecha
        re.IGNORECASE
    )

    try:
        with fitz.open(ruta_pdf_obj) as doc:
            texto_concatenado_primeras_paginas = ""
            for page_num in range(min(1, doc.page_count)): # Enfocarse principalmente en la primera página
                page = doc.load_page(page_num)
                
                # Opción 1: Intentar con bloques de texto
                blocks = page.get_text("blocks") # Obtiene (x0, y0, x1, y1, texto, block_no, block_type)
                for b_idx, block in enumerate(blocks):
                    texto_bloque = block[4].replace("\n", " ").strip() # texto del bloque, reemplazar saltos de línea
                    # print(f"      DEBUG PDF Texto Bloque {b_idx}: '{texto_bloque}'") # Para depuración intensa
                    coincidencia = patron_fecha.search(texto_bloque)
                    if coincidencia:
                        fecha_extraida_str = coincidencia.group(1) # El grupo capturado
                        partes = fecha_extraida_str.lower().split(" de ")
                        if len(partes) == 3:
                            # print(f"    + Fecha encontrada en bloque (Oficio: {ruta_pdf_obj.stem}): {fecha_extraida_str}")
                            return f"{partes[0].strip()} de {partes[1].strip().upper()} {partes[2].strip()}"
                        else: # Debería tener 3 partes
                            # print(f"    - Formato de fecha inesperado en partes (Oficio: {ruta_pdf_obj.stem}): {partes}")
                            return fecha_extraida_str.upper() # Fallback
                
                # Opción 2: Si no se encontró en bloques, intentar con el texto completo de la página
                # Esto es un fallback por si la fecha está dividida entre bloques o el regex necesita más contexto.
                if page_num == 0: # Solo para la primera página para este fallback
                    texto_completo_pagina = page.get_text("text").replace("\n", " ")
                    texto_concatenado_primeras_paginas += texto_completo_pagina + " "

            # Intentar con el texto concatenado de las primeras páginas si no se encontró antes
            if texto_concatenado_primeras_paginas:
                # print(f"      DEBUG PDF Texto Concatenado: '{texto_concatenado_primeras_paginas[:500]}...'") # Para depuración
                coincidencia_concat = patron_fecha.search(texto_concatenado_primeras_paginas)
                if coincidencia_concat:
                    fecha_extraida_concat_str = coincidencia_concat.group(1)
                    partes_concat = fecha_extraida_concat_str.lower().split(" de ")
                    if len(partes_concat) == 3:
                        # print(f"    + Fecha encontrada en texto concatenado (Oficio: {ruta_pdf_obj.stem}): {fecha_extraida_concat_str}")
                        return f"{partes_concat[0].strip()} de {partes_concat[1].strip().upper()} {partes_concat[2].strip()}"
                    else:
                        return fecha_extraida_concat_str.upper() # Fallback

            # print(f"    - Patrón de fecha no encontrado en {ruta_pdf_obj.name} después de revisar bloques y texto concatenado.")
            return None
            
    except ImportError: # Si fitz no estuviera disponible (aunque ya lo comprobamos antes)
        print(f"    - Error Crítico (extracción fecha): Módulo fitz (PyMuPDF) no disponible para {ruta_pdf_obj.name}.")
        return None
    except Exception as e:
        print(f"    - Error procesando PDF {ruta_pdf_obj.name} para extraer fecha: {type(e).__name__} - {e}")
        # traceback.print_exc(limit=2) # Descomentar para depuración más profunda
        return None

def buscar_pdf_oficio(numero_oficio_str, directorio_base_pdfs_obj, nombre_contribuyente_opcional=None):
    """
    Busca un archivo PDF para un número de oficio dado dentro de un directorio base
    que puede contener subcarpetas (como las de CP), EXCLUYENDO carpetas que contengan 'ESCANEADOS'.
    El nombre del PDF se espera como PREFIJO_OFICIO_PDF + numero_oficio_str + "_" + nombre_contribuyente + ".pdf"
    o variantes que contengan PREFIJO_OFICIO_PDF + numero_oficio_str.
    """
    if not directorio_base_pdfs_obj or not directorio_base_pdfs_obj.is_dir():
        print(f"    - Advertencia (búsqueda PDF): Directorio base de PDFs no válido o no encontrado: {directorio_base_pdfs_obj}")
        return None

    oficio_numerico_para_patron = numero_oficio_str.upper().replace(PREFIJO_OFICIO_PDF, "").strip()

    patron_con_prefijo_y_numero = f"{PREFIJO_OFICIO_PDF}{oficio_numerico_para_patron}*.pdf"
    patron_general_con_prefijo = f"*{PREFIJO_OFICIO_PDF}{oficio_numerico_para_patron}*.pdf"
    patron_solo_numero_en_nombre = f"*{oficio_numerico_para_patron}*.pdf"

    archivos_encontrados_validos = []
    carpetas_a_ignorar_keywords = ["ESCANEADOS", "ESCANEADAS", "SCAN", "RESULTADOS_SCAN"] # Añadido "RESULTADOS_SCAN" por si acaso

    patrones_de_busqueda_ordenados = []

    if nombre_contribuyente_opcional:
        nombre_contrib_limpio = "".join(filter(str.isalnum, nombre_contribuyente_opcional)).upper()
        if nombre_contrib_limpio:
            patron_oficiocompleto_nombre = f"{PREFIJO_OFICIO_PDF}{oficio_numerico_para_patron}_{nombre_contrib_limpio}*.pdf"
            patrones_de_busqueda_ordenados.append(patron_oficiocompleto_nombre)
            patron_oficiocompleto_flexible_nombre = f"{PREFIJO_OFICIO_PDF}{oficio_numerico_para_patron}*_{nombre_contrib_limpio}*.pdf"
            patrones_de_busqueda_ordenados.append(patron_oficiocompleto_flexible_nombre)

    patrones_de_busqueda_ordenados.extend([
        patron_con_prefijo_y_numero,
        patron_general_con_prefijo,
        patron_solo_numero_en_nombre
    ])
    
    patrones_de_busqueda_ordenados_unicos = []
    for p in patrones_de_busqueda_ordenados:
        if p not in patrones_de_busqueda_ordenados_unicos:
            patrones_de_busqueda_ordenados_unicos.append(p)
    
    for idx_patron, patron_actual in enumerate(patrones_de_busqueda_ordenados_unicos):
        if not patron_actual: continue

        for pdf_path in directorio_base_pdfs_obj.rglob(patron_actual):
            ignorar_este_pdf = False
            try:
                # --- CORRECCIÓN: Definir partes_ruta_relativa_actual aquí ---
                partes_ruta_relativa_actual = [part.upper() for part in pdf_path.relative_to(directorio_base_pdfs_obj).parts]
                # --- FIN CORRECCIÓN ---
                for keyword_ignorar in carpetas_a_ignorar_keywords:
                    # Comprobar si alguna parte de la ruta contiene la palabra clave a ignorar
                    if any(keyword_ignorar in parte_ruta_upper for parte_ruta_upper in partes_ruta_relativa_actual):
                        ignorar_este_pdf = True
                        break 
            except ValueError: 
                pass 

            if ignorar_este_pdf:
                continue 
            
            if pdf_path not in archivos_encontrados_validos:
                 archivos_encontrados_validos.append(pdf_path)
        
        if archivos_encontrados_validos and nombre_contribuyente_opcional and idx_patron < 2 : 
            break
            
    if not archivos_encontrados_validos:
        return None

    if len(archivos_encontrados_validos) == 1:
        return archivos_encontrados_validos[0]
    else: 
        pdfs_en_carpeta_cp = [
            p for p in archivos_encontrados_validos 
            if len(p.relative_to(directorio_base_pdfs_obj).parts) > 1 and 
               p.relative_to(directorio_base_pdfs_obj).parts[0].upper() == "CP"
        ]
        
        if pdfs_en_carpeta_cp:
            if len(pdfs_en_carpeta_cp) == 1:
                return pdfs_en_carpeta_cp[0]
            else: 
                print(f"    - Advertencia (búsqueda PDF): Múltiples PDFs válidos encontrados para oficio {numero_oficio_str}, incluso en carpetas CP. Usando el primero de CP: {pdfs_en_carpeta_cp[0].name}")
                for f_idx, f_path_valido in enumerate(pdfs_en_carpeta_cp): 
                    print(f"        (CP) {f_idx+1}. {f_path_valido.relative_to(directorio_base_pdfs_obj)}")
                return pdfs_en_carpeta_cp[0]
        else: 
            print(f"    - Advertencia (búsqueda PDF): Múltiples PDFs válidos encontrados para oficio {numero_oficio_str} (ninguno priorizado en carpeta CP). Usando el primero general: {archivos_encontrados_validos[0].name}")
            for f_idx, f_path_valido in enumerate(archivos_encontrados_validos): 
                print(f"        {f_idx+1}. {f_path_valido.relative_to(directorio_base_pdfs_obj)}")
            return archivos_encontrados_validos[0]


def obtener_datos_placas_montos_de_bd_multas(oficio_id_reporte, df_bd_multas_maestro):
    """
    Obtiene datos de placas y montos del DataFrame de la BD de multas maestra.
    Usa las constantes COL_BD_MAESTRA_OFICIO, COL_BD_MAESTRA_PLACA, COL_BD_MAESTRA_MONTO.
    Devuelve una lista de diccionarios: [{'PLACA': 'ABC', 'MONTO': 100.0}, ...]
    """
    if df_bd_multas_maestro.empty:
        print(f"    - Advertencia (BD Maestra): DataFrame de BD multas maestro está vacío. No se pueden obtener placas/montos para {oficio_id_reporte}.")
        return [{'PLACA': 'BD MAESTRA VACÍA', 'MONTO': 0.00}]

    try:
        oficio_id_reporte_normalizado = normalizar_oficio_multas(str(oficio_id_reporte)) # Normaliza el ID del reporte

        # Filtrar por el oficio_id_reporte (comparando formas normalizadas)
        data_oficio_en_bd = df_bd_multas_maestro[ # df_bd_multas_maestro es en realidad df_base_de_datos_original
            df_bd_multas_maestro[COL_BD_MAESTRA_OFICIO].apply(normalizar_oficio_multas) == oficio_id_reporte_normalizado
        ]
    

        if data_oficio_en_bd.empty:
            # print(f"    - Info (BD Maestra): Oficio {oficio_id_reporte} no encontrado en la BD Maestra (col: '{COL_BD_MAESTRA_OFICIO}').")
            return [{'PLACA': 'OFICIO NO EN BD MAESTRA', 'MONTO': 0.00}]
        
        resultados = []
        for _, row_bd in data_oficio_en_bd.iterrows():
            placa_bd = row_bd[COL_BD_MAESTRA_PLACA]
            monto_bd = row_bd[COL_BD_MAESTRA_MONTO]
            # Asegurar que el monto sea numérico, si no, intentar convertir o poner 0.0
            if not isinstance(monto_bd, (int, float)):
                try:
                    monto_bd = float(str(monto_bd).replace("$","").replace(",","").strip())
                except ValueError:
                    print(f"    - Advertencia (BD Maestra): Monto '{monto_bd}' para oficio {oficio_id_reporte}, placa {placa_bd} no es numérico. Usando 0.00.")
                    monto_bd = 0.00
            
            resultados.append({'PLACA': str(placa_bd).strip(), 'MONTO': monto_bd})
        
        if not resultados:
             print(f"    - Advertencia (BD Maestra): Oficio {oficio_id_reporte} encontrado, pero sin datos de placa/monto (cols: '{COL_BD_MAESTRA_PLACA}', '{COL_BD_MAESTRA_MONTO}').")
             return [{'PLACA': 'DATOS FALTANTES EN BD', 'MONTO': 0.00}]

        return resultados

    except KeyError as ke:
        print(f"    - Error GRAVE (BD Maestra): Columna '{ke}' no existe en la BD Maestra de multas. Verifica las constantes COL_BD_MAESTRA_*.")
        return [{'PLACA': 'ERROR COLUMNA BD MAESTRA', 'MONTO': 0.00}]
    except Exception as e:
        print(f"    - Error inesperado al obtener placas/montos de BD Maestra para {oficio_id_reporte}: {e}")
        return [{'PLACA': 'ERROR INESPERADO BD MAESTRA', 'MONTO': 0.00}]

def configurar_locale_es():
    """Intenta configurar el locale a español para formateo de fechas y moneda."""
    locales_a_intentar = ['es_MX.UTF-8', 'es-MX', 'es_MX', 'Spanish_Mexico.1252', 
                          'es_ES.UTF-8', 'es-ES', 'es_ES', 'Spanish_Spain.1252', 
                          'Spanish', 'es', ''] # '' para el default del sistema
    locale_configurado_exitoso = False
    for loc_str in locales_a_intentar:
        try:
            locale.setlocale(locale.LC_ALL, loc_str)
            # Verificar si realmente se estableció a algo parecido a español
            test_date = date(2024, 1, 1) # Enero
            nombre_mes_local = test_date.strftime('%B').lower()
            if "enero" in nombre_mes_local or "january" in nombre_mes_local: # Aceptar inglés como fallback si es el default
                print(f"    - Locale configurado a: {locale.getlocale()} (Mes de prueba: {nombre_mes_local})")
                locale_configurado_exitoso = True
                return True
        except locale.Error:
            continue # Probar el siguiente
        except Exception as e_loc:
            print(f"    - (!) Advertencia configurando locale '{loc_str}': {e_loc}")
            continue
            
    if not locale_configurado_exitoso:
        print("    - (!) ADVERTENCIA CRÍTICA: No se pudo configurar un locale en español o inglés.")
        print("      El formato de fechas y moneda podría ser incorrecto.")
        print("      Se usará el formato por defecto del sistema.")
        try:
            locale.setlocale(locale.LC_ALL, '') # Reset al default del sistema
            print(f"    - Locale reseteado a default del sistema: {locale.getlocale()}")
            return True # Permitir continuar con el default
        except Exception as e_reset_loc:
            print(f"    - (!) Error reseteando locale a default: {e_reset_loc}")
            return False # Falla crítica si ni el default funciona
    return False # No debería llegar aquí si uno de los locales funcionó o el default

def numero_a_texto_moneda_mx(numero):
    """Convierte un número a su representación en texto como moneda mexicana."""
    global failed_operations_log
    try:
        if numero is None or pd.isna(numero):
            numero = 0.0
        
        numero_f = float(str(numero).replace("$", "").replace(",", "").strip())
        
        parte_entera = int(numero_f)
        parte_decimal = int(round((numero_f - parte_entera) * 100))

        if parte_decimal >= 100: # Ajuste por redondeo
            parte_entera += 1
            parte_decimal = 0
            
        texto_entero_min = num2words(parte_entera, lang='es')
        texto_entero_capitalizado = texto_entero_min.capitalize()
        
        texto_moneda = f"({texto_entero_capitalizado} pesos {parte_decimal:02d}/100 M.N.)"
        return texto_moneda
    except Exception as e:
        print(f"    - (!) Error convirtiendo número '{numero}' a texto moneda: {e}")
        failed_operations_log["conversion_numero_texto"].append({'numero': numero, 'error': str(e)})
        return "(Error en conversión a texto)"

def formatear_valor_celda(valor, es_moneda=True):
    """Formatea un valor numérico para mostrarlo en celdas, con o sin símbolo de moneda."""
    try:
        if valor is None or pd.isna(valor):
            valor_limpio_str = "0"
        else:
            valor_limpio_str = str(valor).replace("$", "").replace(",", "").strip()

        if not valor_limpio_str:
            return "$ 0.00" if es_moneda else "0.00"
            
        numero = float(valor_limpio_str)
        formato_locale = locale.getlocale(locale.LC_NUMERIC)
        # Usar locale.format_string para separadores de miles correctos si el locale está bien configurado
        if formato_locale != (None, None) and 'es' in (formato_locale[0] or '').lower() : # Si es un locale español
            return f"${locale.format_string('%.2f', numero, grouping=True)}" if es_moneda else locale.format_string('%.2f', numero, grouping=True)
        else: # Fallback a formato simple si el locale no es español o no está bien configurado
            return f"${numero:,.2f}" if es_moneda else f"{numero:.2f}"
    except (ValueError, TypeError):
        return "$ 0.00" if es_moneda else "0.00"

def formatear_fecha_corta(valor_fecha):
    """Formatea una fecha a DD/MM/AAAA. Devuelve string vacío si la fecha es inválida."""
    if valor_fecha is None or pd.isna(valor_fecha) or str(valor_fecha).strip() == "":
        return ""
    try:
        # Intentar convertir a datetime, probando diferentes formatos si es necesario
        dt_obj = pd.to_datetime(valor_fecha, errors='coerce')
        if pd.isna(dt_obj): # Si falla el primer intento, probar con dayfirst=True
            dt_obj = pd.to_datetime(valor_fecha, dayfirst=True, errors='coerce')
        
        return dt_obj.strftime("%d/%m/%Y") if not pd.isna(dt_obj) else str(valor_fecha)
    except Exception:
        return str(valor_fecha) # Devolver el valor original como string si todo falla

def extraer_cp_de_direccion(direccion_texto):
    """Extrae un código postal de 5 dígitos del final de un texto de dirección."""
    if direccion_texto is None or pd.isna(direccion_texto) or not isinstance(direccion_texto, str) or not direccion_texto.strip():
        return "SIN_CP"
    match = re.search(r'\b(\d{5})\b$', direccion_texto.strip())
    return match.group(1) if match else "SIN_CP"

def crear_hoja_reporte_impresion(writer, datos_para_reporte, titulo_principal, cabecera_columna):
    """
    Crea una segunda hoja en un archivo Excel con un formato de impresión específico,
    llenando las columnas de forma vertical.
    """
    try:
        workbook = writer.book
        sheet_name = "Reporte para Impresión"
        if sheet_name in workbook.sheetnames:
            del workbook[sheet_name]
        
        ws = workbook.create_sheet(sheet_name)

        # --- 1. Configuración de Estilo y Altura de Fila ---
        font_style = Font(name='Calibri', size=26)
        row_height = 29

        # --- 2. Lógica de Relleno por Páginas ---
        max_filas_por_bloque = 43
        columnas_bloques = [('A', 'B'), ('D', 'E'), ('G', 'H')]
        puntero_datos = 0
        num_datos = len(datos_para_reporte)
        fila_actual_excel = 1

        while puntero_datos < num_datos:
            # --- Iniciar una nueva "página" ---
            # Dejar una fila en blanco si no es la primera página
            if fila_actual_excel > 1:
                fila_actual_excel += 1

            # Título Principal de la página
            ws.merge_cells(f'A{fila_actual_excel}:H{fila_actual_excel}')
            titulo_cell = ws[f'A{fila_actual_excel}']
            
            if titulo_principal == "PREDIAL":
                titulo_cell.value = "PREDIAL - DESPACHO MALDONADO"
            else: # Asumir MULTAS u otro
                titulo_cell.value = "MULTAS DE TRANSITO - DESPACHO MALDONADO"
            
            titulo_cell.font = font_style
            titulo_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[fila_actual_excel].height = row_height
            fila_actual_excel += 1

            # Cabeceras de las columnas para esta página
            fila_inicio_segmento = fila_actual_excel
            for col_fila, col_dato in columnas_bloques:
                ws[f'{col_fila}{fila_inicio_segmento}'] = "FILA"
                ws[f'{col_dato}{fila_inicio_segmento}'] = cabecera_columna
                ws[f'{col_fila}{fila_inicio_segmento}'].font = font_style
                ws[f'{col_dato}{fila_inicio_segmento}'].font = font_style
            ws.row_dimensions[fila_inicio_segmento].height = row_height

            # --- Lógica de llenado VERTICAL ---
            # Bucle exterior: por cada bloque de columnas
            for col_fila, col_dato in columnas_bloques:
                # Bucle interior: por cada fila DENTRO de ese bloque
                for i_fila_en_bloque in range(max_filas_por_bloque):
                    if puntero_datos >= num_datos:
                        break # Salir si ya no hay datos
                    
                    fila_absoluta = fila_inicio_segmento + 1 + i_fila_en_bloque
                    ws.row_dimensions[fila_absoluta].height = row_height
                    
                    # Escribir el número de fila
                    cell_fila = ws[f'{col_fila}{fila_absoluta}']
                    cell_fila.value = puntero_datos + 1
                    cell_fila.font = font_style
                    
                    # Escribir el dato (Oficio/Expediente)
                    cell_dato = ws[f'{col_dato}{fila_absoluta}']
                    valor_original = datos_para_reporte[puntero_datos]
                    
                    if titulo_principal == "MULTAS":
                        cell_dato.value = str(valor_original)[-6:]
                    else:
                        cell_dato.value = str(valor_original)
                    
                    cell_dato.font = font_style
                    puntero_datos += 1
                
                if puntero_datos >= num_datos:
                    break
            
            # Preparar la fila para la siguiente "página"
            fila_actual_excel = fila_inicio_segmento + 1 + max_filas_por_bloque

        print(f"    - Hoja '{sheet_name}' creada y formateada con {num_datos} registros.")

    except Exception as e:
        print(f"    - (!) Error crítico al crear la hoja de reporte para impresión: {e}")
        import traceback
        traceback.print_exc()

# --- Funciones para la Base de Datos Maestra (Excel) ---

def cargar_bd_maestra_multas(ruta_excel_bd_maestra):
    """
    Carga la Base de Datos Maestra desde el archivo Excel especificado.
    Si el archivo no existe, o la hoja 'BD_Maestra_Multas' no existe,
    crea un DataFrame vacío con la estructura definida.
    Asegura que todas las columnas definidas en COLUMNAS_BD_MAESTRA_MULTAS existan
    y tengan los tipos de datos definidos en TIPOS_BD_MAESTRA_MULTAS.
    """
    with excel_lock_multas:
        sheet_name_bd = "BD_Maestra_Multas"
        if ruta_excel_bd_maestra.exists():
            try:
                print(f"    - Cargando BD Maestra de Multas desde: {ruta_excel_bd_maestra}")
                # Leer sin especificar dtypes inicialmente para manejar mejor las conversiones
                df = pd.read_excel(ruta_excel_bd_maestra, sheet_name=sheet_name_bd)
                
                # Verificar y ajustar columnas y tipos
                for col_name in COLUMNAS_BD_MAESTRA_MULTAS:
                    expected_dtype_str = TIPOS_BD_MAESTRA_MULTAS[col_name]
                    
                    if col_name not in df.columns:
                        print(f"      -> Columna '{col_name}' no encontrada. Añadiendo con tipo '{expected_dtype_str}'.")
                        if expected_dtype_str == 'Int64':
                            df[col_name] = pd.NA # Para columnas Int64, el valor por defecto es NA
                        else: # Para str y otros
                            df[col_name] = "" 
                    
                    # Intentar convertir al tipo esperado
                    try:
                        current_dtype_name = df[col_name].dtype.name
                        if current_dtype_name == expected_dtype_str: # Si ya tiene el tipo correcto
                            if expected_dtype_str == 'str': # Asegurar que los NaN en str sean ""
                                df[col_name] = df[col_name].fillna("").astype(str)
                            # Para Int64 y otros tipos, si ya coincide, no hacer nada extra aquí.
                            continue # Saltar a la siguiente columna

                        # Si el tipo no coincide, intentar la conversión
                        if expected_dtype_str == 'Int64':
                            df[col_name] = pd.to_numeric(df[col_name], errors='coerce').astype(expected_dtype_str)
                        elif expected_dtype_str == 'str':
                             df[col_name] = df[col_name].fillna("").astype(str) # Convertir NaN a "" antes de astype(str)
                        else: 
                            df[col_name] = df[col_name].astype(expected_dtype_str) # Para otros tipos
                    
                    except Exception as e_type_conv:
                        print(f"      - (!) Advertencia: No se pudo convertir la columna '{col_name}' (tipo actual: {df[col_name].dtype}) al tipo esperado '{expected_dtype_str}'. Se usará string como fallback. Error: {e_type_conv}")
                        df[col_name] = df[col_name].astype(str).fillna("") # Fallback a string
                
                # Reordenar y asegurar que todas las columnas de COLUMNAS_BD_MAESTRA_MULTAS estén presentes
                df = df.reindex(columns=COLUMNAS_BD_MAESTRA_MULTAS)
                # Aplicar tipos una vez más después del reindex por si se crearon nuevas columnas
                for col_name_reindex in COLUMNAS_BD_MAESTRA_MULTAS:
                    current_dtype_reindex = df[col_name_reindex].dtype.name
                    expected_dtype_reindex = TIPOS_BD_MAESTRA_MULTAS[col_name_reindex]
                    if current_dtype_reindex != expected_dtype_reindex:
                        try:
                            if expected_dtype_reindex == 'Int64':
                                df[col_name_reindex] = pd.to_numeric(df[col_name_reindex], errors='coerce').astype(expected_dtype_reindex)
                            elif expected_dtype_reindex == 'str':
                                df[col_name_reindex] = df[col_name_reindex].fillna("").astype(str)
                            else:
                                df[col_name_reindex] = df[col_name_reindex].astype(expected_dtype_reindex)
                        except Exception as e_reindex_astype_final:
                             print(f"      - (!) Advertencia (post-reindex): No se pudo convertir '{col_name_reindex}' a '{expected_dtype_reindex}'. Usando string. Error: {e_reindex_astype_final}")
                             df[col_name_reindex] = df[col_name_reindex].astype(str).fillna("")
                
                print(f"    - BD Maestra de Multas cargada y procesada: {len(df)} registros.")
                return df
            
            except (InvalidFileException, ValueError, KeyError) as e_excel_read:
                print(f"    - (!) ADVERTENCIA: No se pudo leer o procesar '{ruta_excel_bd_maestra}' (Hoja: '{sheet_name_bd}'). Se creará/usará una BD vacía. Error: {e_excel_read}")
                # traceback.print_exc() # Descomentar para depuración detallada del error
            except Exception as e_general_load:
                print(f"    - (!) ADVERTENCIA GENERAL: Error inesperado al cargar BD Maestra '{ruta_excel_bd_maestra}'. Se creará/usará una BD vacía. Error: {e_general_load}")

        else:
            print(f"    - Archivo BD Maestra de Multas '{ruta_excel_bd_maestra}' no encontrado. Se creará uno nuevo.")

        # Crear DataFrame vacío con la estructura correcta si el archivo no existe o falla la carga
        df_vacio = pd.DataFrame(columns=COLUMNAS_BD_MAESTRA_MULTAS)
        # Aplicar tipos y valores por defecto
        for col_name_empty, dtype_empty in TIPOS_BD_MAESTRA_MULTAS.items():
            if dtype_empty == 'str':
                df_vacio[col_name_empty] = pd.Series(dtype='object').fillna("") # Iniciar con strings vacíos
            elif dtype_empty == 'Int64':
                df_vacio[col_name_empty] = pd.Series(dtype='Int64') # Iniciar con NA para Int64
            else:
                df_vacio[col_name_empty] = pd.Series(dtype=dtype_empty)
        
        df_vacio = df_vacio.astype(TIPOS_BD_MAESTRA_MULTAS) # Asegurar tipos finales
        return df_vacio



def actualizar_o_agregar_registro_bd_multas(df_bd_maestra, datos_registro_actualizar):
    """
    Actualiza un registro existente en el DataFrame de la BD Maestra basado en COL_OFICIO_MULTAS,
    o agrega uno nuevo si no existe.
    Solo actualiza las columnas presentes en datos_registro_actualizar.
    Preserva los valores de otras columnas si el registro ya existe.
    Maneja correctamente la asignación a columnas Int64.
    """
    oficio_a_buscar = str(datos_registro_actualizar.get(COL_OFICIO_MULTAS, "")).strip()
    if not oficio_a_buscar:
        print("    - (!) Error: Se intentó actualizar/agregar un registro sin OFICIO en la BD Maestra.")
        failed_operations_log["actualizacion_bd_maestra"].append({"error": "Oficio faltante", "datos": datos_registro_actualizar})
        return df_bd_maestra

    if COL_OFICIO_MULTAS in df_bd_maestra.columns:
        df_bd_maestra[COL_OFICIO_MULTAS] = df_bd_maestra[COL_OFICIO_MULTAS].astype(str).fillna("")
    else: 
        print(f"    - (!) Advertencia: Columna '{COL_OFICIO_MULTAS}' no encontrada en BD Maestra al actualizar. Se añadirá.")
        df_bd_maestra[COL_OFICIO_MULTAS] = ""
        df_bd_maestra[COL_OFICIO_MULTAS] = df_bd_maestra[COL_OFICIO_MULTAS].astype(TIPOS_BD_MAESTRA_MULTAS[COL_OFICIO_MULTAS])

    indices_existentes = df_bd_maestra.index[df_bd_maestra[COL_OFICIO_MULTAS] == oficio_a_buscar].tolist()

    if indices_existentes:
        idx_actualizar = indices_existentes[-1] 
        for columna, valor_nuevo in datos_registro_actualizar.items():
            if columna in df_bd_maestra.columns:
                tipo_esperado = TIPOS_BD_MAESTRA_MULTAS.get(columna)
                try:
                    if pd.isna(valor_nuevo):
                        valor_convertido = pd.NA if tipo_esperado == 'Int64' else ''
                    elif tipo_esperado == 'Int64':
                        # Convertir a numérico, luego si no es NA, a entero de Python antes de asignar.
                        # Pandas se encarga de la conversión a Int64 (nullable integer) si la columna ya tiene ese tipo.
                        # Si es pd.NA, se asigna directamente.
                        num_val = pd.to_numeric(valor_nuevo, errors='coerce')
                        valor_convertido = pd.NA if pd.isna(num_val) else int(num_val)
                    elif tipo_esperado == 'str':
                        valor_convertido = str(valor_nuevo).strip() if valor_nuevo is not None else ""
                    else: 
                        valor_convertido = valor_nuevo 
                    
                    df_bd_maestra.loc[idx_actualizar, columna] = valor_convertido
                except Exception as e_conv_update:
                    # CORRECCIÓN: Si la conversión a Int64 falla, no intentar asignar string directamente a una columna Int64.
                    # Asignar pd.NA para Int64 o string vacío para otras.
                    print(f"      - (!) Advertencia (Actualización): Error convirtiendo valor '{valor_nuevo}' para columna '{columna}' (Oficio: {oficio_a_buscar}). Error: {e_conv_update}")
                    if tipo_esperado == 'Int64':
                        df_bd_maestra.loc[idx_actualizar, columna] = pd.NA
                        print(f"          -> Se asignó pd.NA a '{columna}' para Oficio {oficio_a_buscar}.")
                    else:
                        df_bd_maestra.loc[idx_actualizar, columna] = str(valor_nuevo) if valor_nuevo is not None else ""
                        print(f"          -> Se asignó valor como string a '{columna}' para Oficio {oficio_a_buscar}.")
    else:
        nuevo_registro_dict = {}
        for col_maestra in COLUMNAS_BD_MAESTRA_MULTAS:
            tipo_col_maestra = TIPOS_BD_MAESTRA_MULTAS[col_maestra]
            if col_maestra in datos_registro_actualizar:
                valor_entrante = datos_registro_actualizar[col_maestra]
                try:
                    if pd.isna(valor_entrante):
                        nuevo_registro_dict[col_maestra] = pd.NA if tipo_col_maestra == 'Int64' else ""
                    elif tipo_col_maestra == 'Int64':
                        num_val_nuevo = pd.to_numeric(valor_entrante, errors='coerce')
                        nuevo_registro_dict[col_maestra] = pd.NA if pd.isna(num_val_nuevo) else int(num_val_nuevo)
                    elif tipo_col_maestra == 'str':
                        nuevo_registro_dict[col_maestra] = str(valor_entrante).strip() if valor_entrante is not None else ""
                    else:
                        nuevo_registro_dict[col_maestra] = valor_entrante
                except Exception as e_conv_new_reg:
                    print(f"      - (!) Advertencia (Nuevo Registro): Error convirtiendo valor '{valor_entrante}' para '{col_maestra}'. Error: {e_conv_new_reg}")
                    if tipo_col_maestra == 'Int64':
                        nuevo_registro_dict[col_maestra] = pd.NA
                    else:
                        nuevo_registro_dict[col_maestra] = str(valor_entrante) if valor_entrante is not None else ""
            else: 
                nuevo_registro_dict[col_maestra] = pd.NA if tipo_col_maestra == 'Int64' else ""
        
        df_nueva_fila = pd.DataFrame([nuevo_registro_dict]) # Crea el DF sin restringir columnas
        
        # Aseguramos que los tipos de datos sean correctos antes de unir
        for col, tipo in TIPOS_BD_MAESTRA_MULTAS.items():
            if col in df_nueva_fila.columns:
                try:
                    df_nueva_fila[col] = df_nueva_fila[col].astype(tipo)
                except (ValueError, TypeError):
                    if tipo == 'Int64': # Si falla la conversión a Int64, probamos con float y luego Int
                        df_nueva_fila[col] = pd.to_numeric(df_nueva_fila[col], errors='coerce').astype(tipo)
                    else: # Si todo falla, lo dejamos como string
                        df_nueva_fila[col] = df_nueva_fila[col].astype(str)

        df_bd_maestra = pd.concat([df_bd_maestra, df_nueva_fila], ignore_index=True)
        # Re-aplicar tipos al DataFrame completo después de concat para asegurar Int64
        for col_final_concat, tipo_final_concat in TIPOS_BD_MAESTRA_MULTAS.items():
            if col_final_concat in df_bd_maestra.columns:
                if tipo_final_concat == 'Int64':
                    df_bd_maestra[col_final_concat] = pd.to_numeric(df_bd_maestra[col_final_concat], errors='coerce').astype('Int64')
                # Para 'str', ya debería estar bien, pero una pasada extra no daña
                elif tipo_final_concat == 'str':
                     df_bd_maestra[col_final_concat] = df_bd_maestra[col_final_concat].fillna("").astype(str)


    return df_bd_maestra
# --- FIN DE LA FUNCIÓN CORREGIDA ---


def guardar_bd_maestra_multas(df_bd_a_guardar, ruta_excel_bd_maestra):
    """
    Guarda el DataFrame en la hoja 'BD_Maestra_Multas' de un archivo Excel.
    Si el archivo o la hoja no existen, los crea.
    Sobrescribe completamente el contenido de la hoja 'BD_Maestra_Multas'.
    Intenta preservar otras hojas si existen.
    """
    with excel_lock_multas:
        sheet_name_bd = "BD_Maestra_Multas"
        print(f"    - Guardando BD Maestra de Multas en: '{ruta_excel_bd_maestra}', Hoja: '{sheet_name_bd}' ({len(df_bd_a_guardar)} registros)")

        try:
            # Asegurar que el DataFrame tenga las columnas en el orden definido y los tipos correctos
            df_para_escribir = df_bd_a_guardar.reindex(columns=COLUMNAS_BD_MAESTRA_MULTAS).copy()
            for col, tipo_esperado in TIPOS_BD_MAESTRA_MULTAS.items():
                try:
                    if df_para_escribir[col].dtype.name != tipo_esperado:
                        if tipo_esperado == 'Int64':
                            df_para_escribir[col] = pd.to_numeric(df_para_escribir[col], errors='coerce').astype('Int64')
                        elif tipo_esperado == 'str':
                            df_para_escribir[col] = df_para_escribir[col].fillna("").astype(str)
                        else:
                            df_para_escribir[col] = df_para_escribir[col].astype(tipo_esperado)
                except Exception as e_final_astype:
                    print(f"      - (!) Advertencia (Guardado): No se pudo convertir columna '{col}' al tipo '{tipo_esperado}'. Se guardará como string. Error: {e_final_astype}")
                    df_para_escribir[col] = df_para_escribir[col].fillna("").astype(str)


            # Usar ExcelWriter para más control sobre las hojas
            # Esto permite reemplazar una hoja sin afectar otras (si el archivo ya existe)
            # o crear un archivo nuevo si no existe.
            mode_write = 'a' if ruta_excel_bd_maestra.exists() else 'w'
            if_sheet_exists_action = 'replace' if ruta_excel_bd_maestra.exists() else None

            with pd.ExcelWriter(ruta_excel_bd_maestra, engine='openpyxl', mode=mode_write, if_sheet_exists=if_sheet_exists_action) as writer:
                df_para_escribir.to_excel(writer, sheet_name=sheet_name_bd, index=False)
            
            print(f"    -> BD Maestra de Multas guardada exitosamente en '{ruta_excel_bd_maestra}', hoja '{sheet_name_bd}'.")
            return True
        except PermissionError:
            print(f"    - (!) ERROR CRÍTICO: Permiso denegado al intentar guardar '{ruta_excel_bd_maestra}'. ¿El archivo está abierto por otro programa?")
            failed_operations_log["guardado_bd_maestra"].append({"error": "Permiso denegado", "ruta": str(ruta_excel_bd_maestra)})
            return False
        except Exception as e:
            print(f"    - (!) ERROR CRÍTICO al guardar BD Maestra de Multas '{ruta_excel_bd_maestra}' con openpyxl/pandas: {e}")
            # traceback.print_exc() # Descomentar para depuración detallada
            failed_operations_log["guardado_bd_maestra"].append({"error": str(e), "ruta": str(ruta_excel_bd_maestra)})
            return False

# --- Funciones de Manejo de Archivos PDF (Reutilizadas y Adaptadas) ---

def contar_paginas_pdf(ruta_pdf_a_contar):
    """Cuenta las páginas de un archivo PDF. Devuelve None si hay error."""
    nombre_base_pdf = Path(ruta_pdf_a_contar).name
    try:
        with open(ruta_pdf_a_contar, 'rb') as f_pdf:
            reader = PyPDF2Reader(f_pdf, strict=False) # strict=False para más tolerancia
            count = len(reader.pages)
            return count
    except FileNotFoundError:
        # print(f"      - (!) Error contar_paginas: Archivo no encontrado '{nombre_base_pdf}'")
        failed_operations_log["conteo_paginas_pdf"].append({"archivo": nombre_base_pdf, "error": "FileNotFoundError"})
        return None
    except PyPDF2ReadError as e_pdf_read:
        print(f"      - (!) Error PyPDF2 al leer '{nombre_base_pdf}' para conteo: {e_pdf_read}")
        failed_operations_log["conteo_paginas_pdf"].append({"archivo": nombre_base_pdf, "error": f"PyPDF2ReadError: {e_pdf_read}"})
        return None
    except Exception as e_conteo:
        print(f"      - (!) Error inesperado al contar páginas de '{nombre_base_pdf}': {e_conteo}")
        failed_operations_log["conteo_paginas_pdf"].append({"archivo": nombre_base_pdf, "error": f"Exception: {e_conteo}"})
        return None

def extraer_paginas_pdf_multas(ruta_pdf_original, ruta_pdf_salida, especificacion_paginas="TODAS"):
    """
    Extrae páginas específicas de un PDF y las guarda en un nuevo archivo.
    Especificacion_paginas puede ser:
     - "TODAS": Copia todas las páginas.
     - "ULTIMA": Extrae solo la última página.
     - "RESTO": Extrae todas menos la última.
    Devuelve True si la extracción fue exitosa, False en caso contrario.
    """
    nombre_base_original = Path(ruta_pdf_original).name
    try:
        with open(ruta_pdf_original, 'rb') as f_in_pdf:
            reader = PyPDF2Reader(f_in_pdf, strict=False)
            writer = PyPDF2Writer()
            total_paginas_original = len(reader.pages)

            if total_paginas_original == 0:
                print(f"      - (!) Error extraer_paginas: '{nombre_base_original}' no tiene páginas.")
                failed_operations_log["extraccion_paginas_pdf"].append({"archivo_origen": nombre_base_original, "error": "PDF original sin páginas"})
                return False

            if especificacion_paginas == "TODAS":
                for page in reader.pages:
                    writer.add_page(page)
            elif especificacion_paginas == "ULTIMA":
                writer.add_page(reader.pages[total_paginas_original - 1])
            elif especificacion_paginas == "RESTO":
                if total_paginas_original <= 1:
                    # print(f"      - (*) Info extraer_paginas: '{nombre_base_original}' tiene 1 página o menos. 'RESTO' resulta en PDF vacío.")
                    # Se crea un PDF vacío intencionalmente en este caso, lo cual es un éxito técnico.
                    pass 
                else:
                    for i in range(total_paginas_original - 1): # Todas excepto la última
                        writer.add_page(reader.pages[i])
            else:
                print(f"      - (!) Error extraer_paginas: Especificación '{especificacion_paginas}' no reconocida.")
                failed_operations_log["extraccion_paginas_pdf"].append({"archivo_origen": nombre_base_original, "error": f"Especificación desconocida: {especificacion_paginas}"})
                return False

            # Guardar el PDF resultante
            Path(ruta_pdf_salida).parent.mkdir(parents=True, exist_ok=True) # Asegurar que la carpeta de salida exista
            with open(ruta_pdf_salida, 'wb') as f_out_pdf:
                writer.write(f_out_pdf)
            
            # Verificar si el archivo de salida se creó y tiene contenido (si se esperaban páginas)
            if especificacion_paginas == "RESTO" and total_paginas_original <= 1:
                return True # PDF vacío es el resultado esperado y correcto
            elif Path(ruta_pdf_salida).exists() and Path(ruta_pdf_salida).stat().st_size > 0:
                return True
            elif Path(ruta_pdf_salida).exists(): # Existe pero está vacío
                print(f"      - (!) Advertencia extraer_paginas: PDF de salida '{Path(ruta_pdf_salida).name}' se creó pero está vacío.")
                return True # Considerar éxito si se creó, aunque esté vacío y no debería.
            else: # No existe
                print(f"      - (!) Error extraer_paginas: PDF de salida '{Path(ruta_pdf_salida).name}' no se encontró después de la escritura.")
                return False

    except FileNotFoundError:
        print(f"      - (!) Error extraer_paginas: Archivo PDF original no encontrado '{nombre_base_original}'")
        failed_operations_log["extraccion_paginas_pdf"].append({"archivo_origen": nombre_base_original, "error": "FileNotFoundError (origen)"})
        return False
    except PyPDF2ReadError as e_pdf_read_ext:
        print(f"      - (!) Error PyPDF2 al leer '{nombre_base_original}' para extracción: {e_pdf_read_ext}")
        failed_operations_log["extraccion_paginas_pdf"].append({"archivo_origen": nombre_base_original, "error": f"PyPDF2ReadError: {e_pdf_read_ext}"})
        return False
    except Exception as e_ext:
        print(f"      - (!) Error inesperado extrayendo páginas de '{nombre_base_original}': {e_ext}")
        # traceback.print_exc() # Descomentar para depuración
        failed_operations_log["extraccion_paginas_pdf"].append({"archivo_origen": nombre_base_original, "error": f"Exception: {e_ext}"})
        # Intentar eliminar salida parcial si existe y falló la escritura
        if Path(ruta_pdf_salida).exists():
            try: Path(ruta_pdf_salida).unlink()
            except: pass
        return False

def preprocesar_csv_multas(csv_file_path, output_file_path, delimiter='|'):
    """Preprocesa el CSV para limpiar caracteres problemáticos antes de leer con pandas."""
    print(f"    - Preprocesando CSV: '{Path(csv_file_path).name}' -> '{Path(output_file_path).name}'")
    line_count = 0
    try:
        with open(csv_file_path, 'r', encoding='utf-8', errors='replace') as infile, \
             open(output_file_path, 'w', encoding='utf-8', newline='') as outfile:
            
            reader = csv.reader(infile, delimiter=delimiter, escapechar='\\', quotechar='"', dialect=csv.excel)
            writer = csv.writer(outfile, delimiter=delimiter, quoting=csv.QUOTE_MINIMAL, escapechar='\\')
            
            try:
                header = next(reader)
                writer.writerow(header)
                line_count = 1
                for row_idx, row in enumerate(reader, start=2):
                    line_count = row_idx
                    try:
                        writer.writerow(row)
                    except csv.Error as row_err:
                        print(f"        - (!) Error escritura CSV línea {line_count}: {row_err}. Fila (inicio): {str(row)[:100]}. Saltando.")
                        failed_operations_log["preprocesamiento_csv"].append({'linea': line_count, 'error': str(row_err), 'fila_inicio': str(row)[:100]})
            except StopIteration:
                print(f"        - (!) Advertencia: CSV '{Path(csv_file_path).name}' vacío o solo con encabezado.")
            except csv.Error as e_csv_read:
                print(f"        - (!) Error CSV (general) durante lectura en '{Path(csv_file_path).name}': {e_csv_read}")
                failed_operations_log["preprocesamiento_csv"].append({'archivo': Path(csv_file_path).name, 'error': f"Lectura CSV: {e_csv_read}"})
            except Exception as e_proc:
                print(f"        - (!) Error inesperado preprocesando CSV '{Path(csv_file_path).name}': {e_proc}")
                failed_operations_log["preprocesamiento_csv"].append({'archivo': Path(csv_file_path).name, 'error': f"Inesperado: {e_proc}"})
        
        print(f"    - Preprocesamiento CSV completado para '{Path(csv_file_path).name}'. {max(0, line_count-1)} líneas de datos procesadas.")
    except FileNotFoundError:
        print(f"    - (!) Error Fatal: Archivo CSV no encontrado: {csv_file_path}")
        failed_operations_log["preprocesamiento_csv"].append({'archivo': Path(csv_file_path).name, 'error': "FileNotFoundError"})
        raise # Re-lanzar para detener el proceso si el CSV principal falta
    except Exception as e_open:
        print(f"    - (!) Error Fatal abriendo/creando archivos para preprocesamiento de '{Path(csv_file_path).name}': {e_open}")
        failed_operations_log["preprocesamiento_csv"].append({'archivo': Path(csv_file_path).name, 'error': f"Apertura/Creación: {e_open}"})
        raise

def crear_reporte_estado_flotillas(config_multas_actual, df_bd_maestra):
    """
    Lee la hoja 'Flotillas', la cruza con la BD Maestra (YA CARGADA) para verificar el estado
    y genera un nuevo archivo Excel con los resultados.
    """
    print("--- Generando Reporte de Estado de Flotillas ---")
    
    # 1. Cargar la lista de oficios de Flotillas (esto se queda igual)
    try:
        filepath = Path(config_multas_actual["data_file_path"])
        # ... (código para cargar df_flotillas se mantiene igual que antes) ...
        sheet_name = "Flotillas"
        excel_file = pd.ExcelFile(filepath)
        if sheet_name not in excel_file.sheet_names:
            print(f"  (!) Error: Hoja '{sheet_name}' no encontrada en '{filepath.name}'.")
            return
        df_flotillas = pd.read_excel(filepath, sheet_name=sheet_name, dtype=str)
        df_flotillas.columns = [str(col).strip().upper() for col in df_flotillas.columns]
        if "OFICIO" not in df_flotillas.columns:
            print(f"  (!) Error: La hoja '{sheet_name}' no contiene la columna 'OFICIO'.")
            return
        df_flotillas['OFICIO_NORM'] = df_flotillas['OFICIO'].apply(normalizar_oficio_multas)
    except Exception as e:
        print(f"  (!) Error al cargar la hoja de Flotillas: {e}")
        return

    # 2. USAR la Base de Datos Maestra que fue pasada como argumento
    #    (SE ELIMINA EL BLOQUE try/except que la cargaba desde aquí)
    if df_bd_maestra.empty:
        print("  (*) Advertencia: La Base de Datos Maestra está vacía. Todos los estados se reportarán como 'No encontrado en BD'.")
    else:
        # Normalizar el oficio en la BD Maestra para el cruce
        col_oficio_bd = config_multas_actual["col_expediente"]
        if col_oficio_bd in df_bd_maestra.columns:
            df_bd_maestra['OFICIO_NORM'] = df_bd_maestra[col_oficio_bd].apply(normalizar_oficio_multas)
        else:
            print(f"  (!) Error: Columna de oficio '{col_oficio_bd}' no encontrada en la BD Maestra provista.")
            # Crear columna vacía para que el resto del código no falle
            df_bd_maestra['OFICIO_NORM'] = ''

    # 3. Cruzar la información
    def determinar_estado_impresion(estado):
        if pd.isna(estado) or estado == "":
            return "No encontrado en BD Maestra"
        
        estado_upper = str(estado).upper()
        if "IMPRESION COMPLETADA" in estado_upper or "ULTIMA PAG IMPRESA" in estado_upper:
            return "Impreso"
        elif "NO GENERADO" in estado_upper:
             return estado # Devuelve el estado específico como "NO GENERADO (MENOR A 179)"
        else:
            return "Pendiente de Impresión"

    # Preparar el DataFrame de la BD Maestra para el merge
    df_bd_estados = df_bd_maestra[['OFICIO_NORM', config_multas_actual["col_estado_bd_maestra"]]].drop_duplicates(subset=['OFICIO_NORM'], keep='last')

    # Unir df_flotillas con los estados de la BD Maestra
    df_reporte = pd.merge(
        df_flotillas,
        df_bd_estados,
        on='OFICIO_NORM',
        how='left'
    )
    
    col_estado_bd = config_multas_actual["col_estado_bd_maestra"]
    df_reporte['Estado Impresión'] = df_reporte[col_estado_bd].apply(determinar_estado_impresion)
    
    # 4. Preparar y guardar el archivo final
    # Eliminar columnas temporales y reordenar
    df_reporte_final = df_reporte[['OFICIO', 'NOMBRE', 'MONTO', 'Estado Impresión']]
    
    # Crear carpeta de reportes si no existe
    reportes_path = Path(config_multas_actual["base_path"]) / "REPORTES"
    reportes_path.mkdir(exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = reportes_path / f"Reporte_Estado_Flotillas_{timestamp}.xlsx"
    
    try:
        df_reporte_final.to_excel(output_filename, index=False)
        print("\n  --- Reporte de Estado de Flotillas generado exitosamente ---")
        print(f"  Archivo guardado en: {output_filename}")
    except Exception as e:
        print(f"  (!) Error al guardar el archivo de reporte: {e}")

def normalizar_oficio_multas(oficio_str):
    if not oficio_str or pd.isna(oficio_str):
        return "OFICIO_VACIO_O_INVALIDO"
    s = str(oficio_str).upper().strip()
    match_slashes = re.search(r'DI\s*/\s*DCF\s*/\s*MT\s*/\s*(\d+)', s)
    if match_slashes:
        return f"DIDCFMT{match_slashes.group(1)}"
    match_prefix = re.search(r'DIDCFMT\s*(\d+)', s)
    if match_prefix:
        return f"DIDCFMT{match_prefix.group(1)}"
    if s.isdigit():
        return f"DIDCFMT{s}"
    numeros_encontrados = re.findall(r'\d+', s)
    if numeros_encontrados:
        return f"DIDCFMT{numeros_encontrados[-1]}"
    # print(f"    - ADVERTENCIA: Oficio '{oficio_str}' no pudo ser normalizado a un formato conocido. Se devuelve: '{s}'") # Opcional: habilitar para depuración
    return s

def convertir_a_formato_con_barras(oficio_input_str):
    """
    Convierte un oficio en cualquier formato reconocible (ej. DIDCFMT12345 o DI/DCF/MT/12345)
    al formato estándar con barras (ej. DI/DCF/MT/12345).
    Si el formato de entrada no es un oficio reconocible, devuelve el string original.
    """
    if not isinstance(oficio_input_str, str):
        return oficio_input_str

    oficio_didcfmt = normalizar_oficio_multas(oficio_input_str) # Primero normaliza a DIDCFMT

    if oficio_didcfmt.startswith("DIDCFMT") and oficio_didcfmt[7:].isdigit():
        # Si se pudo normalizar a DIDCFMTXXXXXX, entonces convertir ESTE a formato con barras
        return f"DI/DCF/MT/{oficio_didcfmt[7:]}"
    
    # Si no se pudo normalizar a un formato DIDCFMT (ej. ya era "OFICIO_VACIO_O_INVALIDO")
    # o si el input original ya estaba con barras y normalizar_oficio_multas no lo cambió (porque falló la normalización),
    # se devuelve el input original.
    # O si el input original ya era con barras y normalizar_oficio_multas SÍ lo cambió a DIDCFMT,
    # la condición anterior ya lo manejó.
    return oficio_input_str

def leer_csv_datos_multas(ruta_csv_preprocesado):
    """Lee el CSV preprocesado y lo carga en un DataFrame de pandas."""
    print(f"    - Leyendo CSV de datos preprocesado: '{Path(ruta_csv_preprocesado).name}'")
    try:
        # Leer todas las columnas como string inicialmente para evitar problemas de inferencia de tipos
        df_datos = pd.read_csv(ruta_csv_preprocesado, encoding='utf-8', sep='|', dtype=str, low_memory=False, quoting=csv.QUOTE_MINIMAL, escapechar='\\')
        
        # Limpiar espacios en blanco de todas las celdas de string
        for col in df_datos.columns:
            if df_datos[col].dtype == 'object': # Solo para columnas de tipo string/object
                df_datos[col] = df_datos[col].str.strip().fillna('') # Llenar NaN con string vacío después de strip
            else: # Para columnas no-string, solo llenar NaN si es apropiado (ej. con 0 para numéricos si se desea)
                df_datos[col] = df_datos[col].fillna('') # Llenar NaN con string vacío por ahora, se convertirán después

        print(f"    - CSV '{Path(ruta_csv_preprocesado).name}' leído. {len(df_datos)} filas cargadas.")
        return df_datos
    except FileNotFoundError:
        print(f"    - (!) Error: CSV preprocesado no encontrado: {ruta_csv_preprocesado}")
        failed_operations_log["lectura_csv_datos"].append({'archivo': Path(ruta_csv_preprocesado).name, 'error': "FileNotFoundError"})
        return pd.DataFrame() # Devolver DataFrame vacío
    except pd.errors.EmptyDataError:
        print(f"    - (!) Error: CSV preprocesado vacío: {ruta_csv_preprocesado}")
        failed_operations_log["lectura_csv_datos"].append({'archivo': Path(ruta_csv_preprocesado).name, 'error': "EmptyDataError"})
        return pd.DataFrame()
    except Exception as e:
        print(f"    - (!) Error crítico leyendo CSV preprocesado '{Path(ruta_csv_preprocesado).name}': {e}")
        failed_operations_log["lectura_csv_datos"].append({'archivo': Path(ruta_csv_preprocesado).name, 'error': str(e)})
        return pd.DataFrame() # Devolver DataFrame vacío

# --- Funciones de Procesamiento de Documentos DOCX ---

def reemplazar_placeholders_en_runs_multas(paragraphs, replacements_dict):
    """
    Reemplaza placeholders en los runs de una lista de párrafos de forma segura,
    preservando campos como los números de página.
    """
    if not paragraphs: return

    for p in paragraphs:
        if p is None: continue

        # Se necesita obtener el texto completo del párrafo para encontrar placeholders
        # que podrían estar divididos entre diferentes "runs" (segmentos de texto).
        full_text = "".join(run.text for run in p.runs)

        # Solo si se detecta un placeholder, se procede a modificar el párrafo.
        if any(key in full_text for key in replacements_dict):

            # Se realiza el reemplazo en la cadena de texto completa.
            for placeholder, value in replacements_dict.items():
                if placeholder in full_text:
                    text_value = str(value) if value is not None and not pd.isna(value) else ""
                    full_text = full_text.replace(placeholder, text_value)

            # A continuación, la lógica "quirúrgica":
            # 1. Se borra el texto de todos los segmentos ('runs') del párrafo.
            # 2. Se coloca TODO el texto ya modificado en el PRIMER segmento.
            # 3. Los demás segmentos se dejan vacíos.
            # Esto conserva el párrafo y sus campos (como el número de página),
            # pero actualiza su contenido de texto de forma segura.
            for i, run in enumerate(p.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ""

def reemplazar_placeholders_multas(document_obj, replacements_dict):
    """Reemplaza placeholders en todo el documento (párrafos, tablas, encabezados, pies)."""
    if document_obj is None: return

    # Reemplazar en párrafos principales
    reemplazar_placeholders_en_runs_multas(document_obj.paragraphs, replacements_dict)

    # Reemplazar en tablas
    for table in document_obj.tables:
        if table is None: continue
        for row_idx, row in enumerate(table.rows):
            if row is None: continue
            is_header_row = (row_idx == 0) # Asumir que la primera fila es cabecera
            for cell in row.cells:
                if cell is None: continue
                if not is_header_row: # Aplicar a celdas de datos
                    reemplazar_placeholders_en_runs_multas(cell.paragraphs, replacements_dict)
                # else: # Lógica para cabeceras si también pueden tener placeholders (opcional)
                #     pass 

    # Reemplazar en encabezados y pies de página
    for section in document_obj.sections:
        if section is None: continue
        if section.header:
            reemplazar_placeholders_en_runs_multas(section.header.paragraphs, replacements_dict)
            for table_header in section.header.tables:
                for row_header in table_header.rows:
                    for cell_header in row_header.cells:
                        reemplazar_placeholders_en_runs_multas(cell_header.paragraphs, replacements_dict)
        if section.footer:
            reemplazar_placeholders_en_runs_multas(section.footer.paragraphs, replacements_dict)
            for table_footer in section.footer.tables:
                for row_footer in table_footer.rows:
                    for cell_footer in row_footer.cells:
                        reemplazar_placeholders_en_runs_multas(cell_footer.paragraphs, replacements_dict)

def procesar_tabla_dinamica_multas(tabla_obj, registros_df_oficio, importe_total_formateado_tabla, oficio_actual_str):
    """Procesa y llena la tabla dinámica de multas en el documento Word."""
    global failed_operations_log
    if tabla_obj is None:
        print(f"        - (!) Error Interno: La tabla proporcionada para Oficio {oficio_actual_str} es None.")
        failed_operations_log[f"procesamiento_tabla_oficio_{oficio_actual_str}"].append({'error': 'Objeto tabla es None'})
        return

    # print(f"        - Procesando tabla para Oficio {oficio_actual_str}...")
    tabla_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Identificar etiqueta de la fila de total en la plantilla (si existe)
    etiqueta_total_plantilla = "TOTAL DE MULTAS DE TRANSITO:" # Default
    num_filas_plantilla_original = len(tabla_obj.rows)
    if num_filas_plantilla_original > 1:
        for i in range(num_filas_plantilla_original -1, 0, -1): # Iterar desde el final
            try:
                row_text_concat = "".join(cell.text.strip() for cell in tabla_obj.rows[i].cells).upper()
                if "[MULTA]" in row_text_concat or "TOTAL DE MULTAS" in row_text_concat:
                    if tabla_obj.rows[i].cells:
                         etiqueta_total_plantilla = tabla_obj.rows[i].cells[0].text.strip()
                    break 
            except Exception: continue

    # Eliminar todas las filas de la tabla EXCEPTO la primera (cabecera)
    if len(tabla_obj.rows) > 1:
        for i in range(len(tabla_obj.rows) - 1, 0, -1):
            try:
                row_element = tabla_obj.rows[i]._element
                row_element.getparent().remove(row_element)
            except Exception as e_del_row:
                print(f"            - (!) Adv: Error eliminando fila preexistente de plantilla (índice {i}, Oficio {oficio_actual_str}): {e_del_row}")

    # Formatear la cabecera (primera fila)
    if len(tabla_obj.rows) > 0:
        try:
            header_row = tabla_obj.rows[0]
            header_row.height = Inches(0.20)
            for cell_idx, cell in enumerate(header_row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                original_text_header = cell.text.strip() 
                
                for p_old in cell.paragraphs: p_old._element.getparent().remove(p_old._element)
                
                p_header = cell.add_paragraph()
                p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_header = p_header.add_run(original_text_header)
                run_header.font.name = FONT_NAME_TABLE_MULTAS
                run_header.font.size = FONT_SIZE_TABLE_MULTAS
                run_header.bold = True
        except Exception as e_format_header:
            print(f"            - (!) Adv. al aplicar formato a cabecera (Oficio {oficio_actual_str}): {e_format_header}")
    else:
        print(f"        - (!) Advertencia: Tabla para Oficio {oficio_actual_str} quedó sin filas. No se puede procesar.")
        return

    # Columnas esperadas en registros_df_oficio para la tabla. El orden importa.
    # Estas deben coincidir con las columnas del DataFrame que se pasa.
    columnas_datos_tabla = ["PLACAS", "FECHA", "BOLETA", "CONCEPTO", "INFRACCION", "SANCION", "MOTIVACION", "Uma/SM_numeric", "CUOTAS_numeric", "IMPORTE_numeric"]
    num_columnas_esperadas_datos = len(columnas_datos_tabla)
    
    for idx, fila_datos_oficio in registros_df_oficio.iterrows():
        try:
            nueva_fila_tabla = tabla_obj.add_row()
            celdas_nueva_fila = nueva_fila_tabla.cells
            
            if len(celdas_nueva_fila) != num_columnas_esperadas_datos:
                 print(f"            - (!) Adv (Oficio {oficio_actual_str}, Fila datos {idx}): Discrepancia de columnas. Creadas: {len(celdas_nueva_fila)}, Esperadas: {num_columnas_esperadas_datos}.")
            
            for j, nombre_col_datos in enumerate(columnas_datos_tabla):
                if j >= len(celdas_nueva_fila): break 
                
                celda_a_llenar = celdas_nueva_fila[j]
                valor_dato = fila_datos_oficio.get(nombre_col_datos, '')
                texto_celda_final = ''
                
                if nombre_col_datos == "FECHA":
                    texto_celda_final = formatear_fecha_corta(valor_dato)
                elif nombre_col_datos == "IMPORTE_numeric": # Esta columna ya es numérica
                    texto_celda_final = formatear_valor_celda(valor_dato, es_moneda=True) # Añadir $
                elif nombre_col_datos in ["CUOTAS_numeric", "Uma/SM_numeric"]: # Estas ya son numéricas
                    texto_celda_final = formatear_valor_celda(valor_dato, es_moneda=False) # Sin $
                else:
                    texto_celda_final = str(valor_dato if pd.notna(valor_dato) else '')

                for p_old_cell in celda_a_llenar.paragraphs: p_old_cell._element.getparent().remove(p_old_cell._element)
                p_cell = celda_a_llenar.add_paragraph()
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_cell = p_cell.add_run(texto_celda_final)
                run_cell.font.name = FONT_NAME_TABLE_MULTAS
                run_cell.font.size = FONT_SIZE_TABLE_MULTAS
                run_cell.bold = False
                celda_a_llenar.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            nueva_fila_tabla.height = Inches(0.20)
        except Exception as e_row_tabla:
            print(f"        - (!) Error GRAVE añadiendo fila de datos a tabla (Índice DF: {idx}, Oficio: {oficio_actual_str}): {e_row_tabla}")
            # traceback.print_exc(limit=1) # Descomentar para depuración
            failed_operations_log[f"error_fila_tabla_oficio_{oficio_actual_str}"].append({'idx_df': idx, 'error': str(e_row_tabla)})

    # Añadir fila final de total
    num_columnas_reales_tabla = len(tabla_obj.columns)
    if num_columnas_reales_tabla >= 2:
        try:
            fila_total_tabla = tabla_obj.add_row()
            celdas_fila_total = fila_total_tabla.cells
            fila_total_tabla.height = Inches(0.20)

            celda_etiqueta_total_fusionada = celdas_fila_total[0].merge(celdas_fila_total[num_columnas_reales_tabla - 2])
            
            for p_old_label in celda_etiqueta_total_fusionada.paragraphs: p_old_label._element.getparent().remove(p_old_label._element)
            p_label_total = celda_etiqueta_total_fusionada.add_paragraph()
            p_label_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_label_total = p_label_total.add_run(etiqueta_total_plantilla)
            r_label_total.font.name = FONT_NAME_TABLE_MULTAS; r_label_total.font.size = FONT_SIZE_TABLE_MULTAS; r_label_total.bold = True
            celda_etiqueta_total_fusionada.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            celda_valor_total = celdas_fila_total[num_columnas_reales_tabla - 1]
            for p_old_valor in celda_valor_total.paragraphs: p_old_valor._element.getparent().remove(p_old_valor._element)
            p_valor_total = celda_valor_total.add_paragraph()
            # El importe_total_formateado_tabla ya viene con $ y formateado
            p_valor_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_valor_total = p_valor_total.add_run(importe_total_formateado_tabla)
            r_valor_total.font.name = FONT_NAME_TABLE_MULTAS; r_valor_total.font.size = FONT_SIZE_TABLE_MULTAS; r_valor_total.bold = True
            celda_valor_total.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception as e_fila_total:
            print(f"        - (!) Error crítico añadiendo fila total para Oficio {oficio_actual_str}: {e_fila_total}")
            # traceback.print_exc(limit=1) # Descomentar para depuración
            failed_operations_log[f"error_fila_total_oficio_{oficio_actual_str}"].append({'error': str(e_fila_total)})
    elif num_columnas_reales_tabla == 1: # Manejo si solo hay una columna
        try:
            fila_total_unica_col = tabla_obj.add_row()
            celda_unica_total = fila_total_unica_col.cells[0]
            fila_total_unica_col.height = Inches(0.20)
            for p_old_single in celda_unica_total.paragraphs: p_old_single._element.getparent().remove(p_old_single._element)
            p_total_single_col = celda_unica_total.add_paragraph()
            texto_total_single_col = f"{etiqueta_total_plantilla} {importe_total_formateado_tabla}"
            p_total_single_col.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_total_single_col = p_total_single_col.add_run(texto_total_single_col)
            r_total_single_col.font.name = FONT_NAME_TABLE_MULTAS; r_total_single_col.font.size = FONT_SIZE_TABLE_MULTAS; r_total_single_col.bold = True
            celda_unica_total.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception as e_total_single_col:
            print(f"        - (!) Error añadiendo fila total (1 columna) para Oficio {oficio_actual_str}: {e_total_single_col}")
    else:
        print(f"        - (!) No se añade fila total para Oficio {oficio_actual_str}, la tabla no tiene columnas suficientes (columnas: {num_columnas_reales_tabla}).")

def buscar_pdf_oficial_generado(oficio_norm, config_modo, nombre_contrib, cp):
    """
    Busca activamente un PDF generado oficial basado en la información del oficio.
    Devuelve la ruta (Path object) si lo encuentra, o None si no.
    """
    print(f"        - Buscando PDF oficial para Oficio (Norm): {oficio_norm}...")
    
    directorio_documentos = Path(config_modo["output_docs_path"])
    oficio_con_barras = convertir_a_formato_con_barras(oficio_norm)
    nombre_base_archivo = limpiar_texto(f"{oficio_con_barras}_{nombre_contrib}")
    nombre_archivo_pdf = f"{nombre_base_archivo}.pdf"
    
    rutas_a_buscar = []
    
    if cp and cp not in ["SIN_CP", "SIN_CP_VALIDO"]:
        ruta_en_cp = directorio_documentos / "CP" / f"CP_{cp}" / nombre_archivo_pdf
        rutas_a_buscar.append(ruta_en_cp)

    ruta_en_vacias = directorio_documentos / "VACIAS" / nombre_archivo_pdf
    rutas_a_buscar.append(ruta_en_vacias)
    
    ruta_en_raiz = directorio_documentos / nombre_archivo_pdf
    rutas_a_buscar.append(ruta_en_raiz)

    for ruta in rutas_a_buscar:
        if ruta.is_file():
            print(f"          -> ¡Éxito! PDF oficial encontrado en: {ruta}")
            return ruta
            
    print(f"        - No se encontró un PDF oficial pre-generado para el oficio {oficio_norm}.")
    return None

# --- Función Principal para Generar un Documento Individual ---

def generar_un_documento_multas_v2(
    datos_oficio_df, 
    plantilla_path_obj, 
    ruta_pdf_destino_obj, 
    modo_extraccion_paginas, 
    oficio_slashed_para_placeholder
    ):
    """
    [VERSIÓN CON MEJOR MANEJO DE ERRORES]
    Genera un documento DOCX, lo convierte a PDF y opcionalmente extrae páginas.
    """
    global failed_operations_log, Docx2PdfFileNotFoundError

    pdf_final_generado_ok = False
    num_paginas_pdf_completo = None
    
    carpeta_destino_final = ruta_pdf_destino_obj.parent
    nombre_base_archivo_sin_ext = ruta_pdf_destino_obj.stem
    timestamp_unico = str(time.time_ns())
    ruta_docx_temp = carpeta_destino_final / f"~TEMP_{nombre_base_archivo_sin_ext}_{timestamp_unico}.docx"
    ruta_pdf_completo_intermedio_temp = carpeta_destino_final / f"~TEMP_FULL_{nombre_base_archivo_sin_ext}_{timestamp_unico}.pdf"

    try:
        # --- Generación del DOCX (esta parte no cambia) ---
        # Crear una copia profunda del elemento XML de la plantilla
        document = Document(plantilla_path_obj)      
        try:
            for section in document.sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(13.4)
        except Exception as e_ajuste_pagina:
            print(f"            - (!) Adv. ajustando tamaño página (Oficio Placeholder: {oficio_slashed_para_placeholder}): {e_ajuste_pagina}")

        if datos_oficio_df.empty:
            raise ValueError(f"No hay datos en datos_oficio_df para el Oficio (Placeholder: {oficio_slashed_para_placeholder})")
        
        fila_representativa = datos_oficio_df.iloc[0]
        fecha_actual_doc = ""
        if configurar_locale_es():
            try: fecha_actual_doc = date.today().strftime("%d de %B de %Y")
            except Exception as e_fecha_doc: print(f"            - (!) Adv. formateando fecha para doc (Oficio Placeholder: {oficio_slashed_para_placeholder}): {e_fecha_doc}")
        
        if 'IMPORTE_numeric' not in datos_oficio_df.columns:
            raise KeyError(f"La columna 'IMPORTE_numeric' es necesaria en datos_oficio_df para Oficio (Placeholder: {oficio_slashed_para_placeholder})")
        
        suma_importes_oficio = datos_oficio_df['IMPORTE_numeric'].sum()
        monto_total_texto_oficio = numero_a_texto_moneda_mx(suma_importes_oficio)
        monto_total_formateado_celda = formatear_valor_celda(suma_importes_oficio, es_moneda=True)

        replacements_doc = {
            "[NOMBRE_COMPLETO]": fila_representativa.get("NOMBRE_COMPLETO_CALC", "CONTRIBUYENTE DESCONOCIDO"),
            "[DIRECCION]": fila_representativa.get("DIRECCION_CSV", "DIRECCIÓN DESCONOCIDA"),
            "[MULTA]": monto_total_formateado_celda,
            "[MULTA_TEXTO]": monto_total_texto_oficio,
            "[OFICIO]": oficio_slashed_para_placeholder,
            "[FECHA]": fecha_actual_doc
        }
        reemplazar_placeholders_multas(document, replacements_doc)

        if document.tables:
            procesar_tabla_dinamica_multas(
                document.tables[0], 
                datos_oficio_df,
                monto_total_formateado_celda,
                oficio_slashed_para_placeholder
            )
        
        document.save(ruta_docx_temp)

        ### CAMBIO 1: Se aísla la llamada a 'convert' para capturar el error exacto.
        try:
            convert(str(ruta_docx_temp), str(ruta_pdf_completo_intermedio_temp))
        except Exception as e_convert:
            print(f"        - (!) FALLO DIRECTO EN CONVERSIÓN. Error: {e_convert}")
            print(f"        - (!) REVISA EL ARCHIVO DOCX TEMPORAL PARA PISTAS: {ruta_docx_temp}")
            raise # Re-lanza la excepción para que sea capturada por el bloque principal

        ### CAMBIO 2: Se mejora el mensaje de error si el PDF no se crea o está vacío.
        if not (ruta_pdf_completo_intermedio_temp.exists() and ruta_pdf_completo_intermedio_temp.stat().st_size > 0):
            error_msg_pdf_intermedio = (f"Fallo al generar PDF completo intermedio. El archivo no se creó o está vacío. "
                                        f"Verifica el DOCX temporal: {ruta_docx_temp}")
            raise RuntimeError(error_msg_pdf_intermedio)
        
        num_paginas_pdf_completo = contar_paginas_pdf(ruta_pdf_completo_intermedio_temp)
        if num_paginas_pdf_completo is None:
            print(f"            - (!) Advertencia: No se pudo contar páginas del PDF completo intermedio para Oficio (Placeholder: {oficio_slashed_para_placeholder}).")
            failed_operations_log[f"conteo_paginas_pdf_completo_{oficio_slashed_para_placeholder}"].append({"error": "Conteo fallido en PDF intermedio"})

        if modo_extraccion_paginas == "TODAS":
            shutil.copy2(ruta_pdf_completo_intermedio_temp, ruta_pdf_destino_obj)
        else: 
            if not extraer_paginas_pdf_multas(ruta_pdf_completo_intermedio_temp, ruta_pdf_destino_obj, modo_extraccion_paginas):
                raise RuntimeError(f"Fallo al extraer páginas '{modo_extraccion_paginas}' del PDF intermedio para Oficio (Placeholder: {oficio_slashed_para_placeholder}).")
        
        if ruta_pdf_destino_obj.exists() and ruta_pdf_destino_obj.stat().st_size > 0:
            pdf_final_generado_ok = True
        elif modo_extraccion_paginas == "RESTO" and num_paginas_pdf_completo is not None and num_paginas_pdf_completo <= 1 and ruta_pdf_destino_obj.exists():
            pdf_final_generado_ok = True
        else:
            error_msg_pdf_final = f"El archivo PDF final ({modo_extraccion_paginas}) no se generó o está vacío: {ruta_pdf_destino_obj.name} para Oficio (Placeholder: {oficio_slashed_para_placeholder})"
            raise RuntimeError(error_msg_pdf_final)

    except Docx2PdfFileNotFoundError as e_docx2pdf_fnf:
        error_msg = f"MS Word/LibreOffice no encontrado/accesible o ruta de archivo inválida durante conversión para Oficio (Placeholder: {oficio_slashed_para_placeholder}). Detalles: {e_docx2pdf_fnf}"
        print(f"        - (!) Error FATAL conversión PDF: {error_msg}")
        failed_operations_log[f"error_conversion_app_no_encontrada_{oficio_slashed_para_placeholder}"].append({'error': error_msg, 'docx_temp': str(ruta_docx_temp)})
    except ImportError as e_imp_fatal:
         print(f"        - (!) Error FATAL: Módulo requerido no importado/funcional (Oficio Placeholder: {oficio_slashed_para_placeholder}). {e_imp_fatal}")
         failed_operations_log[f"error_importacion_critica_{oficio_slashed_para_placeholder}"].append({'error': str(e_imp_fatal)})
    except Exception as e_gen_doc:
        ### CAMBIO 3: El mensaje de error que se imprime ahora será mucho más específico gracias a los cambios anteriores.
        error_msg_gen = f"Fallo en generación DOCX, conversión PDF o extracción de página ({modo_extraccion_paginas}) para Oficio (Placeholder: {oficio_slashed_para_placeholder}): {e_gen_doc}"
        print(f"        - (!) Error Crítico: {error_msg_gen}")
        failed_operations_log[f"error_generacion_documento_{modo_extraccion_paginas}_{oficio_slashed_para_placeholder}"].append({'error': str(e_gen_doc), 'docx_temp': str(ruta_docx_temp)})
    finally:
        if ruta_docx_temp.exists():
            try: ruta_docx_temp.unlink()
            except Exception as e_del_docx_temp: print(f"                - (!) Adv. no se pudo eliminar DOCX temp '{ruta_docx_temp.name}': {e_del_docx_temp}")
        
        if modo_extraccion_paginas != "TODAS" and ruta_pdf_completo_intermedio_temp.exists():
            try: ruta_pdf_completo_intermedio_temp.unlink()
            except Exception as e_del_pdf_full_temp: print(f"                - (!) Adv. no se pudo eliminar PDF completo intermedio temp '{ruta_pdf_completo_intermedio_temp.name}': {e_del_pdf_full_temp}")
        elif modo_extraccion_paginas == "TODAS" and ruta_pdf_completo_intermedio_temp.exists() and ruta_pdf_completo_intermedio_temp != ruta_pdf_destino_obj:
            try: ruta_pdf_completo_intermedio_temp.unlink()
            except Exception as e_del_pdf_full_temp_extra: print(f"                - (!) Adv. limpieza extra PDF intermedio '{ruta_pdf_completo_intermedio_temp.name}': {e_del_pdf_full_temp_extra}")

    return pdf_final_generado_ok, num_paginas_pdf_completo



def extract_oficio_from_page_scan_multas(page_fitz_obj, page_num_debug):
    """
    Extrae el OFICIO buscando únicamente en la mitad inferior de la página,
    haciéndolo más rápido y preciso.
    """
    try:
        # Aumentar el zoom para mejor calidad de imagen
        zoom = 3.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page_fitz_obj.get_pixmap(colorspace=fitz.csGRAY, matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # --- MEJORA CLAVE: Recortar la imagen a la mitad inferior ---
        img_width, img_height = img.size
        area_recorte = (0, int(img_height * 0.50), img_width, img_height) # Desde el 50% hacia abajo
        img_recortada = img.crop(area_recorte)
        
        # Realizar OCR solo en la imagen recortada
        texto_pagina = pytesseract.image_to_string(img_recortada, lang='spa', config='--psm 4')

        if not texto_pagina.strip():
            return [], "OCR_MULTAS_PAGINA_VACIA"

        # El patrón busca "DIDCFMT" seguido de 6 o más dígitos
        patron_oficio = re.compile(r"D[I1L]DCFMT\s*(\d{6,})", re.IGNORECASE)
        coincidencias = patron_oficio.findall(texto_pagina)

        if not coincidencias:
            return [], "OCR_MULTAS_SIN_PATRON"

        # Devolver el o los oficios encontrados (normalmente será uno)
        candidatos_normalizados = [f"DIDCFMT{match}" for match in coincidencias]
        return list(set(candidatos_normalizados)), "OCR_MULTAS_CANDIDATOS_OK"

    except Exception as e_ocr_page:
        print(f"        - (!) Error inesperado en OCR de Multas en página {page_num_debug}: {e_ocr_page}")
        return [], "OCR_MULTAS_EXCEPCION"

def save_pdf_group_scan_multas(doc_fitz_original_obj, page_indices_list_0based, oficio_str_grupo, temp_dir_path_obj):
    """
    Guarda un grupo de páginas (especificadas por sus índices 0-based) de un documento fitz
    en un nuevo PDF temporal dentro del directorio temporal especificado.
    Retorna el Path object al PDF temporal creado, o None si falla.
    """
    if not page_indices_list_0based or not oficio_str_grupo:
        return None
    
    nombre_archivo_temporal_grupo = f"TEMP_GRUPO_{limpiar_texto(oficio_str_grupo)}_{time.time_ns()}.pdf"
    ruta_pdf_temporal_grupo = temp_dir_path_obj / nombre_archivo_temporal_grupo
    
    try:
        doc_fitz_nuevo_grupo = fitz.open() # Crear PDF vacío
        # Insertar las páginas seleccionadas del documento original
        # PyMuPDF espera una secuencia de números de página 0-based para insert_pdf
        doc_fitz_nuevo_grupo.insert_pdf(doc_fitz_original_obj, from_page=page_indices_list_0based[0], 
                                        to_page=page_indices_list_0based[-1], start_at=0)

        if doc_fitz_nuevo_grupo.page_count > 0:
            doc_fitz_nuevo_grupo.save(str(ruta_pdf_temporal_grupo), garbage=4, deflate=True) # Guardar optimizado
            doc_fitz_nuevo_grupo.close()

            if ruta_pdf_temporal_grupo.is_file() and ruta_pdf_temporal_grupo.stat().st_size > 0:
                return ruta_pdf_temporal_grupo
            else:
                print(f"        - (!) Error: PDF temporal de grupo para Oficio {oficio_str_grupo} no se creó o está vacío en {ruta_pdf_temporal_grupo}.")
                failed_operations_log[f"guardar_grupo_pdf_scan_{oficio_str_grupo}"].append({"error": "PDF temporal vacío o no creado"})
                return None
        else:
            print(f"        - (!) No se añadieron páginas al PDF temporal de grupo para Oficio {oficio_str_grupo}.")
            doc_fitz_nuevo_grupo.close()
            return None
    except Exception as e_save_grupo:
        print(f"        - (!) Error al crear PDF temporal de grupo para Oficio {oficio_str_grupo}: {e_save_grupo}")
        failed_operations_log[f"guardar_grupo_pdf_scan_{oficio_str_grupo}"].append({"error": str(e_save_grupo)})
        if 'doc_fitz_nuevo_grupo' in locals() and doc_fitz_nuevo_grupo:
            try: doc_fitz_nuevo_grupo.close()
            except: pass
        return None


def group_and_split_pdf_dynamically_scan_multas(pdf_fuente_path_obj, temp_dir_para_grupos_obj, df_bd_maestra_para_verificar, config_multas_actual, df_csv_principal_preparado):

    """
    [VERSIÓN MEJORADA] Analiza un PDF, agrupa páginas por OFICIO validado contra la BD Maestra
    y crea PDFs temporales para cada grupo.
    Devuelve una lista de tuplas de grupos válidos y un log de los no encontrados.
    """
    resultados_validos = []
    oficios_no_encontrados_log = {}
    doc_fitz_pdf_fuente = None

    # Crear un conjunto de oficios válidos (normalizados) de la BD Maestra para una búsqueda rápida
    col_oficio_bd = config_multas_actual["col_expediente"]
    oficios_validos_set = set()
    if not df_bd_maestra_para_verificar.empty and col_oficio_bd in df_bd_maestra_para_verificar.columns:
        oficios_validos_set = set(df_bd_maestra_para_verificar[col_oficio_bd].apply(normalizar_oficio_multas))
        print(f"    - Creado set de validación de Multas con {len(oficios_validos_set)} oficios únicos de la BD Maestra.")
    else:
        print("    - (!) ADVERTENCIA (Scan Multas): BD Maestra vacía o sin columna de oficio. No se podrá validar ningún OCR.")

    try:
        oficios_y_sus_paginas = {}
        oficios_invalidos_ya_reportados = set()
        doc_fitz_pdf_fuente = fitz.open(str(pdf_fuente_path_obj))
        total_paginas_pdf_fuente = doc_fitz_pdf_fuente.page_count
        print(f"    - Paso 1 (Multas): Mapeando y validando {total_paginas_pdf_fuente} páginas...")

        for i_pagina in range(total_paginas_pdf_fuente):
            num_pagina_para_mostrar = i_pagina + 1
            try:
                pagina_obj = doc_fitz_pdf_fuente.load_page(i_pagina)
                # La nueva función devuelve una LISTA de candidatos
                posibles_oficios, metodo_extraccion = extract_oficio_from_page_scan_multas(pagina_obj, num_pagina_para_mostrar)

                if not posibles_oficios:
                    continue

                # --- Lógica de Validación ---
                # 1. Si ya conocemos el oficio de una página anterior, lo agrupamos
                oficio_ya_conocido = None
                for candidato in posibles_oficios:
                    if candidato in oficios_y_sus_paginas:
                        oficios_y_sus_paginas[candidato]["paginas"].append(i_pagina)
                        oficio_ya_conocido = candidato
                        break
                    if candidato in oficios_invalidos_ya_reportados:
                        oficio_ya_conocido = candidato
                        break
                
                if oficio_ya_conocido:
                    continue

                # 2. Si son nuevos, intentar validar cada candidato contra la BD
                oficio_validado = None
                for candidato in posibles_oficios:
                    if candidato in oficios_validos_set:
                        oficio_validado = candidato
                        break
                
                # 3. Decidir qué hacer
                if oficio_validado:
                    print(f"    [+] VALIDADO: Oficio '{oficio_validado}' encontrado en BD Maestra (pág. {num_pagina_para_mostrar}).")
                    oficios_y_sus_paginas[oficio_validado] = {"paginas": [i_pagina], "metodo_ocr": metodo_extraccion}
                else:
                    # No se encontró en la BD Maestra. Ahora, verificar si es un "huérfano" válido en el CSV.
                    oficio_huerfano_validado = None
                    for candidato in posibles_oficios:
                        datos_csv = buscar_datos_oficio_en_csv_multas(candidato, df_csv_principal_preparado)
                        if datos_csv.get('encontrado_en_csv'):
                            oficio_huerfano_validado = candidato
                            break

                    if oficio_huerfano_validado:
                        # SÍ es un huérfano válido. Tratarlo como un oficio a procesar.
                        print(f"    [+] HUÉRFANO VALIDADO: Oficio '{oficio_huerfano_validado}' encontrado en CSV (pág. {num_pagina_para_mostrar}). Se generará.")
                        if oficio_huerfano_validado not in oficios_y_sus_paginas:
                            oficios_y_sus_paginas[oficio_huerfano_validado] = {"paginas": [], "metodo_ocr": metodo_extraccion}
                        oficios_y_sus_paginas[oficio_huerfano_validado]["paginas"].append(i_pagina)
                    else:
                        # NO está en la BD Maestra NI en el CSV. Es un error real.
                        oficio_a_reportar = posibles_oficios[0]
                        print(f"        [!] OFICIO NO ENCONTRADO en ninguna fuente: '{oficio_a_reportar}' (pág. {num_pagina_para_mostrar})")
                        for candidato in posibles_oficios:
                            oficios_invalidos_ya_reportados.add(candidato)
                        oficios_no_encontrados_log[oficio_a_reportar] = {"primera_pagina": num_pagina_para_mostrar}
            except Exception as e_proc_pagina:
                print(f"        - (!) Error mapeando página de Multas {num_pagina_para_mostrar}: {e_proc_pagina}.")
                continue

        print(f"\n    - Paso 1 (Multas) completo. {len(oficios_y_sus_paginas)} oficios válidos y {len(oficios_no_encontrados_log)} no encontrados.")
        print(f"    - Paso 2 (Multas): Creando PDFs de grupo para oficios válidos...")

        # --- INICIO: Bloque para ordenar por aparición en el PDF ---
        print(f"    - Ordenando oficios encontrados por orden de aparición en el PDF...")
        oficios_ordenados = sorted(
            oficios_y_sus_paginas.items(),
            key=lambda item: min(item[1]['paginas'])
        )
        # --- FIN DEL BLOQUE DE ORDENAMIENTO ---

        for oficio_str, data in oficios_ordenados:
            # El oficio ya está normalizado (DIDCFMTXXXXX)
            ruta_temp_grupo = save_pdf_group_scan_multas(doc_fitz_pdf_fuente, data["paginas"], oficio_str, temp_dir_para_grupos_obj)
            if ruta_temp_grupo:
                # El "oficio original" para el log es el mismo que el normalizado en este caso
                resultados_validos.append((oficio_str, oficio_str, ruta_temp_grupo, len(data["paginas"]), data["metodo_ocr"]))

    except Exception as e_open_split_pdf:
        print(f"    - (!) Error Crítico durante el mapeo de Multas: {e_open_split_pdf}")
    finally:
        if doc_fitz_pdf_fuente:
            try: doc_fitz_pdf_fuente.close()
            except: pass

    return resultados_validos, oficios_no_encontrados_log

def buscar_datos_oficio_en_csv_multas(oficio_norm_buscado, df_csv_principal_completo):
    """
    Busca un oficio normalizado en el DataFrame del CSV principal.
    Devuelve un diccionario con: 'nombre_contribuyente', 'direccion', 'cp', 'monto_formateado', 'encontrado_en_csv' (bool).
    Los valores son None si no se encuentran.
    """
    datos_encontrados = {
        'nombre_contribuyente': None, 
        'direccion': None, 
        'cp': None, 
        'monto_formateado': None,
        'encontrado_en_csv': False
    }
    if df_csv_principal_completo is None or df_csv_principal_completo.empty:
        return datos_encontrados

    # Asumir que df_csv_principal_completo ya tiene 'OFICIO_NORM_CSV', 'NOMBRE_COMPLETO_CALC',
    # 'DIRECCION_CSV', 'CP_CALC_CSV', 'IMPORTE_numeric'.
    
    registros_oficio_en_csv = df_csv_principal_completo[df_csv_principal_completo['OFICIO_NORM_CSV'] == oficio_norm_buscado]
    
    if not registros_oficio_en_csv.empty:
        fila_representativa_csv_oficio = registros_oficio_en_csv.iloc[0]
        datos_encontrados['nombre_contribuyente'] = fila_representativa_csv_oficio.get('NOMBRE_COMPLETO_CALC')
        datos_encontrados['direccion'] = fila_representativa_csv_oficio.get('DIRECCION_CSV')
        datos_encontrados['cp'] = fila_representativa_csv_oficio.get('CP_CALC_CSV')
        
        # Para el monto, sumar todos los importes de este oficio y formatear
        suma_importes_oficio_csv = registros_oficio_en_csv['IMPORTE_numeric'].sum()
        datos_encontrados['monto_formateado'] = formatear_valor_celda(suma_importes_oficio_csv, es_moneda=True)
        datos_encontrados['encontrado_en_csv'] = True
        
    return datos_encontrados

def crear_reporte_despachos_main(directorio_multas_cliente_str, config_multas_actual):
    """
    Función principal para crear el reporte de despachos.
    directorio_multas_cliente_str: Ruta (string) a la carpeta del modo MULTAS actual.
    config_multas_actual: Diccionario de configuración del modo actual (mode_config).
    """
    print("\n--- Iniciando Creación de Reporte de Despachos (Versión Actualizada) ---")
    
    # --- CAMBIO 1: Se definen nuevas constantes para las columnas del archivo de entrega ---
    COL_ENTREGA_EXPEDIENTE = "EXPEDIENTE"
    COL_ENTREGA_FECHA_NOTIF = "FECHA-DE NOTIFICACION"
    COL_ENTREGA_ESTATUS = "ESTATUS" # Se asume que esta columna se usará para filtrar
    
    # !!! IMPORTANTE: Confirma el valor exacto que indica una entrega exitosa !!!
    VALOR_ESTATUS_NOTIFICADO = "NOTIFICADO" # <--- ¡¡¡REVISA Y AJUSTA ESTE VALOR!!!

    if not config_multas_actual:
        print("Error Crítico: El diccionario de configuración del modo (config_multas_actual) no fue proporcionado.")
        return

    ruta_base_modo_actual_str = config_multas_actual.get('base_path')
    if not ruta_base_modo_actual_str or not Path(ruta_base_modo_actual_str).is_dir():
        print(f"Error Crítico: 'config_multas_actual['base_path']' ('{ruta_base_modo_actual_str or 'NO DEFINIDO'}') no es un directorio válido.")
        return
        
    carpeta_multas_lote_actual = Path(ruta_base_modo_actual_str) 
    carpeta_datos_despacho = carpeta_multas_lote_actual / NOMBRE_CARPETA_DATOS_DESPACHO
    
    if not carpeta_datos_despacho.is_dir():
        print(f"Error: No se encontró la carpeta '{NOMBRE_CARPETA_DATOS_DESPACHO}' dentro de '{carpeta_multas_lote_actual}'.")
        return

    # Selección del archivo de entrega (sin cambios en esta parte)
    archivos_excel_en_datos_despacho = [
        f for f in carpeta_datos_despacho.iterdir() 
        if f.is_file() and f.suffix.lower() == '.xlsx'
    ]
    if not archivos_excel_en_datos_despacho:
        print(f"Error: No se encontraron archivos Excel en '{carpeta_datos_despacho}' para seleccionar.")
        return

    print("\n  Seleccione el archivo de entrega (ej. 'ENTREGA...xlsx'):")
    for i, f_excel in enumerate(archivos_excel_en_datos_despacho):
        print(f"    {i+1}. {f_excel.name}")
    
    ruta_reporte_oficios_excel = None
    while True:
        try:
            opcion_str = input(f"  Ingrese el número del archivo (1-{len(archivos_excel_en_datos_despacho)}): ")
            opcion = int(opcion_str) - 1
            if 0 <= opcion < len(archivos_excel_en_datos_despacho):
                ruta_reporte_oficios_excel = archivos_excel_en_datos_despacho[opcion]
                break
            else:
                print("  Selección inválida. Intente de nuevo.")
        except ValueError:
            print("  Entrada no válida. Ingrese un número.")        
    print(f"  - Archivo de entrega seleccionado: {ruta_reporte_oficios_excel.name}")

    # Carga de BASE_DE_DATOS.xlsx (sin cambios en esta parte)
    ruta_base_de_datos_original_str = config_multas_actual.get('data_file_path') 
    if not ruta_base_de_datos_original_str or not Path(ruta_base_de_datos_original_str).exists():
        print(f"Error Crítico: La ruta al archivo 'BASE_DE_DATOS.xlsx' no es válida.")
        return
        
    ruta_base_de_datos_original_obj = Path(ruta_base_de_datos_original_str)
    df_base_de_datos_original = pd.read_excel(ruta_base_de_datos_original_obj, dtype=str, keep_default_na=False, na_filter=False)
    df_base_de_datos_original.columns = df_base_de_datos_original.columns.str.strip()
    print(f"  - Archivo fuente de datos ('{ruta_base_de_datos_original_obj.name}') cargado.")

    # --- CAMBIO 2: Se pide al usuario la FECHA DE OFICIO una sola vez ---
    print("\n" + "="*70)
    fecha_de_oficio_global = input("--> Ingrese la FECHA DE OFICIO que se usará para todos los registros: ").strip()
    if not fecha_de_oficio_global:
        print("  Advertencia: No se ingresó fecha de oficio. La columna quedará vacía.")
        fecha_de_oficio_global = ""
    print(f"  - Se usará '{fecha_de_oficio_global}' como FECHA DE OFICIO para todo el reporte.")
    print("="*70 + "\n")

    # Leer y procesar el archivo de entrega seleccionado
    df_notificados = pd.DataFrame()
    try:
        # --- CAMBIO 3: Se ajusta `header=1` para leer los encabezados desde la fila 2 de Excel ---
        df_entrega = pd.read_excel(ruta_reporte_oficios_excel, sheet_name=0, header=0)
        df_entrega.columns = df_entrega.columns.str.strip()
        print(f"  - Columnas leídas del archivo de entrega: {df_entrega.columns.tolist()}")

        columnas_entrega_requeridas = [COL_ENTREGA_EXPEDIENTE, COL_ENTREGA_FECHA_NOTIF, COL_ENTREGA_ESTATUS]
        for col_req in columnas_entrega_requeridas:
            if col_req not in df_entrega.columns:
                print(f"  Error GRAVE: La columna requerida '{col_req}' NO existe en '{ruta_reporte_oficios_excel.name}'.")
                print(f"  Asegúrese de que el archivo tenga los encabezados correctos en la fila 2.")
                return
        
        # --- CAMBIO 4: El filtro ahora se basa en la columna ESTATUS y el valor definido ---
        df_notificados = df_entrega[
            df_entrega[COL_ENTREGA_ESTATUS].astype(str).str.strip().str.upper() == VALOR_ESTATUS_NOTIFICADO.upper()
        ].copy()
        
        if df_notificados.empty:
            print(f"  - No se encontraron registros con {COL_ENTREGA_ESTATUS}='{VALOR_ESTATUS_NOTIFICADO}' en '{ruta_reporte_oficios_excel.name}'.")
            print("  - No se generará el reporte.")
            return
        print(f"  - {len(df_notificados)} oficios encontrados como '{VALOR_ESTATUS_NOTIFICADO}'.")

    except Exception as e:
        print(f"Error al leer o procesar el archivo de entrega ('{ruta_reporte_oficios_excel.name}'): {e}")
        return

    # Preparar datos para el nuevo Excel
    fecha_actual_sistema = datetime.now()
    fecha_reporte_formateada_str = formatear_fecha_reporte(fecha_actual_sistema)
    
    datos_para_nuevo_excel = []
    contador_filas_excel_salida = 0

    print("\n  Procesando expedientes notificados para el nuevo reporte:")
    for _, fila_notif in df_notificados.iterrows():
        # --- CAMBIO 5: Se obtienen los datos de las nuevas columnas ---
        expediente_actual = str(fila_notif[COL_ENTREGA_EXPEDIENTE]).strip()
        fecha_notif_actual = str(fila_notif[COL_ENTREGA_FECHA_NOTIF]).strip()

        print(f"  -> Procesando Expediente: {expediente_actual}")
        
        # La función para obtener placas y montos ahora usa el EXPEDIENTE como ID.
        # Se renombra el primer argumento de la función para mayor claridad.
        lista_placas_montos_oficio = obtener_datos_placas_montos_de_bd_multas(
            expediente_actual, 
            df_base_de_datos_original 
        )

        if not lista_placas_montos_oficio or \
           (len(lista_placas_montos_oficio) == 1 and lista_placas_montos_oficio[0]['PLACA'] in ['OFICIO NO EN BD MAESTRA', 'BD MAESTRA VACÍA', 'DATOS FALTANTES EN BD']):
            print(f"    - Advertencia: No se obtuvieron placas/montos válidos para el expediente {expediente_actual}. Se creará una fila placeholder.")
            contador_filas_excel_salida += 1
            datos_para_nuevo_excel.append({
                "NUMERO": contador_filas_excel_salida,
                "FECHA DE REPORTE": fecha_reporte_formateada_str,
                "DESPACHO": TEXTO_DESPACHO_FIJO,
                "NO. DE CONTROL Ó DOCUMENTO": expediente_actual,
                "FECHA DE OFICIO": fecha_de_oficio_global, # <-- Usa la fecha global
                "FECHA DE NOTIFICACION": fecha_notif_actual,
                "PLACA": 'VERIFICAR BASE_DE_DATOS.xlsx',
                "MONTO": 0.00
            })
        else:
            df_placas_montos_temp = pd.DataFrame(lista_placas_montos_oficio)
            df_placas_montos_temp['MONTO'] = pd.to_numeric(df_placas_montos_temp['MONTO'], errors='coerce').fillna(0.0)
            df_agrupado_por_placa = df_placas_montos_temp.groupby('PLACA', as_index=False)['MONTO'].sum()

            for _, fila_placa_agrupada in df_agrupado_por_placa.iterrows():
                contador_filas_excel_salida += 1
                placa_unica_actual = fila_placa_agrupada['PLACA']
                monto_sumado_placa = fila_placa_agrupada['MONTO']
                
                datos_para_nuevo_excel.append({
                    "NUMERO": contador_filas_excel_salida,
                    "FECHA DE REPORTE": fecha_reporte_formateada_str,
                    "DESPACHO": TEXTO_DESPACHO_FIJO,
                    "NO. DE CONTROL Ó DOCUMENTO": expediente_actual,
                    "FECHA DE OFICIO": fecha_de_oficio_global, # <-- Usa la fecha global
                    "FECHA DE NOTIFICACION": fecha_notif_actual,
                    "PLACA": placa_unica_actual,
                    "MONTO": monto_sumado_placa
                })

    if not datos_para_nuevo_excel:
        print("\nNo se generaron datos válidos para el reporte final.")
        return

    df_reporte_final_despachos = pd.DataFrame(datos_para_nuevo_excel)
    
    # Guardado del archivo final (sin cambios en esta parte)
    carpeta_salida_final_reportes = carpeta_datos_despacho / NOMBRE_CARPETA_REPORTES_GENERADOS
    carpeta_salida_final_reportes.mkdir(parents=True, exist_ok=True)
    nombre_archivo_salida_final = f"RELACION_DESPACHOS_GENERADA_{fecha_actual_sistema.strftime('%Y%m%d_%H%M%S')}.xlsx"
    ruta_salida_excel_final = carpeta_salida_final_reportes / nombre_archivo_salida_final

    try:
        df_reporte_final_despachos.to_excel(ruta_salida_excel_final, index=False, engine='openpyxl')
        print(f"\n--- Reporte de Despachos Generado Exitosamente ---")
        print(f"  Guardado en: {ruta_salida_excel_final}")
    except Exception as e:
        print(f"Error al guardar el nuevo archivo Excel '{nombre_archivo_salida_final}': {e}")



def generar_documento_faltante_para_escaneo_multas(
    oficio_norm_faltante, # Este es DIDCFMT...
    df_csv_principal_preparado, 
    directorio_temporal_para_generado_obj, 
    config_multas_actual 
    ):
    global failed_operations_log
    
    print(f"            - [Generar Faltante Scan] Intentando generar documento para Oficio (Norm): {oficio_norm_faltante}")

    if df_csv_principal_preparado is None or df_csv_principal_preparado.empty:
        # ... (manejo de error igual) ...
        return None
    
    if 'OFICIO_NORM_CSV' not in df_csv_principal_preparado.columns:
        # ... (manejo de error igual) ...
        return None
        
    registros_csv_oficio_faltante = df_csv_principal_preparado[
        df_csv_principal_preparado['OFICIO_NORM_CSV'] == oficio_norm_faltante # Comparar DIDCFMT con DIDCFMT
    ].copy()

    if registros_csv_oficio_faltante.empty:
        # ... (manejo de error igual) ...
        return None

    try:
        plantilla_path_del_modo_str = config_multas_actual.get("template_file_path")
        # ... (resto de la lógica de plantilla igual) ...
        if not plantilla_path_del_modo_str: raise ValueError("Ruta plantilla no definida.")
        plantilla_path_del_modo = Path(plantilla_path_del_modo_str)
        if not plantilla_path_del_modo.is_file(): raise FileNotFoundError(f"Plantilla no encontrada: {plantilla_path_del_modo}")
    except Exception as e_plantilla:
        print(f"            - (!) Error CRÍTICO (Generar Faltante Scan): Problema con la plantilla para Oficio {oficio_norm_faltante}. Detalle: {e_plantilla}")
        failed_operations_log[f"generar_faltante_scan_{oficio_norm_faltante}"].append({"error": f"Error de plantilla: {e_plantilla}", "timestamp": datetime.now().isoformat()})
        return None

    # ---- MODIFICACIÓN CLAVE para el placeholder y nombre de archivo ----
    # Convertir el oficio normalizado (DIDCFMT...) al formato con barras para el placeholder y el nombre del archivo temporal
    oficio_faltante_con_barras = convertir_a_formato_con_barras(oficio_norm_faltante)
    
    nombre_contrib_para_nombre = registros_csv_oficio_faltante.iloc[0].get('NOMBRE_COMPLETO_CALC', 'CONTRIBUYENTE_DESCONOCIDO')
    
    # Usar el formato con barras para el nombre del archivo temporal
    nombre_base_pdf_temp_generado = limpiar_texto(f"TEMP_GEN_SCAN_{oficio_faltante_con_barras}_{nombre_contrib_para_nombre}")
    # ---- FIN MODIFICACIÓN ----
    ruta_pdf_generado_temporal_obj = directorio_temporal_para_generado_obj / f"{nombre_base_pdf_temp_generado}_{time.time_ns()}.pdf"

    print(f"            - Preparando para generar PDF para Oficio (Norm: {oficio_norm_faltante}, Placeholder: {oficio_faltante_con_barras}) usando plantilla: {plantilla_path_del_modo.name}")
    
    pdf_generado_ok_scan, num_paginas_generado = generar_un_documento_multas_v2(
        datos_oficio_df=registros_csv_oficio_faltante,
        plantilla_path_obj=plantilla_path_del_modo, 
        ruta_pdf_destino_obj=ruta_pdf_generado_temporal_obj,
        modo_extraccion_paginas="TODAS", 
        oficio_slashed_para_placeholder=oficio_faltante_con_barras # MODIFICADO: Pasar formato con barras
    )

    if pdf_generado_ok_scan:
        print(f"            -> ÉXITO (Generar Faltante Scan): Documento temporal ({num_paginas_generado if num_paginas_generado is not None else 'N/A'} págs) generado para {oficio_faltante_con_barras} en {ruta_pdf_generado_temporal_obj.name}.")
        return ruta_pdf_generado_temporal_obj
    else:
        print(f"            -> FALLO (Generar Faltante Scan): No se pudo generar el documento temporal para {oficio_faltante_con_barras}.")
        failed_operations_log[f"generar_faltante_scan_FALLO_FINAL_{oficio_faltante_con_barras}"].append({"error": "generar_un_documento_multas_v2 reportó fallo.", "timestamp": datetime.now().isoformat()})
        return None


def merge_pdfs_multas(ruta_pdf_generado_obj, ruta_pdf_escaneado_obj, ruta_pdf_fusionado_salida_obj):
    """
    Fusiona un PDF generado (puede ser None) y un PDF escaneado en un único archivo de salida.
    El PDF generado (si se proporciona y existe) va primero.
    Retorna True si la fusión fue exitosa y se creó el archivo, False en caso contrario.
    """
    writer = PyPDF2Writer()
    fusion_realizada_con_contenido = False

    # 1. Añadir páginas del PDF generado (si existe y tiene contenido)
    if ruta_pdf_generado_obj and ruta_pdf_generado_obj.exists() and ruta_pdf_generado_obj.stat().st_size > 0:
        try:
            with open(ruta_pdf_generado_obj, 'rb') as f_gen:
                reader_gen = PyPDF2Reader(f_gen, strict=False)
                if reader_gen.pages: 
                    for page in reader_gen.pages:
                        writer.add_page(page)
                    fusion_realizada_con_contenido = True
                    # print(f"            [+] Páginas del generado '{ruta_pdf_generado_obj.name}' añadidas a la fusión.")
        except PyPDF2ReadError as e_merge_gen_read:
             print(f"            [!] Advertencia (PyPDF2): No se pudieron leer páginas del PDF generado '{ruta_pdf_generado_obj.name}': {e_merge_gen_read}")
             failed_operations_log[f"fusion_pdf_error_lectura_generado"].append({'archivo': ruta_pdf_generado_obj.name, 'error': str(e_merge_gen_read)})
        except Exception as e_merge_gen_other:
            print(f"            [!] Advertencia (Otro): No se pudieron añadir páginas del PDF generado '{ruta_pdf_generado_obj.name}': {e_merge_gen_other}")
            failed_operations_log[f"fusion_pdf_error_añadir_generado"].append({'archivo': ruta_pdf_generado_obj.name, 'error': str(e_merge_gen_other)})
    
    # 2. Añadir páginas del PDF escaneado (debe existir y tener contenido)
    if ruta_pdf_escaneado_obj and ruta_pdf_escaneado_obj.exists() and ruta_pdf_escaneado_obj.stat().st_size > 0:
        try:
            with open(ruta_pdf_escaneado_obj, 'rb') as f_scan:
                reader_scan = PyPDF2Reader(f_scan, strict=False)
                if reader_scan.pages:
                    for page in reader_scan.pages:
                        writer.add_page(page)
                    fusion_realizada_con_contenido = True
                    # print(f"            [+] Páginas del escaneado '{ruta_pdf_escaneado_obj.name}' añadidas a la fusión.")
                else: # Problema si el escaneado está vacío
                    print(f"            [!] Error: PDF escaneado '{ruta_pdf_escaneado_obj.name}' no tiene páginas.")
                    failed_operations_log[f"fusion_pdf_escaneado_sin_paginas"].append({'archivo': ruta_pdf_escaneado_obj.name})
                    if not writer.pages: return False # Si no hay nada que escribir
        except PyPDF2ReadError as e_merge_scan_read:
            print(f"            [!] Error (PyPDF2): No se pudieron leer páginas del PDF escaneado '{ruta_pdf_escaneado_obj.name}': {e_merge_scan_read}")
            failed_operations_log[f"fusion_pdf_error_lectura_escaneado"].append({'archivo': ruta_pdf_escaneado_obj.name, 'error': str(e_merge_scan_read)})
            return False # Error crítico si el escaneado principal no se puede leer
        except Exception as e_merge_scan_other:
            print(f"            [!] Error (Otro): No se pudieron añadir páginas del PDF escaneado '{ruta_pdf_escaneado_obj.name}': {e_merge_scan_other}")
            failed_operations_log[f"fusion_pdf_error_añadir_escaneado"].append({'archivo': ruta_pdf_escaneado_obj.name, 'error': str(e_merge_scan_other)})
            return False
    else: # Si el PDF escaneado no existe o está vacío
        print(f"            [!] Error CRÍTICO: PDF escaneado fuente '{ruta_pdf_escaneado_obj}' no encontrado o vacío. No se puede fusionar.")
        failed_operations_log[f"fusion_pdf_escaneado_no_encontrado"].append({'archivo': str(ruta_pdf_escaneado_obj)})
        if not writer.pages: return False

    # Guardar el PDF fusionado solo si se añadieron páginas
    if writer.pages:
        try:
            ruta_pdf_fusionado_salida_obj.parent.mkdir(parents=True, exist_ok=True)
            with open(ruta_pdf_fusionado_salida_obj, 'wb') as f_out_merge:
                writer.write(f_out_merge)
            return True
        except Exception as e_write_merge:
            print(f"            [!] Error al escribir PDF fusionado '{ruta_pdf_fusionado_salida_obj.name}': {e_write_merge}")
            failed_operations_log[f"fusion_pdf_error_escritura"].append({'archivo_salida': ruta_pdf_fusionado_salida_obj.name, 'error': str(e_write_merge)})
            return False
    else: # No se añadieron páginas
        print(f"            [!] No hay páginas para escribir en el PDF fusionado '{ruta_pdf_fusionado_salida_obj.name}'. No se creó el archivo.")
        return False


def process_single_oficio_scan_multas(
    oficio_escaneado_str_original,
    oficio_escaneado_str_norm,
    ruta_pdf_temporal_grupo_escaneado_obj,
    num_paginas_grupo_escaneado,
    metodo_extraccion_ocr,
    df_bd_maestra_global,
    df_csv_principal_preparado,
    nombre_pdf_fuente_original_str,
    directorio_temporal_principal_obj,
    directorio_salida_escaneos_oficio_obj,
    config_multas_actual,
    funcion_de_subida
):
    """
    Procesa un único oficio extraído del escaneo, implementando la lógica para
    oficios existentes y "huérfanos" con todas las actualizaciones de BD requeridas.
    """
    print(f"        - Procesando Oficio Escaneado (Norm: '{oficio_escaneado_str_norm}'), Fuente: {nombre_pdf_fuente_original_str}")
    
    # --- 1. PREPARACIÓN INICIAL ---
    oficio_slashed_para_reporte = convertir_a_formato_con_barras(oficio_escaneado_str_norm)
    datos_para_excel = {
        "Oficio": oficio_slashed_para_reporte,
        "NombreArchivo": "", "Paginas": num_paginas_grupo_escaneado,
        "Direccion": "", "CP": "", "Monto": "", "Error": ""
    }
    df_bd_maestra_actualizada_local = df_bd_maestra_global.copy()
    ruta_pdf_oficial_para_fusion = None
    nombre_contrib_final = "CONTRIBUYENTE_DESCONOCIDO"

    # --- 2. LÓGICA CENTRAL: ¿EXISTE O ES HUÉRFANO? ---
    entrada_bd_maestra = df_bd_maestra_actualizada_local[
        df_bd_maestra_actualizada_local['OFICIO'].apply(normalizar_oficio_multas) == oficio_escaneado_str_norm
    ]

    if not entrada_bd_maestra.empty:
        # --- CAMINO A: OFICIO ENCONTRADO EN BD MAESTRA ---
        print(f"            - Oficio encontrado en BD Maestra. Buscando PDF generado...")
        fila_bd = entrada_bd_maestra.iloc[0]
        nombre_contrib_final = fila_bd.get("NOMBRE CONTRIBUYENTE", "DESCONOCIDO_EN_BD")
        cp = fila_bd.get("CP", "SIN_CP")
        
        datos_para_excel.update({
            "Direccion": fila_bd.get("DIRECCION", ""),
            "CP": cp, "Monto": formatear_valor_celda(fila_bd.get("MONTO", 0.0))
        })
        
        ruta_pdf_oficial_para_fusion = buscar_pdf_oficial_generado(oficio_escaneado_str_norm, config_multas_actual, nombre_contrib_final, cp)
        
        if not ruta_pdf_oficial_para_fusion:
            print(f"            - PDF oficial no encontrado. Generando sobre la marcha como respaldo...")
            ruta_pdf_oficial_para_fusion = generar_documento_faltante_para_escaneo_multas(
                oficio_escaneado_str_norm, df_csv_principal_preparado, directorio_temporal_principal_obj, config_multas_actual
            )
            
    else:
        # --- CAMINO B: OFICIO "HUÉRFANO" ---
        print(f"            - (!) Oficio no encontrado en BD Maestra. Buscando en BASE_DE_DATOS.xlsx...")
        datos_csv = buscar_datos_oficio_en_csv_multas(oficio_escaneado_str_norm, df_csv_principal_preparado)
        
        if datos_csv.get('encontrado_en_csv'):
            print(f"            - Oficio huérfano encontrado en CSV. Generando y añadiendo a Lote A...")
            nombre_contrib_final = datos_csv.get('nombre_contribuyente', 'DESCONOCIDO_EN_CSV')
            
            # Generar el documento que falta
            ruta_pdf_oficial_para_fusion = generar_documento_faltante_para_escaneo_multas(
                oficio_escaneado_str_norm, df_csv_principal_preparado, directorio_temporal_principal_obj, config_multas_actual
            )
            
            if not ruta_pdf_oficial_para_fusion:
                datos_para_excel["Error"] = "FALLO_GENERACION_HUERFANO"
            else:
                # Calcular todos los datos necesarios para la nueva fila
                registros = df_csv_principal_preparado[df_csv_principal_preparado['OFICIO_NORM_CSV'] == oficio_escaneado_str_norm]
                num_registros = len(registros)
                monto_total = registros['IMPORTE_numeric'].sum()
                paginas_generadas = contar_paginas_pdf(ruta_pdf_oficial_para_fusion)
                hojas_calculadas = math.ceil(paginas_generadas / 2) if paginas_generadas is not None else pd.NA

                # Asignar nuevo ID de Lote A
                ultimo_numero_A = obtener_ultimo_id_de_lote_especifico(df_bd_maestra_actualizada_local, "A")
                nuevo_id = f"A-{(ultimo_numero_A + 1):03d}"
                print(f"            - Asignando nuevo ID de Lote A: {nuevo_id}")

                # Preparar el nuevo registro COMPLETO
                nuevo_registro = {
                    "ID": nuevo_id,
                    "OFICIO": oficio_slashed_para_reporte,
                    "NOMBRE CONTRIBUYENTE": nombre_contrib_final,
                    "DIRECCION": datos_csv.get('direccion', ''),
                    "CP": datos_csv.get('cp', ''),
                    "MONTO": monto_total,
                    "ESTADO": "IMPRESION COMPLETADA",
                    "HOJAS POR DOCUMENTO": hojas_calculadas,
                    "REGISTROS EN BD GENERACION": num_registros,
                    "FECHA IMPRESION": "XXX",
                    "BASE DE DATOS ESCANEO": f"BD_Escaneo_{Path(nombre_pdf_fuente_original_str).stem}_Multas.xlsx",
                    "Ruta PDF Generado": str(ruta_pdf_oficial_para_fusion.relative_to(config_multas_actual["base_path"])).replace('\\', '/'),
                    "MOVIMIENTO": "ESCANEADO (EN REVISION)",
                }
                df_bd_maestra_actualizada_local = actualizar_o_agregar_registro_bd_multas(df_bd_maestra_actualizada_local, nuevo_registro)
                datos_para_excel["Error"] = "GENERADO Y AÑADIDO A BD"
        else:
            print(f"            - (!) Error: Oficio huérfano '{oficio_escaneado_str_norm}' tampoco fue encontrado en BASE_DE_DATOS.xlsx.")
            datos_para_excel["Error"] = "OFICIO NO ENCONTRADO EN NINGUNA FUENTE"
            return datos_para_excel, df_bd_maestra_global

    # --- 3. FUSIÓN Y ACTUALIZACIÓN FINAL ---
    if not ruta_pdf_oficial_para_fusion:
        print(f"            - (!) Error Crítico: No se pudo encontrar ni generar el PDF oficial. Abortando fusión.")
        datos_para_excel["Error"] += "; FALLO FUSION"
        return datos_para_excel, df_bd_maestra_actualizada_local

    nombre_archivo_fusionado_base = limpiar_texto(f"{oficio_slashed_para_reporte}_{nombre_contrib_final}")
    ruta_pdf_fusionado_final = directorio_salida_escaneos_oficio_obj / f"{nombre_archivo_fusionado_base}.pdf"

    if merge_pdfs_multas(ruta_pdf_oficial_para_fusion, ruta_pdf_temporal_grupo_escaneado_obj, ruta_pdf_fusionado_final):
        print(f"            -> Fusión exitosa: {ruta_pdf_fusionado_final.name}")
        datos_para_excel["NombreArchivo"] = ruta_pdf_fusionado_final.name
        
        ruta_relativa_pdf_escaneado = str(ruta_pdf_fusionado_final.relative_to(Path(config_multas_actual["output_docs_path"]))).replace('\\', '/')
        
        # Actualizar la BD con la ruta del PDF FUSIONADO
        indices_actualizar = df_bd_maestra_actualizada_local.index[df_bd_maestra_actualizada_local['OFICIO'].apply(normalizar_oficio_multas) == oficio_escaneado_str_norm].tolist()
        if indices_actualizar:
            df_bd_maestra_actualizada_local.loc[indices_actualizar[-1], 'Ruta PDF Escaneado'] = ruta_relativa_pdf_escaneado
        
        if funcion_de_subida:
            funcion_de_subida(ruta_pdf_fusionado_final, config_multas_actual)
    else:
        datos_para_excel["Error"] += "; FALLO FUSION"
    
    return datos_para_excel, df_bd_maestra_actualizada_local


def generar_documentos_multas_core(
    df_datos_principales,
    df_bd_maestra_actual,
    config_multas_actual,
    modo_generacion_menu_solicitado,
    max_archivos_a_generar,
    letra_lote,
    contador_inicial_lote
    ):
    """
    Función principal para generar documentos de Multas, llamada por el generador maestro.
    Asegura que el OFICIO se maneje en formato DI/DCF/MT/XXXXXX para la generación
    de contenido y para la Base de Datos Maestra.
    """
    global failed_operations_log # Si usas un log global de errores
    if 'failed_operations_log' not in globals(): # Asegurar que exista, sino crear uno local
        from collections import defaultdict
        failed_operations_log = defaultdict(list)
        
    print(f"\n--- (Multas Logic Core) Iniciando Generación en Modo Menú: {modo_generacion_menu_solicitado} ---")

    plantilla_path = Path(config_multas_actual["template_file_path"])
    carpeta_principal_salida = Path(config_multas_actual["output_docs_path"])
    # failed_operations_log.clear() # Decide si limpiar el log global aquí o manejarlo externamente

    if not configurar_locale_es(): # Asume que esta función existe y configura el locale
        print("(!) ADVERTENCIA (Multas Logic): No se pudo configurar el locale a español.")

    # Constantes de estado y modo desde config_multas_actual
    EST_PENDIENTE = config_multas_actual["ESTADO_PENDIENTE"]
    EST_ERROR_GEN = config_multas_actual["ESTADO_ERROR_GENERACION"]
    EST_GEN_COMPLETO = config_multas_actual["ESTADO_GEN_COMPLETO"]
    EST_GEN_ULTIMA = config_multas_actual["ESTADO_GEN_ULTIMA"]
    EST_GEN_RESTO = config_multas_actual["ESTADO_GEN_RESTO"]
    EST_IMP_ULTIMA = config_multas_actual["ESTADO_IMP_ULTIMA"] # Usado para filtrar en modo RESTO
    # EST_IMP_COMPLETO = config_multas_actual["ESTADO_IMP_COMPLETO"] # Usado para excluir

    # Modos de generación definidos en generador.py (asegúrate que las claves coincidan)
    MODO_COMPLETO_INTERNO = config_multas_actual.get("MODO_GENERACION_COMPLETO", "COMPLETO")
    MODO_ULTIMA_INTERNO = config_multas_actual.get("MODO_GENERACION_ULTIMA", "ULTIMA")
    MODO_RESTO_INTERNO = config_multas_actual.get("MODO_GENERACION_RESTO", "RESTO")
    MODO_ESPECIFICOS_INTERNO = config_multas_actual.get("MODO_GENERACION_ESPECIFICOS", "ESPECIFICOS")


    modo_extraccion_pdf_logica = ""
    estado_objetivo_bd_si_exito_logica = ""
    estados_validos_para_procesar_filtro_logica = []
    estados_a_excluir_filtro_bd_logica = [config_multas_actual.get("ESTADO_IMP_COMPLETO")] # Siempre excluir los ya impresos completos

    if modo_generacion_menu_solicitado == MODO_COMPLETO_INTERNO:
        modo_extraccion_pdf_logica = "TODAS"
        estado_objetivo_bd_si_exito_logica = EST_GEN_COMPLETO
        estados_validos_para_procesar_filtro_logica = [EST_PENDIENTE, EST_ERROR_GEN, EST_GEN_ULTIMA, EST_IMP_ULTIMA, EST_GEN_RESTO]
        estados_a_excluir_filtro_bd_logica.append(EST_GEN_COMPLETO) # No regenerar si ya está completo
    elif modo_generacion_menu_solicitado == MODO_ULTIMA_INTERNO:
        modo_extraccion_pdf_logica = "ULTIMA"
        estado_objetivo_bd_si_exito_logica = EST_GEN_ULTIMA
        estados_validos_para_procesar_filtro_logica = [EST_PENDIENTE, EST_ERROR_GEN, EST_GEN_COMPLETO] # Puede generar ultima desde pendiente, error o completo
        estados_a_excluir_filtro_bd_logica.extend([EST_GEN_ULTIMA, EST_IMP_ULTIMA, EST_GEN_RESTO])
    elif modo_generacion_menu_solicitado == MODO_RESTO_INTERNO:
        modo_extraccion_pdf_logica = "RESTO"
        estado_objetivo_bd_si_exito_logica = EST_GEN_RESTO
        estados_validos_para_procesar_filtro_logica = [EST_IMP_ULTIMA, EST_GEN_ULTIMA] # Solo procesar los que ya tienen última generada o impresa
        estados_a_excluir_filtro_bd_logica.extend([EST_GEN_RESTO, EST_GEN_COMPLETO, EST_PENDIENTE, EST_ERROR_GEN])
    elif modo_generacion_menu_solicitado == MODO_ESPECIFICOS_INTERNO:
        modo_extraccion_pdf_logica = "TODAS" # Para específicos, usualmente se regenera completo
        estado_objetivo_bd_si_exito_logica = EST_GEN_COMPLETO
        estados_validos_para_procesar_filtro_logica = [] # Sin filtro de estado inicial, se procesan los IDs dados
        estados_a_excluir_filtro_bd_logica = [] # No excluir ninguno por estado (se procesan los especificados)
    else:
        print(f"     - (!) Error (Multas Logic): Modo de generación '{modo_generacion_menu_solicitado}' no reconocido.")
        return df_bd_maestra_actual

    # 2. Preparar df_datos_principales (el DataFrame que viene de BASE_DE_DATOS.xlsx)
    try:
        col_importe_datos = config_multas_actual.get("col_importe_csv", "IMPORTE") # Nombre de columna en CSV
        col_cuotas_datos = config_multas_actual.get("col_cuotas_csv", "CUOTAS")    # Nombre de columna en CSV
        col_uma_sm_datos = config_multas_actual.get("col_uma_sm_csv", "Uma/SM")    # Nombre de columna en CSV
        col_nombre_datos = config_multas_actual.get("col_nombre_base_csv", "NOMBRE")
        col_apaterno_datos = config_multas_actual.get("col_apaterno_csv", "APELLIDO PATERNO")
        col_amaterno_datos = config_multas_actual.get("col_amaterno_csv", "APELLIDO MATERNO")
        col_direccion_datos_multas = config_multas_actual.get("col_direccion_completa_csv", "DIRECCION")
        # Esta es la columna del OFICIO ORIGINAL tal como viene en BASE_DE_DATOS.xlsx
        col_oficio_original_csv = config_multas_actual.get("id_col_csv", "OFICIO")

        df_datos_preparados = df_datos_principales.copy()

        # Crear columnas numéricas limpias
        for target_col_numeric, source_col_name_csv in zip(
            ['IMPORTE', 'CUOTAS', 'Uma/SM'], # Nombres estándar para las columnas _numeric
            [col_importe_datos, col_cuotas_datos, col_uma_sm_datos]
        ):
            if source_col_name_csv in df_datos_preparados.columns:
                df_datos_preparados[f'{target_col_numeric}_numeric'] = pd.to_numeric(
                    df_datos_preparados[source_col_name_csv].astype(str).str.replace(r'[$,]', '', regex=True).str.strip(),
                    errors='coerce'
                ).fillna(0.0)
            else: # Si la columna fuente no existe en el CSV, crear la _numeric con ceros
                df_datos_preparados[f'{target_col_numeric}_numeric'] = 0.0
                print(f"  (*) Advertencia (Multas Logic): Columna fuente '{source_col_name_csv}' no encontrada para {target_col_numeric}_numeric. Usando 0.0.")

        # Crear NOMBRE_COMPLETO_CALC
        nombre_s = df_datos_preparados.get(col_nombre_datos, pd.Series(dtype=str, index=df_datos_preparados.index)).astype(str).str.strip()
        ap_s = df_datos_preparados.get(col_apaterno_datos, pd.Series(dtype=str, index=df_datos_preparados.index)).astype(str).str.strip()
        am_s = df_datos_preparados.get(col_amaterno_datos, pd.Series(dtype=str, index=df_datos_preparados.index)).astype(str).str.strip()
        df_datos_preparados['NOMBRE_COMPLETO_CALC'] = (nombre_s + ' ' + ap_s + ' ' + am_s).str.replace(r'\s+', ' ', regex=True).str.strip().replace('', 'CONTRIBUYENTE_DESCONOCIDO')
        
        # Crear OFICIO_NORM_CSV (versión normalizada a DIDCFMT... del oficio del CSV, para joins y comparaciones)
        if col_oficio_original_csv in df_datos_preparados.columns:
            df_datos_preparados['OFICIO_NORM_CSV'] = df_datos_preparados[col_oficio_original_csv].apply(normalizar_oficio_multas)
        else:
            df_datos_preparados['OFICIO_NORM_CSV'] = pd.Series([normalizar_oficio_multas(None)] * len(df_datos_preparados), dtype=str, index=df_datos_preparados.index)
            print(f"  (*) Advertencia (Multas Logic): Columna de OFICIO original '{col_oficio_original_csv}' no encontrada en datos de entrada. 'OFICIO_NORM_CSV' tendrá valores por defecto.")

        # Crear CP_CALC_CSV (extraído de la dirección) y DIRECCION_CSV (limpia)
        if col_direccion_datos_multas in df_datos_preparados.columns:
            df_datos_preparados['DIRECCION_CSV'] = df_datos_preparados[col_direccion_datos_multas].astype(str).str.strip()
            df_datos_preparados['CP_CALC_CSV'] = df_datos_preparados['DIRECCION_CSV'].apply(lambda x: extraer_cp_de_direccion(x) if pd.notna(x) else "SIN_CP").astype(str)
        else:
            df_datos_preparados['DIRECCION_CSV'] = "DIRECCION_DESCONOCIDA"
            df_datos_preparados['CP_CALC_CSV'] = "SIN_CP"
            print(f"  (*) Advertencia (Multas Logic): Columna de dirección '{col_direccion_datos_multas}' no encontrada en datos de entrada. Usando valores por defecto.")

    except KeyError as e_key:
        print(f"  (!) Error GRAVE (Multas Logic): Falta columna clave para preparación de datos CSV: {e_key}. Columnas disponibles: {df_datos_principales.columns.tolist()}.")
        return df_bd_maestra_actual
    except Exception as e_prep_csv:
        print(f"  (!) Error (Multas Logic) durante la preparación de columnas del CSV: {e_prep_csv}")
        return df_bd_maestra_actual

    # Calcular conteo de registros por oficio (basado en el OFICIO_NORM_CSV)
    col_id_normalizado_csv_para_conteo = 'OFICIO_NORM_CSV' 
    df_conteo_registros_por_oficio = pd.DataFrame()
    if col_id_normalizado_csv_para_conteo in df_datos_preparados.columns:
        df_conteo_registros_por_oficio = df_datos_preparados.groupby(col_id_normalizado_csv_para_conteo, as_index=False, observed=False).size().rename(columns={'size': 'CONTEO_CALCULADO'})
    else: # No debería ocurrir si la preparación anterior funcionó
        print(f"  (!) ADVERTENCIA (Multas Logic - Conteo): Columna '{col_id_normalizado_csv_para_conteo}' no encontrada en datos preparados. El conteo de registros en BD Maestra será NA.")

    # 3. Combinar df_datos_preparados con df_bd_maestra_actual para filtrar por estado
    col_id_en_bd_maestra = config_multas_actual["col_expediente"] # Nombre de la columna OFICIO en la BD Maestra
    col_estado_en_bd_maestra = config_multas_actual["col_estado_bd_maestra"]
    
    df_bd_maestra_para_merge = df_bd_maestra_actual.copy()
    col_id_bd_maestra_norm_temp = col_id_en_bd_maestra + "_NORM_BDM_TEMP" # Columna temporal para el ID normalizado de la BDM

    if col_id_en_bd_maestra in df_bd_maestra_para_merge.columns:
        df_bd_maestra_para_merge[col_id_en_bd_maestra] = df_bd_maestra_para_merge[col_id_en_bd_maestra].astype(str) # Asegurar que es string antes de .apply
        df_bd_maestra_para_merge[col_id_bd_maestra_norm_temp] = df_bd_maestra_para_merge[col_id_en_bd_maestra].apply(normalizar_oficio_multas) # Normalizar ID de BDM a DIDCFMT...
    else:
        print(f"  (!) Error Crítico (Multas Logic): Columna ID '{col_id_en_bd_maestra}' no existe en BD Maestra. El merge para filtrar por estado fallará.")
        return df_bd_maestra_actual

    # Seleccionar solo las columnas necesarias de la BD Maestra para el merge
    columnas_necesarias_bdm_para_merge = [col_id_en_bd_maestra, col_id_bd_maestra_norm_temp, col_estado_en_bd_maestra]
    columnas_necesarias_bdm_para_merge = [col for col in columnas_necesarias_bdm_para_merge if col in df_bd_maestra_para_merge.columns] # Solo las que existan

    df_combinado = pd.merge(
        df_datos_preparados, 
        df_bd_maestra_para_merge[columnas_necesarias_bdm_para_merge],
        left_on='OFICIO_NORM_CSV', # Este es DIDCFMT... del CSV preparado
        right_on=col_id_bd_maestra_norm_temp, # Este es DIDCFMT... de la BDM
        how='left', # Mantener todos los registros del CSV, añadir info de BDM si hay match
        suffixes=('_csv', '_bdm') # Sufijos para columnas con el mismo nombre
    )

    # Limpiar columnas temporales o con sufijos si es necesario
    if col_id_bd_maestra_norm_temp in df_combinado.columns:
        df_combinado.drop(columns=[col_id_bd_maestra_norm_temp], inplace=True)
    
    # Determinar la columna de estado final para el filtrado
    col_estado_final_para_filtrar = 'ESTADO_FINAL_PARA_FILTRAR'
    col_estado_bdm_post_merge = col_estado_en_bd_maestra # Si no hubo colisión de nombres
    if f"{col_estado_en_bd_maestra}_bdm" in df_combinado.columns: # Si hubo colisión y se añadió sufijo
        col_estado_bdm_post_merge = f"{col_estado_en_bd_maestra}_bdm"
    
    if col_estado_bdm_post_merge in df_combinado.columns:
        df_combinado[col_estado_final_para_filtrar] = df_combinado[col_estado_bdm_post_merge].fillna(EST_PENDIENTE)
    else: # Si la columna de estado no vino de la BDM (ej. oficio nuevo no en BDM)
        print(f"  (*) Advertencia (Multas Logic): Columna de estado '{col_estado_bdm_post_merge}' no encontrada después del merge. Asumiendo estado '{EST_PENDIENTE}' para todos los registros.")
        df_combinado[col_estado_final_para_filtrar] = EST_PENDIENTE
    
    # Filtrar registros según el modo de generación y el estado
    df_registros_a_procesar = df_combinado
    if modo_generacion_menu_solicitado != MODO_ESPECIFICOS_INTERNO: # Para ESPECIFICOS, se filtran por ID en generador.py
        condicion_filtro_estado = df_registros_a_procesar[col_estado_final_para_filtrar].isin(estados_validos_para_procesar_filtro_logica)
        if estados_a_excluir_filtro_bd_logica: 
            condicion_filtro_estado &= ~df_registros_a_procesar[col_estado_final_para_filtrar].isin(estados_a_excluir_filtro_bd_logica)
        df_registros_a_procesar = df_registros_a_procesar[condicion_filtro_estado].copy() # .copy() para evitar SettingWithCopyWarning

    if df_registros_a_procesar.empty:
        print(f"     - (Multas Logic) No hay oficios que cumplan los criterios de estado para el modo '{modo_generacion_menu_solicitado}'.")
        return df_bd_maestra_actual

    # Obtener lista de OFICIOS únicos (normalizados) a procesar, respetando el límite de max_archivos_a_generar
    # 'OFICIO_NORM_CSV' es la columna con los IDs normalizados (DIDCFMT...)
    col_id_normalizado_para_unicos = 'OFICIO_NORM_CSV' 
    
    if col_id_normalizado_para_unicos not in df_registros_a_procesar.columns:
        print(f"     - (!) Error Crítico (Multas Logic): Columna ID normalizada '{col_id_normalizado_para_unicos}' no presente en los registros filtrados. No se puede continuar.")
        return df_bd_maestra_actual

    oficios_unicos_normalizados = df_registros_a_procesar[col_id_normalizado_para_unicos].drop_duplicates().tolist()
    
    if max_archivos_a_generar > 0:
        oficios_unicos_a_procesar_final = oficios_unicos_normalizados[:max_archivos_a_generar]
    else: # Si es 0 o negativo, procesar todos los elegibles
        oficios_unicos_a_procesar_final = oficios_unicos_normalizados
        
    if not oficios_unicos_a_procesar_final:
        print(f"     - (Multas Logic) No quedaron oficios únicos para procesar después de aplicar el límite de cantidad.")
        return df_bd_maestra_actual
        
    print(f"     - (Multas Logic) {len(oficios_unicos_a_procesar_final)} oficios únicos (normalizados) se procesarán para el modo '{modo_generacion_menu_solicitado}'.")

    # Iniciar el proceso de generación para los oficios seleccionados
    documentos_generados_exitosamente_multas = 0
    df_bd_maestra_para_actualizar_logica = df_bd_maestra_actual.copy() # Trabajar sobre una copia de la BDM
    processed_oficio_ids_in_master_db_format = [] # <--- AÑADE ESTA LÍNEA (NUEVO)

    
    # Asegurar que las carpetas de salida existan
    sub_cp_modo = carpeta_principal_salida / "CP"; sub_vacias_modo = carpeta_principal_salida / "VACIAS"
    sub_cp_modo.mkdir(parents=True, exist_ok=True); sub_vacias_modo.mkdir(parents=True, exist_ok=True)

    letra_actual_para_ciclo = letra_lote
    contador_actual_para_ciclo = contador_inicial_lote

    for i, oficio_norm_actual_iteracion in enumerate(oficios_unicos_a_procesar_final):
        
        # --- 1. Obtener y preparar los datos para el oficio actual ---
        registros_csv_para_este_oficio = df_registros_a_procesar[
            df_registros_a_procesar[col_id_normalizado_para_unicos] == oficio_norm_actual_iteracion
        ].copy()

        if registros_csv_para_este_oficio.empty:
            continue

        monto_total_oficio = registros_csv_para_este_oficio['IMPORTE_numeric'].sum()
        fila_representativa_csv = registros_csv_para_este_oficio.iloc[0]
        oficio_original_del_input_csv = fila_representativa_csv.get(col_oficio_original_csv, oficio_norm_actual_iteracion)
        oficio_final_formato_barras = convertir_a_formato_con_barras(oficio_original_del_input_csv)
        
        processed_oficio_ids_in_master_db_format.append(oficio_final_formato_barras)
        
        datos_para_bd = {} # Diccionario para los datos que se escribirán en la BD

        # --- 2. Decidir si se genera el documento o se salta por monto bajo ---
        if monto_total_oficio < 179:
            # Camino A: El monto es muy bajo. No se genera PDF, solo se actualiza la BD.
            print(f"\n     --- (Multas Logic) SALTANDO Oficio (Norm: {oficio_norm_actual_iteracion}) por monto bajo ({monto_total_oficio:.2f}) ---")
            
            datos_para_bd = {
                "ESTADO": config_multas_actual.get("ESTADO_NO_GENERADO_MONTO_BAJO", "NO GENERADO (MENOR A 179)"),
                "MONTO": monto_total_oficio,
                "REGISTROS EN BD GENERACION": len(registros_csv_para_este_oficio),
                "OFICIO": oficio_final_formato_barras,
                "NOMBRE CONTRIBUYENTE": fila_representativa_csv.get('NOMBRE_COMPLETO_CALC', 'DESCONOCIDO'),
                "DIRECCION": fila_representativa_csv.get('DIRECCION_CSV', ''),
                "CP": fila_representativa_csv.get('CP_CALC_CSV', 'SIN_CP'),
                "HOJAS POR DOCUMENTO": pd.NA, # No hay documento, no hay hojas
                "Ruta PDF Generado": "",
                "MOVIMIENTO": "EN DESPACHO (GENERADO)"
            }
        else:
            # Camino B: El monto es suficiente. Se procede con la generación del documento.
            print(f"\n     --- (Multas Logic Core) ({i+1}/{len(oficios_unicos_a_procesar_final)}) Procesando Oficio: {oficio_final_formato_barras} ---")
            
            nombre_contrib_actual = fila_representativa_csv.get('NOMBRE_COMPLETO_CALC', 'DESCONOCIDO')
            cp_actual_csv = fila_representativa_csv.get('CP_CALC_CSV', 'SIN_CP')
            
            nombre_base_pdf_salida = limpiar_texto(f"{oficio_final_formato_barras}_{nombre_contrib_actual}")
            carpeta_salida_pdf_final = sub_vacias_modo
            if cp_actual_csv and cp_actual_csv not in ["SIN_CP", "Desconocido", "DIRECCION_DESCONOCIDA", "SIN_CP_SCAN"]:
                carpeta_salida_pdf_final = sub_cp_modo / f"CP_{cp_actual_csv}"
            carpeta_salida_pdf_final.mkdir(parents=True, exist_ok=True)
            ruta_pdf_destino_final_obj = carpeta_salida_pdf_final / f"{nombre_base_pdf_salida}.pdf"

            pdf_generado_ok, paginas_pdf_original = generar_un_documento_multas_v2(
                datos_oficio_df=registros_csv_para_este_oficio, 
                plantilla_path_obj=plantilla_path,
                ruta_pdf_destino_obj=ruta_pdf_destino_final_obj, 
                modo_extraccion_paginas=modo_extraccion_pdf_logica,
                oficio_slashed_para_placeholder=oficio_final_formato_barras
            )
            
            if pdf_generado_ok:
                documentos_generados_exitosamente_multas += 1

            estado_final_para_bd = EST_ERROR_GEN if not pdf_generado_ok else estado_objetivo_bd_si_exito_logica
            hojas_calculadas_para_bd = pd.NA
            if pdf_generado_ok and paginas_pdf_original is not None and paginas_pdf_original > 0:
                hojas_calculadas_para_bd = math.ceil(paginas_pdf_original / 2)
            
            ruta_relativa_pdf_para_bd = ""
            if pdf_generado_ok and ruta_pdf_destino_final_obj.exists():
                try:
                    ruta_relativa_pdf_para_bd = str(ruta_pdf_destino_final_obj.relative_to(Path(config_multas_actual["output_docs_path"])))
                except ValueError:
                    ruta_relativa_pdf_para_bd = ruta_pdf_destino_final_obj.name

            datos_para_bd = {
                "OFICIO": oficio_final_formato_barras,
                "NOMBRE CONTRIBUYENTE": nombre_contrib_actual,
                "DIRECCION": fila_representativa_csv.get('DIRECCION_CSV', ''),
                "CP": cp_actual_csv,
                "MONTO": monto_total_oficio,
                "ESTADO": estado_final_para_bd,
                "HOJAS POR DOCUMENTO": hojas_calculadas_para_bd, # <--- LÍNEA CORREGIDA
                "REGISTROS EN BD GENERACION": len(registros_csv_para_este_oficio),
                "Ruta PDF Generado": ruta_relativa_pdf_para_bd,
                "MOVIMIENTO": "EN DESPACHO (GENERADO)"
            }

        # --- 3. Lógica Unificada para Actualizar o Añadir en la BD Maestra ---
        col_id_en_bd_maestra = config_multas_actual["col_expediente"]
        indices_existentes_en_bdm = df_bd_maestra_para_actualizar_logica.index[
            df_bd_maestra_para_actualizar_logica[col_id_en_bd_maestra].apply(normalizar_oficio_multas) == oficio_norm_actual_iteracion
        ].tolist()

        if indices_existentes_en_bdm:
            # El registro YA EXISTE: Se actualizan los campos, respetando el ID existente.
            idx_actualizar_bdm = indices_existentes_en_bdm[-1]
            id_existente = df_bd_maestra_para_actualizar_logica.loc[idx_actualizar_bdm, 'ID']
            print(f"     -> (Multas Logic) RESPETANDO ID de lote existente '{id_existente}' para Oficio: {oficio_final_formato_barras}")
            for col_nombre, valor_nuevo in datos_para_bd.items():
                if col_nombre in df_bd_maestra_para_actualizar_logica.columns:
                    df_bd_maestra_para_actualizar_logica.loc[idx_actualizar_bdm, col_nombre] = valor_nuevo
        else:
            # El registro es NUEVO: Se genera un nuevo ID y se añade la fila completa.
            contador_actual_para_ciclo += 1
            if contador_actual_para_ciclo > 100:
                contador_actual_para_ciclo = 1
                letra_actual_para_ciclo = obtener_siguiente_letra_lote(letra_actual_para_ciclo)
            
            nuevo_id_generado = f"{letra_actual_para_ciclo}-{contador_actual_para_ciclo:03d}"
            datos_para_bd['ID'] = nuevo_id_generado
            print(f"     -> (Multas Logic) Asignando NUEVO ID de lote '{nuevo_id_generado}' para Oficio: {oficio_final_formato_barras}")
            
            for col_maestra in config_multas_actual["db_master_columns"]:
                if col_maestra not in datos_para_bd:
                    tipo_col = config_multas_actual["db_master_types"].get(col_maestra)
                    datos_para_bd[col_maestra] = pd.NA if tipo_col == 'Int64' else ""
            
            df_nueva_fila = pd.DataFrame([datos_para_bd], columns=config_multas_actual["db_master_columns"])
            df_bd_maestra_para_actualizar_logica = pd.concat([df_bd_maestra_para_actualizar_logica, df_nueva_fila], ignore_index=True)
            
        # --- 4. Re-asegurar tipos de datos al final de CADA iteración ---
        # Esto previene errores de tipos mixtos en el DataFrame para la siguiente iteración.
        try:
            tipos_a_aplicar = {k: v for k, v in config_multas_actual["db_master_types"].items() if k in df_bd_maestra_para_actualizar_logica.columns}
            if tipos_a_aplicar:
                df_bd_maestra_para_actualizar_logica = df_bd_maestra_para_actualizar_logica.astype(tipos_a_aplicar)
        except Exception as e_astype:
            print(f"    (*) Advertencia: Fallo menor en re-astype para Oficio {oficio_norm_actual_iteracion}. Error: {e_astype}")


    print(f"\n     --- (Multas Logic) Resumen del Lote (Modo Menú: {modo_generacion_menu_solicitado}) ---")
    print(f"     - Oficios únicos solicitados para procesar: {len(oficios_unicos_a_procesar_final)}")
    print(f"     - Documentos PDF generados exitosamente: {documentos_generados_exitosamente_multas}")

        # Ahora el "return" se ejecuta después de que el ciclo "for" ha terminado.
    return df_bd_maestra_para_actualizar_logica, processed_oficio_ids_in_master_db_format
    
    


def run_scan_and_process_multas(df_bd_maestra_global, df_csv_principal_global, config_multas_actual, funcion_de_subida=None):
    """
    Función principal que orquesta el escaneo de PDFs de Multas, procesamiento y registro.
    Utiliza rutas específicas del modo desde config_multas_actual y prepara los datos del CSV.
    """
    global failed_operations_log
    print("\n--- Iniciando Modo Escaneo de Documentos de Multas ---")

    # 1. Definir rutas de escaneo basadas en config_multas_actual
    try:
        base_path_modo = Path(config_multas_actual["base_path"])
        output_docs_path_modo = Path(config_multas_actual["output_docs_path"])
    except KeyError as e_path_config:
        print(f"    - (!) ERROR CRÍTICO: Falta configuración de ruta esencial ('base_path' o 'output_docs_path') en config_multas_actual. Error: {e_path_config}")
        return df_bd_maestra_global
        
    source_pdf_folder_scan_multas_modo = base_path_modo / "PDFs_A_Escanear_Multas"
    output_scan_artifacts_multas_modo_root = output_docs_path_modo / "DOCUMENTOS_ESCANEADOS_MULTAS"

    print(f"    - Carpeta de PDFs a Escanear (Modo Multas): {source_pdf_folder_scan_multas_modo}")
    print(f"    - Carpeta Raíz de Salida para Escaneos (Modo Multas): {output_scan_artifacts_multas_modo_root}")

    # 2. Validar dependencias críticas
    try:
        pytesseract.get_tesseract_version() # Forzará error si Tesseract no está bien
        if 'fitz' not in sys.modules: # Chequeo simbólico
            print("    - (!) Advertencia: Módulo 'fitz' (PyMuPDF) no detectado en sys.modules.")
    except (ImportError, NameError, pytesseract.TesseractNotFoundError, AttributeError) as e_dep_ocr_scan:
        print(f"    - (!) ERROR CRÍTICO: Falta dependencia de OCR o Tesseract no está configurado: {e_dep_ocr_scan}")
        print("        El modo Escaneo de Multas no puede continuar.")
        return df_bd_maestra_global

    # 3. Validar y crear carpetas
    if not source_pdf_folder_scan_multas_modo.is_dir():
        print(f"    - (!) ERROR CRÍTICO: Carpeta PDFs fuente '{source_pdf_folder_scan_multas_modo}' NO existe para modo '{config_multas_actual['mode_name']}'.")
        return df_bd_maestra_global
    try:
        output_scan_artifacts_multas_modo_root.mkdir(parents=True, exist_ok=True)
    except OSError as e_mkdir:
        print(f"    - (!) ERROR CRÍTICO: No se pudo crear carpeta de salida escaneos '{output_scan_artifacts_multas_modo_root}'. Error: {e_mkdir}"); return df_bd_maestra_global

    lista_pdfs_a_escanear = list(source_pdf_folder_scan_multas_modo.glob("*.pdf"))
    if not lista_pdfs_a_escanear:
        print(f"    - No se encontraron PDFs en '{source_pdf_folder_scan_multas_modo}'."); return df_bd_maestra_global
    print(f"    - Se encontraron {len(lista_pdfs_a_escanear)} PDFs para escanear.")

    # 4. PREPARACIÓN DEL DATAFRAME CSV PRINCIPAL (df_csv_principal_global)
    df_csv_preparado_para_scan = pd.DataFrame() 
    if df_csv_principal_global is not None and not df_csv_principal_global.empty:
        print("    - Preparando datos del CSV principal para la lógica de escaneo...")
        try:
            df_csv_preparado_para_scan = df_csv_principal_global.copy()
            
            col_oficio_csv = config_multas_actual.get("id_col_csv", "OFICIO") 
            col_nombre_csv = config_multas_actual.get("col_nombre_base_csv", "NOMBRE")
            col_apaterno_csv = config_multas_actual.get("col_apaterno_csv", "APELLIDO PATERNO")
            col_amaterno_csv = config_multas_actual.get("col_amaterno_csv", "APELLIDO MATERNO")
            col_direccion_csv = config_multas_actual.get("col_direccion_completa_csv", "DIRECCION")
            col_importe_csv = config_multas_actual.get("col_importe_csv", "IMPORTE") # Usar la misma clave que en generación
            col_cuotas_csv = config_multas_actual.get("col_cuotas_csv", "CUOTAS")    
            col_uma_sm_csv = config_multas_actual.get("col_uma_sm_csv", "Uma/SM") 

            # Crear columnas numéricas
            for target_col_numeric, source_col_name in zip(
                ['IMPORTE', 'CUOTAS', 'Uma/SM'], [col_importe_csv, col_cuotas_csv, col_uma_sm_csv]
            ):
                if source_col_name in df_csv_preparado_para_scan.columns:
                    df_csv_preparado_para_scan[f'{target_col_numeric}_numeric'] = pd.to_numeric(
                        df_csv_preparado_para_scan[source_col_name].astype(str).str.replace(r'[$,]', '', regex=True).str.strip(),
                        errors='coerce').fillna(0.0)
                else:
                    print(f"      (*) Advertencia (Prep Scan CSV): Columna fuente '{source_col_name}' no en CSV. '{target_col_numeric}_numeric' será 0.0.")
                    df_csv_preparado_para_scan[f'{target_col_numeric}_numeric'] = 0.0
            
            # Crear NOMBRE_COMPLETO_CALC
            nombre_s = df_csv_preparado_para_scan.get(col_nombre_csv, pd.Series(dtype=str, index=df_csv_preparado_para_scan.index)).astype(str).str.strip()
            ap_s = df_csv_preparado_para_scan.get(col_apaterno_csv, pd.Series(dtype=str, index=df_csv_preparado_para_scan.index)).astype(str).str.strip()
            am_s = df_csv_preparado_para_scan.get(col_amaterno_csv, pd.Series(dtype=str, index=df_csv_preparado_para_scan.index)).astype(str).str.strip()
            df_csv_preparado_para_scan['NOMBRE_COMPLETO_CALC'] = (nombre_s + ' ' + ap_s + ' ' + am_s).str.replace(r'\s+', ' ', regex=True).str.strip().replace('', 'CONTRIBUYENTE_DESCONOCIDO_SCAN')

            # Crear OFICIO_NORM_CSV (el OFICIO normalizado del CSV de entrada)
            if col_oficio_csv in df_csv_preparado_para_scan.columns:
                df_csv_preparado_para_scan['OFICIO_NORM_CSV'] = df_csv_preparado_para_scan[col_oficio_csv].apply(normalizar_oficio_multas)
            else:
                print(f"      (*) Advertencia (Prep Scan CSV): Columna de oficio '{col_oficio_csv}' no en CSV. 'OFICIO_NORM_CSV' con valores por defecto.")
                df_csv_preparado_para_scan['OFICIO_NORM_CSV'] = pd.Series([normalizar_oficio_multas(None)] * len(df_csv_preparado_para_scan), dtype=str, index=df_csv_preparado_para_scan.index)

            # Crear CP_CALC_CSV y DIRECCION_CSV
            if col_direccion_csv in df_csv_preparado_para_scan.columns:
                df_csv_preparado_para_scan['DIRECCION_CSV'] = df_csv_preparado_para_scan[col_direccion_csv].astype(str).str.strip()
                df_csv_preparado_para_scan['CP_CALC_CSV'] = df_csv_preparado_para_scan['DIRECCION_CSV'].apply(lambda x: extraer_cp_de_direccion(x) if pd.notna(x) else "SIN_CP").astype(str)
            else:
                print(f"      (*) Advertencia (Prep Scan CSV): Columna de dirección '{col_direccion_csv}' no en CSV. Usando defaults.")
                df_csv_preparado_para_scan['DIRECCION_CSV'] = "DIRECCION_DESCONOCIDA_SCAN"
                df_csv_preparado_para_scan['CP_CALC_CSV'] = "SIN_CP_SCAN"
            
            print("    - Datos del CSV principal preparados para la lógica de escaneo.")
        except Exception as e_prep_scan:
            print(f"    - (!) Error GRAVE durante la preparación de datos CSV para escaneo: {e_prep_scan}")
            failed_operations_log["preparacion_csv_scan_multas"].append({"error": str(e_prep_scan), "timestamp": datetime.now().isoformat()})
            return df_bd_maestra_global 
    elif df_csv_principal_global is None:
         print("    - El DataFrame CSV principal (df_csv_principal_global) es None. No se puede preparar para escaneo.")
    else: # es vacío
         print("    - El DataFrame CSV principal está vacío. No hay datos que preparar para escaneo.")
    # --- FIN: PREPARACIÓN DEL DATAFRAME CSV PRINCIPAL ---

    df_bd_maestra_modificada_total_scan = df_bd_maestra_global.copy()

    for i_pdf_fuente, pdf_fuente_path in enumerate(lista_pdfs_a_escanear):
        print(f"\n    --- Procesando PDF Fuente ({i_pdf_fuente + 1}/{len(lista_pdfs_a_escanear)}): {pdf_fuente_path.name} ---")
        
        nombre_base_pdf_fuente = pdf_fuente_path.stem
        # Los resultados de cada PDF fuente van a una subcarpeta específica para ellos
        carpeta_salida_para_este_pdf_fuente = output_scan_artifacts_multas_modo_root / f"Resultados_Scan_{nombre_base_pdf_fuente}"
        carpeta_salida_para_este_pdf_fuente.mkdir(parents=True, exist_ok=True) # Asegurar que la carpeta para los PDFs fusionados exista
        nombre_excel_escaneo_individual = f"BD_Escaneo_{nombre_base_pdf_fuente}_Multas.xlsx"
        ruta_excel_escaneo_individual = carpeta_salida_para_este_pdf_fuente / nombre_excel_escaneo_individual # Guardar Excel dentro de su carpeta de resultados

        log_oficios_no_encontrados = []

        if ruta_excel_escaneo_individual.exists():
            print(f"        - El archivo Excel de escaneo '{nombre_excel_escaneo_individual}' ya existe. Saltando '{pdf_fuente_path.name}'.")
            continue

        registros_para_excel_fuente_actual = []
        df_bd_maestra_para_este_pdf_fuente = df_bd_maestra_modificada_total_scan.copy() 

        with tempfile.TemporaryDirectory(prefix=f"scan_multas_temp_{nombre_base_pdf_fuente}_") as temp_dir_str:
            temp_dir_path = Path(temp_dir_str)
            
            # Asumimos que group_and_split_pdf_dynamically_scan_multas ahora devuelve:
            # (texto_original_del_ocr_que_coincidio, oficio_normalizado_del_ocr, ruta_pdf_temporal, num_pags, metodo_ocr)
            lista_grupos_oficio_escaneado, oficios_no_encontrados_map = group_and_split_pdf_dynamically_scan_multas(
                pdf_fuente_path, 
                temp_dir_path,
                df_bd_maestra_para_este_pdf_fuente,
                config_multas_actual,
                df_csv_principal_preparado=df_csv_preparado_para_scan # <-- NUEVO PARÁMETRO
            )
            
            # --- NUEVO: Poblar el log de no encontrados ---
            for oficio_nf, data_nf in oficios_no_encontrados_map.items():
                 log_oficios_no_encontrados.append({
                    "OficioNoEncontrado": oficio_nf,
                    "PrimeraPaginaDetectada": data_nf["primera_pagina"],
                    "Error": "Oficio no existe en la Base de Datos Maestra"
                 })

            if not lista_grupos_oficio_escaneado:
                print(f"        - No se encontraron grupos de oficios válidos en '{pdf_fuente_path.name}'.")
                df_excel_vacio = pd.DataFrame(columns=COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL) # COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL debe estar definida
                try:
                    df_excel_vacio.to_excel(ruta_excel_escaneo_individual, index=False, engine='openpyxl')
                except Exception as e_save_empty_excel:
                    print(f"        - (!) Error guardando Excel de escaneo vacío '{nombre_excel_escaneo_individual}': {e_save_empty_excel}")
                continue

            print(f"        - Se encontraron {len(lista_grupos_oficio_escaneado)} grupos de oficios en '{pdf_fuente_path.name}'. Procesando...")
            
            df_bd_maestra_actualizada_iteracion_fuente = df_bd_maestra_para_este_pdf_fuente.copy() 
            
            for datos_grupo in lista_grupos_oficio_escaneado:
                # Desempaquetar la tupla (asumiendo el nuevo formato de devolución)
                # Si group_and_split... aún no devuelve el original, ajusta aquí.
                if len(datos_grupo) == 5: # (oficio_original_ocr, oficio_norm_ocr, ruta_pdf, num_pags, metodo_ocr)
                    oficio_original_ocr, oficio_norm_ocr, ruta_pdf_grupo_g, num_pags_g, metodo_ocr_g = datos_grupo
                elif len(datos_grupo) == 4: # (oficio_norm_ocr, ruta_pdf, num_pags, metodo_ocr) - Fallback si no se devuelve el original
                    print(f"        (*) Advertencia: group_and_split... no devolvió oficio original del OCR. Usando normalizado como original para el grupo que inicia con (Norm: {datos_grupo[0]}).")
                    oficio_normalizado_ocr = datos_grupo[0]
                    oficio_original_ocr = oficio_normalizado_ocr # Usar el normalizado como placeholder del original
                    ruta_pdf_grupo_g, num_pags_g, metodo_ocr_g = datos_grupo[1], datos_grupo[2], datos_grupo[3]
                else:
                    print(f"        (!) Error: Formato de datos de grupo inesperado: {datos_grupo}. Saltando este grupo.")
                    continue

                datos_excel_un_oficio, df_bd_actualizada_un_oficio_iter = process_single_oficio_scan_multas(
                    oficio_escaneado_str_original=oficio_original_ocr, 
                    oficio_escaneado_str_norm=oficio_norm_ocr,    
                    ruta_pdf_temporal_grupo_escaneado_obj=ruta_pdf_grupo_g,
                    num_paginas_grupo_escaneado=num_pags_g,
                    metodo_extraccion_ocr=metodo_ocr_g,
                    df_bd_maestra_global=df_bd_maestra_actualizada_iteracion_fuente, 
                    df_csv_principal_preparado=df_csv_preparado_para_scan, 
                    nombre_pdf_fuente_original_str=pdf_fuente_path.name,
                    directorio_temporal_principal_obj=temp_dir_path,
                    directorio_salida_escaneos_oficio_obj=carpeta_salida_para_este_pdf_fuente,
                    config_multas_actual=config_multas_actual,
                    funcion_de_subida=funcion_de_subida # <-- AÑADIR ESTA LÍNEA
                )
                registros_para_excel_fuente_actual.append(datos_excel_un_oficio)
                df_bd_maestra_actualizada_iteracion_fuente = df_bd_actualizada_un_oficio_iter
            
            df_bd_maestra_modificada_total_scan = df_bd_maestra_actualizada_iteracion_fuente

        if registros_para_excel_fuente_actual or log_oficios_no_encontrados:
            # 1. Crear el DataFrame con los registros procesados exitosamente (si los hay)
            if registros_para_excel_fuente_actual:
                df_exitosos = pd.DataFrame(registros_para_excel_fuente_actual, columns=COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL)
            else:
                df_exitosos = pd.DataFrame(columns=COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL)

            # 2. Crear el DataFrame con los oficios no encontrados (si los hay)
            if log_oficios_no_encontrados:
                df_no_encontrados_log = pd.DataFrame(log_oficios_no_encontrados)
                # Renombrar columnas para que coincidan y poder concatenar
                df_no_encontrados_log.rename(columns={
                    "OficioNoEncontrado": "Oficio", 
                    "PrimeraPaginaDetectada": "Paginas"
                }, inplace=True)
                # Asegurarse de que tenga todas las columnas del Excel final, rellenando con ""
                for col in COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL:
                    if col not in df_no_encontrados_log.columns:
                        df_no_encontrados_log[col] = ""
            else:
                df_no_encontrados_log = pd.DataFrame()

            # 3. Concatenar ambos DataFrames
            df_excel_individual = pd.concat([df_exitosos, df_no_encontrados_log], ignore_index=True)
            
            # 4. Guardar el archivo Excel unificado
            try:
                print(f"        - Guardando bitácora de escaneo (éxitos y errores): {ruta_excel_escaneo_individual.name}")

                # --- INICIA BLOQUE DE GUARDADO DE EXCEL CON REPORTE (YA NO HAY 'if' EXTRA) ---
                lista_de_oficios = df_excel_individual["Oficio"].tolist()

                with pd.ExcelWriter(ruta_excel_escaneo_individual, engine='openpyxl') as writer:
                    # Escribir la primera hoja con el log de escaneo
                    df_excel_individual.to_excel(writer, sheet_name="Log_Escaneo", index=False)

                    # Escribir la segunda hoja con el reporte para impresión
                    crear_hoja_reporte_impresion(
                        writer,
                        datos_para_reporte=lista_de_oficios,
                        titulo_principal="MULTAS",
                        cabecera_columna="OFICIO"
                    )

                print(f"        - Bitácora y Reporte de Impresión guardados en: {ruta_excel_escaneo_individual.name}")

                # Lógica de subida del archivo Excel de bitácora
                if funcion_de_subida:
                    print(f"        - Subiendo bitácora de escaneo '{ruta_excel_escaneo_individual.name}' al servidor...")
                    subida_ok = funcion_de_subida(ruta_excel_escaneo_individual, config_multas_actual)
                    if not subida_ok:
                        print(f"        - (!) ADVERTENCIA: La subida de la bitácora de escaneo Excel falló.")

            except Exception as e_save_excel_scan:
                print(f"        - (!) Error guardando o subiendo Excel de escaneo individual '{nombre_excel_escaneo_individual}': {e_save_excel_scan}")
                failed_operations_log[f"guardar_excel_escaneo_{nombre_excel_escaneo_individual}"].append({"error": str(e_save_excel_scan)})

                # <<< Lógica de subida de la bitácora (ya está bien) >>>
                if funcion_de_subida:
                    print(f"        - Subiendo bitácora de escaneo '{ruta_excel_escaneo_individual.name}' al servidor...")
                    subida_ok = funcion_de_subida(ruta_excel_escaneo_individual, config_multas_actual)
                    if not subida_ok:
                        print(f"        - (!) ADVERTENCIA: La subida de la bitácora de escaneo Excel falló.")
            except Exception as e_save_excel_scan:
                print(f"        - (!) Error guardando o subiendo Excel de escaneo individual '{nombre_excel_escaneo_individual}': {e_save_excel_scan}")
                failed_operations_log[f"guardar_excel_escaneo_{nombre_excel_escaneo_individual}"].append({"error": str(e_save_excel_scan)})
        else:
            print(f"        - No se generaron registros para el Excel de escaneo de '{pdf_fuente_path.name}'.")
        # --- FIN DEL BLOQUE CORREGIDO ---

    print("\n--- Modo Escaneo de Documentos de Multas Finalizado ---")
    return df_bd_maestra_modificada_total_scan