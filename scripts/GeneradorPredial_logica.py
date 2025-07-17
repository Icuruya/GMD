# -*- coding: utf-8 -*-
import os
import re
import locale
import pandas as pd
import math
import unicodedata
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from datetime import datetime
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import nsmap
import traceback
import time # <--- AÑADIR ESTA LÍNEA
from datetime import datetime # Asegurarse que está importado
from threading import Lock # NUEVO: Para acceso seguro al Excel
from num2words import num2words
from pathlib import Path # <--- AÑADE ESTA LÍNEA AQUÍ
import tempfile
import shutil
import sys
import paramiko
import copy
from openpyxl.styles import Font, Alignment


excel_lock = Lock()


# --- RUTA BASE DE DOCUMENTOS PREDIAL EN EL SERVIDOR ---
RUTA_BASE_PDF_SERVIDOR_PREDIAL = Path(r"\\asesorescloud.ddns.net\datos_gobierno\PREDIAL - AFC\DOCUMENTOS")

# --- NUEVO: Imports para Excel y PDF ---

# Imports para OCR y manipulación de PDF (escaneo)
try:
    import fitz # PyMuPDF
    print("    - Módulo 'fitz' (PyMuPDF) importado para Escaneo Predial.")
except ImportError:
    print("¡ERROR FATAL! PyMuPDF (fitz) no encontrado. Necesario para el modo Escaneo Predial.")
    # Considera si el script debe detenerse si falta esta dependencia crítica para escaneo.
    # podrías lanzar la excepción: raise
    pass # O manejarlo de otra forma si el script puede funcionar parcialmente sin esto

try:
    import pytesseract
    from PIL import Image
    print("    - Módulos 'pytesseract' y 'PIL' importados para OCR Predial.")
    # Si necesitas configurar la ruta de Tesseract (si no está en el PATH):
    # tesseract_paths = [
    #     r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    #     r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    # ]
    # pytesseract.pytesseract.tesseract_cmd = '' 
    # for tess_path in tesseract_paths:
    #     if Path(tess_path).is_file():
    #         pytesseract.pytesseract.tesseract_cmd = tess_path
    #         print(f"    - Tesseract OCR encontrado en: {tess_path}")
    #         break
    # if not pytesseract.pytesseract.tesseract_cmd:
    #     print("    - (!) Advertencia: Tesseract OCR no encontrado en rutas comunes. Se intentará usar el PATH.")
except ImportError:
    print("¡ERROR FATAL! No se encontraron 'pytesseract' o 'Pillow'. Necesarios para OCR Predial.")
    # Considera cómo manejar esto; podrías lanzar la excepción: raise
    pass
# --- Imports para Excel y PDF ---
try:
    import openpyxl
    from openpyxl import load_workbook, Workbook 
    from openpyxl.utils.dataframe import dataframe_to_rows
    print("Biblioteca 'openpyxl' encontrada.")
except ImportError:
    print("ERROR CRITICO: La biblioteca 'openpyxl' no está instalada.")
    print("Puedes instalarla con: pip install openpyxl")
    exit() 

try:
    from PyPDF2 import PdfReader, PdfWriter
    from PyPDF2.errors import PdfReadError 
    print("Biblioteca 'PyPDF2' encontrada.")
except ImportError:
    print("ERROR CRITICO: La biblioteca 'PyPDF2' no está instalada.")
    print("Puedes instalarla con: pip install pypdf2")
    exit() 

try:
    from docx2pdf import convert
    DOCX2PDF_INSTALLED = True
    print("Biblioteca 'docx2pdf' encontrada.")
except ImportError:
    DOCX2PDF_INSTALLED = False
    print("ADVERTENCIA: La biblioteca 'docx2pdf' no está instalada.")
    # ... (mensajes de advertencia)

try:
    from num2words import num2words
    NUM2WORDS_INSTALLED = True
except ImportError:
    NUM2WORDS_INSTALLED = False
    print("Advertencia: La biblioteca 'num2words' no está instalada.")
    print("Puedes instalarla con: pip install num2words")
    print("El placeholder [FECHA_TEXTO] no funcionará correctamente.")

# Patrón para "EXP NNNNNNNN" (sin punto, para la parte superior)
# Asume que el número de expediente tiene 7 u 8 dígitos. Ajusta \d{7,8} si es diferente (ej. \d{8})
# LÍNEA NUEVA Y CORREGIDA
EXPEDIENTE_REGEX_SCAN_PREDIAL_NODOT = re.compile(r"(?:Expediente Catastral:|EXP)\s*:?\s*(\d{7,8})\b", re.IGNORECASE)
# Patrón para "EXP. NNNNNNNN" (con punto, para la parte inferior)
EXPEDIENTE_REGEX_SCAN_PREDIAL_DOT = re.compile(r"EXP\.\s+(\d{7,8})\b", re.IGNORECASE) # (modificado)

# Columnas para los Excel individuales de escaneo de Predial (asegúrate que esté)
COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL_PREDIAL = ["EXPEDIENTE", "NOMBRE CONTRIBUYENTE", "DIRECCION", "COLONIA", "BIMESTRE", "AÑOS", "PaginasEscaneadas", "NombreArchivoFusionado", "RutaArchivoFusionado", "PaginasDetectadas"]# --- Configuración ---
CARPETA_PRINCIPAL = "DOCUMENTOS"
NOMBRE_BD_MAESTRA = "base_datos_maestra.xlsx"
RUTA_BD_MAESTRA = os.path.join(CARPETA_PRINCIPAL, NOMBRE_BD_MAESTRA)

#os.makedirs(CARPETA_PRINCIPAL, exist_ok=True)

# --- Columnas de la BD Maestra (Esquema COMPLETO) ---
COL_EXPEDIENTE = "EXPEDIENTE"
COL_NOMBRE_CONTRIBUYENTE = "NOMBRE CONTRIBUYENTE"
COL_DIRECCION = "DIRECCION" 
COL_COLONIA = "COLONIA"
COL_BIMESTRE = "BIMESTRE" 
COL_ANOS = "AÑOS" 
COL_MONTO = "MONTO" 
COL_ESTADO = "ESTADO" 
COL_ID_LOTE = "ID Lote"  # Corregido para coincidir con el input del usuario
COL_PRIORIDAD = "Prioridad" # Corregido para coincidir con el input del usuario
COL_HOJAS_DOC = "HOJAS POR DOCUMENTO" 
COL_IMPRESORA = "Impresora" # Nueva columna
COL_FECHA_HORA_IMPRESION = "Fecha Hora Impresion" # Nueva columna


# Lista ordenada de TODAS las columnas para el DataFrame y Excel
COLUMNAS_BD_MAESTRA = [
    COL_EXPEDIENTE, COL_NOMBRE_CONTRIBUYENTE, COL_DIRECCION, COL_COLONIA,
    COL_BIMESTRE, COL_ANOS, COL_MONTO, COL_ESTADO,
    COL_ID_LOTE, COL_PRIORIDAD, COL_HOJAS_DOC,
    COL_IMPRESORA, COL_FECHA_HORA_IMPRESION, "MOVIMIENTO" # Nuevas columnas añadidas
]

# Tipos de datos para pandas para TODAS las columnas
TIPOS_BD_MAESTRA = {
    COL_EXPEDIENTE: 'str', COL_NOMBRE_CONTRIBUYENTE: 'str', COL_DIRECCION: 'str', COL_COLONIA: 'str',
    COL_BIMESTRE: 'str', COL_ANOS: 'str', COL_MONTO: 'str', COL_ESTADO: 'str',
    COL_ID_LOTE: 'str', COL_PRIORIDAD: 'Int64', COL_HOJAS_DOC: 'Int64',
    COL_IMPRESORA: 'str', COL_FECHA_HORA_IMPRESION: 'str',
     "MOVIMIENTO": 'str' # Tipos para nuevas columnas
}

# --- Constantes de Estado SINCRONIZADAS ---
ESTADO_GEN_ULTIMA = "Generado ultima hoja"
ESTADO_GEN_COMPLETO = "Generado COMPLETO"
ESTADO_IMP_ULTIMA = "Ultima PAG impresa"
ESTADO_IMP_COMPLETO = "Impresión Completada"
ESTADO_ERROR_GENERACION = "Error Generacion" 
ESTADO_PENDIENTE = "Pendiente" 

# --- Modos de Generación ---
MODO_COMPLETO = "COMPLETO"
MODO_ULTIMA = "ULTIMA"
MODO_RESTO = "RESTO"

# --- Archivos de Entrada (Mantener los originales) ---
csv_data_file = "BASE_DE_DATOS.csv"
config_file = "config_columnas.csv"
pm_csv_file = "PM.csv"
plantilla_word = "PLANTILLA.docx"
eliminar_docx_intermedio = True


# --- Texto de la Leyenda Final (Página 17) ---
LEYENDA_PAG17_TEXTO = """
Siendo las _________ horas con _________ minutos del día __________ de ___________ de 20_____, ubicado en el domicilio citado al rubro, con fundamento en los artículos 2, 7 último párrafo 13, 14, 34 fracciones V, VI, y VII, 40, 42, 133, 134, 136 y 137 del Código Fiscal del Estado de Nuevo León; y 30, 31, inciso A) fracciones I y XV y 35 fracciones I, II, VIII, XXII, XXV, XXVI, XXVII, XXIX, XXXII, XLII, XLIII, XLIV, XLV, XLVII y XLIX, del Reglamento de la Administración Pública Municipal de Monterrey; Artículo 6 de la Ley de Ingresos de los Municipios del Estado de Nuevo León para el ejercicio fiscal 2023, 96, 99 y 100 de la Ley de Gobierno Municipal de Nuevo León, en relación con el artículo 11 del Reglamento de la Administración Pública Municipal de Monterrey, el suscrito cuyo nombre y cargo aparece al calce de la presente diligencia, identificándome mediante oficio número ______________________ expedido por la C. Directora de Recaudación Inmobiliaria de la Secretaría de Finanzas y Administración del Municipio de Monterrey en fecha ______________________________________, con vigencia desde su expedición hasta el 31 de diciembre de 2025; con quien atiende la diligencia de notificación, el (la) C. cuyo nombre aparece al calce de la presente diligencia y quien se identifica con _______________________________, número ____________________________, persona a quien se le hace lectura integra y se le entrega documento en original que contiene la (s) DETERMINACIÓN (ES) de CRÉDITO (S) FISCAL (ES) emitida (s) por la Directora de Recaudación Inmobiliaria de la Secretaría de Finanzas y Administración del Municipio de Monterrey, en fecha señalada al rubro, que contiene (n) firma (s) autógrafa (s) de la Directora antes aludida, relativo al (los) expediente (s) catastral (es) citados al rubro, firmando quienes participan en ésta diligencia y así quisieron hacerlo.

EL MINISTRO EJECUTOR                                           PERSONA CON QUIEN SE ENTIENDE LA DILIGENCIA

NOMBRE: _____________________                                          NOMBRE: ____________________________

FIRMA: _______________________                                           FIRMA:_______________________________
"""

PROTECTED_PLACEHOLDERS = ["[LOGO]", "[TITULO]"]



# --- Funciones de utilidad (sin cambios) ---
def _generar_citatorio_pdf_interno(expediente, config_predial, temp_dir):
    """
    Función interna para generar el PDF del citatorio para un expediente específico.
    Devuelve la ruta (Path object) al PDF temporal del citatorio o None si falla.
    """
    try:
        expediente_str = str(expediente).strip()
        if len(expediente_str) == 7:
            expediente_str = '0' + expediente_str

        # Construir la ruta a la plantilla del citatorio desde la configuración del modo
        ruta_plantilla_citatorio = Path(config_predial["base_path"]) / 'CITATORIO.docx'
        
        if not ruta_plantilla_citatorio.is_file():
            print(f"    (!) ERROR: No se encontró la plantilla del citatorio en: {ruta_plantilla_citatorio}")
            return None

        doc = Document(ruta_plantilla_citatorio)
        # Reemplazar el placeholder de forma segura
        for paragraph in doc.paragraphs:
            if '[EXPEDIENTE]' in paragraph.text:
                # Esta técnica reemplaza el texto sin perder el formato del párrafo
                texto_original = paragraph.text
                nuevo_texto = texto_original.replace('[EXPEDIENTE]', expediente_str)
                paragraph.text = nuevo_texto
                # Rompemos el bucle si ya lo encontramos para ser eficientes
                break 
        
        temp_docx_path = temp_dir / f"temp_citatorio_{expediente}.docx"
        temp_pdf_path = temp_dir / f"citatorio_{expediente}.pdf"
        
        doc.save(temp_docx_path)
        convert(str(temp_docx_path), str(temp_pdf_path))
        
        if temp_pdf_path.exists() and temp_pdf_path.stat().st_size > 0:
            print(f"      - PDF de citatorio generado para {expediente}.")
            return temp_pdf_path
        else:
            print(f"      - (!) Fallo al convertir a PDF el citatorio para {expediente}.")
            return None

    except Exception as e:
        print(f"      - (!) Error crítico al generar PDF de citatorio para {expediente}: {e}")
        traceback.print_exc()
        return None

def is_page_blank_scan_predial(page_fitz_obj, umbral_caracteres=20):
    """
    [VERSIÓN CORREGIDA Y MEJORADA]
    Verifica si una página está esencialmente en blanco.
    Ahora revisa la presencia de texto, imágenes y dibujos para tomar la decisión.
    """
    # 1. Comprobar si hay texto suficiente en la página.
    # Esto funciona para documentos generados digitalmente.
    if len(page_fitz_obj.get_text("text").strip()) > umbral_caracteres:
        return False

    # 2. Comprobar si la página contiene imágenes.
    # Esto es CRUCIAL para detectar páginas escaneadas que no tienen texto digital.
    if page_fitz_obj.get_images(full=True):
        return False

    # 3. Como respaldo, comprobar si hay dibujos vectoriales (líneas, tablas, etc.).
    if page_fitz_obj.get_drawings():
        return False

    # 4. Si no hay texto, ni imágenes, ni dibujos, la página se considera en blanco.
    return True

def limpiar_texto(texto):
    if pd.isna(texto):
        return ""
    texto_str = str(texto)
    nfkd_form = unicodedata.normalize('NFKD', texto_str)
    texto_sin_acentos = "".join([c for c in nfkd_form if not unicodedata.combining(c)])
    texto_limpio = re.sub(r'[^\w\s-]', '_', texto_sin_acentos)
    return re.sub(r'\s+', '_', texto_limpio).strip().lower()

def page_contains_keywords_scan_predial(page_fitz_obj):
    """Revisa si una página contiene palabras clave de citatorio."""
    keywords = ["CITATORIO", "DILIGENCIA", "NOTIFICACIÓN", "ACTA CIRCUNSTANCIADA"]
    texto_pagina = page_fitz_obj.get_text("text").upper()
    for keyword in keywords:
        if keyword in texto_pagina:
            return True
    return False

def save_pdf_group_scan_predial(doc_fitz_original_obj, page_indices_list_0based, expediente_str_grupo, temp_dir_path_obj, config_predial_actual): # (verificado)
    """
    Guarda un grupo de páginas (especificadas por sus índices 0-based) de un documento fitz
    en un nuevo PDF temporal dentro del directorio temporal especificado.
    Retorna el Path object al PDF temporal creado, o None si falla.
    """
    if not page_indices_list_0based or not expediente_str_grupo:
        print(f"        - (save_pdf_group_scan_predial) Error: Lista de páginas o expediente no proporcionado. Exp: {expediente_str_grupo}")
        return None

    nombre_archivo_temporal_grupo = f"TEMP_GRUPO_PREDIAL_{limpiar_texto(expediente_str_grupo)}_{time.time_ns()}.pdf" # limpiar_texto es de GeneradorPredial_logica.py
    ruta_pdf_temporal_grupo = temp_dir_path_obj / nombre_archivo_temporal_grupo

    doc_fitz_nuevo_grupo = None 
    try:
        doc_fitz_nuevo_grupo = fitz.open() 

        for page_index in page_indices_list_0based: # (iteración robusta)
            doc_fitz_nuevo_grupo.insert_pdf(doc_fitz_original_obj, from_page=page_index, to_page=page_index, start_at=-1)

        if doc_fitz_nuevo_grupo.page_count > 0:
            doc_fitz_nuevo_grupo.save(str(ruta_pdf_temporal_grupo), garbage=4, deflate=True, clean=True) 

            if ruta_pdf_temporal_grupo.is_file() and ruta_pdf_temporal_grupo.stat().st_size > 0:
                return ruta_pdf_temporal_grupo
            else:
                print(f"        - (!) Error (save_pdf_group_scan_predial): PDF temporal para Exp {expediente_str_grupo} no se creó o está vacío en {ruta_pdf_temporal_grupo}.")
                return None
        else:
            print(f"        - (!) (save_pdf_group_scan_predial): No se añadieron páginas al PDF temporal para Exp {expediente_str_grupo}.")
            return None
    except Exception as e_save_grupo:
        print(f"        - (!) Error (save_pdf_group_scan_predial) al crear PDF temporal para Exp {expediente_str_grupo}: {e_save_grupo}")
        return None
    finally:
        if doc_fitz_nuevo_grupo:
            try: doc_fitz_nuevo_grupo.close()
            except: pass

def obtener_siguiente_letra_lote(letra_actual):
    """
    Obtiene la siguiente letra del alfabeto para el lote.
    Ej: 'A' -> 'B', 'Z' -> 'AA', 'AZ' -> 'BA'.
    """
    if not letra_actual or not letra_actual.isalpha():
        return 'A'
    
    letra_actual = letra_actual.upper()
    if letra_actual == 'Z' * len(letra_actual):
        return 'A' * (len(letra_actual) + 1)
    
    letras = list(letra_actual)
    i = len(letras) - 1
    while i >= 0:
        if letras[i] == 'Z':
            letras[i] = 'A'
            i -= 1
        else:
            letras[i] = chr(ord(letras[i]) + 1)
            break
    return "".join(letras)

def group_and_split_pdf_dynamically_scan_predial(pdf_fuente_path_obj, temp_dir_para_grupos_obj, config_predial_actual, df_bd_maestra_para_verificar, df_csv_principal_predial):
    """
    [VERSIÓN CORREGIDA]
    Analiza un PDF, agrupa páginas por expediente validado y agrupa correctamente
    las páginas de continuación (citatorios, etc.) al último expediente válido encontrado.
    """
    resultados_validos = []
    doc_fitz = None
    
    # --- INICIO DE LA MODIFICACIÓN: Lógica de Estado ---
    # Esta variable recordará el último expediente válido mientras recorremos el PDF.
    ultimo_expediente_valido = None
    # --- FIN DE LA MODIFICACIÓN ---

    # Pre-cargar los sets de expedientes válidos para búsquedas rápidas (sin cambios)
    col_exp_bd = config_predial_actual["col_expediente"]
    expedientes_validos_bd_set = set(df_bd_maestra_para_verificar[col_exp_bd].astype(str).str.strip().str.lstrip('0'))
    expedientes_validos_csv_set = set(df_csv_principal_predial[col_exp_bd].astype(str).str.strip().str.lstrip('0'))

    try:
        expedientes_y_sus_paginas = {}
        doc_fitz = fitz.open(str(pdf_fuente_path_obj))

        for i_pagina in range(doc_fitz.page_count):
            pagina_obj = doc_fitz.load_page(i_pagina)
            # Extraemos los posibles expedientes de la página (sin cambios)
            posibles_exp, _ = extract_expediente_from_page_scan_predial(pagina_obj, i_pagina + 1, config_predial_actual)

            expediente_encontrado_y_validado = None

            # Intentar validar los expedientes extraídos (sin cambios)
            if posibles_exp:
                for exp in posibles_exp:
                    if exp in expedientes_validos_bd_set or exp in expedientes_validos_csv_set:
                        expediente_encontrado_y_validado = exp
                        break
            
            # --- INICIO DE LA MODIFICACIÓN: Nueva Lógica de Decisión ---
            if expediente_encontrado_y_validado:
                # CASO 1: Encontramos un número de expediente válido en la página.
                # Actualizamos nuestro "último expediente válido" y lo asignamos.
                ultimo_expediente_valido = expediente_encontrado_y_validado
                
                if ultimo_expediente_valido not in expedientes_y_sus_paginas:
                    expedientes_y_sus_paginas[ultimo_expediente_valido] = {"paginas": []}
                
                expedientes_y_sus_paginas[ultimo_expediente_valido]["paginas"].append(i_pagina)
            
            else:
                # CASO 2: NO encontramos un número de expediente válido en la página.
                # Ahora verificamos si es una página de continuación o si está en blanco.
                # Para ser de continuación, la página NO debe estar en blanco Y ya debemos tener un expediente anterior.
                if not is_page_blank_scan_predial(pagina_obj) and ultimo_expediente_valido:
                    # ¡Es una hoja de continuación! La asignamos al último expediente válido.
                    print(f"    - Página {i_pagina + 1} sin expediente, pero con contenido. Asignada a Exp. {ultimo_expediente_valido}.")
                    expedientes_y_sus_paginas[ultimo_expediente_valido]["paginas"].append(i_pagina)
                # Si la página está en blanco, o si aún no hemos encontrado ningún expediente válido,
                # simplemente la ignoramos y no hacemos nada.
            # --- FIN DE LA MODIFICACIÓN ---

        print("    - Ordenando expedientes encontrados por orden de aparición en el PDF...")
        # El resto de la función para ordenar y crear los archivos PDF no necesita cambios.
        expedientes_ordenados = sorted(
            expedientes_y_sus_paginas.items(),
            key=lambda item: min(item[1]['paginas'])
        )

        for exp_str, data in expedientes_ordenados:
            ruta_temp_grupo = save_pdf_group_scan_predial(doc_fitz, data["paginas"], exp_str, temp_dir_para_grupos_obj, config_predial_actual)
            if ruta_temp_grupo:
                resultados_validos.append((exp_str, ruta_temp_grupo, len(data["paginas"]), "METODO_CORREGIDO", data["paginas"]))

    finally:
        if doc_fitz:
            doc_fitz.close()

    # Se devuelve un log de errores vacío porque ahora los "no encontrados" se manejan dentro.
    return resultados_validos, {}

def crear_hoja_reporte_impresion(writer, datos_para_reporte, titulo_principal, cabecera_columna):
    """
    Crea una segunda hoja en un archivo Excel con un formato de impresión específico,
    llenando las columnas de forma vertical.
    """
    try:
        workbook = writer.book
        sheet_name = "Reporte para Impresión"
        if sheet_name in workbook.sheetnames:
            del workbook[sheet_name]

        ws = workbook.create_sheet(sheet_name)

        font_style = Font(name='Calibri', size=26) # <-- TAMAÑO 26 CORREGIDO
        row_height = 29

        max_filas_por_bloque = 43
        columnas_bloques = [('A', 'B'), ('D', 'E'), ('G', 'H')]
        puntero_datos = 0
        num_datos = len(datos_para_reporte)
        fila_actual_excel = 1

        while puntero_datos < num_datos:
            if fila_actual_excel > 1:
                fila_actual_excel += 1

            ws.merge_cells(f'A{fila_actual_excel}:H{fila_actual_excel}')
            titulo_cell = ws[f'A{fila_actual_excel}']

            # Título dinámico para el reporte
            titulo_final = f"{titulo_principal} - DESPACHO MALDONADO"
            # Esta parte es un placeholder, ya que el lote se determina en la función que llama.
            # Pero la dejamos flexible.
            if "Lote" in titulo_principal:
                 titulo_final = titulo_principal.replace("Lote", "LOTE") + " - DESPACHO MALDONADO"

            titulo_cell.value = titulo_final
            titulo_cell.font = font_style
            titulo_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[fila_actual_excel].height = row_height
            fila_actual_excel += 1

            fila_inicio_segmento = fila_actual_excel
            for col_fila, col_dato in columnas_bloques:
                ws[f'{col_fila}{fila_inicio_segmento}'] = "FILA"
                ws[f'{col_dato}{fila_inicio_segmento}'] = cabecera_columna
                ws[f'{col_fila}{fila_inicio_segmento}'].font = font_style
                ws[f'{col_dato}{fila_inicio_segmento}'].font = font_style
            ws.row_dimensions[fila_inicio_segmento].height = row_height

            for col_fila, col_dato in columnas_bloques:
                for i_fila_en_bloque in range(max_filas_por_bloque):
                    if puntero_datos >= num_datos:
                        break

                    fila_absoluta = fila_inicio_segmento + 1 + i_fila_en_bloque
                    ws.row_dimensions[fila_absoluta].height = row_height

                    ws[f'{col_fila}{fila_absoluta}'] = puntero_datos + 1
                    ws[f'{col_fila}{fila_absoluta}'].font = font_style

                    ws[f'{col_dato}{fila_absoluta}'] = str(datos_para_reporte[puntero_datos])
                    ws[f'{col_dato}{fila_absoluta}'].font = font_style

                    puntero_datos += 1
                if puntero_datos >= num_datos:
                    break

            fila_actual_excel = fila_inicio_segmento + 1 + max_filas_por_bloque

        print(f"    - Hoja '{sheet_name}' creada y formateada con {num_datos} registros.")

    except Exception as e:
        print(f"    - (!) Error crítico al crear la hoja de reporte para impresión: {e}")
        import traceback
        traceback.print_exc()
def buscar_datos_expediente_en_csv_predial(expediente_buscado_str, df_csv_principal_predial_completo, config_predial_actual): # (verificado)
    datos_encontrados = {
        'nombre_contribuyente': None, 'colonia': None, 'monto_formateado': None, 'encontrado_en_csv': False
    }
    if df_csv_principal_predial_completo is None or df_csv_principal_predial_completo.empty:
        return datos_encontrados

    col_exp_csv = config_predial_actual.get("col_expediente", "EXPEDIENTE") 
    col_nombre_csv = config_predial_actual.get("col_nombre_csv_original", "NOMBRE CONTRIBUYENTE") 
    col_colonia_csv = config_predial_actual.get("col_colonia_csv_original", "COLONIA") 

    if col_exp_csv not in df_csv_principal_predial_completo.columns:
         print(f"    (!) Advertencia (buscar_datos_csv_predial): Columna expediente '{col_exp_csv}' no encontrada en CSV. No se pueden buscar datos.")
         return datos_encontrados

    df_csv_principal_predial_completo['EXPEDIENTE_NORM_CSV_TEMP'] = df_csv_principal_predial_completo[col_exp_csv].astype(str).str.lstrip('0')

    registros_expediente_en_csv = df_csv_principal_predial_completo[
        df_csv_principal_predial_completo['EXPEDIENTE_NORM_CSV_TEMP'] == expediente_buscado_str
    ].copy() 

    df_csv_principal_predial_completo.drop(columns=['EXPEDIENTE_NORM_CSV_TEMP'], inplace=True, errors='ignore')

    if not registros_expediente_en_csv.empty:
        fila_representativa_csv_exp = registros_expediente_en_csv.iloc[0]
        datos_encontrados['nombre_contribuyente'] = fila_representativa_csv_exp.get(col_nombre_csv)
        datos_encontrados['colonia'] = fila_representativa_csv_exp.get(col_colonia_csv)

        monto_total_calculado_csv = 0.0
        try:
            # --- LÓGICA DE CÁLCULO DE MONTO TOTAL DESDE CSV (ADAPTAR A TU CASO) ---
            # Esta es una simplificación. Debes adaptar la lógica de suma de montos
            # de tu función generar_documentos_predial_core para que opere sobre fila_representativa_csv_exp.
            # Si tienes una columna ya calculada con el total en df_csv_principal_predial_completo:
            col_monto_total_directo_csv = config_predial_actual.get("col_monto_calculado_csv", "MONTO_CALCULADO_TOTAL") # Inventa una clave para config
            if col_monto_total_directo_csv in fila_representativa_csv_exp:
                monto_total_calculado_csv = safe_float(fila_representativa_csv_exp.get(col_monto_total_directo_csv, 0.0)) # safe_float de GeneradorPredial_logica.py
            else:
                # Si no, implementa la suma de las columnas de adeudo, recargos, etc., de la fila.
                # Ejemplo placeholder (¡DEBES COMPLETAR ESTO CON TU LÓGICA DE SUMAS!):
                # sum_impuesto_csv = safe_float(fila_representativa_csv_exp.get("ADEUDO_TOTAL_CSV", 0.0)) 
                # sum_actualizacion_csv = safe_float(fila_representativa_csv_exp.get("ACTUALIZACION_TOTAL_CSV", 0.0))
                # sum_recargos_csv = safe_float(fila_representativa_csv_exp.get("RECARGOS_TOTAL_CSV", 0.0))
                # monto_total_calculado_csv = sum_impuesto_csv + sum_actualizacion_csv + sum_recargos_csv
                print(f"    (*) (buscar_datos_csv_predial) Cálculo de monto total desde CSV detallado no implementado completamente. Monto será 0.")
                monto_total_calculado_csv = 0.0

            datos_encontrados['monto_formateado'] = formatear_valor_v6(monto_total_calculado_csv, "MONTO_TOTAL_CSV_PREDIAL_SCAN", {"MONTO_TOTAL_CSV_PREDIAL_SCAN": "moneda"}) # formatear_valor_v6 de GeneradorPredial_logica.py
        except Exception as e_monto_csv:
            print(f"    (!) Error calculando monto desde CSV para Exp {expediente_buscado_str}: {e_monto_csv}")
            datos_encontrados['monto_formateado'] = "$0.00"
        datos_encontrados['encontrado_en_csv'] = True

    return datos_encontrados


def generar_un_documento_predial_para_escaneo(
    datos_expediente_df, # DataFrame con la fila (o filas) del CSV para este expediente
    plantilla_path_obj,
    ruta_pdf_destino_obj, # Ruta final donde se guardará el PDF completo
    expediente_str_para_log, # Expediente ya normalizado
    config_predial_actual_local, # El mode_config
    nombres_columnas_csv_local, # Nombres de columnas como están en datos_expediente_df
    tipos_columnas_csv_local    # Tipos de datos definidos en config_columnas.xlsx para esas columnas
):
    pdf_final_generado_ok = False
    num_paginas_generado = None # Para el conteo de páginas del PDF generado

    # Crear un nombre temporal para el DOCX en la misma carpeta del PDF de destino
    temp_docx_name = f"~TEMP_DOCX_FOR_SCAN_PREDIAL_{expediente_str_para_log}_{time.time_ns()}.docx"
    ruta_temp_docx_predial_scan = ruta_pdf_destino_obj.parent / temp_docx_name

    if datos_expediente_df.empty:
        print(f"        - (!) (Generar Faltante Predial Scan): No hay datos en datos_expediente_df para Exp {expediente_str_para_log}")
        return False, None

    # Asumimos que si hay múltiples filas para un expediente en el CSV de generación,
    # la primera es representativa para datos generales, y el resto son para tablas (esto es raro en Predial).
    # Usualmente, Predial tiene una sola fila muy ancha por expediente en su BASE_DE_DATOS.xlsx.
    row = datos_expediente_df.iloc[0] 

    try:
        # --- INICIO: COPIA Y ADAPTACIÓN DE LÓGICA DE GENERACIÓN DE DOCX ---
        # Esta sección es una réplica de la lógica que está DENTRO del bucle `for index, row in df_datos_ordenados.iterrows():`
        # en tu función `generar_documentos_predial_core`. Debes copiarla aquí y adaptarla.
        # Variables a adaptar:
        #   - Usa 'row' (la que se obtiene de datos_expediente_df.iloc[0]).
        #   - Usa 'plantilla_path_obj' en lugar de 'plantilla_path'.
        #   - Usa 'config_predial_actual_local' para acceder a configuraciones del modo si es necesario.
        #   - Usa 'nombres_columnas_csv_local' y 'tipos_columnas_csv_local' para formatear_valor_v6.
        #   - Carga el 'pm_set' usando config_predial_actual_local["pm_file_path"].
        #   - El DOCX se guarda en 'ruta_temp_docx_predial_scan'.
        #   - Asegúrate que todas las funciones auxiliares que llama (ej. safe_float, formatear_valor_v6,
        #     reemplazar_en_documento_v2, eliminar_elementos_inactivos_v_usuario, etc.)
        #     estén definidas y accesibles en GeneradorPredial_logica.py.

        print(f"        - (Generar Doc. Escaneo Predial) Iniciando generación DOCX para Exp: {expediente_str_para_log}")

        # 1. Cargar PM Set (necesario para la leyenda)
        pm_set_local_scan = cargar_lista_pm(config_predial_actual_local["pm_file_path"]) # cargar_lista_pm debe estar definida

        # 2. Lógica de años activos/inactivos (copiar de generar_documentos_predial_core)
        anos_activos = []
        anos_inactivos = []
        # ... (tu lógica completa para llenar anos_activos y anos_inactivos usando 'row',
        #      nombres_columnas_csv_local y tipos_columnas_csv_local) ...
        # Ejemplo simplificado (debes usar tu lógica completa):
        possible_year_cols_scan = {2022: "2022", 2023: "2023", 2024: "2024"} # Asume estos nombres en config_cols
        impuesto_year_cols_scan = {year: f"IMPUESTO PREDIAL DEL AÑO {year}" for year in [2022, 2023, 2024]}
        for year_scan in [2022, 2023, 2024]:
            activo_scan = False; valor_check_scan = None
            year_col_name_scan = possible_year_cols_scan.get(year_scan)
            imp_col_name_scan = impuesto_year_cols_scan.get(year_scan)

            # Verificar si la columna existe con el nombre directo del año o como "IMPUESTO PREDIAL DEL AÑO XXXX"
            if year_col_name_scan in row.index: # row es una Serie de Pandas
                valor_check_scan = str(row.get(year_col_name_scan, '0')).strip()
            elif imp_col_name_scan in row.index:
                valor_check_scan = str(row.get(imp_col_name_scan, '0')).strip()

            if valor_check_scan and valor_check_scan != '0' and valor_check_scan.lower() not in ['nan', 'na', '', '-']:
                if safe_float(valor_check_scan, 0.0) != 0.0: # safe_float debe estar definida
                    activo_scan = True
            if activo_scan: anos_activos.append(year_scan)
            else: anos_inactivos.append(year_scan)


        # 3. Placeholders de fecha, año base (copiar de generar_documentos_predial_core)
        now_scan = datetime.now() # datetime debe estar importado
        fecha_hoy_str_scan = "(Error Locale)"
        fecha_texto_str_scan = "(Error Locale/num2words)"
        # Asumiendo que configurar_locale() ya fue llamado y NUM2WORDS_INSTALLED está definido
        if locale.getlocale(locale.LC_TIME)[0]: # Verifica si el locale está configurado
             try: fecha_hoy_str_scan = now_scan.strftime("%d de %B de %Y").lower()
             except Exception: pass
             if NUM2WORDS_INSTALLED: # NUM2WORDS_INSTALLED debe estar definido
                 try: # ... (lógica para fecha_texto_str_scan) ...
                    dia_num_scan = now_scan.day; ano_num_scan = now_scan.year
                    dia_palabra_scan = num2words(dia_num_scan, lang='es') # num2words debe estar importado
                    ano_palabra_scan = num2words(ano_num_scan, lang='es')
                    medio_fecha_scan = now_scan.strftime("de %B de %Y").lower()
                    fecha_texto_str_scan = f"{dia_num_scan}-{dia_palabra_scan} {medio_fecha_scan}-{ano_palabra_scan}"
                 except Exception: pass
             else: fecha_texto_str_scan = "(num2words no instalado para escaneo)"
        ano_placeholder_scan = str(min(anos_activos)) if anos_activos else str(now_scan.year)


        # 4. Cargar plantilla
        doc_scan = Document(plantilla_path_obj)

        # 5. Extraer PERIODO y calcular bimestre/año para BD (copiar de generar_documentos_predial_core)
        # La columna "PERIODO" debe estar en nombres_columnas_csv_local y en `row`
        periodo_col_name_scan = "PERIODO" # Asegúrate que este es el nombre en tu config_columnas.xlsx
        periodo_str_scan = str(row.get(periodo_col_name_scan, "")).strip()
        periodo_año_scan = 0; periodo_bim_scan = 0
        if re.match(r"^\d{6}$", periodo_str_scan): # re debe estar importado
             try: 
                 periodo_año_scan = int(periodo_str_scan[:4])
                 periodo_bim_scan = int(periodo_str_scan[4:6])
             except: pass

        # 6. Reemplazos Iniciales (copiar y adaptar de generar_documentos_predial_core)
        #    Usa nombres_columnas_csv_local y tipos_columnas_csv_local para formatear_valor_v6
        initial_replacements_scan = {
             "[FECHA]": fecha_hoy_str_scan,
             "[FECHA_TEXTO]": fecha_texto_str_scan,
             "[AÑO]": ano_placeholder_scan,
             # ... otros placeholders directos y los que se iteran desde nombres_columnas_csv_local ...
        }
        # Debes replicar aquí el bucle que itera sobre nombres_columnas_csv_local
        # y que llama a formatear_valor_v6, incluyendo la lógica de años inactivos y bimestres.
        # Ejemplo (debes copiar tu lógica completa):
        placeholders_sumas_total_scan = ["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", "[TOTAL]", "[TOTAL PESOS CORREGIDOS]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]","[SUMA DE MONTOS ACTUALIZADOS]", "[SUMA DE MONTOS ACTUALIZADOS TEXTO]","[SUMA DE LOS MONTOS DE RECARGO TOTAL]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]", "[SUMA DE LA SANCION POR BIMESTRE]", "[SUMA DE LA SANCION POR BIMESTRE TEXTO]","[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]","[LEYENDA_FINAL_PAG17]"] 
        for col_name_iter_scan in nombres_columnas_csv_local:
            placeholder_iter_scan = f"[{col_name_iter_scan}]"
            if placeholder_iter_scan in placeholders_sumas_total_scan or placeholder_iter_scan in initial_replacements_scan or col_name_iter_scan.startswith("IMPUESTO PREDIAL DEL AÑO"):
                continue
            # ... (tu lógica completa de formateo y manejo de años/bimestres para este placeholder) ...
            # Esto incluye la llamada a formatear_valor_v6(row.get(col_name_iter_scan), col_name_iter_scan, tipos_columnas_csv_local)
            # y la lógica especial para BIMESTRE X YYYY, MONTO ACTUALIZADO..., etc.
            # Y la corrección para forzar SANCION POR BIMESTRE a ser igual que MONTO ACTUALIZADO
            valor_original_scan = row.get(col_name_iter_scan)
            valor_formateado_scan = formatear_valor_v6(valor_original_scan, col_name_iter_scan, tipos_columnas_csv_local) # formatear_valor_v6 debe estar definida
            # Aplicar aquí la lógica de ceros para bimestres pasados, y espacios para bimestres futuros con valor cero.
            # ...
            initial_replacements_scan[placeholder_iter_scan] = valor_formateado_scan # Valor final después de la lógica

        # Forzar SANCION POR BIMESTRE = MONTO ACTUALIZADO (como en generar_documentos_predial_core)
        for year_s_override_scan in [2022, 2023, 2024]:
            for bim_s_override_scan in range(1, 7):
                sancion_ph_scan = f"[SANCION POR BIMESTRE {bim_s_override_scan} {year_s_override_scan}]"
                monto_act_ph_scan = f"[MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim_s_override_scan} {year_s_override_scan}]"
                valor_monto_act_formateado_scan = initial_replacements_scan.get(monto_act_ph_scan)
                if valor_monto_act_formateado_scan is not None:
                    initial_replacements_scan[sancion_ph_scan] = valor_monto_act_formateado_scan
                else: # Si monto actualizado no existe, poner sanción a 0 o espacio
                     initial_replacements_scan[sancion_ph_scan] = formatear_valor_v6("0", sancion_ph_scan, tipos_columnas_csv_local)


        reemplazar_en_documento_v2(doc_scan, initial_replacements_scan, fase="inicial_escaneo") # reemplazar_en_documento_v2 debe estar definida

        # 7. Eliminar elementos inactivos, procesar tablas (copiar de generar_documentos_predial_core)
        tablas_protegidas_scan = [] # Debe ser una lista vacía para cada documento
        eliminar_elementos_inactivos_v_usuario(doc_scan, anos_inactivos, periodo_año_scan, periodo_bim_scan, tablas_protegidas_scan) # esta función debe estar definida
        procesar_tablas_suelo_construccion(doc_scan, anos_activos) # esta función debe estar definida

        # 8. Calcular Sumas (copiar de generar_documentos_predial_core)
        #    Esto es crucial para [TOTAL], [TOTAL PESOS CORREGIDOS], etc.
        #    Usa 'row' y los nombres de columna de nombres_columnas_csv_local.
        # ... (tu lógica completa de cálculo de sum_impuesto_anual_total, sum_monto_actualizado_total, etc.) ...
        #    total_redondeado_bd_scan = round(total_general_recalculado_scan)
        # Ejemplo simplificado (DEBES USAR TU LÓGICA COMPLETA):
        impuesto_anual_por_año_scan = {2022:0.0, 2023:0.0, 2024:0.0} #...llenar esto
        sum_impuesto_anual_total_scan = 0.0 # ...calcular esto
        sum_monto_actualizado_total_scan = 0.0 # ...calcular esto
        sum_recargos_total_scan = 0.0 # ...calcular esto
        sum_sancion_total_scan = sum_monto_actualizado_total_scan # Importante: Sanción visible = Monto Actualizado total

        for year_iter_scan in [2022, 2023, 2024]:
            for bim_iter_scan in range(1, 7):
                if (year_iter_scan < periodo_año_scan) or (year_iter_scan == periodo_año_scan and bim_iter_scan < periodo_bim_scan): continue
                campo_bim_scan_iter = f"BIMESTRE {bim_iter_scan} {year_iter_scan}"
                campo_act_scan_iter = f"MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim_iter_scan} {year_iter_scan}"
                campo_rec_scan_iter = f"RECARGOS POR BIMESTRE {bim_iter_scan} {year_iter_scan}"
                # Sanción original no se usa para la suma visible, pero podría existir en row
                # campo_san_scan_iter = f"SANCION POR BIMESTRE {bim_iter_scan} {year_iter_scan}"

                valor_bim_actual_scan_iter = safe_float(row.get(campo_bim_scan_iter, 0.0))
                impuesto_anual_por_año_scan[year_iter_scan] += valor_bim_actual_scan_iter
                sum_impuesto_anual_total_scan += valor_bim_actual_scan_iter
                sum_monto_actualizado_total_scan += safe_float(row.get(campo_act_scan_iter, 0.0))
                sum_recargos_total_scan += safe_float(row.get(campo_rec_scan_iter, 0.0))

        monto_actualizado_restados_scan = sum_monto_actualizado_total_scan - sum_impuesto_anual_total_scan
        # Total general recalculado usando sum_monto_actualizado_total_scan como la "sanción visible"
        total_general_recalculado_scan = sum_impuesto_anual_total_scan + monto_actualizado_restados_scan + sum_recargos_total_scan + sum_monto_actualizado_total_scan
        total_redondeado_bd_scan = round(total_general_recalculado_scan)


        # 9. Reemplazos Finales (copiar de generar_documentos_predial_core)
        #    Asegúrate que [TOTAL] usa total_redondeado_bd_scan
        #    y que [LEYENDA_FINAL_PAG17] usa pm_set_local_scan.
        final_replacements_scan = {
            # ... todos tus placeholders finales como [SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS], [TOTAL], [TOTAL PESOS CORREGIDOS], etc.
            # usando las sumas calculadas en el paso 8 y formatear_valor_v6 / formatear_texto_moneda
        }
        final_replacements_scan["[TOTAL]"] = formatear_valor_v6(total_redondeado_bd_scan, "[TOTAL]", {"[TOTAL]":"moneda"})
        final_replacements_scan["[TOTAL PESOS CORREGIDOS]"] = formatear_texto_moneda(total_redondeado_bd_scan) # formatear_texto_moneda debe estar definida
        final_replacements_scan["[LEYENDA_FINAL_PAG17]"] = "" if expediente_str_para_log in pm_set_local_scan else LEYENDA_PAG17_TEXTO # LEYENDA_PAG17_TEXTO debe estar definida

        # ... (el resto de los placeholders para sumas, bimestres ordinales, BM[año], etc.)
        # Ejemplo para [SUMA DE LA SANCION POR BIMESTRE] y su texto, usando sum_monto_actualizado_total_scan
        final_replacements_scan["[SUMA DE LA SANCION POR BIMESTRE]"] = formatear_valor_v6(sum_monto_actualizado_total_scan, "[SUMA DE LA SANCION POR BIMESTRE]", {"[SUMA DE LA SANCION POR BIMESTRE]": "moneda"})
        final_replacements_scan["[SUMA DE LA SANCION POR BIMESTRE TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total_scan)

        final_replacements_scan["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]"] = formatear_valor_v6(sum_impuesto_anual_total_scan, "[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", {"[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]":"moneda"})
        final_replacements_scan["[SUMA DE MONTOS ACTUALIZADOS]"] = formatear_valor_v6(sum_monto_actualizado_total_scan, "[SUMA DE MONTOS ACTUALIZADOS]", {"[SUMA DE MONTOS ACTUALIZADOS]":"moneda"})
        final_replacements_scan["[SUMA DE MONTOS ACTUALIZADOS TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total_scan)
        final_replacements_scan["[SUMA DE MONTOS ACTUALIZADOS RESTADOS]"] = formatear_valor_v6(monto_actualizado_restados_scan, "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]", {"[SUMA DE MONTOS ACTUALIZADOS RESTADOS]":"moneda"})
        final_replacements_scan["[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"] = formatear_texto_moneda(monto_actualizado_restados_scan)
        final_replacements_scan["[SUMA DE LOS MONTOS DE RECARGO TOTAL]"] = formatear_valor_v6(sum_recargos_total_scan, "[SUMA DE LOS MONTOS DE RECARGO TOTAL]", {"[SUMA DE LOS MONTOS DE RECARGO TOTAL]":"moneda"})
        final_replacements_scan["[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]"] = formatear_texto_moneda(sum_recargos_total_scan)
        final_replacements_scan["[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]"] = formatear_texto_moneda(sum_impuesto_anual_total_scan)

        # Placeholders para [IMPUESTO PREDIAL DEL AÑO X]
        for año_ip_scan in [2022, 2023, 2024]:
            valor_ip_anual_scan = round(impuesto_anual_por_año_scan[año_ip_scan], 2)
            placeholder_ip_scan = f"[IMPUESTO PREDIAL DEL AÑO {año_ip_scan}]"
            if valor_ip_anual_scan > 0:
                final_replacements_scan[placeholder_ip_scan] = formatear_valor_v6(valor_ip_anual_scan, placeholder_ip_scan, {placeholder_ip_scan: 'moneda'})
            else:
                final_replacements_scan[placeholder_ip_scan] = " " # O formatear_valor_v6(0.0, ...) si prefieres "$ 0.00"

        reemplazar_en_documento_v2(doc_scan, final_replacements_scan, fase="final_escaneo")

        # 10. Eliminar "CONSTRUCCION" si aplica (copiar de generar_documentos_predial_core)
        if verificar_ausencia_info_construccion_v2(row, anos_activos): # verificar_ausencia_info_construccion_v2 debe estar definida
            eliminar_palabra_especifica_del_documento(doc_scan, "CONSTRUCCION") # eliminar_palabra_especifica_del_documento debe estar definida

        # --- FIN: COPIA Y ADAPTACIÓN DE LÓGICA DE GENERACIÓN DE DOCX ---

        doc_scan.save(ruta_temp_docx_predial_scan)
        print(f"        - (Generar Doc. Escaneo Predial) DOCX temporal guardado: {ruta_temp_docx_predial_scan.name}")

        # --- Conversión a PDF (Siempre completo para esta función) ---
        if DOCX2PDF_INSTALLED: # Constante global definida en GeneradorPredial_logica.py
            print(f"        - (Generar Doc. Escaneo Predial) Convirtiendo a PDF: {ruta_pdf_destino_obj.name}")
            convert(str(ruta_temp_docx_predial_scan), str(ruta_pdf_destino_obj)) # convert de docx2pdf
            if ruta_pdf_destino_obj.exists() and ruta_pdf_destino_obj.stat().st_size > 0:
                pdf_final_generado_ok = True
                num_paginas_generado = contar_paginas_pdf(ruta_pdf_destino_obj) # contar_paginas_pdf debe estar definida
                print(f"        - (Generar Doc. Escaneo Predial) PDF generado ({num_paginas_generado if num_paginas_generado is not None else 'N/A'} págs).")
            else:
                print(f"        - (!) (Generar Doc. Escaneo Predial): PDF en destino '{ruta_pdf_destino_obj.name}' no creado o vacío tras conversión.")
        else:
            print(f"        - (!) (Generar Doc. Escaneo Predial): docx2pdf no instalado. No se pudo generar PDF para {expediente_str_para_log}.")

    except Exception as e_gen_doc_scan_predial:
        print(f"        - (!) Error GRAVE (Generar Doc. Escaneo Predial) para Exp {expediente_str_para_log}: {e_gen_doc_scan_predial}")
        import traceback
        traceback.print_exc() # Para depuración detallada
        pdf_final_generado_ok = False
        num_paginas_generado = None
    finally:
        if ruta_temp_docx_predial_scan.exists():
            try: ruta_temp_docx_predial_scan.unlink()
            except Exception as e_del_temp_docx:
                 print(f"        - (!) Advertencia: No se pudo eliminar DOCX temporal '{ruta_temp_docx_predial_scan.name}': {e_del_temp_docx}")

    return pdf_final_generado_ok, num_paginas_generado


def generar_documento_faltante_para_escaneo_predial(
    expediente_faltante_str,
    df_csv_principal_predial_completo, # <--- Este es el nombre correcto
    directorio_temporal_para_generado_obj,
    config_predial_actual
):
    """
    Genera un documento completo de Predial sobre la marcha para ser usado en el proceso de escaneo.
    Devuelve la ruta (Path object) al PDF temporal generado, o None si falla.
    """
    # Filtrar el DataFrame principal para obtener los datos de este expediente
    registros_csv_expediente = df_csv_principal_predial_completo[
    df_csv_principal_predial_completo['EXPEDIENTE'].astype(str).str.lstrip('0') == expediente_faltante_str
    ]

    if registros_csv_expediente.empty:
        print(f"            - (!) Error Interno: No se encontraron datos en CSV para generar el expediente faltante {expediente_faltante_str}.")
        return None

    # Lógica de generación (simplificada, ya que no necesitamos actualizar la BD aquí)
    try:
        # Aquí iría una versión adaptada de tu lógica de `generar_un_documento_predial`
        # Por simplicidad, asumimos que tienes una función que puede generar el DOCX
        # y lo convierte a PDF. Usaremos una función placeholder aquí.
        # En tu caso real, aquí deberías llamar a la lógica que crea el DOCX.
        print(f"            - (Placeholder) Generando DOCX y PDF para expediente {expediente_faltante_str}...")
        
        # Simulación de la creación de un PDF
        nombre_contrib = registros_csv_expediente.iloc[0].get("NOMBRE CONTRIBUYENTE", "desconocido")
        nombre_temp_pdf = f"TEMP_GEN_SCAN_PREDIAL_{expediente_faltante_str}_{limpiar_texto(nombre_contrib)}.pdf"
        ruta_pdf_generado_temporal = directorio_temporal_para_generado_obj / nombre_temp_pdf

        # Aquí llamarías a tu función real que convierte de DOCX a PDF.
        # Por ejemplo: convert(ruta_docx_temp, ruta_pdf_generado_temporal)
        # Como placeholder, crearemos un archivo vacío para que el flujo continúe.
        ruta_pdf_generado_temporal.touch() # Placeholder
        
        if ruta_pdf_generado_temporal.exists():
             return ruta_pdf_generado_temporal
        else:
            return None

    except Exception as e:
        print(f"            - (!) FALLO CRÍTICO generando documento faltante para Exp {expediente_faltante_str}: {e}")
        return None

def generar_documento_faltante_para_escaneo_predial(
    expediente_faltante_str, # Ya normalizado
    df_csv_principal_predial_completo,
    directorio_temporal_para_generado_obj, # Path object al directorio temporal general para esta ejecución de escaneo
    config_predial_actual # El mode_config
):
    if df_csv_principal_predial_completo is None or df_csv_principal_predial_completo.empty:
        print(f"            - (!) Error (Generar Faltante Scan Predial): DataFrame CSV principal no disponible para Exp {expediente_faltante_str}.")
        return None

    # Identificar la columna de expediente en el CSV de datos
    # Esta info viene de config_columnas.xlsx, mapeada por generador.py a config_predial_actual["col_expediente"]
    # pero para el CSV de datos, puede ser el mismo nombre o uno diferente si config_columnas.xlsx lo mapea.
    # Asumimos que generador.py ha puesto el nombre de la columna ID del CSV de datos en config_predial_actual
    # bajo una clave como "col_expediente_en_datos_csv" o similar.
    # Si no, usamos el nombre estándar de la BD "EXPEDIENTE" como fallback para buscar en el CSV.
    col_exp_csv_datos = config_predial_actual.get("col_expediente_en_datos_csv", config_predial_actual.get("col_expediente", "EXPEDIENTE"))

    if col_exp_csv_datos not in df_csv_principal_predial_completo.columns:
         print(f"            - (!) Error (Generar Faltante Scan Predial): Columna ID '{col_exp_csv_datos}' no encontrada en el DataFrame de datos CSV (BASE_DE_DATOS.xlsx).")
         return None

    # Filtrar el DataFrame CSV principal para obtener solo los registros de este expediente.
    # Necesitamos normalizar la columna del CSV para la comparación (quitar ceros a la izquierda).
    df_csv_principal_predial_completo['EXP_NORM_TEMP_PARA_FILTRO_SCAN'] = df_csv_principal_predial_completo[col_exp_csv_datos].astype(str).str.lstrip('0')

    registros_csv_expediente_faltante = df_csv_principal_predial_completo[
        df_csv_principal_predial_completo['EXP_NORM_TEMP_PARA_FILTRO_SCAN'] == expediente_faltante_str # expediente_faltante_str ya está normalizado
    ].copy()

    df_csv_principal_predial_completo.drop(columns=['EXP_NORM_TEMP_PARA_FILTRO_SCAN'], inplace=True, errors='ignore')

    if registros_csv_expediente_faltante.empty:
        print(f"            - (!) Error (Generar Faltante Scan Predial): No se encontraron datos en CSV principal para Exp {expediente_faltante_str}.")
        return None

    # Crear nombre y ruta para el PDF generado TEMPORALMENTE
    # Necesitamos el nombre del contribuyente desde el CSV
    col_nombre_csv_datos = config_predial_actual.get("col_nombre_csv_original", "NOMBRE CONTRIBUYENTE") # El nombre en config_columnas.xlsx para el nombre del contribuyente
    nombre_contrib_para_nombre_scan = registros_csv_expediente_faltante.iloc[0].get(col_nombre_csv_datos, 'DESCONOCIDO_CSV')

    nombre_base_pdf_temp_generado = limpiar_texto(f"TEMP_GEN_SCAN_PREDIAL_{expediente_faltante_str}_{nombre_contrib_para_nombre_scan}")
    ruta_pdf_generado_temporal_obj = directorio_temporal_para_generado_obj / f"{nombre_base_pdf_temp_generado}_{time.time_ns()}.pdf"

    # Obtener nombres y tipos de columnas del config_columnas.xlsx para pasarlos a la función de generación.
    # generador.py debería haber cargado estos y puestos en config_predial_actual.
    nombres_cols_csv_para_gen_faltante = config_predial_actual.get("nombres_columnas_csv_cargadas")
    tipos_cols_csv_para_gen_faltante = config_predial_actual.get("tipos_columnas_csv_cargadas")

    if not nombres_cols_csv_para_gen_faltante or not tipos_cols_csv_para_gen_faltante:
        print(f"            - (!) Advertencia (Generar Faltante Scan Predial): Nombres/Tipos de columnas CSV no encontrados en config. Intentando cargar desde '{config_predial_actual['config_cols_file_path']}'.")
        # cargar_config_columnas es de GeneradorPredial_logica.py (o debería ser importada/definida)
        # Esta función devuelve (nombres, tipos) o (None, None)
        nombres_temp, tipos_temp = cargar_config_columnas(config_predial_actual["config_cols_file_path"])
        if nombres_temp:
            nombres_cols_csv_para_gen_faltante = nombres_temp
            tipos_cols_csv_para_gen_faltante = tipos_temp
        else:
            print(f"            - (!) Error CRITICO (Generar Faltante Scan Predial): No se pudo cargar la configuración de columnas. Abortando generación de faltante para {expediente_faltante_str}.")
            return None

    pdf_generado_ok_scan, _ = generar_un_documento_predial_para_escaneo(
        datos_expediente_df=registros_csv_expediente_faltante,
        plantilla_path_obj=Path(config_predial_actual["template_file_path"]),
        ruta_pdf_destino_obj=ruta_pdf_generado_temporal_obj, # Se guarda en el directorio temporal
        expediente_str_para_log=expediente_faltante_str,
        config_predial_actual_local=config_predial_actual,
        nombres_columnas_csv_local=nombres_cols_csv_para_gen_faltante,
        tipos_columnas_csv_local=tipos_cols_csv_para_gen_faltante
    )

    if pdf_generado_ok_scan:
        print(f"            -> ÉXITO (Generar Faltante Scan Predial): Documento temporal generado para {expediente_faltante_str} en {ruta_pdf_generado_temporal_obj.name}.")
        return ruta_pdf_generado_temporal_obj
    else:
        print(f"            -> FALLO (Generar Faltante Scan Predial): No se pudo generar el documento temporal para {expediente_faltante_str}.")
        return None

def merge_pdfs_predial(ruta_pdf_oficial_obj, ruta_pdf_escaneado_obj, ruta_pdf_fusionado_salida_obj, config_predial_actual):
    """
    Fusiona un PDF oficial (puede ser None si no se pudo generar/encontrar) y un PDF escaneado.
    El PDF oficial (si existe) va primero.
    Retorna True si la fusión fue exitosa y se creó el archivo, False en caso contrario.
    """
    writer = PdfWriter()
    fusion_realizada_con_contenido = False

    # 1. Añadir páginas del PDF "oficial" (generado o de la BD)
    if ruta_pdf_oficial_obj and ruta_pdf_oficial_obj.exists() and ruta_pdf_oficial_obj.stat().st_size > 0:
        try:
            with open(ruta_pdf_oficial_obj, 'rb') as f_oficial:
                reader_oficial = PdfReader(f_oficial, strict=False)
                if reader_oficial.pages: 
                    for page in reader_oficial.pages:
                        writer.add_page(page)
                    fusion_realizada_con_contenido = True
                    print(f"            [+] Páginas del PDF oficial '{ruta_pdf_oficial_obj.name}' añadidas a la fusión.")
        except PdfReadError as e_merge_oficial_read:
             print(f"            [!] Advertencia (Merge Predial - PyPDF2): No se pudieron leer páginas del PDF oficial '{ruta_pdf_oficial_obj.name}': {e_merge_oficial_read}")
        except Exception as e_merge_oficial_other:
            print(f"            [!] Advertencia (Merge Predial - Otro): No se pudieron añadir páginas del PDF oficial '{ruta_pdf_oficial_obj.name}': {e_merge_oficial_other}")
    elif ruta_pdf_oficial_obj:
         print(f"            [*] Info (Merge Predial): PDF oficial '{ruta_pdf_oficial_obj.name}' no encontrado o vacío. No se añadirá a la fusión.")


    # 2. Añadir páginas del PDF escaneado (este DEBE existir)
    if ruta_pdf_escaneado_obj and ruta_pdf_escaneado_obj.exists() and ruta_pdf_escaneado_obj.stat().st_size > 0:
        try:
            with open(ruta_pdf_escaneado_obj, 'rb') as f_scan:
                reader_scan = PdfReader(f_scan, strict=False)
                if reader_scan.pages:
                    for page in reader_scan.pages:
                        writer.add_page(page)
                    fusion_realizada_con_contenido = True
                    print(f"            [+] Páginas del PDF escaneado '{ruta_pdf_escaneado_obj.name}' añadidas a la fusión.")
                else: 
                    print(f"            [!] Error (Merge Predial): PDF escaneado '{ruta_pdf_escaneado_obj.name}' no tiene páginas (inesperado).")
                    # Si no hay nada del oficial Y el escaneado está vacío, no hay nada que escribir.
                    if not writer.pages: return False 
        except PdfReadError as e_merge_scan_read:
            print(f"            [!] Error CRITICO (Merge Predial - PyPDF2): No se pudieron leer páginas del PDF escaneado '{ruta_pdf_escaneado_obj.name}': {e_merge_scan_read}")
            return False 
        except Exception as e_merge_scan_other:
            print(f"            [!] Error CRITICO (Merge Predial - Otro): No se pudieron añadir páginas del PDF escaneado '{ruta_pdf_escaneado_obj.name}': {e_merge_scan_other}")
            return False
    else: 
        print(f"            [!] Error CRÍTICO (Merge Predial): PDF escaneado fuente '{ruta_pdf_escaneado_obj}' no encontrado o vacío. No se puede fusionar.")
        # Si no hay nada del oficial Y el escaneado no existe/está vacío.
        if not writer.pages: return False

    # Guardar el PDF fusionado solo si se añadieron páginas
    if writer.pages:
        try:
            ruta_pdf_fusionado_salida_obj.parent.mkdir(parents=True, exist_ok=True) # Asegurar que la carpeta de salida exista
            with open(ruta_pdf_fusionado_salida_obj, 'wb') as f_out_merge:
                writer.write(f_out_merge)
            print(f"            -> PDF Fusionado (Predial) guardado en: {ruta_pdf_fusionado_salida_obj.name}")
            return True
        except Exception as e_write_merge_predial:
            print(f"            [!] Error (Merge Predial) al escribir PDF fusionado '{ruta_pdf_fusionado_salida_obj.name}': {e_write_merge_predial}")
            return False
    else: 
        print(f"            [!] (Merge Predial) No hay páginas para escribir en el PDF fusionado '{ruta_pdf_fusionado_salida_obj.name}'. No se creó el archivo.")
        return False

def reemplazar_parrafo_con_negritas(doc, placeholder, texto_completo, texto_negritas):
    """
    [VERSIÓN MEJORADA] Busca un párrafo que contenga un placeholder y lo reemplaza con texto
    que tiene una parte específica en negritas, preservando la fuente y alineación originales.
    """
    for p in doc.paragraphs:
        if placeholder in p.text:
            # Guardar formato original del primer 'run' y del párrafo
            original_run = p.runs[0] if p.runs else None
            original_font_name = original_run.font.name if original_run and original_run.font.name else None
            original_font_size = original_run.font.size if original_run and original_run.font.size else None
            p_alignment = p.alignment
            p_style = p.style

            # Dividir el texto
            partes = texto_completo.split(texto_negritas)
            if len(partes) != 2:
                p.text = texto_completo
                continue

            parte1, parte2 = partes[0], partes[1]

            # Limpiar el párrafo
            for run in p.runs:
                p._element.remove(run._element)

            # Añadir las nuevas partes con el formato correcto
            run1 = p.add_run(parte1)
            run2 = p.add_run(texto_negritas)
            run2.bold = True
            run3 = p.add_run(parte2)

            # Reaplicar fuente y tamaño a todas las partes nuevas
            for new_run in [run1, run2, run3]:
                if original_font_name:
                    new_run.font.name = original_font_name
                if original_font_size:
                    new_run.font.size = original_font_size

            # Reaplicar formato del párrafo
            p.alignment = p_alignment
            p.style = p.style
            return

def process_single_expediente_scan_predial(
    expediente_escaneado_str,
    ruta_pdf_temporal_grupo_escaneado_obj,
    num_paginas_grupo_escaneado,
    metodo_extraccion_ocr,
    paginas_detectadas_list,
    df_bd_maestra_global_predial,
    df_csv_principal_global_predial,
    nombre_pdf_fuente_original_str,
    directorio_temporal_principal_obj,
    directorio_salida_escaneos_exp_obj,
    config_predial_actual,
    funcion_de_subida
):
    """
    Procesa un único expediente de PREDIAL extraído del escaneo.
    """
    print(f"        --- Procesando Expediente Predial: {expediente_escaneado_str} ---")

    # --- 1. PREPARACIÓN INICIAL ---
    paginas_str = ", ".join([str(p + 1) for p in paginas_detectadas_list])
    datos_para_excel = {
        "EXPEDIENTE": expediente_escaneado_str, "NOMBRE CONTRIBUYENTE": "", "DIRECCION": "",
        "COLONIA": "", "BIMESTRE": "", "AÑOS": "", "PaginasEscaneadas": num_paginas_grupo_escaneado,
        "NombreArchivoFusionado": "", "RutaArchivoFusionado": "", "PaginasDetectadas": paginas_str
    }
    df_bd_maestra_actualizada_local = df_bd_maestra_global_predial.copy()
    ruta_pdf_oficial_para_fusion = None
    nombre_contrib_final = "CONTRIBUYENTE_DESCONOCIDO"

    # --- 2. LÓGICA CENTRAL: BUSCAR EL EXPEDIENTE ---
    # La validación ya se hizo en group_and_split, aquí solo recuperamos los datos
    
    # Primero buscar en la BD Maestra
    entrada_bd_maestra = df_bd_maestra_actualizada_local[
        df_bd_maestra_actualizada_local['EXPEDIENTE'].astype(str).str.lstrip('0') == expediente_escaneado_str
    ]

    if not entrada_bd_maestra.empty:
        # --- CAMINO A: EXPEDIENTE ENCONTRADO EN BD MAESTRA ---
        print(f"            - Expediente encontrado en BD Maestra. Verificando PDF...")
        fila_bd = entrada_bd_maestra.iloc[0]
        nombre_contrib_final = fila_bd.get("NOMBRE CONTRIBUYENTE", "DESCONOCIDO_EN_BD")
        # Poblar datos para el log de Excel con info de la BD
        datos_para_excel.update({
            "NOMBRE CONTRIBUYENTE": nombre_contrib_final,
            "DIRECCION": fila_bd.get("DIRECCION", ""),
            "COLONIA": fila_bd.get("COLONIA", ""),
            "BIMESTRE": fila_bd.get("BIMESTRE", ""),
            "AÑOS": fila_bd.get("AÑOS", "")
        })

        ruta_relativa_pdf = fila_bd.get("Ruta PDF Generado", "")
        if ruta_relativa_pdf:
            ruta_pdf_oficial_para_fusion = Path(config_predial_actual["output_docs_path"]) / ruta_relativa_pdf
            if not ruta_pdf_oficial_para_fusion.is_file():
                print(f"            - (!) ADVERTENCIA: La ruta en BD no existe en disco: {ruta_pdf_oficial_para_fusion}")
                ruta_pdf_oficial_para_fusion = None # Forzar regeneración
    else:
        # --- CAMINO B: EXPEDIENTE HUÉRFANO (Solo existe en BASE_DE_DATOS.xlsx) ---
        print(f"            - Expediente no en BD Maestra. Recuperando datos de BASE_DE_DATOS.xlsx...")
        datos_csv = buscar_datos_expediente_en_csv_predial(expediente_escaneado_str, df_csv_principal_global_predial, config_predial_actual)
        if datos_csv.get('encontrado_en_csv'):
            nombre_contrib_final = datos_csv.get('nombre_contribuyente', 'DESCONOCIDO_EN_CSV')
            datos_para_excel.update({
                "NOMBRE CONTRIBUYENTE": nombre_contrib_final,
                "COLONIA": datos_csv.get('colonia', '')
            })
        else:
             # Este caso es improbable si group_and_split funciona bien, pero es un buen seguro
             print(f"            - (!) ERROR: Expediente '{expediente_escaneado_str}' no se pudo encontrar en ninguna fuente de datos.")
             datos_para_excel["Error"] = "EXPEDIENTE_SIN_DATOS"
             return datos_para_excel, df_bd_maestra_global_predial
             
    # Si en cualquier camino el PDF no se encontró, se genera como respaldo
    if not ruta_pdf_oficial_para_fusion:
        print(f"            - PDF oficial no encontrado o no válido, generando sobre la marcha...")
        ruta_pdf_oficial_para_fusion = generar_documento_faltante_para_escaneo_predial(
            expediente_escaneado_str, df_csv_principal_global_predial, directorio_temporal_principal_obj, config_predial_actual
        )
    
    # --- 3. FUSIÓN Y ACTUALIZACIÓN FINAL ---
    if not ruta_pdf_oficial_para_fusion:
        datos_para_excel["Error"] = "FALLO_GENERACION/FUSION"
        return datos_para_excel, df_bd_maestra_actualizada_local

    nombre_archivo_fusionado = limpiar_texto(f"PREDIAL_SCAN_{expediente_escaneado_str}_{nombre_contrib_final}")
    ruta_pdf_fusionado_final = directorio_salida_escaneos_exp_obj / f"{nombre_archivo_fusionado}.pdf"

    if merge_pdfs_predial(ruta_pdf_oficial_para_fusion, ruta_pdf_temporal_grupo_escaneado_obj, ruta_pdf_fusionado_final, config_predial_actual):
        ruta_relativa_fusionado = str(ruta_pdf_fusionado_final.relative_to(Path(config_predial_actual["output_docs_path"]))).replace('\\', '/')
        datos_para_excel.update({
            "NombreArchivoFusionado": ruta_pdf_fusionado_final.name,
            "RutaArchivoFusionado": ruta_relativa_fusionado
        })

        # --- Actualización Final de la Base de Datos ---
        # Prepara los datos que se van a guardar. Esta vez no se crea un registro nuevo aquí,
        # se asume que la función de generación se encargará de ello si es necesario.
        # Aquí solo actualizamos los datos del escaneo.
        datos_registro_bd = {
            "EXPEDIENTE": expediente_escaneado_str,
            "Ruta PDF Escaneado": ruta_relativa_fusionado,
            "BASE DE DATOS ESCANEO": f"BD_Escaneo_{Path(nombre_pdf_fuente_original_str).stem}_Predial.xlsx",
            "MOVIMIENTO": "ESCANEADO"
        }
        # actualizar_o_agregar_registro_bd se encargará de encontrar la fila y actualizar estos campos.
        df_bd_maestra_actualizada_local = actualizar_o_agregar_registro_bd(df_bd_maestra_actualizada_local, datos_registro_bd)
        
        if funcion_de_subida:
            funcion_de_subida(ruta_pdf_fusionado_final, config_predial_actual)
    else:
        datos_para_excel["Error"] = "FALLO_FUSION"
    
    if ruta_pdf_oficial_para_fusion and ruta_pdf_oficial_para_fusion.name.startswith("TEMP_GEN_SCAN_"):
        ruta_pdf_oficial_para_fusion.unlink(missing_ok=True)
        
    return datos_para_excel, df_bd_maestra_actualizada_local

def run_scan_and_process_predial(df_bd_maestra_global, df_csv_principal_global, config_predial_actual, funcion_de_subida=None):
    """
    Función principal que orquesta el escaneo de PDFs de Predial, procesamiento y registro.
    df_bd_maestra_global: DataFrame de la BD Maestra de Predial (ya cargado por generador.py).
    df_csv_principal_global: DataFrame de BASE_DE_DATOS.xlsx de Predial (ya cargado por generador.py).
    config_predial_actual: Diccionario de configuración del modo (de generador.py).
    """
    print("\n--- Iniciando Modo Escaneo de Documentos de PREDIAL ---")

    # Rutas (obtenidas de config_predial_actual que es mode_config)
    # Carpeta de entrada para PDFs a escanear. El usuario debe crearla y poner los PDFs allí.
    # Ej: [Directorio_Modo_Predial]/PDFs_A_Escanear_Predial/
    # config_predial_actual["base_path"] apunta al directorio del modo (ej. .../PREDIAL_FERNANDO/)
    source_pdf_folder_scan_predial = Path(config_predial_actual["base_path"]) / "PDFs_A_Escanear_Predial"

    # Carpeta de salida para los artefactos del escaneo (PDFs fusionados, Excel logs individuales)
    # config_predial_actual["output_docs_path"] apunta a .../DOCUMENTOS/ dentro del modo.
    # generador.py ya crea una subcarpeta "ESCANEADOS_PREDIAL" dentro de output_docs_path.
    output_scan_artifacts_predial_root = Path(config_predial_actual["output_docs_path"]) / "ESCANEADOS_PREDIAL"

    print(f"    - Carpeta de PDFs a Escanear (Predial): {source_pdf_folder_scan_predial}")
    print(f"    - Carpeta Raíz de Salida para Artefactos de Escaneo (Predial): {output_scan_artifacts_predial_root}")

    # Validar dependencias críticas (Tesseract, fitz, Pillow)
    try:
        pytesseract.get_tesseract_version() # Lanza error si Tesseract no está o no se encuentra
        # fitz y PIL se verifican en los imports iniciales, pero una verificación aquí no está de más.
        if 'fitz' not in sys.modules: raise ImportError("PyMuPDF (fitz) no está disponible.")
    except Exception as e_dep_ocr_scan_predial:
        print(f"    - (!) ERROR CRÍTICO (Scan Predial): Falta una dependencia de OCR o Tesseract no está configurado: {e_dep_ocr_scan_predial}")
        print("        El modo Escaneo de Predial no puede continuar. Verifique las instalaciones.")
        return df_bd_maestra_global # Devolver la BD Maestra sin cambios

    if not source_pdf_folder_scan_predial.is_dir():
        print(f"    - (!) ERROR CRÍTICO (Scan Predial): La carpeta de PDFs fuente para escanear NO existe: {source_pdf_folder_scan_predial}")
        print(f"          Por favor, cree esta carpeta y coloque los PDFs a escanear allí.")
        # Opcionalmente, crearla: source_pdf_folder_scan_predial.mkdir(parents=True, exist_ok=True)
        return df_bd_maestra_global

    # La carpeta output_scan_artifacts_predial_root ya debería haber sido creada por generador.py
    output_scan_artifacts_predial_root.mkdir(parents=True, exist_ok=True) # Asegurar por si acaso

    lista_pdfs_a_escanear_predial = list(source_pdf_folder_scan_predial.glob("*.pdf"))
    if not lista_pdfs_a_escanear_predial:
        print(f"    - No se encontraron archivos PDF en '{source_pdf_folder_scan_predial}'. Nada que escanear para Predial.")
        return df_bd_maestra_global

    df_bd_maestra_modificada_total_scan_predial = df_bd_maestra_global.copy()

    for i_pdf_fuente, pdf_fuente_path_predial in enumerate(lista_pdfs_a_escanear_predial):
        print(f"\n    --- Procesando PDF Fuente Predial ({i_pdf_fuente + 1}/{len(lista_pdfs_a_escanear_predial)}): {pdf_fuente_path_predial.name} ---")

        nombre_base_pdf_fuente_predial = pdf_fuente_path_predial.stem

        # --- CAMBIO 1: Definir la ruta de la carpeta de resultados ---
        carpeta_salida_lote_actual = output_scan_artifacts_predial_root / f"Resultados_Scan_{nombre_base_pdf_fuente_predial}"
        carpeta_salida_lote_actual.mkdir(parents=True, exist_ok=True)

        # --- CAMBIO 2: La bitácora de escaneo ahora se guardará DENTRO de la carpeta del lote ---
        nombre_excel_escaneo_individual_predial = f"BD_Escaneo_{nombre_base_pdf_fuente_predial}_Predial.xlsx"
        ruta_excel_escaneo_individual_predial = carpeta_salida_lote_actual / nombre_excel_escaneo_individual_predial

        # --- CAMBIO 3: Inicializar listas para los nuevos logs ---
        registros_para_excel_fuente_actual_predial = []
        log_expedientes_no_encontrados = []

        if ruta_excel_escaneo_individual_predial.exists():
            print(f"        - El archivo Excel de escaneo Predial '{nombre_excel_escaneo_individual_predial}' ya existe. Saltando procesamiento de '{pdf_fuente_path_predial.name}'.")
            continue

        with tempfile.TemporaryDirectory(prefix=f"scan_predial_temp_{nombre_base_pdf_fuente_predial}_") as temp_dir_str_predial:
            temp_dir_path_predial = Path(temp_dir_str_predial)

            # En la nueva versión, esta función ahora devuelve también los no encontrados
            lista_grupos_exp_escaneado, expedientes_no_encontrados_map = group_and_split_pdf_dynamically_scan_predial(
                pdf_fuente_path_predial,
                temp_dir_path_predial,
                config_predial_actual,
                df_bd_maestra_modificada_total_scan_predial, # Argumento 1
                df_csv_principal_global                      # Argumento 2 (EL QUE FALTABA)
            )

            # Poblar la lista de no encontrados para el log
            for exp, data in expedientes_no_encontrados_map.items():
                log_expedientes_no_encontrados.append({
                    "ExpedienteNoEncontrado": exp,
                    "PrimeraPaginaDetectada": data["primera_pagina"]
                })

            if not lista_grupos_exp_escaneado:
                print(f"        - No se encontraron grupos de expedientes válidos en '{pdf_fuente_path_predial.name}'.")
            else:
                print(f"        - Se encontraron {len(lista_grupos_exp_escaneado)} expedientes válidos. Procesando cada grupo...")
                for exp_str_g, ruta_pdf_grupo_g, num_pags_g, metodo_ocr_g, paginas_list in lista_grupos_exp_escaneado:
                    datos_excel_exp, df_bd_maestra_modificada_total_scan_predial = process_single_expediente_scan_predial(
                        expediente_escaneado_str=exp_str_g,
                        ruta_pdf_temporal_grupo_escaneado_obj=ruta_pdf_grupo_g,
                        num_paginas_grupo_escaneado=num_pags_g,
                        metodo_extraccion_ocr=metodo_ocr_g,
                        df_bd_maestra_global_predial=df_bd_maestra_modificada_total_scan_predial,
                        df_csv_principal_global_predial=df_csv_principal_global,
                        nombre_pdf_fuente_original_str=pdf_fuente_path_predial.name,
                        directorio_temporal_principal_obj=temp_dir_path_predial,
                        directorio_salida_escaneos_exp_obj=carpeta_salida_lote_actual, # Usar la nueva ruta
                        config_predial_actual=config_predial_actual,
                        funcion_de_subida=funcion_de_subida,
                        paginas_detectadas_list=paginas_list,
                    )
                    if datos_excel_exp:
                        registros_para_excel_fuente_actual_predial.append(datos_excel_exp)

        # --- Guardar la bitácora de escaneos exitosos ---
        # --- Guardar la bitácora de escaneos exitosos ---
        if registros_para_excel_fuente_actual_predial or log_expedientes_no_encontrados:
            # Unificar los expedientes exitosos y los no encontrados en un solo DataFrame
            df_exitosos = pd.DataFrame(registros_para_excel_fuente_actual_predial, columns=COLUMNAS_EXCEL_ESCANEO_INDIVIDUAL_PREDIAL)
            df_no_encontrados = pd.DataFrame(log_expedientes_no_encontrados)
            if not df_no_encontrados.empty:
                df_no_encontrados.rename(columns={"ExpedienteNoEncontrado": "EXPEDIENTE", "PrimeraPaginaDetectada": "PaginasDetectadas"}, inplace=True)
            
            df_excel_individual_predial = pd.concat([df_exitosos, df_no_encontrados], ignore_index=True)

            try:
                # --- Lógica para el título dinámico del reporte ---
                titulo_reporte = "PREDIAL - DESPACHO MALDONADO"
                match_lote = re.search(r'Lote\s*([A-Z])', pdf_fuente_path_predial.name, re.IGNORECASE)
                if match_lote:
                    lote = match_lote.group(1).upper()
                    titulo_reporte = f"PREDIAL - LOTE {lote} - DESPACHO MALDONADO"

                # Obtener la lista de expedientes para el reporte
                lista_de_expedientes = df_excel_individual_predial["EXPEDIENTE"].tolist()

                with pd.ExcelWriter(ruta_excel_escaneo_individual_predial, engine='openpyxl') as writer:
                    # Escribir la primera hoja (el log detallado)
                    df_excel_individual_predial.to_excel(writer, sheet_name="Log_Escaneo", index=False)
                    
                    # Escribir la segunda hoja (el reporte para impresión)
                    crear_hoja_reporte_impresion(
                        writer,
                        datos_para_reporte=lista_de_expedientes,
                        titulo_principal="PREDIAL", # <-- Título para la lógica interna
                        cabecera_columna="EXPEDIENTE" # <-- Cabecera de la columna
                    )
                
                print(f"        - Bitácora y Reporte de Impresión guardados en: {ruta_excel_escaneo_individual_predial.name}")
                
                if funcion_de_subida:
                    print(f"        - Subiendo bitácora de escaneo al servidor...")
                    funcion_de_subida(ruta_excel_escaneo_individual_predial, config_predial_actual)

            except Exception as e_save_excel:
                print(f"        - (!) Error guardando el archivo Excel de escaneo con reporte: {e_save_excel}")
        # --- FIN DEL BLOQUE DE GUARDADO ---
        # --- CAMBIO 4: Guardar el NUEVO log de expedientes no encontrados ---
        if log_expedientes_no_encontrados:
            ruta_log_no_encontrados = carpeta_salida_lote_actual / "log_expedientes_no_encontrados.xlsx"
            df_no_encontrados = pd.DataFrame(log_expedientes_no_encontrados)
            try:
                print(f"        - ¡Atención! Se encontraron expedientes no registrados. Guardando log en: {ruta_log_no_encontrados}")
                df_no_encontrados.to_excel(ruta_log_no_encontrados, index=False, engine='openpyxl')
            except Exception as e_save_log_nf:
                print(f"        - (!) Error guardando el log de no encontrados: {e_save_log_nf}")

    print("\n--- Modo Escaneo de Documentos de PREDIAL Finalizado ---")
    return df_bd_maestra_modificada_total_scan_predial

def extract_expediente_from_page_scan_predial(page_fitz_obj, page_num_debug, config_predial_actual):
    """
    Extrae expedientes de la parte SUPERIOR e INFERIOR, los normaliza
    (quitando ceros iniciales) y los compara para validación.
    """
    exp_top_norm = None
    exp_bottom_norm = None

    try:
        zoom = float(config_predial_actual.get("ocr_zoom_factor_predial", 2.5))
        mat = fitz.Matrix(zoom, zoom)
        pix = page_fitz_obj.get_pixmap(colorspace=fitz.csGRAY, matrix=mat, alpha=False)
        img = Image.frombytes("L", [pix.width, pix.height], pix.samples)
        img_width, img_height = img.size

        # --- Búsqueda en la parte SUPERIOR ---
        top_crop_end_y = int(img_height * 0.35) # Buscar en el 35% superior
        area_top = (0, 0, img_width, top_crop_end_y)
        texto_top = pytesseract.image_to_string(img.crop(area_top), lang='spa', config="--psm 4")
        
        match_top = EXPEDIENTE_REGEX_SCAN_PREDIAL_NODOT.search(texto_top)
        if match_top:
            # Normalización: quitar ceros a la izquierda
            exp_top_norm = match_top.group(1).lstrip('0')

        # --- Búsqueda en la parte INFERIOR ---
        bottom_crop_start_y = int(img_height * 0.80) # Buscar en el 20% inferior
        area_bottom = (0, bottom_crop_start_y, img_width, img_height)
        texto_bottom = pytesseract.image_to_string(img.crop(area_bottom), lang='spa', config="--psm 6")
        
        match_bottom = EXPEDIENTE_REGEX_SCAN_PREDIAL_DOT.search(texto_bottom)
        if match_bottom:
            # Normalización: quitar ceros a la izquierda
            exp_bottom_norm = match_bottom.group(1).lstrip('0')
        
        # --- Lógica de Decisión ---
        if exp_top_norm and exp_bottom_norm:
            if exp_top_norm == exp_bottom_norm:
                return [exp_top_norm], "OCR_MATCH"
            else:
                # Si hay discrepancia, devuelve ambos para que la siguiente función decida
                return [exp_top_norm, exp_bottom_norm], "OCR_DISCREPANCIA"
        elif exp_top_norm:
            return [exp_top_norm], "OCR_SOLO_TOP"
        elif exp_bottom_norm:
            return [exp_bottom_norm], "OCR_SOLO_BOTTOM"
        else:
            return [], "OCR_NO_ENCONTRADO"

    except Exception as e:
        print(f"        - (!) Error inesperado en OCR Predial página {page_num_debug}: {e}")
        return [], "OCR_EXCEPCION"


def configurar_locale():
    locales_a_intentar = ['es_MX.UTF-8', 'es-MX', 'es_MX', 'Spanish_Mexico', 'es_ES.UTF-8', 'es-ES', 'es_ES', 'Spanish', '']
    locale_configurado = ""
    found = False
    for loc in locales_a_intentar:
        try:
            locale.setlocale(locale.LC_ALL, loc)
            locale_configurado = locale.getlocale(locale.LC_TIME)
            try:
                test_date = datetime(2024, 4, 1)
                month_name = test_date.strftime('%B')
                if "bril" in month_name.lower() or "apr" in month_name.lower() or "avr" in month_name.lower(): # Consider Spanish variations
                    print(f"Locale de tiempo configurado y verificado: {locale_configurado} (Ej: {month_name})")
                    found = True
                    break
                else:
                    print(f"Advertencia: Locale {loc} ({locale_configurado}) no parece ser español (Mes: {month_name}). Intentando siguiente.")
            except Exception as e_verify:
                print(f"Advertencia: No se pudo verificar locale {loc} ({locale_configurado}): {e_verify}")
        except locale.Error:
            continue
        except Exception as e_setlocale:
            print(f"Advertencia: Error inesperado configurando locale {loc}: {e_setlocale}")
            continue
    if not found:
        print("Error CRITICO: No se pudo configurar un locale en español para formatear fechas.")
        print("Asegúrate de que un locale como 'es_MX.UTF-8' o 'es_ES.UTF-8' esté instalado en tu sistema.")
        return False
    try:
        print(f"Localización general configurada: {locale.getlocale()}")
    except Exception as e:
        print(f"Advertencia: Locale configurado pero no se pudo obtener el nombre general: {e}")
    return True

def cargar_config_columnas(filepath):
    encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1']
    delimiters_to_try = ['|', ',', ';']
    df_config = None # Inicializar fuera de los bucles internos para mantener el último leído con éxito
    error_msg = "No se pudo leer el archivo de configuración." 

    for delim in delimiters_to_try:
        for enc in encodings_to_try:
            try:
                # Intentar leer con la combinación actual
                df_config_temp = pd.read_csv(filepath, delimiter=delim, encoding=enc, dtype=str, skipinitialspace=True)
                df_config_temp.columns = df_config_temp.columns.str.strip()
                
                # Verificar si las columnas necesarias están presentes
                if "NombreEncabezado" in df_config_temp.columns and "TipoDato" in df_config_temp.columns:
                    print(f"      Configuración leída (Codificación:{enc}, Delimitador:'{delim}')")
                    df_config = df_config_temp # Guardar la configuración válida
                    
                    # Procesar la configuración válida
                    df_config = df_config.fillna('')
                    df_config["NombreEncabezado"] = df_config["NombreEncabezado"].str.strip()
                    if df_config["NombreEncabezado"].duplicated().any():
                        print("Advertencia: Nombres de encabezado duplicados en el archivo de configuración. Se usarán los primeros encontrados.")
                        df_config = df_config.drop_duplicates(subset="NombreEncabezado", keep='first')
                    
                    nombres = df_config["NombreEncabezado"].tolist()
                    tipos = df_config.set_index('NombreEncabezado')['TipoDato'].str.strip().str.lower().to_dict()
                    print(f"Configuración de {len(nombres)} columnas cargada.")
                    return nombres, tipos # Retornar inmediatamente si es exitoso
                else:
                    # Actualizar mensaje de error si las columnas esperadas no se encuentran
                    error_msg = f"Encabezados 'NombreEncabezado' o 'TipoDato' no encontrados (Codificación:{enc}, Delimitador:'{delim}'). Detectados: {df_config_temp.columns.tolist()}"
            except FileNotFoundError:
                print(f"Error: Archivo de configuración '{filepath}' no encontrado.")
                return None, None # Retornar inmediatamente si el archivo no existe
            except Exception as e:
                error_msg = f"Error leyendo el archivo de configuración '{filepath}' (Codificación:{enc}, Delimitador:'{delim}'): {e}"
                # No es necesario asignar df_config = None aquí, ya que df_config_temp no se asignó a df_config
                continue # Continuar con la siguiente codificación/delimitador
        # Si se encontró una configuración válida en el bucle interno de codificaciones,
        # df_config no será None y el return ya habrá ocurrido.
        # Si se completa el bucle de codificaciones sin éxito para el delimitador actual,
        # df_config podría seguir siendo None o tener el valor de un intento anterior fallido de otro delimitador.
        # Por eso es importante que el return esté dentro del if que valida las columnas.

    # Si se sale de todos los bucles sin haber retornado, significa que falló.
    print(f"Error: No se pudo leer la configuración de columnas desde '{filepath}'. Último error: {error_msg}")
    return None, None

# In GeneradorPredial_logica.py

def formatear_valor_v6(valor, nombre_columna, tipos_columnas):
    if pd.isna(valor):
        return ""
    valor_str = str(valor).strip()
    nombre_col_limpio = nombre_columna.strip()
    
    tipo_dato = 'texto' # Default
    if tipos_columnas is not None:
        tipo_dato = tipos_columnas.get(nombre_col_limpio, 'texto').lower()

    # --- Override tipo_dato for specific placeholder names (for sums) ---
    if nombre_col_limpio.startswith(("[SUMA", "[TOTAL", "[IMPUESTO PREDIAL DEL AÑO")):
        tipo_dato = 'moneda'

    if not valor_str or valor_str == '-' or valor_str.isspace():
        if tipo_dato == 'moneda': # Return "$ 0.00" for empty monetary values
            return "$ 0.00"
        return ""

    # --- Handle non-numeric types first that should return before 'numero' is defined ---
    if tipo_dato == 'texto':
        return valor_str
    if tipo_dato == 'texto_preformateado':
        return valor_str 
    if tipo_dato == 'texto_sin_formato_num': # e.g., EXPEDIENTE, CP
        return valor_str
    if tipo_dato == 'fecha':
        return valor_str if valor_str else "Sin fecha"
    
    if tipo_dato == 'porcentaje':
        if valor_str.endswith('%'):
            return valor_str # Already formatted as "X%"
        try:
            # Convertir a float. Asumimos que el valor de entrada es un decimal (ej., "0.32" para 32%)
            numero_decimal = float(valor_str.replace(",", "")) 
            numero_porcentaje = numero_decimal * 100 # Convertir 0.32 a 32.0
            # Formatear como entero (sin decimales) y añadir el símbolo %
            return f"{locale.format_string('%.0f', numero_porcentaje, grouping=False)}%"
        except (ValueError, TypeError):
            print(f"    Advertencia (formatear_valor_v6): Valor '{valor_str}' para columna '{nombre_columna}' con TipoDato 'PORCENTAJE' no es numérico. Se devuelve como está.")
            return valor_str

    # --- At this point, we expect a numeric type that needs 'numero' ---
    # --- Define 'numero' for all subsequent numeric processing ---
    try:
        # Standardize numeric conversion
        numero = float(valor_str.replace(",", "").replace("$", ""))
    except (ValueError, TypeError):
        # If conversion to float fails for a supposedly numeric type (that wasn't caught above)
        print(f"    Advertencia (formatear_valor_v6): Valor '{valor_str}' para columna '{nombre_columna}' (TipoDato: {tipo_dato}) no es numérico válido. Se devuelve como está.")
        return valor_str # Return original string if not a valid number for numeric types

    # --- Specific formatting for numeric TipoDato (now 'numero' is defined) ---
    if tipo_dato == 'moneda':
        # Esta es la Opción B que te sugerí, asegurando que 'numero' ya está definido.
        formatted_value_monetary = locale.format_string('%.2f', numero, grouping=True, monetary=True)
        currency_symbol = locale.localeconv().get('currency_symbol', '$').strip()
        
        # Verificar si el símbolo de moneda ya fue agregado por monetary=True
        # Considerar números negativos, ej. -$10.00
        is_negative = numero < 0
        temp_check_val = formatted_value_monetary
        if is_negative and temp_check_val.startswith('-'):
            temp_check_val = temp_check_val[1:] # Quitar el '-' para la comprobación del símbolo

        if not temp_check_val.strip().startswith(currency_symbol):
            # Si no está, añadirlo manualmente.
            sign = "-" if is_negative else ""
            # Usamos abs(numero) para el formateo numérico y anteponemos el signo y símbolo
            return f"{sign}{currency_symbol}{locale.format_string('%.2f', abs(numero), grouping=True)}"
        return formatted_value_monetary # Ya tiene el símbolo (o no se necesita añadir)

    elif tipo_dato == 'decimal_3':
        return locale.format_string("%.3f", numero, grouping=True)
    elif tipo_dato == 'decimal_4':
        return locale.format_string("%.4f", numero, grouping=True)
    elif tipo_dato == 'numero_general': # E.g., number with 2 decimals and grouping
        return locale.format_string("%.2f", numero, grouping=True)
    
    # --- Special handling for column names (can override TipoDato) ---
    sin_decimales_campos = [
        f"MESES DE RETARDO EN EL PAGO {i} {a}" for a in [2022, 2023, 2024] for i in range(1, 7)
    ]
    campos_cierre_redondeado = [ 
        f"VALOR CATASTRAL {i} {a}" for a in [2022, 2023, 2024] for i in range(1, 6)
    ] + [
        f"VALOR CATASTRAL CONSTRUCCION {i} {a}" for a in [2022, 2023, 2024] for i in range(1, 13)
    ]

    if nombre_col_limpio in campos_cierre_redondeado:
        numero_entero = round(numero)
        # Asegurar que también tenga el símbolo de moneda si es necesario (consistente con 'moneda')
        currency_symbol_special = locale.localeconv().get('currency_symbol', '$').strip()
        return f"{currency_symbol_special} {locale.format_string('%.2f', numero_entero, grouping=True)}"
    
    if nombre_col_limpio in sin_decimales_campos:
        return locale.format_string("%.0f", numero, grouping=True)

    # Fallback para otros casos numéricos no explícitamente manejados
    return locale.format_string("%.2f", numero, grouping=True)

# --- Nueva Función Core para ser llamada por generador_maestro.py ---
def generar_documentos_predial_core(
    df_datos_para_procesar,
    pm_set_actual,
    config_predial_actual,
    modo_generacion_solicitado,
    max_docs_a_generar,
    df_bd_maestra_actualizada,
    nombres_columnas_csv, # Nombres de columnas como están en df_datos_para_procesar
    tipos_columnas_csv,   # Tipos de datos definidos en config_columnas.xlsx
    letra_lote,
    contador_inicial_lote,
    funcion_de_subida
):
    """
    Función principal para generar documentos de Predial, llamada por el generador maestro.
    Utiliza DataFrames pre-cargados.
    """
    print(f"\n--- (Predial Logic Core) Iniciando Generación en Modo: {modo_generacion_solicitado} ---")
    print(f"--- (Predial Logic Core) Máximo a generar: {'Todos los aplicables' if max_docs_a_generar == 0 or max_docs_a_generar >= len(df_datos_para_procesar) else max_docs_a_generar} ---")

    plantilla_path = config_predial_actual["template_file_path"]
    carpeta_principal_salida = config_predial_actual["output_docs_path"]
    col_expediente_bd = config_predial_actual["col_expediente"]
    col_estado_bd = config_predial_actual["col_estado_bd_maestra"]

    # Acceder a constantes de estado desde la configuración del modo
    ESTADO_PENDIENTE_MODO = config_predial_actual["ESTADO_PENDIENTE"]
    ESTADO_ERROR_GENERACION_MODO = config_predial_actual["ESTADO_ERROR_GENERACION"]
    ESTADO_GEN_COMPLETO_MODO = config_predial_actual["ESTADO_GEN_COMPLETO"]
    ESTADO_GEN_ULTIMA_MODO = config_predial_actual["ESTADO_GEN_ULTIMA"]
    ESTADO_IMP_COMPLETO_MODO = config_predial_actual["ESTADO_IMP_COMPLETO"]
    ESTADO_IMP_ULTIMA_MODO = config_predial_actual["ESTADO_IMP_ULTIMA"]
    ESTADO_PDF_NO_ENCONTRADO_MODO = config_predial_actual["ESTADO_PDF_NO_ENCONTRADO"] # <--- AÑADE ESTA LÍNEA


    MODO_COMPLETO_INTERNO = config_predial_actual["MODO_GENERACION_COMPLETO"]
    MODO_ULTIMA_INTERNO = config_predial_actual["MODO_GENERACION_ULTIMA"]
    MODO_RESTO_INTERNO = config_predial_actual["MODO_GENERACION_RESTO"]
    MODO_ESPECIFICOS_INTERNO = config_predial_actual["MODO_GENERACION_ESPECIFICOS"]


    locale_ok = configurar_locale() # Usa la de GeneradorPredial_logica.py
    if not locale_ok:
        print("Error CRITICO (Predial Logic): No se pudo configurar el locale español. Abortando lógica.")
        return df_bd_maestra_actualizada

    # Determinar el nombre de la columna de ID en df_datos_para_procesar
    # Esta es la columna que contiene el "EXPEDIENTE" en el archivo BASE_DE_DATOS.xlsx
    # Su nombre real viene de config_columnas.xlsx (y está en nombres_columnas_csv)
    col_expediente_datos = ""
    # Opción 1: El generador maestro podría pasar el nombre exacto de la columna ID de los datos
    # si se añade a mode_config, ej. mode_config["col_expediente_en_datos_csv"]
    # Opción 2: Buscar un nombre estándar como "EXPEDIENTE" entre nombres_columnas_csv
    # Opción 3: Asumir que es la primera columna de nombres_columnas_csv
    
    # Vamos con una combinación de Opción 2 y 3 como fallback:
    if "EXPEDIENTE" in nombres_columnas_csv: # Si "EXPEDIENTE" es uno de los nombres de config_columnas.xlsx
        # Necesitamos encontrar cuál de los nombres_columnas_csv corresponde a "EXPEDIENTE"
        # Si config_columnas.xlsx dice "ID_PREDIO" en lugar de "EXPEDIENTE", esto fallará.
        # Sería mejor si generador.py identificara y pasara el nombre exacto de la columna ID de los datos.
        # Por ahora, si "EXPEDIENTE" está en la lista de nombres de columna del df_datos_para_procesar, lo usamos.
        if "EXPEDIENTE" in df_datos_para_procesar.columns:
             col_expediente_datos = "EXPEDIENTE"
    
    if not col_expediente_datos: # Si no se encontró "EXPEDIENTE" directamente
        if nombres_columnas_csv: # Tomar el primer nombre de columna definido en config_columnas.xlsx
            col_expediente_datos = nombres_columnas_csv[0]
            print(f"    (*) Advertencia (Predial Logic): Usando la primera columna '{col_expediente_datos}' de config_columnas.xlsx como ID del expediente en los datos de entrada.")
        else: # No debería pasar si nombres_columnas_csv se cargó bien
            print("    (!) Error (Predial Logic): No se pudieron determinar los nombres de columna de los datos de entrada. Abortando lógica.")
            return df_bd_maestra_actualizada
            
    if col_expediente_datos not in df_datos_para_procesar.columns:
        print(f"    (!) Error (Predial Logic): La columna de expediente para datos '{col_expediente_datos}' no se encuentra en el DataFrame de entrada. Columnas disponibles: {df_datos_para_procesar.columns.tolist()}")
        return df_bd_maestra_actualizada


    # --- Filtrar df_datos_para_procesar (ya está hecho en generador.py, pero podemos re-validar o refinar) ---
    # La variable df_datos_para_procesar que llega aquí YA ESTÁ FILTRADA por generador.py
    # para el modo_generacion_solicitado y por los expedientes específicos si aplica.
    # El max_docs_a_generar también ya se aplicó en generador.py
    
    df_datos_listos_para_bucle = df_datos_para_procesar # Usar el DataFrame ya filtrado y limitado
    
    print(f"  (Predial Logic Core) Total de registros a procesar (ya filtrados por maestro): {len(df_datos_listos_para_bucle)}")
    if df_datos_listos_para_bucle.empty:
        print("  (Predial Logic Core) No hay registros para procesar después del filtrado del maestro.")
        return df_bd_maestra_actualizada

    generados_count_logica = 0
    df_bd_maestra_para_iteracion = df_bd_maestra_actualizada.copy() # Trabajar sobre una copia

    # Rutas de salida específicas del modo
    carpeta_colonias_out_modo = carpeta_principal_salida / "COLONIAS"
    carpeta_vacias_out_modo = carpeta_principal_salida / "VACIAS"
    # El generador_maestro ya debería haber creado estas carpetas base.

    # Nombres de columnas de datos (de config_columnas.xlsx)
    col_nombre_contrib_datos = config_predial_actual.get("col_nombre_csv_original", "NOMBRE")
    col_direccion_contrib_datos = "DIRECCION" # Asumir que así se llama en BASE_DE_DATOS.xlsx / config_columnas.xlsx
    col_colonia_datos = config_predial_actual.get("col_colonia_csv_original", "COLONIA")
    periodo_col_datos = "PERIODO" # Asumir

    generados_count_logica = 0
    df_bd_maestra_para_iteracion = df_bd_maestra_actualizada.copy()

    letra_actual_para_ciclo = letra_lote
    contador_actual_para_ciclo = contador_inicial_lote

    for index, row in df_datos_listos_para_bucle.iterrows():
        if max_docs_a_generar > 0 and generados_count_logica >= max_docs_a_generar:
            print(f"  (Predial Logic Core) Límite de {max_docs_a_generar} documentos alcanzado. Deteniendo procesamiento de este lote.")
            break # Salir del bucle for
        expediente_actual = str(row.get(col_expediente_datos, "")).strip()
        nombre_contribuyente = str(row.get(col_nombre_contrib_datos, "SIN_NOMBRE")).strip()
        # -- INICIO BLOQUE PARA OMITIR POR CLAVE DE EXTENSION --
        clave_extension = str(row.get("CLAVE DE EXTENSION", "0")).strip()
        if clave_extension != "0":
            print(f"  - OMITIENDO: Expediente {expediente_actual} tiene CLAVE DE EXTENSION {clave_extension}. No se generará.")
            # Opcional: Podrías actualizar la BD aquí con un estado "Omitido por Clave" si lo necesitas.
            continue # Salta al siguiente registro del bucle
        # -- FIN BLOQUE PARA OMITIR --
        direccion_contribuyente = str(row.get(col_direccion_contrib_datos, "DIRECCION_DESCONOCIDA")).strip()
        colonia_actual_datos = str(row.get(col_colonia_datos, "")).strip()

        if not expediente_actual:
            print(f"  (!) Advertencia (Predial Logic): Registro {index+1} sin ID de expediente válido ('{col_expediente_datos}'). Saltando.")
            continue
        
        nombre_base_limpio = limpiar_texto(f"{expediente_actual}_{nombre_contribuyente}")
        if not nombre_base_limpio or len(nombre_base_limpio) < len(expediente_actual):
            nombre_base_limpio = limpiar_texto(expediente_actual) if limpiar_texto(expediente_actual) else f"exp_{expediente_actual}"

        if not colonia_actual_datos or colonia_actual_datos == "0" or colonia_actual_datos.upper() == "VACIA":
            ruta_salida_exp_actual_modo = carpeta_vacias_out_modo
            colonia_para_bd_actualizada = "VACIAS"
        else:
            colonia_nombre_limpio_modo = limpiar_texto(colonia_actual_datos)
            ruta_salida_exp_actual_modo = carpeta_colonias_out_modo / colonia_nombre_limpio_modo
            ruta_salida_exp_actual_modo.mkdir(parents=True, exist_ok=True) # La lógica de Predial crea la subcarpeta de la colonia.
            colonia_para_bd_actualizada = colonia_actual_datos
        
        # REEMPLAZO PARA LAS RUTAS
        ruta_final_pdf = ruta_salida_exp_actual_modo / f"{nombre_base_limpio}.pdf"

        # Archivos intermedios que se crearán y borrarán en cada iteración
        ruta_temp_docx_principal = ruta_salida_exp_actual_modo / f"~{nombre_base_limpio}_main.docx"
        ruta_temp_pdf_principal = ruta_salida_exp_actual_modo / f"~{nombre_base_limpio}_main_TEMP.pdf"
        ruta_pdf_final_unido_temp = ruta_salida_exp_actual_modo / f"~{nombre_base_limpio}_UNIDO_TEMP.pdf"

        pdf_generado_final_ok = False
        estado_final_bd_exp = ESTADO_ERROR_GENERACION_MODO
        paginas_doc_completo_contadas = pd.NA 
        monto_final_para_bd = ""
        bimestre_final_para_bd = ""
        anos_final_para_bd = ""

        try:
            anos_activos = []
            anos_inactivos = []
            possible_year_cols = {2022: "2022", 2023: "2023", 2024: "2024"}
            impuesto_year_cols = {year: f"IMPUESTO PREDIAL DEL AÑO {year}" for year in [2022, 2023, 2024]}

            for year_val in [2022, 2023, 2024]:
                activo = False; valor_check = None
                year_col_data_name = str(year_val) 
                impuesto_year_col_data_name = f"IMPUESTO PREDIAL DEL AÑO {year_val}"

                if year_col_data_name in row.index:
                    valor_check = str(row.get(year_col_data_name, '0')).strip()
                elif impuesto_year_col_data_name in row.index:
                    valor_check = str(row.get(impuesto_year_col_data_name, '0')).strip()
                
                if valor_check and valor_check != '0' and valor_check.lower() not in ['nan', 'na', '', '-']:
                    if safe_float(valor_check, 0.0) != 0.0: activo = True
                if activo: anos_activos.append(year_val)
                else: anos_inactivos.append(year_val)
            
            if anos_activos: anos_final_para_bd = f"{min(anos_activos)}-{max(anos_activos)}"
            else: anos_final_para_bd = ""

            now = datetime.now(); fecha_hoy_str_calc = "(Error Locale)"; fecha_texto_str_calc = "(Error Locale/num2words)"
            if locale_ok:
                 try: fecha_hoy_str_calc = now.strftime("%d de %B de %Y").lower()
                 except Exception: pass
                 if NUM2WORDS_INSTALLED: # NUM2WORDS_INSTALLED es de GeneradorPredial_logica.py
                     try:
                         dia_num = now.day; ano_num_calc = now.year
                         dia_palabra = num2words(dia_num, lang='es'); ano_palabra_calc = num2words(ano_num_calc, lang='es')
                         medio_fecha = now.strftime("de %B de %Y").lower()
                         fecha_texto_str_calc = f"{dia_num}-{dia_palabra} {medio_fecha}-{ano_palabra_calc}"
                     except Exception: pass
                 else: fecha_texto_str_calc = "(num2words no instalado)"
            ano_placeholder_calc = str(min(anos_activos)) if anos_activos else str(now.year)

            
            doc = Document(plantilla_path)

            periodo_str_calc = str(row.get(periodo_col_datos, "")).strip()
            periodo_año_calc = 0; periodo_bim_calc = 0
            if re.match(r"^\d{6}$", periodo_str_calc):
                 try: 
                     periodo_año_calc = int(periodo_str_calc[:4])
                     periodo_bim_calc = int(periodo_str_calc[4:6])
                     bimestre_final_para_bd = f"{periodo_bim_calc}-{periodo_año_calc}"
                 except: bimestre_final_para_bd = ""
            else: bimestre_final_para_bd = ""
            
            initial_replacements_calc = {} # Asegúrate que se inicializa antes del bucle
            placeholders_sumas_total_calc = [
                "[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", "[SUMA DE MONTOS ACTUALIZADOS]", 
                "[SUMA DE LOS MONTOS DE RECARGO TOTAL]", "[SUMA DE LA SANCION POR BIMESTRE]", "[TOTAL]",
                "[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]", 
                "[SUMA DE MONTOS ACTUALIZADOS TEXTO]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]", 
                "[SUMA DE LA SANCION POR BIMESTRE TEXTO]", "[TOTAL PESOS CORREGIDOS]", 
                "[LEYENDA_FINAL_PAG17]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]", 
                "[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"
            ]
            initial_replacements_calc["[FECHA]"] = fecha_hoy_str_calc
            initial_replacements_calc["[FECHA_TEXTO]"] = fecha_texto_str_calc
            initial_replacements_calc["[AÑO]"] = ano_placeholder_calc

            # BUCLE ORIGINAL DONDE SE LLENA initial_replacements_calc DESDE LA FILA (row)
            for col_data_name in nombres_columnas_csv: # Iterar sobre nombres de config_columnas.xlsx
                placeholder = f"[{col_data_name}]"
                if placeholder in placeholders_sumas_total_calc or placeholder in initial_replacements_calc or col_data_name.startswith("IMPUESTO PREDIAL DEL AÑO"):
                    continue
                
                valor_final_calc = ""
                col_year_match = re.search(r"(?:^|\b|_|\s)(\d{4})$", col_data_name.strip())
                col_year_parsed = None
                if col_year_match:
                    year_str_parsed = col_year_match.group(1)
                    if year_str_parsed in ["2022", "2023", "2024"]:
                        col_year_parsed = int(year_str_parsed)
                
                if col_year_parsed is not None and col_year_parsed in anos_inactivos:
                    valor_final_calc = ""
                elif col_data_name in row.index: # row usa los nombres de nombres_columnas_csv
                    valor_original_row = row[col_data_name]
                    # formatear_valor_v6 usa tipos_columnas_csv que tiene los tipos de config_columnas.xlsx
                    valor_formateado_calc = formatear_valor_v6(valor_original_row, col_data_name, tipos_columnas_csv)
                    
                    patrones_bimestre_calc = [r"^BIMESTRE\s+(\d+)\s+(\d{4})$", r"^MONTO ACTUALIZADO PREDIAL POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^RECARGOS POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^SANCION POR BIMESTRE\s+(\d+)\s+(\d{4})$"]
                    for patron_bim_calc in patrones_bimestre_calc:
                        match_bim_calc = re.match(patron_bim_calc, col_data_name.upper())
                        if match_bim_calc:
                            bim_num_calc = int(match_bim_calc.group(1))
                            bim_year_calc = int(match_bim_calc.group(2))
                            if (bim_year_calc < periodo_año_calc) or (bim_year_calc == periodo_año_calc and bim_num_calc < periodo_bim_calc):
                                valor_formateado_calc = formatear_valor_v6("0", col_data_name, tipos_columnas_csv)
                            else:
                                valor_limpio_calc = str(valor_original_row).replace(",", "").replace("$", "").strip()
                                try: 
                                    numero_bim_calc = float(valor_limpio_calc)
                                except: 
                                    numero_bim_calc = 0
                                if numero_bim_calc > 0: 
                                    valor_formateado_calc = valor_formateado_calc
                                else: 
                                    valor_formateado_calc = " "
                            break
                    valor_final_calc = valor_formateado_calc
                    if col_data_name == col_expediente_datos: # Si la columna actual es la del expediente
                        valor_str_exp_calc = str(valor_final_calc).strip()
                        valor_final_calc = '0' + valor_str_exp_calc if len(valor_str_exp_calc) == 7 else valor_str_exp_calc
                initial_replacements_calc[placeholder] = valor_final_calc
            
            # --- PUNTO DE CORRECCIÓN IMPORTANTE ---
            # MUEVE EL BLOQUE DE CÓDIGO PARA FORZAR SANCIONES AQUÍ:
            # ANTES de reemplazar_en_documento_v2(doc, initial_replacements_calc, fase="inicial")

            print("      (Predial Logic Core) FORZANDO placeholders SANCION BIMESTRE con valores de MONTO ACTUALIZADO...")
            for year_s_override in [2022, 2023, 2024]: # Itera sobre los años relevantes
                for bim_s_override in range(1, 7): # Itera sobre los bimestres (1 al 6)
                    
                    sancion_ph = f"[SANCION POR BIMESTRE {bim_s_override} {year_s_override}]"
                    monto_act_ph = f"[MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim_s_override} {year_s_override}]"

                    valor_monto_actualizado_formateado = initial_replacements_calc.get(monto_act_ph)

                    if valor_monto_actualizado_formateado is not None:
                        initial_replacements_calc[sancion_ph] = valor_monto_actualizado_formateado
                        # print(f"DEBUG: Forzando {sancion_ph} = {valor_monto_actualizado_formateado}") # Descomenta para depurar si es necesario
                    else:
                        valor_cero_formateado_para_sancion = " " 
                        if tipos_columnas_csv:
                             valor_cero_formateado_para_sancion = formatear_valor_v6("0", sancion_ph, tipos_columnas_csv)
                        initial_replacements_calc[sancion_ph] = valor_cero_formateado_para_sancion
                        # print(f"DEBUG: {monto_act_ph} no encontrado o None. Forzando {sancion_ph} = {valor_cero_formateado_para_sancion}") # Descomenta para depurar

            # Ahora sí, se llama a la función de reemplazo inicial con initial_replacements_calc ya modificado
            reemplazar_en_documento_v2(doc, initial_replacements_calc, fase="inicial") # Usa la de Predial_logica
            
            tablas_protegidas_calc = []
            eliminar_elementos_inactivos_v_usuario(doc, anos_inactivos, periodo_año_calc, periodo_bim_calc, tablas_protegidas_calc) # Usa la de Predial_logica
            procesar_tablas_suelo_construccion(doc, anos_activos) # Usa la de Predial_logica

            # --- Cálculos de sumas (estos ya estaban bien por tus cambios anteriores) ---
            impuesto_anual_calc = {2022: 0.0, 2023: 0.0, 2024: 0.0}
            sum_impuesto_anual_total_calc = 0.0
            sum_monto_actualizado_total_calc = 0.0 # Suma de los MONTOS ACTUALIZADOS del Excel
            sum_recargos_total_calc = 0.0
            sum_sancion_total_calc = 0.0 # Esta es la suma de las SANCIONES ORIGINALES del Excel
                                         # pero no la usaremos para el total general ni para la suma de sanción visible

            for year_s in [2022, 2023, 2024]:
                for bim_s in range(1, 7):
                    if (year_s < periodo_año_calc) or (year_s == periodo_año_calc and bim_s < periodo_bim_calc): continue
                    campo_bim_data = f"BIMESTRE {bim_s} {year_s}"
                    campo_act_data = f"MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim_s} {year_s}"
                    campo_rec_data = f"RECARGOS POR BIMESTRE {bim_s} {year_s}"
                    # campo_san_data = f"SANCION POR BIMESTRE {bim_s} {year_s}" # Ya no es crítico para el cálculo final de sanción visible
                    
                    valor_bim_actual_calc = safe_float(row.get(campo_bim_data, 0.0))
                    impuesto_anual_calc[year_s] += valor_bim_actual_calc
                    sum_impuesto_anual_total_calc += valor_bim_actual_calc
                    sum_monto_actualizado_total_calc += safe_float(row.get(campo_act_data, 0.0))
                    sum_recargos_total_calc += safe_float(row.get(campo_rec_data, 0.0))
                    # sum_sancion_total_calc += safe_float(row.get(campo_san_data, 0.0)) # Se puede omitir si no se usa

            monto_actualizado_restados_calc = sum_monto_actualizado_total_calc - sum_impuesto_anual_total_calc
            
            # Total general recalculado usando sum_monto_actualizado_total_calc para la parte de "sanciones"
            # Este cambio ya lo tenías correcto:
            total_general_recalculado_calc = sum_impuesto_anual_total_calc + monto_actualizado_restados_calc + sum_recargos_total_calc + sum_monto_actualizado_total_calc
            total_redondeado_bd_calc = round(total_general_recalculado_calc)
            monto_final_para_bd = f"$ {locale.format_string('%.2f', total_redondeado_bd_calc, grouping=True)}"

            # --- Populado de final_replacements_calc (tus cambios aquí ya eran correctos) ---
            final_replacements_calc = {}
            bimestres_ordinales_calc = {1: "PRIMER", 2: "SEGUNDO", 3: "TERCERO", 4: "CUARTO", 5: "QUINTO", 6: "SEXTO"}
            ordinal_texto_calc = bimestres_ordinales_calc.get(periodo_bim_calc, "")
            final_replacements_calc["[BIMESTRE_ORDINAL]"] = ordinal_texto_calc
            final_replacements_calc["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]"] = formatear_valor_v6(sum_impuesto_anual_total_calc, "[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", tipos_columnas_csv)
            final_replacements_calc["[SUMA DE MONTOS ACTUALIZADOS]"] = formatear_valor_v6(sum_monto_actualizado_total_calc, "[SUMA DE MONTOS ACTUALIZADOS]", tipos_columnas_csv)
            final_replacements_calc["[SUMA DE LOS MONTOS DE RECARGO TOTAL]"] = formatear_valor_v6(sum_recargos_total_calc, "[SUMA DE LOS MONTOS DE RECARGO TOTAL]", tipos_columnas_csv)
            
            # SUMA DE SANCION usa sum_monto_actualizado_total_calc (este cambio ya lo tenías):
            final_replacements_calc["[SUMA DE LA SANCION POR BIMESTRE]"] = formatear_valor_v6(sum_monto_actualizado_total_calc, "[SUMA DE LA SANCION POR BIMESTRE]", tipos_columnas_csv)
            final_replacements_calc["[SUMA DE MONTOS ACTUALIZADOS RESTADOS]"] = formatear_valor_v6(monto_actualizado_restados_calc, "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]", tipos_columnas_csv)
            final_replacements_calc["[TOTAL]"] = formatear_valor_v6(total_redondeado_bd_calc, "[TOTAL]", tipos_columnas_csv)

            for año_ip_calc in [2022, 2023, 2024]:
                valor_ip_anual_calc = round(impuesto_anual_calc[año_ip_calc], 2) 
                placeholder_ip_calc = f"[IMPUESTO PREDIAL DEL AÑO {año_ip_calc}]"
                if valor_ip_anual_calc > 0:
                    final_replacements_calc[placeholder_ip_calc] = formatear_valor_v6(valor_ip_anual_calc, placeholder_ip_calc, tipos_columnas_csv)
                else:
                    final_replacements_calc[placeholder_ip_calc] = formatear_valor_v6(0.0, placeholder_ip_calc, tipos_columnas_csv) 

            final_replacements_calc["[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]"] = formatear_texto_moneda(sum_impuesto_anual_total_calc)
            final_replacements_calc["[SUMA DE MONTOS ACTUALIZADOS TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total_calc)
            final_replacements_calc["[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]"] = formatear_texto_moneda(sum_recargos_total_calc)
            # SUMA DE SANCION TEXTO usa sum_monto_actualizado_total_calc (este cambio ya lo tenías):
            final_replacements_calc["[SUMA DE LA SANCION POR BIMESTRE TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total_calc)
            final_replacements_calc["[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"] = formatear_texto_moneda(monto_actualizado_restados_calc)
            final_replacements_calc["[TOTAL PESOS CORREGIDOS]"] = formatear_texto_moneda(total_redondeado_bd_calc)
            
            es_pm_calc = expediente_actual in pm_set_actual 
            final_replacements_calc["[LEYENDA_FINAL_PAG17]"] = "" if es_pm_calc else LEYENDA_PAG17_TEXTO

            bm_por_anio = {}
            for anio in [2022, 2023, 2024]:
                if anio == periodo_año_calc:
                    bm_por_anio[anio] = periodo_bim_calc
                elif anio > periodo_año_calc:
                    bm_por_anio[anio] = 1
            for anio_bm, bim_val in bm_por_anio.items():
                final_replacements_calc[f"[BM{anio_bm}]"] = f"{bim_val}°"

            # EL BLOQUE DE FORZAR SANCIONES YA NO VA AQUÍ, SE MOVIÓ ARRIBA
            # ANTES DE LA PRIMERA LLAMADA A reemplazar_en_documento_v2

            reemplazar_en_documento_v2(doc, final_replacements_calc, fase="final")

            # --- INICIO DE LA NUEVA LÓGICA PARA [PARRAFO_MILLAR] Y [MILLAR_TEXTO] ---
            print("      Determinando párrafo y texto de millar...")
            
            # Obtener el nombre real de la columna 'MILLAR' desde la configuración
            millar_col_name = "MILLAR" # Nombre que pusimos en config_columnas.xlsx
            parrafo_final_millar = "Párrafo por defecto o error: El valor de millar no fue reconocido."
            millar_texto_final = "X-X" # Valor por defecto

            if millar_col_name in row.index:
                # Usar safe_float para convertir el valor de forma segura
                millar_valor = safe_float(row.get(millar_col_name, 0.0), 0.0)

                if math.isclose(millar_valor, 0.002):
                    millar_texto_final = "2-DOS"
                    parrafo_final_millar = (
                        "De conformidad con el artículo 21 bis-8 primer párrafo de la Ley de Hacienda para los Municipios del Estado de Nuevo León, "
                        "el valor catastral del “INMUEBLE” de cada año descrito, se le aplica la tasa del 2-dos al millar, es decir, "
                        "se multiplica dicho valor catastral por .002, por lo que, tratándose de predios con uso de casa habitación, como en el presente caso, "
                        "conforme a lo establecido en el segundo párrafo del artículo ya mencionado, se multiplica el valor catastral mencionado por un factor de .002, "
                        "de lo que resulta el Impuesto Predial adeudado, el cual queda de la siguiente manera:"
                    )
                elif math.isclose(millar_valor, 0.003):
                    millar_texto_final = "3-TRES"
                    parrafo_final_millar = (
                        "De conformidad con el artículo 21 bis-8 primer párrafo de la Ley de Hacienda para los Municipios del Estado de Nuevo León, "
                        "el valor catastral del “INMUEBLE” de cada año descrito, se le aplica la tasa del 2-dos al millar, es decir, "
                        "se multiplica dicho valor catastral por .002 y, en el caso de predios con cualquier uso distinto al de casa habitación, "
                        "como en el presente caso, dado que se trata de un predio con una edificación comercial, se pagará el impuesto predial adicionando 1-uno al millar a la tasa mencionada, "
                        "conforme a lo establecido en el tercer párrafo del artículo ya mencionado, por lo que una vez adicionada la tasa de 2-dos al millar más 1-uno al millar, "
                        "da como resultado por un factor de .003, de lo que resulta en el Impuesto Predial adeudado, el cual queda de la siguiente manera:"
                    )
                elif math.isclose(millar_valor, 0.004):
                    millar_texto_final = "4-CUATRO"
                    parrafo_final_millar = (
                        "De conformidad con el artículo 21 bis-8 primer párrafo de la Ley de Hacienda para los Municipios del Estado de Nuevo León, "
                        "el valor catastral del “INMUEBLE” de cada año descrito, se le aplica la tasa del 2-dos al millar, es decir, "
                        "se multiplica dicho valor catastral por .002 y en el caso de predios con cualquier uso distinto al de casa habitación, "
                        "como en el presente caso, dado que se trata de un predio baldío, se pagará el impuesto predial adicionando 2-dos al millar a la tasa mencionada, "
                        "conforme a lo establecido en el segundo párrafo del artículo ya mencionado, por lo que una vez adicionada la tasa de 2-dos al millar más 2-dos al millar, "
                        "da como resultado una tasa del 4-cautro al millar, es decir, se multiplica el valor catastral mencionado por un factor de .004, "
                        "de lo que resulta el Impuesto Predial adeudado, el cual queda de la siguiente manera:"
                    )
                else:
                    # Este es el párrafo original que me diste, lo usaré como fallback
                    millar_texto_final = f"({millar_valor:.3f})-ERROR"
                    parrafo_final_millar = (
                         "De conformidad con el artículo 21 bis-8 primer párrafo de la Ley de Hacienda para los Municipios del Estado de Nuevo León, el valor catastral del “INMUEBLE” de cada año descrito, "
                         "se le aplica la tasa del 2-dos al millar, es decir, se multiplica dicho valor catastral por .002  y, en el caso de predios con cualquier uso distinto al de casa habitacion, "
                         "como en el presente caso, dado que se trata de un predio con una edificacion comercial, se pagara el impuesto predial adicionando 1-uno al millar a la tasa mencionada, "
                         "conforme a lo establecido en el tercer parrafo del articulo ya mencionado, por lo que una vez adicionada la tasa de 2-dos al millar mas 1- al millar, da como resultado una tasa del 3-tres al millar, "
                         "es decir se multiplica el valor catastral mencionado por un factor de .003, de lo que resulta en el Impuesto Predial Adeudado, el cual queda de la siguiente manera:"
                    )
                    print(f"      (*) Advertencia: Valor de Millar '{millar_valor}' no reconocido para Exp {expediente_actual}. Usando párrafo por defecto.")

            # Añadir el texto del millar al diccionario de reemplazos finales
            # El formato de este placeholder se tomará del que le diste en la plantilla.
            final_replacements_calc["[MILLAR_TEXTO]"] = millar_texto_final

            # --- FIN DE LA NUEVA LÓGICA ---


            # Llamada a la función de reemplazo de placeholders finales (esta línea ya existe)
            reemplazar_en_documento_v2(doc, final_replacements_calc, fase="final")
            
            # --- NUEVA LLAMADA PARA REEMPLAZAR EL PÁRRAFO COMPLETO ---
            # Esto se hace DESPUÉS de los reemplazos normales
            print("      Reemplazando párrafo de millar con formato...")
            reemplazar_parrafo_con_negritas(doc, "[PARRAFO_MILLAR]", parrafo_final_millar, "“INMUEBLE”")
            # --- FIN DE LA NUEVA LLAMADA ---

            if verificar_ausencia_info_construccion_v2(row, anos_activos): # Usa la de Predial_logica
                eliminar_palabra_especifica_del_documento(doc, "CONSTRUCCION") # Usa la de Predial_logica
            
            # Guardar el DOCX principal en su nueva ruta temporal
            doc.save(ruta_temp_docx_principal)

            # --- INICIO: NUEVA LÓGICA DE GENERACIÓN Y UNIÓN ---

            # 1. Convertir el documento principal a su PDF temporal
            if DOCX2PDF_INSTALLED:
                convert(str(ruta_temp_docx_principal), str(ruta_temp_pdf_principal))
                if not (ruta_temp_pdf_principal.exists() and ruta_temp_pdf_principal.stat().st_size > 0):
                    raise RuntimeError(f"Fallo al generar el PDF principal para el expediente {expediente_actual}")
            else:
                raise ImportError("docx2pdf no está instalado, no se puede continuar.")

            # 2. Generar el PDF del citatorio
            with tempfile.TemporaryDirectory() as temp_dir_citatorio:
                ruta_temp_pdf_citatorio = _generar_citatorio_pdf_interno(expediente_actual, config_predial_actual, Path(temp_dir_citatorio))
                if not ruta_temp_pdf_citatorio:
                    raise RuntimeError(f"Fallo al generar el PDF del citatorio para el expediente {expediente_actual}")

                # 3. Unir el PDF principal y el del citatorio
                print(f"    (Predial Logic Core) Uniendo PDF principal y citatorio...")
                reader_principal = PdfReader(str(ruta_temp_pdf_principal))
                reader_citatorio = PdfReader(str(ruta_temp_pdf_citatorio))
                writer_final = PdfWriter()
                
                paginas_principal = list(reader_principal.pages)
                total_paginas_principal = len(paginas_principal)
                
                # Calcular el punto de inserción (antes de la última "hoja")
                if total_paginas_principal <= 2:
                    punto_insercion = total_paginas_principal
                elif total_paginas_principal % 2 == 0:
                    punto_insercion = total_paginas_principal - 2
                else: # Impar
                    punto_insercion = total_paginas_principal - 1
                
                for i in range(punto_insercion):
                    writer_final.add_page(paginas_principal[i])
                
                writer_final.add_page(reader_citatorio.pages[0])

                for i in range(punto_insercion, total_paginas_principal):
                    writer_final.add_page(paginas_principal[i])

                with open(ruta_pdf_final_unido_temp, 'wb') as f_out:
                    writer_final.write(f_out)
            
            paginas_doc_completo_contadas = len(writer_final.pages)
            print(f"    (Predial Logic Core) PDF final unido creado con {paginas_doc_completo_contadas} páginas.")

            # 4. Aplicar LÓGICA DE EXTRACCIÓN AL PDF UNIDO
            if modo_generacion_solicitado == MODO_COMPLETO_INTERNO or modo_generacion_solicitado == MODO_ESPECIFICOS_INTERNO:
                shutil.copy2(ruta_pdf_final_unido_temp, ruta_final_pdf)
                estado_final_bd_exp = ESTADO_GEN_COMPLETO_MODO
            elif modo_generacion_solicitado == MODO_ULTIMA_INTERNO:
                if extraer_paginas_pdf(ruta_pdf_final_unido_temp, ruta_final_pdf, "ULTIMA"):
                    estado_final_bd_exp = ESTADO_GEN_ULTIMA_MODO
            elif modo_generacion_solicitado == MODO_RESTO_INTERNO:
                if extraer_paginas_pdf(ruta_pdf_final_unido_temp, ruta_final_pdf, "RESTO"):
                    estado_final_bd_exp = config_predial_actual["ESTADO_GEN_RESTO"]
            
            # Verificación final
            if ruta_final_pdf.exists() and ruta_final_pdf.stat().st_size > 0:
                pdf_generado_final_ok = True
                print(f"    (Predial Logic Core) PDF final ({modo_generacion_solicitado}) creado exitosamente: {ruta_final_pdf.name}")


        except Exception as main_loop_error_predial_core:
            print(f"    (!) Error INESPERADO (Predial Logic Core) procesando expediente {expediente_actual}: {main_loop_error_predial_core}")
            import traceback
            traceback.print_exc()
            estado_final_bd_exp = ESTADO_ERROR_GENERACION_MODO
            paginas_doc_completo_contadas = pd.NA

        finally:
            # --- Limpieza de TODOS los archivos temporales ---
            if eliminar_docx_intermedio and 'ruta_temp_docx_principal' in locals() and ruta_temp_docx_principal.exists():
                try: ruta_temp_docx_principal.unlink()
                except Exception: pass
            if 'ruta_temp_pdf_principal' in locals() and ruta_temp_pdf_principal.exists():
                try: ruta_temp_pdf_principal.unlink()
                except Exception: pass
            if 'ruta_pdf_final_unido_temp' in locals() and ruta_pdf_final_unido_temp.exists():
                try: ruta_pdf_final_unido_temp.unlink()
                except Exception: pass

            tipo_de_suelo_valor = str(row.get("TIPO DE SUELO", "")).strip()

            # --- Lógica de Actualización de BD Maestra (Modificada) ---
            # 1. Preparar el diccionario de datos a actualizar SIN el ID.
            hojas_calculadas_predial = pd.NA
            if paginas_doc_completo_contadas is not None and pd.notna(paginas_doc_completo_contadas) and paginas_doc_completo_contadas > 0:
                hojas_calculadas_predial = math.ceil(paginas_doc_completo_contadas / 2)

            # Primero, lee el valor de la fila de datos
            tipo_de_suelo_valor = str(row.get("TIPO DE SUELO", "")).strip()

            datos_para_actualizar = {
                "EXPEDIENTE": expediente_actual,
                "NOMBRE CONTRIBUYENTE": nombre_contribuyente,
                "DIRECCION": direccion_contribuyente,
                "COLONIA": colonia_para_bd_actualizada,
                "BIMESTRE": bimestre_final_para_bd,
                "AÑOS": anos_final_para_bd,
                "TIPO": tipo_de_suelo_valor,
                "MONTO": total_redondeado_bd_calc if 'total_redondeado_bd_calc' in locals() and pd.notna(total_redondeado_bd_calc) else 0.0,
                "ESTADO": estado_final_bd_exp,
                "HOJAS POR DOCUMENTO": hojas_calculadas_predial, # <--- LÍNEA CORREGIDA
                "FECHA IMPRESION": "",
                "BASE DE DATOS ESCANEO": "",
                "MOVIMIENTO": "EN DESPACHO (GENERADO)"
            }
            
            # Añadir la ruta del PDF si se generó exitosamente
            if pdf_generado_final_ok:
                try:
                    ruta_docs_base = Path(config_predial_actual["output_docs_path"])
                    ruta_relativa_pdf = ruta_final_pdf.relative_to(ruta_docs_base).as_posix()
                    datos_para_actualizar["Ruta PDF Generado"] = ruta_relativa_pdf
                except (ValueError, NameError):
                    datos_para_actualizar["Ruta PDF Generado"] = ruta_final_pdf.name if 'ruta_final_pdf' in locals() else ""
            else:
                datos_para_actualizar["Ruta PDF Generado"] = ""

            # 2. Verificar si el expediente ya existe en la BD Maestra
            indices_existentes_bd = df_bd_maestra_para_iteracion.index[
                df_bd_maestra_para_iteracion[col_expediente_bd].astype(str) == str(expediente_actual)
            ].tolist()

            if indices_existentes_bd:
                # El registro YA EXISTE. Recuperar su ID y usarlo en el log.
                id_existente = df_bd_maestra_para_iteracion.loc[indices_existentes_bd[-1], 'ID']

                # --- AHORA IMPRIMIMOS EL LOG CON EL ID CORRECTO ---
                print(f"\n  --- (Predial Logic Core) ({generados_count_logica + 1}/{len(df_datos_listos_para_bucle)}) ID: {id_existente} | Exp: {expediente_actual} ---")
                print(f"    (Predial Logic Core) RESPETANDO ID de lote existente '{id_existente}'...")

                idx_actualizar_bd = indices_existentes_bd[-1]
                # Actualizar los campos del registro existente
                for col_nombre, valor_nuevo in datos_para_actualizar.items():
                    if col_nombre in df_bd_maestra_para_iteracion.columns:
                        df_bd_maestra_para_iteracion.loc[idx_actualizar_bd, col_nombre] = valor_nuevo
            else:
                # El registro es NUEVO. Generar, asignar e imprimir el nuevo ID

                # --- INICIO LÓGICA DE ASIGNACIÓN DE ID ---
                contador_actual_para_ciclo += 1
                if contador_actual_para_ciclo > 100:
                    contador_actual_para_ciclo = 1
                    letra_actual_para_ciclo = obtener_siguiente_letra_lote(letra_actual_para_ciclo)
                
                nuevo_id_generado = f"{letra_actual_para_ciclo}-{contador_actual_para_ciclo:03d}"
                datos_para_actualizar['ID'] = nuevo_id_generado
                # --- FIN LÓGICA DE ASIGNACIÓN DE ID ---

                # --- AHORA IMPRIMIMOS EL LOG CON EL ID NUEVO ---
                print(f"\n  --- (Predial Logic Core) ({generados_count_logica + 1}/{len(df_datos_listos_para_bucle)}) NUEVO ID: {nuevo_id_generado} | Exp: {expediente_actual} ---")
                print(f"    (Predial Logic Core) Asignando nuevo ID de lote '{nuevo_id_generado}'...")


                # Asegurarse de que todas las columnas de la BD Maestra estén presentes antes de añadir
                for col_maestra in config_predial_actual["db_master_columns"]:
                    if col_maestra not in datos_para_actualizar:
                        tipo_col = config_predial_actual["db_master_types"].get(col_maestra)
                        datos_para_actualizar[col_maestra] = pd.NA if tipo_col == 'Int64' else ""

                df_nueva_fila = pd.DataFrame([datos_para_actualizar], columns=config_predial_actual["db_master_columns"])
                df_bd_maestra_para_iteracion = pd.concat([df_bd_maestra_para_iteracion, df_nueva_fila], ignore_index=True)
# ...

            # Re-asegurar tipos de datos al final de cada iteración
            try:
                tipos_a_aplicar = {k: v for k, v in config_predial_actual["db_master_types"].items() if k in df_bd_maestra_para_iteracion.columns}
                df_bd_maestra_para_iteracion = df_bd_maestra_para_iteracion.astype(tipos_a_aplicar)
            except Exception as e_astype:
                 print(f"    (*) Advertencia: Fallo menor en re-astype para Exp {expediente_actual}. Error: {e_astype}")

            if pdf_generado_final_ok:
                generados_count_logica += 1
                print(f"    (Predial Logic Core) PDF generado para Exp {expediente_actual} es OK. Procediendo a registrar ruta y subir...")

                # 1. Calcular ruta relativa del PDF para la BD
                try:
                    ruta_docs_base = Path(config_predial_actual["output_docs_path"])
                    ruta_relativa_pdf = ruta_final_pdf.relative_to(ruta_docs_base).as_posix()
                    datos_para_actualizar["Ruta PDF Generado"] = ruta_relativa_pdf # <-- CORREGIDO
                    print(f"       - Ruta relativa para BD: {ruta_relativa_pdf}")
                except (ValueError, NameError):
                    print(f"       - (!) Advertencia: No se pudo calcular la ruta relativa para {ruta_final_pdf}. Se guardará solo el nombre.")
                    datos_para_actualizar["Ruta PDF Generado"] = ruta_final_pdf.name if 'ruta_final_pdf' in locals() else "" # <-- CORREGIDO

                # 2. Subir el archivo PDF generado al servidor
                print(f"      - Intentando subir PDF '{ruta_final_pdf.name}' al servidor...")
                # La función subir_archivo_al_servidor ahora está en este mismo archivo
                subida_exitosa = funcion_de_subida(ruta_final_pdf, config_predial_actual)
                if not subida_exitosa:
                    print(f"      - (!) ERROR CRÍTICO: La subida del PDF para el expediente {expediente_actual} falló.")
                    # Opcional: podrías cambiar el estado a un error específico de subida
                    # estado_final_bd_exp = "ERROR_SUBIDA_PDF" 
                    # datos_reg_bd_actualizar["ESTADO"] = estado_final_bd_exp

            else:
                # Si la generación del PDF falló, asegúrate de que la ruta quede vacía en la BD.
                datos_para_actualizar["Ruta PDF Generado"] = ""
    
    print(f"\n  --- (Predial Logic Core) Proceso Finalizado ---")
    print(f"  Documentos PDF (según modo) generados/intentados en esta ejecución lógica: {generados_count_logica}")
    
    return df_bd_maestra_para_iteracion

    
def formatear_texto_moneda(valor):
    if not NUM2WORDS_INSTALLED: return "(num2words no instalada)"
    if pd.isna(valor): return "(Monto no especificado)"
    try:
        numero = safe_float(valor, None)
        if numero is None: return "(Monto inválido)"
        if abs(numero) > 999_999_999_999: return "(Monto demasiado grande para convertir a texto)"
        numero_redondeado = round(numero, 2)
        parte_entera = int(numero_redondeado)
        parte_decimal = abs(int(round((numero_redondeado - parte_entera) * 100)))
        texto_entero = num2words(parte_entera, lang='es')
        if texto_entero: texto_entero_capitalizado = texto_entero[0].upper() + texto_entero[1:]
        else: texto_entero_capitalizado = "Cero"
        signo = "Menos " if numero < 0 else ""
        return f"({signo}{texto_entero_capitalizado} pesos {parte_decimal:02d}/100 M.N.)"
    except Exception as e:
        print(f"Error en formatear_texto_moneda para valor '{valor}': {e}")
        return "(Error en conversión a texto)"

def replace_text_in_paragraph(paragraph, replacements):
    if paragraph is None or not hasattr(paragraph, 'runs') or not hasattr(paragraph, '_element') or paragraph._element is None: return
    try:
        parent_tag = paragraph._element.getparent().tag
        is_header_footer = parent_tag.lower().endswith(('hdr', 'ftr'))
    except Exception: 
        is_header_footer = False
        parent_tag = ""
    if is_header_footer:
        if "[EXPEDIENTE]" not in paragraph.text: 
            return
        replacements = {k: v for k, v in replacements.items() if k == "[EXPEDIENTE]"}
        if not replacements: 
            return
    try: 
        full_text = paragraph.text
    except Exception: 
        return
    if '[' not in full_text: 
        return

    processed_text = full_text
    found_placeholders_in_para = False
    sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
    for key in sorted_keys:
        if key in processed_text:
            value = replacements[key]
            str_value = str(value) if value is not None else ""
            if key not in str_value:
                processed_text = processed_text.replace(key, str_value)
                found_placeholders_in_para = True
    if any(phrase in processed_text for phrase in ["DOMICILIO FISCAL:", "UBICACIÓN DEL PREDIO:"]):
        parts = [part.strip() for part in processed_text.split(',')]
        non_empty_parts = [part for part in parts if part]
        cleaned_text = ', '.join(non_empty_parts)
        cleaned_text = re.sub(r',\s*CP\s+(\d+)', r' CP \1', cleaned_text)
        processed_text = cleaned_text
    if not found_placeholders_in_para or full_text == processed_text: 
        return

    try:
        first_run = paragraph.runs[0] if paragraph.runs else None
        p_alignment = paragraph.alignment
        p_style = paragraph.style
        for run in list(paragraph.runs):
            parent_el = run._element.getparent()
            if parent_el is not None:
                try: 
                    parent_el.remove(run._element)
                except ValueError: 
                    pass
        if processed_text:
            new_run = paragraph.add_run(processed_text)
            if first_run and hasattr(new_run, 'font'):
                try:
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                    new_run.font.bold = first_run.font.bold
                    new_run.font.italic = first_run.font.italic
                    new_run.font.underline = first_run.font.underline
                    if first_run.font.color and first_run.font.color.rgb:
                        new_run.font.color.rgb = first_run.font.color.rgb
                except Exception:
                    pass
        paragraph.alignment = p_alignment
        if p_style is not None and hasattr(p_style, 'name') and p_style.name != 'Normal':
            try: 
                paragraph.style = p_style
            except Exception: 
                pass
    except Exception as e: 
        print(f"Error reconstruyendo párrafo: {e}")

def reemplazar_en_documento_v2(document, replacements, fase=""):
    print(f"      Aplicando {len(replacements)} reemplazos {fase}...")
    for paragraph in list(document.paragraphs): 
        replace_text_in_paragraph(paragraph, replacements)
    for table in list(document.tables):
        if table._element is None or table._element.getparent() is None: 
            continue
        for row in list(table.rows):
            if row._element is None or row._element.getparent() is None: 
                continue
            try: 
                cells_in_row = list(row.cells)
            except Exception as e:
                print(f"Warn ({fase}): No get celdas tabla cuerpo: {e}")
                continue
            for cell_idx, cell in enumerate(cells_in_row):
                try:
                    if hasattr(cell, '_element') and cell._element is not None and cell._element.getparent() is not None and hasattr(cell, 'paragraphs'):
                        for paragraph in list(cell.paragraphs):
                            if paragraph._element is not None and paragraph._element.getparent() is not None: 
                                replace_text_in_paragraph(paragraph, replacements)
                except IndexError:
                    print(f"Warn ({fase}): Índice celda {cell_idx} inválido tabla cuerpo.")
                    continue
                except Exception as cell_err:
                    print(f"Warn ({fase}): Error celda {cell_idx} tabla cuerpo: {cell_err}")
                    continue
    for section in document.sections:
        for header_footer_part in [section.header, section.footer, section.first_page_header, section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if header_footer_part is not None:
                for paragraph in list(header_footer_part.paragraphs):
                    if paragraph._element is not None and paragraph._element.getparent() is not None: 
                        replace_text_in_paragraph(paragraph, replacements)
                if hasattr(header_footer_part, 'tables'):
                    for table in list(header_footer_part.tables):
                        if table._element is None or table._element.getparent() is None: 
                            continue
                        for row in list(table.rows):
                            if row._element is None or row._element.getparent() is None: 
                                continue
                            try: 
                                cells_in_row_hf = list(row.cells)
                            except Exception as e: 
                                print(f"Warn ({fase}): No get celdas tabla H/F: {e}")
                                continue
                            for cell_idx_hf, cell in enumerate(cells_in_row_hf):
                                try:
                                    if hasattr(cell, '_element') and cell._element is not None and cell._element.getparent() is not None and hasattr(cell, 'paragraphs'):
                                        for paragraph in list(cell.paragraphs):
                                            if paragraph._element is not None and paragraph._element.getparent() is not None: 
                                                replace_text_in_paragraph(paragraph, replacements)
                                except IndexError:
                                    print(f"Warn ({fase}): Índice celda H/F {cell_idx_hf} inválido.")
                                    continue
                                except Exception as cell_err_hf:
                                    print(f"Warn ({fase}): Error celda H/F {cell_idx_hf}: {cell_err_hf}")
                                    continue

def eliminar_filas_areas_vacias(table):
    try:
        if not table.rows: 
            return
        header_cells = table.rows[0].cells
        indices_superficie = [i for i, cell in enumerate(header_cells) if "SUPERFICIE" in cell.text.upper()]
        if not indices_superficie: 
            return

        filas_a_eliminar = []
        for r_idx, row in enumerate(table.rows[1:], 1):
            try:
                eliminar = True
                if len(row.cells) <= max(indices_superficie): 
                    eliminar = False
                else:
                    for idx in indices_superficie:
                        try: 
                            cell_value = row.cells[idx].text
                        except IndexError: 
                            eliminar = False
                            break
                        normalized = cell_value.replace("$", "").replace("[", "").replace("]", "").replace(",", "").strip()
                        if normalized in ["[ ] 0", "[ 10", "[] 0", "[]0", "[0]"]: 
                            normalized = "0"
                        if normalized != "":
                            try:
                                num = float(normalized)
                                if round(num, 5) not in (0.0, 1.0, 2.0): 
                                    eliminar = False
                                    break
                            except ValueError: 
                                eliminar = False
                                break
                if eliminar:
                    if row._element is not None: 
                        filas_a_eliminar.append(row._element)
            except Exception as ex_row:
                print("Warn: Error al procesar fila %d para SUPERFICIE: %s" % (r_idx, ex_row))

        eliminadas_count_local = 0
        for row_elem in filas_a_eliminar:
            try:
                parent = row_elem.getparent()
                if parent is not None: 
                    parent.remove(row_elem)
                    eliminadas_count_local += 1
            except Exception as ex_del:
                print("Warn: Error eliminando fila de area: %s" % ex_del)
    except IndexError:
        print("Error CRITICO: No se pudo acceder a la fila de encabezado en tabla de areas.")
    except Exception as ex_table:
        print("Error CRITICO procesando tabla de areas: %s" % ex_table)

def eliminar_elementos_inactivos_v_usuario(document, anos_inactivos, periodo_año, periodo_bimestre, tablas_protegidas):
    if not anos_inactivos:
        print("      No hay años inactivos para eliminar secciones.")
    else:
        print("      Eliminando elementos para anios inactivos (%s) y filas invalidas de area..." % anos_inactivos)

    elementos_cuerpo_a_eliminar = []
    filas_a_eliminar_por_tabla = {}
    try:
        parrafos_doc = list(document.paragraphs)
        tablas_doc = list(document.tables)
    except Exception as e_list_elements:
        print("--- Debug CRITICAL: Error al listar parrafos o tablas: %s ---" % e_list_elements)
        return

    if anos_inactivos:
        cuerpo_doc = document.element.body
        elementos_cuerpo = list(cuerpo_doc)
        i = 0
        while i < len(elementos_cuerpo):
            elem = elementos_cuerpo[i]
            if isinstance(elem, CT_P):
                para_obj = None
                for p_doc in parrafos_doc:
                    if hasattr(p_doc, '_element') and p_doc._element == elem:
                        para_obj = p_doc
                        break
                if para_obj:
                    try:
                        para_text_raw = para_obj.text
                        if any(ph in para_text_raw for ph in PROTECTED_PLACEHOLDERS):
                            i += 1
                            continue
                        para_text = para_text_raw.strip().upper()
                        match = re.match(r"AÑO\s+(\d{4})", para_text)
                        if match:
                            year = int(match.group(1))
                            if year in anos_inactivos:
                                for j in range(i, min(i + 5, len(elementos_cuerpo))):
                                    if elementos_cuerpo[j] not in elementos_cuerpo_a_eliminar:
                                        elementos_cuerpo_a_eliminar.append(elementos_cuerpo[j])
                                i += 5
                                continue
                    except Exception as e_p1:
                        print("Warn: Error procesando P ANIO en indice %d: %s" % (i, e_p1))
            i += 1

    tabla7_keywords = ["PERIODO ADEUDADO"]
    tabla9_keywords = ["BIMESTRE", "AÑO"]
    tabla8_keywords = ["FACTOR DE ACTUALIZACIÓN"]
    tabla10_keywords = ["RECARGOS POR BIMESTRE"]

    for table_idx, table in enumerate(tablas_doc):
        if table._element is None or table._element.getparent() is None:
            continue
        if table._element in elementos_cuerpo_a_eliminar:
            continue

        processed_as_area_table = False
        filas_a_eliminar_esta_tabla = []
        table_element = table._element

        try:
            if len(table.rows) > 0:
                header_cells = table.rows[0].cells
                try:
                    header_texts_raw = [cell.text for cell in header_cells]
                except Exception as e_print:
                    print("--- Debug Headers Tabla %d: Error al imprimir (%s)" % (table_idx, e_print))
                header_texts = [t.strip().upper() for t in header_texts_raw]
                texto_encabezado = " ".join(header_texts)

                es_tabla_totales = (
                    "IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in texto_encabezado
                    and "[TOTAL]" in texto_encabezado
                )

                if es_tabla_totales:
                    print("🛡️ Tabla de TOTALES detectada y protegida")
                    continue  # No tocar esta tabla

                # Marcar tabla de totales para que nunca se elimine
                # ————— Detecto tabla de totales sin exigir “[TOTAL]” literal —————
                if "IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in texto_encabezado \
                   and "TOTAL" in texto_encabezado:
                    tablas_protegidas.append(table._element)
                    print(f"🛡️ Protegiendo tabla de totales (idx {table_idx})")
                    continue

                for table_element, filas in filas_a_eliminar_por_tabla.items():
                    # Proteger la tabla de totales incluso si por error se marcaron sus filas
                    if any("IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in cell.text.upper() for row in table_element.xpath(".//w:tr") for cell in row.xpath(".//w:t")):
                        print("🚫 Cancelando eliminación de filas en tabla de TOTALES")
                        continue

                for table_element, filas in filas_a_eliminar_por_tabla.items():
                    if table_element in tablas_protegidas:
                        continue  # 🛑 No eliminar filas de tabla protegida
                    for fila in filas:
                        if fila.getparent() is not None:
                            fila.getparent().remove(fila)



                # 1) Saltar completamente la tabla de totales
                if any("IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in h for h in header_texts):
                    continue

                # 2) Procesar inmediatamente la tabla de ÍNDICE NACIONAL DE PRECIOS
                if header_texts and "BIMESTRE/AÑO" in header_texts[0] \
                and any("ÍNDICE NACIONAL DE PRECIOS" in ht for ht in header_texts):
                    for row in table.rows[1:]:
                        cell0 = row.cells[0].text.strip()
                        m = re.match(r"(\d+)°/(\d{4})", cell0)
                        if m:
                            bim, anio = int(m.group(1)), int(m.group(2))
                            if (anio < periodo_año) or (anio == periodo_año and bim < periodo_bimestre):
                                # Poner 0.00 en todas las demás celdas de esa fila
                                for cell in row.cells[1:]:
                                    for p in cell.paragraphs:
                                        p.text = "0.00"
                    # Ya no queremos tocar ni borrar filas de esta tabla
                    continue

            if any("SUPERFICIE" in txt for txt in header_texts):
                eliminar_filas_areas_vacias(table)
                processed_as_area_table = True

            if anos_inactivos:
                table_type = None
                if not processed_as_area_table and header_texts:
                    header_0 = header_texts[0]
                    if any(ph in header_0 for ph in PROTECTED_PLACEHOLDERS):
                        continue

                    if any(kw in header_0 for kw in tabla7_keywords):
                        table_type = 7
                    elif len(header_texts) > 1 and tabla9_keywords[0] in header_0 and tabla9_keywords[1] in header_texts[1]:
                        table_type = 9
                    elif any(kw in header_0 for kw in tabla8_keywords):
                        table_type = 8
                    elif any(kw in header_0 for kw in tabla10_keywords):
                        table_type = 10
                    elif any('SANCION' in ht for ht in header_texts):
                        table_type = 11
                    elif len(table.rows) > 1 and len(table.columns) > 0 and len(table.rows[1].cells) > 0:
                        if re.match(r"\d+°?/\d{4}", table.rows[1].cells[0].text.strip()):
                            table_type = 81011

                if table_type is not None:
                    for row_idx_rel in range(len(table.rows) - 1, 0, -1):
                        row_idx_abs = row_idx_rel
                        row_year = None
                        try:
                            row = table.rows[row_idx_rel]
                            if row._element is None or row._element.getparent() is None:
                                continue
                            cells = row.cells
                            cell0_text = cells[0].text.strip() if len(cells) > 0 else ""
                            cell1_text = cells[1].text.strip() if len(cells) > 1 else ""

                            if table_type == 7:
                                m = re.search(r"(\bdel\s+|\/)\s*(\d{4})\b", cell0_text, re.IGNORECASE)
                                row_year = int(m.group(2)) if m else None
                            elif table_type == 9:
                                try:
                                    bim_from_row = None
                                    row_year = None
                                    # Leer el bimestre (1°, 2°, etc.)
                                    match_bim = re.match(r"(\d+)°", cell0_text)
                                    if match_bim:
                                        bim_from_row = int(match_bim.group(1))
                                    # Leer el año
                                    if cell1_text.isdigit():
                                        row_year = int(cell1_text)

                                    if (row_year is not None and row_year < periodo_año) or (row_year == periodo_año and bim_from_row is not None and bim_from_row < periodo_bimestre):
                                        filas_a_eliminar_esta_tabla.append(row._element)
                                except Exception as e:
                                    print(f"Error procesando fila de tabla 9: {e}")

                            elif table_type in [8, 10, 11, 81011]:
                                m = re.search(r"/\s*(\d{4})\b", cell0_text)
                                row_year = int(m.group(1)) if m else None

                            if row_year is not None and row_year in anos_inactivos:
                                if row._element not in filas_a_eliminar_esta_tabla:
                                    filas_a_eliminar_esta_tabla.append(row._element)

                            if "/" in cell0_text:
                                parts = cell0_text.split("/")
                                try:
                                    bim_from_row = int(parts[0].replace("°", "").strip())
                                except:
                                    bim_from_row = None

                        except IndexError:
                            print("     Warn: Fila %d(%d) tipo %s con menos celdas." % (row_idx_abs+1, row_idx_rel, table_type))
                            continue
                        except Exception as e_row_check:
                            print("     Warn: Error procesando fila %d(%d) tipo %s: %s" % (row_idx_abs+1, row_idx_rel, table_type, e_row_check))
                            continue

                    if filas_a_eliminar_esta_tabla:
                        filas_a_eliminar_por_tabla[table_element] = filas_a_eliminar_esta_tabla

                    # Si la tabla tiene "BIMESTRE/AÑO" y "ÍNDICE NACIONAL DE PRECIOS..." → se identifica como tabla de índice
                    if "BIMESTRE/AÑO" in header_texts[0] and any("ÍNDICE NACIONAL DE PRECIOS" in ht for ht in header_texts):
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            cells = row.cells
                            if len(cells) < 2:
                                continue
                            cell0 = cells[0].text.strip()
                            match = re.match(r"(\d+)°/(\d{4})", cell0)
                            if match:
                                bim = int(match.group(1))
                                year = int(match.group(2))
                                if (year < periodo_año) or (year == periodo_año and bim < periodo_bimestre):
                                    # El bimestre es anterior al periodo → reemplazar todas las celdas (excepto primera si se quiere) por "0.00"
                                    for i in range(1, len(cells)):
                                        cells[i].text = "0.00"
                   

        except IndexError as e_index:
            print("--- Debug WARN: Error Índice tabla %d: %s" % (table_idx, e_index))
            continue
        except Exception as e_table_proc:
            print("--- Debug WARN: Error general tabla %d: %s" % (table_idx, e_table_proc))
            traceback.print_exc()
            continue

    eliminados_count = 0
    if elementos_cuerpo_a_eliminar:
        elementos_cuerpo_eliminados_realmente = []
        for elem in elementos_cuerpo_a_eliminar:
            try:
                parent = elem.getparent() if hasattr(elem, 'getparent') else None
            except AttributeError:
                parent = None

            # 🔒 Protección de tabla de totales
            contiene_texto_protegido = any(
                t.text and (
                    "IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in t.text.upper() or
                    "[TOTAL]" in t.text.upper()
                )
                for t in elem.iter("w:t")
            )
            if contiene_texto_protegido:
                print("🛑 Protegido: NO se elimina una tabla que contiene [TOTAL] o encabezado de totales.")
                continue

            if parent is not None:
                try:
                    parent.remove(elem)
                    eliminados_count += 1
                    elementos_cuerpo_eliminados_realmente.append(elem)
                except Exception as e_del_sec:
                    print("Warn: Error eliminando seccion: %s" % e_del_sec)

    else:
        elementos_cuerpo_eliminados_realmente = []

    if filas_a_eliminar_por_tabla:
        for table_element, row_elements in filas_a_eliminar_por_tabla.items():
            # 🚫 Protección adicional: si la tabla es de totales, no eliminar NADA
            if table_element in tablas_protegidas:
                print("🛡️ Saltando eliminación de filas de tabla protegida (TOTALES)")
                continue

            if table_element in elementos_cuerpo_eliminados_realmente:
                continue

            try: 
                table_parent = table_element.getparent() if hasattr(table_element, 'getparent') else None
            except AttributeError:
                table_parent = None

            if table_parent is None:
                continue

            for row_element in row_elements:
                try: 
                    row_parent = row_element.getparent() if hasattr(row_element, 'getparent') else None
                except AttributeError:
                    row_parent = None
                if row_parent is not None:
                    try: 
                        row_parent.remove(row_element)
                        eliminados_count += 1
                    except Exception as e_del_row:
                        print("Warn: Error eliminando fila por anio: %s" % e_del_row)

            try: 
                table_parent = table_element.getparent() if hasattr(table_element, 'getparent') else None
            except AttributeError:
                table_parent = None
            if table_parent is None:
                continue

            for row_element in row_elements:
                try: 
                    row_parent = row_element.getparent() if hasattr(row_element, 'getparent') else None
                except AttributeError:
                    row_parent = None
                if row_parent is not None:
                    try: 
                        row_parent.remove(row_element)
                        eliminados_count += 1
                    except Exception as e_del_row:
                        print("Warn: Error eliminando fila por anio: %s" % e_del_row)

    if eliminados_count > 0:
        print("      Se eliminaron %d elementos/filas en total." % eliminados_count)

def safe_float(value, default=0.0):
    if pd.isna(value):
        return default
    try:
        cleaned_value = str(value).replace(",", "").strip()
        return float(cleaned_value) if cleaned_value else default
    except (ValueError, TypeError):
        return default

def cargar_lista_pm(filepath):
    pm_expedientes = set()
    encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1']
    delimiters_to_try = ['|', ',', ';']
    df_pm = None
    exp_col = 'EXPEDIENTE'
    pm_col = 'TIPO PM'
    for delim in delimiters_to_try:
        for enc in encodings_to_try:
            try:
                test_df = pd.read_csv(filepath, delimiter=delim, encoding=enc, dtype=str, nrows=0, skipinitialspace=True)
                test_df.columns = test_df.columns.str.strip().str.upper()
                if exp_col.upper() in test_df.columns and pm_col.upper() in test_df.columns:
                    all_cols = pd.read_csv(filepath, delimiter=delim, encoding=enc, dtype=str, nrows=0, skipinitialspace=True).columns.tolist()
                    real_exp_col = [c for c in all_cols if c.strip().upper() == exp_col.upper()][0]
                    real_pm_col = [c for c in all_cols if c.strip().upper() == pm_col.upper()][0]
                    df_pm = pd.read_csv(filepath, delimiter=delim, encoding=enc, dtype=str, usecols=[real_exp_col, real_pm_col], skipinitialspace=True)
                    df_pm.rename(columns={real_exp_col: exp_col, real_pm_col: pm_col}, inplace=True)
                    print("      ¡Archivo PM leido! (Enc:%s, Delim:'%s')" % (enc, delim))
                    break
                else:
                    df_pm = None
                    continue
            except FileNotFoundError:
                print(f"Error: Archivo PM '{filepath}' no encontrado.")
                return pm_expedientes
            except Exception:
                df_pm = None
                continue
        if df_pm is not None:
            break
    if df_pm is None:
        print(f"Warn: No se pudo leer PM '{filepath}'.")
        return pm_expedientes
    try:
        df_pm = df_pm.fillna('')
        pm_entries = df_pm[df_pm[pm_col].astype(str).str.strip() != '']
        pm_expedientes.update(pm_entries[exp_col].astype(str).str.strip())
        print("      Se encontraron %d expedientes PM." % len(pm_expedientes))
    except Exception as e:
        print("Error procesando datos PM: %s" % e)
    return pm_expedientes

# --- NUEVA FUNCIÓN ---
def procesar_tablas_suelo_construccion(document, anos_activos):
    """
    - Elimina completamente las tablas si solo contienen datos de años inactivos.
    - En tablas de años activos, elimina filas con valores en 0 pero conserva las filas finales de totales.
    """
    tablas_a_eliminar = []

    for table in document.tables:
        # —— PROTEJO LA TABLA GLOBAL DE TOTALES ——
        if table.rows:
            header = " ".join(cell.text.upper() for cell in table.rows[0].cells)
            if "IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in header and "TOTAL" in header:
                print("🛡️ Saltando procesar_tablas_suelo_construccion para tabla de totales")
                continue

        try:
            header_text = " ".join(cell.text.upper() for cell in table.rows[0].cells)
            if ("IMPORTE TOTAL DE CONTRIBUCIONES OMITIDAS" in header_text
                and "[TOTAL]" in header_text):
                print("🛡️ Protegiendo tabla global de totales")
                continue
        except Exception:
            pass
        # ——— Protejo aquí la tabla de totales globales ———
        años_detectados = set()
        filas_a_eliminar = []
        contiene_filas_de_totales = False
        contiene_datos_utiles = False
        total_filas = len(table.rows)
        total_filas_totales = 0

        for i, row in enumerate(table.rows):
            texto_fila = " ".join(cell.text.upper() for cell in row.cells)

            # Detectar años en los placeholders (incluso en filas de totales)
            matches = re.findall(r"202\d", texto_fila)
            for match in matches:
                try:
                    años_detectados.add(int(match))
                except:
                    continue

            # Detectar si es fila de totales
            if re.search(r"VALOR CATASTRAL DE CONSTRUCCI|VALOR CATASTRAL TOTAL DERIVADO", texto_fila):
                contiene_filas_de_totales = True
                total_filas_totales += 1
                continue  # Nunca eliminar

            # Verificar si fila contiene valores válidos
            if len(row.cells) >= 2:
                val1 = row.cells[-2].text.replace("$", "").replace(",", "").strip()
                val2 = row.cells[-1].text.replace("$", "").replace(",", "").strip()
                try:
                    n1 = float(val1) if val1 else 0.0
                    n2 = float(val2) if val2 else 0.0
                    if n1 == 0.0 and n2 == 0.0:
                        filas_a_eliminar.append(row._element)
                    else:
                        contiene_datos_utiles = True
                except:
                    continue

        # ❗ Eliminar si:
        # 1. Se detectaron años y todos son inactivos
        # 2. No se detectaron años y no hay datos útiles
        if (años_detectados and años_detectados.isdisjoint(anos_activos)) or (not años_detectados and not contiene_datos_utiles):
            print(f"❌ Eliminando tabla completa (años inválidos o sin datos útiles)")
            tablas_a_eliminar.append(table._element)
            continue



        # Nueva condición: si solo hay filas de totales (sin ninguna fila de datos útil ni con valores), eliminar tabla
        if not contiene_datos_utiles and contiene_filas_de_totales and total_filas == total_filas_totales:
            print(f"❌ Eliminando tabla con solo filas de totales y sin datos útiles")
            tablas_a_eliminar.append(table._element)
            continue

        # Si la tabla es de años activos, eliminar filas vacías (menos las de totales)
        if not contiene_datos_utiles and contiene_filas_de_totales:
            print(f"⚠️ Tabla sin datos útiles en años activos, conservando solo totales")
        for fila in filas_a_eliminar:
            try:
                parent = fila.getparent()
                if parent is not None:
                    parent.remove(fila)
            except Exception as e:
                print(f"⚠️ Error al eliminar fila: {e}")

    # Eliminar tablas completas
    for tabla_elem in tablas_a_eliminar:
        try:
            tabla_elem.getparent().remove(tabla_elem)
        except Exception as e:
            print(f"⚠️ Error al eliminar tabla: {e}")

registros_base = []

# --- NUEVAS FUNCIONES: Gestión Base de Datos Maestra (Excel) ---

def cargar_bd_maestra(ruta_excel):
    with excel_lock:
        if os.path.exists(ruta_excel):
            try:
                print(f"Cargando BD Maestra desde: {ruta_excel}")
                df = pd.read_excel(ruta_excel, sheet_name="BD_Maestra") # Leer sin dtype inicial
                
                # Asegurar que todas las columnas de COLUMNAS_BD_MAESTRA existan y tengan el tipo correcto
                for col_name in COLUMNAS_BD_MAESTRA:
                    expected_dtype = TIPOS_BD_MAESTRA[col_name]
                    if col_name not in df.columns:
                        print(f"  -> Añadiendo columna faltante a BD Maestra: '{col_name}' con tipo '{expected_dtype}'")
                        if expected_dtype == 'Int64':
                            df[col_name] = pd.NA
                        else: # Para str y otros
                            df[col_name] = "" 
                    
                    # Intentar convertir al tipo esperado
                    try:
                        if df[col_name].dtype.name == expected_dtype: # Si ya tiene el tipo correcto
                            if expected_dtype == 'str': # Asegurar que los NaN en str sean ""
                                df[col_name] = df[col_name].fillna("").astype(str)
                            # Para Int64 y otros tipos, si ya coincide, no hacer nada extra aquí.
                            continue # Saltar a la siguiente columna

                        # Si el tipo no coincide, intentar la conversión
                        if expected_dtype == 'Int64':
                            df[col_name] = pd.to_numeric(df[col_name], errors='coerce').astype(expected_dtype)
                        elif expected_dtype == 'str':
                             df[col_name] = df[col_name].astype(str).fillna("")
                        else: 
                            df[col_name] = df[col_name].astype(expected_dtype)
                    except Exception as e_type:
                        print(f"  (!) Advertencia: No se pudo convertir la columna '{col_name}' (tipo actual: {df[col_name].dtype}) al tipo esperado '{expected_dtype}'. Se usará string. Error: {e_type}")
                        df[col_name] = df[col_name].astype(str).fillna("") # Fallback a string
                
                # Seleccionar y reordenar a COLUMNAS_BD_MAESTRA para consistencia interna.
                # Si alguna columna definida en COLUMNAS_BD_MAESTRA aún no existe (improbable después del bucle anterior),
                # se creará aquí por el reindex.
                df = df.reindex(columns=COLUMNAS_BD_MAESTRA)
                # Aplicar tipos una vez más después del reindex por si se crearon nuevas columnas
                # y para asegurar que los valores por defecto (ej. pd.NA para Int64) tengan el tipo correcto.
                for col_name in COLUMNAS_BD_MAESTRA: # Iterar de nuevo para asegurar tipos después de reindex
                    current_col_dtype_name = df[col_name].dtype.name
                    expected_col_dtype = TIPOS_BD_MAESTRA[col_name]
                    
                    # Solo intentar convertir si el tipo actual no es el esperado
                    if current_col_dtype_name != expected_col_dtype:
                        try:
                            if expected_col_dtype == 'Int64':
                                # Para Int64, convertir NaN a pd.NA antes de astype
                                df[col_name] = pd.to_numeric(df[col_name], errors='coerce').astype(expected_col_dtype)
                            elif expected_col_dtype == 'str':
                                 df[col_name] = df[col_name].fillna("").astype(str) # Asegurar que NaN se convierta a ""
                            else: # Para otros tipos
                                df[col_name] = df[col_name].astype(expected_col_dtype)
                        except Exception as e_reindex_astype: # Fallback final
                             print(f"  (!) Advertencia (post-reindex): No se pudo convertir '{col_name}' a '{expected_col_dtype}'. Usando string. Error: {e_reindex_astype}")
                             df[col_name] = df[col_name].astype(str).fillna("")


                print(f"BD Maestra cargada y procesada: {len(df)} registros. Columnas: {df.columns.tolist()}")
                return df
            except Exception as e:
                print(f"(!) ADVERTENCIA: No se pudo leer o procesar '{ruta_excel}'. Se creará/usará una BD vacía. Error: {e}")
                traceback.print_exc()
        else:
            print(f"Archivo BD Maestra '{ruta_excel}' no encontrado. Se creará uno nuevo.")

        df_vacio = pd.DataFrame(columns=COLUMNAS_BD_MAESTRA)
        df_vacio = df_vacio.astype(TIPOS_BD_MAESTRA)
        for col_name, dtype in TIPOS_BD_MAESTRA.items(): # Bucle para asegurar valores iniciales correctos
            if dtype == 'str':
                df_vacio[col_name] = "" # Asegurar strings vacíos en lugar de NaN para columnas string
            # Para Int64, pd.NA ya se maneja por astype(TIPOS_BD_MAESTRA)
        return df_vacio


def actualizar_o_agregar_registro_bd(df_bd, registro_data):
    """
    Actualiza un registro existente basado en EXPEDIENTE o agrega uno nuevo.
    registro_data es un diccionario con los datos de UNA fila.
    Devuelve el DataFrame actualizado.
    """
    expediente = registro_data.get(COL_EXPEDIENTE)
    if not expediente:
        print("(!) Error: Intento de actualizar/agregar registro sin EXPEDIENTE.")
        return df_bd 

    indices = df_bd.index[df_bd[COL_EXPEDIENTE].astype(str) == str(expediente)].tolist()

    if indices:
        idx = indices[-1]
        # print(f"  -> Actualizando registro existente para Expediente: {expediente} (índice: {idx})") # Log reducido
        for col, valor in registro_data.items():
            if col in df_bd.columns:
                target_dtype = TIPOS_BD_MAESTRA.get(col)
                try:
                    if pd.isna(valor):
                        valor_convertido = pd.NA if target_dtype == 'Int64' else ''
                    elif target_dtype == 'Int64':
                        valor_convertido = pd.to_numeric(valor, errors='coerce').astype('Int64')
                    else:
                        valor_convertido = str(valor) 

                    df_bd.loc[idx, col] = valor_convertido
                except Exception as e_conv:
                    print(f"  (!) Warn: Error convirtiendo valor '{valor}' para columna '{col}' (Exp: {expediente}). Se usará string. Error: {e_conv}")
                    df_bd.loc[idx, col] = str(valor) 
    else:
        # print(f"  -> Agregando nuevo registro para Expediente: {expediente}") # Log reducido
        nueva_fila_df = pd.DataFrame([registro_data])
        for col in COLUMNAS_BD_MAESTRA:
            if col not in nueva_fila_df.columns:
                nueva_fila_df[col] = pd.NA if TIPOS_BD_MAESTRA[col] == 'Int64' else ''
        
        # Asegurar el orden de las columnas y los tipos antes de concatenar
        df_bd = pd.concat([df_bd, nueva_fila_df], ignore_index=True)

    return df_bd

# MODIFICACIÓN PRINCIPAL AQUÍ
def guardar_bd_maestra(df_bd, ruta_excel):
    """
    Guarda el DataFrame en la hoja especificada de un archivo Excel existente,
    preservando otras hojas y formatos tanto como sea posible.
    Si el archivo o la hoja no existen, los crea.
    Reemplaza el contenido de la hoja especificada.
    """
    with excel_lock: # Asegura acceso exclusivo
        try:
            sheet_name = "BD_Maestra" # Nombre de la hoja a actualizar
            print(f"Guardando BD Maestra en: '{ruta_excel}', Hoja: '{sheet_name}' ({len(df_bd)} registros)")

            # Intentar cargar el libro de trabajo existente
            try:
                book = load_workbook(ruta_excel)
            except FileNotFoundError:
                # El archivo no existe, crear uno nuevo
                print(f"  Archivo '{ruta_excel}' no encontrado. Creando uno nuevo.")
                book = Workbook()
                # Si es un libro nuevo, la hoja activa por defecto es 'Sheet'.
                # Si nuestra sheet_name es diferente, crearemos la nuestra y eliminaremos 'Sheet' después si es necesario.
            
            # Asegurarse de que la hoja de destino exista
            if sheet_name in book.sheetnames:
                ws = book[sheet_name]
                # Borrar el contenido existente de la hoja para escribir los nuevos datos
                # Esto es como if_sheet_exists='replace'
                # Advertencia: Esto eliminará cualquier formato, filtro, tabla DENTRO de esta hoja.
                print(f"  Limpiando contenido existente de la hoja '{sheet_name}'...")
                ws.delete_rows(1, ws.max_row + 1) # Borra todas las filas, incluyendo encabezados si los hubiera
                # También es buena idea limpiar columnas por si la nueva data es más angosta
                ws.delete_cols(1, ws.max_column + 1)
            else:
                print(f"  Hoja '{sheet_name}' no encontrada. Creándola.")
                ws = book.create_sheet(title=sheet_name)
                # Si acabamos de crear nuestra hoja y la hoja 'Sheet' por defecto existe (y no es la nuestra),
                # y es un libro que tenía 'Sheet' como única hoja o una hoja no deseada.
                if "Sheet" in book.sheetnames and sheet_name != "Sheet" and len(book.sheetnames) > 1:
                    # Si 'Sheet' no es la hoja activa principal y no es nuestra hoja de trabajo, la eliminamos.
                    # Esto es para evitar tener una hoja 'Sheet' vacía si creamos una nueva.
                    if book.active.title == "Sheet" and len(book.worksheets) > 1 : #Si la hoja activa es Sheet y hay mas de una
                         # Cambiar la hoja activa a nuestra hoja de trabajo antes de borrar 'Sheet'
                         book.active = ws

                    # Solo eliminar 'Sheet' si no es la hoja en la que estamos trabajando
                    if "Sheet" in book.sheetnames and sheet_name != "Sheet": #Verificar de nuevo por si acaso
                        print(f"  Eliminando hoja por defecto 'Sheet'.")
                        del book["Sheet"]


            # Escribir los datos del DataFrame en la hoja
            # Asegurar que el DataFrame tenga las columnas en el orden definido en COLUMNAS_BD_MAESTRA
            df_to_write = df_bd[COLUMNAS_BD_MAESTRA].copy()

            print(f"  Escribiendo {len(df_to_write)} filas en la hoja '{sheet_name}'...")
            for r_idx, row in enumerate(dataframe_to_rows(df_to_write, index=False, header=True), 1):
                for c_idx, value in enumerate(row, 1):
                    # Manejo especial para pd.NA que openpyxl podría no manejar bien directamente para ciertos tipos.
                    # Usualmente, None es mejor para celdas vacías.
                    if pd.isna(value):
                        ws.cell(row=r_idx, column=c_idx, value=None)
                    else:
                        ws.cell(row=r_idx, column=c_idx, value=value)
            
            # Si después de todo, 'Sheet' sigue ahí y no es nuestra hoja principal, y hay más de una hoja.
            if "Sheet" in book.sheetnames and sheet_name != "Sheet" and len(book.sheetnames) > 1:
                print(f"  Verificación final: Eliminando hoja por defecto 'Sheet' si aún existe y no es la principal.")
                del book["Sheet"]


            book.save(ruta_excel)
            print(f"  -> BD Maestra guardada exitosamente en '{ruta_excel}', hoja '{sheet_name}'.")
            return True
        except PermissionError:
            print(f"(!) ERROR CRÍTICO: Permiso denegado al guardar '{ruta_excel}'. ¿Archivo abierto?")
            return False
        except Exception as e:
            print(f"(!) ERROR CRÍTICO al guardar BD Maestra '{ruta_excel}' con openpyxl: {e}")
            traceback.print_exc()
            return False

# --- NUEVA FUNCION: Contar páginas PDF ---
def contar_paginas_pdf(ruta_pdf):
    """Cuenta las páginas de un archivo PDF usando PyPDF2."""
    try:
        with open(ruta_pdf, 'rb') as f:
            reader = PdfReader(f, strict=False) # strict=False para más tolerancia
            count = len(reader.pages)
            return count
    except FileNotFoundError:
        print(f"  (!) Error contar_paginas: Archivo no encontrado '{os.path.basename(ruta_pdf)}'")
        return None
    except PdfReadError as e_pdf:
        # Errores comunes: EOF marker not found, file damaged, password protected
        print(f"  (!) Error PyPDF2 al leer '{os.path.basename(ruta_pdf)}': {e_pdf}")
        return None
    except Exception as e:
        print(f"  (!) Error inesperado al contar páginas de '{os.path.basename(ruta_pdf)}': {e}")
        return None

# --- NUEVA FUNCION: Extraer páginas PDF ---
# En GeneradorPredial_logica.py

def extraer_paginas_pdf(ruta_original, ruta_salida, paginas_a_extraer="TODAS"):
    """
    Extrae páginas específicas de un PDF.
    paginas_a_extraer puede ser:
     - "TODAS": Copia el archivo original.
     - "ULTIMA": Extrae solo la última página.
     - "RESTO": Extrae todas menos la última.
     - "PRIMERAS_DOS": Extrae las primeras dos páginas (página 1 y 2).
    Devuelve True si éxito, False si falla.
    """
    try:
        reader = PdfReader(ruta_original, strict=False)
        writer = PdfWriter()
        total_paginas = len(reader.pages)

        if total_paginas == 0 and paginas_a_extraer not in ["PRIMERAS_DOS", "ULTIMA", "RESTO"]: # Permitir que estos modos creen un PDF vacío si el original está vacío.
            print(f"  (!) Error extraer_paginas: '{os.path.basename(ruta_original)}' no tiene páginas.")
            return False

        if paginas_a_extraer == "TODAS":
            for page_num in range(total_paginas):
                writer.add_page(reader.pages[page_num])
        elif paginas_a_extraer == "ULTIMA":
            if total_paginas > 0:
                writer.add_page(reader.pages[total_paginas - 1])
            else: # total_paginas == 0, se creará un PDF vacío
                print(f"  (*) Info extraer_paginas (ULTIMA): '{os.path.basename(ruta_original)}' no tiene páginas. Se creará PDF vacío.")
        elif paginas_a_extraer == "RESTO":
            if total_paginas <= 1:
                print(f"  (*) Info extraer_paginas (RESTO): '{os.path.basename(ruta_original)}' tiene {total_paginas} pág. 'RESTO' resulta en PDF vacío.")
            else:
                for i in range(total_paginas - 1):
                    writer.add_page(reader.pages[i])
        elif paginas_a_extraer == "PRIMERAS_DOS":  # <--- ASEGÚRATE QUE ESTA LÍNEA ES EXACTA
            if total_paginas == 1:
                print(f"  (*) Info extraer_paginas (PRIMERAS_DOS): '{os.path.basename(ruta_original)}' solo tiene 1 página. Extrayendo esa única página.")
                writer.add_page(reader.pages[0])
            elif total_paginas >= 2:
                # print(f"  (*) Info extraer_paginas (PRIMERAS_DOS): Extrayendo las primeras 2 páginas de '{os.path.basename(ruta_original)}'.") # Log opcional
                writer.add_page(reader.pages[0])
                writer.add_page(reader.pages[1])
            else: # total_paginas == 0
                 print(f"  (*) Info extraer_paginas (PRIMERAS_DOS): '{os.path.basename(ruta_original)}' no tiene páginas. Se creará PDF vacío.")
        else:
            print(f"  (!) Error extraer_paginas: Modo '{paginas_a_extraer}' no reconocido.") # <--- ESTE ES EL MENSAJE QUE ESTÁS VIENDO
            return False

        with open(ruta_salida, 'wb') as f_out:
            writer.write(f_out)
        return True

    except FileNotFoundError:
        print(f"  (!) Error extraer_paginas: Archivo original no encontrado '{os.path.basename(ruta_original)}'")
        return False
    except PdfReadError as e_pdf:
        print(f"  (!) Error PyPDF2 al leer '{os.path.basename(ruta_original)}' para extracción: {e_pdf}")
        # No borramos aquí, dejamos que el llamador decida si el archivo de salida es un problema.
        return False
    except Exception as e:
        print(f"  (!) Error inesperado extrayendo páginas de '{os.path.basename(ruta_original)}': {e}")
        traceback.print_exc()
        # Intentar eliminar el archivo de salida si se creó parcialmente y hubo un error grave.
        if os.path.exists(ruta_salida):
            try:
                os.remove(ruta_salida)
                print(f"    INFO: Archivo de salida parcial '{ruta_salida}' eliminado debido a error en extracción.")
            except Exception as e_del:
                print(f"    WARN: No se pudo eliminar el archivo de salida parcial '{ruta_salida}': {e_del}")
        return False
    
# --- NUEVA FUNCIÓN PARA EXPEDIENTES ESPECÍFICOS ---
def generar_expedientes_especificos(csv_data_path, config_path, pm_csv_path, plantilla_path, lista_expedientes, registrar_bd):
    """
    Genera documentos completos para una lista específica de expedientes.
    Permite elegir si se registra o no en la BD Maestra.
    """
    print(f"\n--- Iniciando Generación para Expedientes Específicos ---")
    print(f"Expedientes a procesar: {', '.join(lista_expedientes)}")
    print(f"Registrar en BD Maestra: {'Sí' if registrar_bd else 'No'}")

    locale_ok = configurar_locale()
    if not locale_ok:
        print("Error CRITICO: No se pudo configurar el locale español. Abortando.")
        return

    # Cargar configuración de columnas y lista PM
    print(f"Cargando config: {config_path}")
    nombres_columnas_reales, tipos_columnas = cargar_config_columnas(config_path)
    if nombres_columnas_reales is None:
        print("Error: No se pudo cargar config. Abortando.")
        return
    if "NOMBRE" not in nombres_columnas_reales:
         print("Error CRITICO: La columna 'NOMBRE' (para el nombre del contribuyente) no está definida en config_columnas.csv. Abortando.")
         return
    # (Puedes agregar más validaciones de columnas necesarias aquí si quieres)

    print(f"Cargando lista PM: {pm_csv_path}")
    pm_set = cargar_lista_pm(pm_csv_path)

    # --- Cargar Base de Datos Maestra (solo si se va a registrar) ---
    df_bd_maestra = None
    if registrar_bd:
        df_bd_maestra = cargar_bd_maestra(RUTA_BD_MAESTRA)

    # --- Leer y Preparar Datos de Entrada (BASE_DE_DATOS.csv) ---
    print(f"Leyendo datos de entrada: {csv_data_path}")
    df_datos = None
    try:
        read_encoding = 'utf-8'
        read_args = {'delimiter': '|', 'engine': 'python', 'dtype': str, 'on_bad_lines': 'warn', 'header': None, 'names': nombres_columnas_reales, 'skiprows': 1}
        try:
            df_datos = pd.read_csv(csv_data_path, **read_args, encoding=read_encoding)
        except UnicodeDecodeError:
            read_encoding = 'latin-1'; read_args['encoding'] = read_encoding
            print(f"Reintentando lectura con {read_encoding}...")
            df_datos = pd.read_csv(csv_data_path, **read_args)
        except ValueError as ve:
             try:
                 with open(csv_data_path, 'r', encoding=read_encoding) as f: first_line = f.readline(); num_cols_real = len(first_line.split('|'))
                 print(f"Error Columnas CSV: {ve}\nEsperaba {len(nombres_columnas_reales)}, archivo tiene {num_cols_real}.")
             except Exception as e_inner: print(f"Error Columnas CSV: {ve}\nEsperaba {len(nombres_columnas_reales)}. No se pudo verificar: {e_inner}")
             return
        except FileNotFoundError: print(f"Error CRITICO: Datos '{csv_data_path}' no encontrado."); return
        except Exception as read_err: print(f"Error CRITICO lectura CSV: {read_err}"); traceback.print_exc(); return

        df_datos = df_datos.fillna("")
        print(f"Lectura de datos OK. Registros iniciales: {len(df_datos)}")
    except Exception as e_setup:
        print(f"Error CRITICO setup lectura datos: {e_setup}")
        traceback.print_exc()
        return

    if df_datos is None or df_datos.empty:
        print("Error: No hay datos en el archivo CSV para procesar.")
        if registrar_bd and df_bd_maestra is not None:
            guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA)
        return
    
    df_datos[COL_EXPEDIENTE] = df_datos[COL_EXPEDIENTE].astype(str)

    # --- Filtrar df_datos para que solo contenga los expedientes especificados ---
    df_datos_filtrados = df_datos[df_datos[COL_EXPEDIENTE].isin(lista_expedientes)].copy()
    
    print(f"Registros encontrados para los expedientes especificados: {len(df_datos_filtrados)}")
    if df_datos_filtrados.empty:
         print(f"No se encontraron datos para los expedientes solicitados en '{csv_data_path}'.")
         if registrar_bd and df_bd_maestra is not None:
            guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA)
         return

    # No es necesario el ordenamiento complejo por colonia aquí, ya que son expedientes específicos.
    # Simplemente iteramos sobre los encontrados.
    df_datos_a_procesar = df_datos_filtrados.reset_index(drop=True)

    # --- Bucle Principal de Generación ---
    generados_count = 0
    registros_actualizados_en_bd = 0

    carpeta_colonias_out = os.path.join(CARPETA_PRINCIPAL, "COLONIAS")
    carpeta_vacias_out = os.path.join(CARPETA_PRINCIPAL, "VACIAS")
    os.makedirs(carpeta_colonias_out, exist_ok=True)
    os.makedirs(carpeta_vacias_out, exist_ok=True)

    for index, row in df_datos_a_procesar.iterrows():
        expediente_actual = str(row.get(COL_EXPEDIENTE, "")).strip()
        nombre_contribuyente = str(row.get("NOMBRE", "")).strip()
        direccion_contribuyente = str(row.get("DIRECCION", "")).strip() # Asumiendo columna 'DIRECCION'
        colonia_actual = str(row.get("COLONIA", "")).strip()

        if not expediente_actual: # Debería estar, ya que filtramos por él
            print(f"(!) Advertencia: Registro {index+1} sin EXPEDIENTE válido. Saltando.")
            continue
        if not nombre_contribuyente:
            print(f"(!) Advertencia: Registro {index+1} (Exp: {expediente_actual}) sin NOMBRE. Usando 'SIN_NOMBRE'.")
            nombre_contribuyente = "SIN_NOMBRE"

        print(f"\n--- ({generados_count + 1}/{len(df_datos_a_procesar)}) Procesando Exp: {expediente_actual} | Nombre: {nombre_contribuyente} ---")

        nombre_base_limpio = limpiar_texto(f"{expediente_actual}_{nombre_contribuyente}")
        if not nombre_base_limpio or len(nombre_base_limpio) < len(expediente_actual): # Fallback
            nombre_base_limpio = limpiar_texto(expediente_actual) if limpiar_texto(expediente_actual) else f"exp_{expediente_actual}"

        if not colonia_actual or colonia_actual == "0":
            ruta_salida_expediente = carpeta_vacias_out
            colonia_para_bd = "VACIAS"
        else:
            colonia_nombre_limpio = limpiar_texto(colonia_actual)
            ruta_salida_expediente = os.path.join(carpeta_colonias_out, colonia_nombre_limpio)
            os.makedirs(ruta_salida_expediente, exist_ok=True)
            colonia_para_bd = colonia_actual
        
        ruta_temp_docx = os.path.join(ruta_salida_expediente, f"~{nombre_base_limpio}.docx")
        ruta_final_pdf = os.path.join(ruta_salida_expediente, f"{nombre_base_limpio}.pdf")
        # No necesitamos ruta_temp_pdf_extraccion ya que siempre generamos completo

        pdf_generado_final = False
        estado_final_bd_para_este_exp = ESTADO_ERROR_GENERACION # Asumir error
        paginas_finales = None
        hojas_calculadas = pd.NA
        monto_para_bd = ""
        bimestre_para_bd = ""
        anos_para_bd = ""

        try:
            # --- LÓGICA DE GENERACIÓN DEL DOCX (copiada y adaptada de generar_documentos) ---
            # Esta parte es casi idéntica a la de `generar_documentos`.
            # La principal diferencia es que el modo de generación es siempre "COMPLETO" implícitamente.

            anos_activos = []
            anos_inactivos = []
            possible_year_cols = {2022: "2022", 2023: "2023", 2024: "2024"}
            impuesto_year_cols = {year: f"IMPUESTO PREDIAL DEL AÑO {year}" for year in [2022, 2023, 2024]}
            for year in [2022, 2023, 2024]:
                activo = False; valor_check = None; year_col_name = possible_year_cols.get(year)
                if year_col_name in row.index: valor_check = str(row.get(year_col_name, '0')).strip()
                elif impuesto_year_cols[year] in row.index: valor_check = str(row.get(impuesto_year_cols[year], '0')).strip()
                if valor_check and valor_check != '0' and valor_check.lower() not in ['nan', 'na', '', '-']:
                    if safe_float(valor_check, 0.0) != 0.0: activo = True
                if activo: anos_activos.append(year)
                else: anos_inactivos.append(year)
            print(f"  Anios activos detectados: {anos_activos}")
            if anos_activos: anos_para_bd = f"{min(anos_activos)}-{max(anos_activos)}"
            else: anos_para_bd = ""

            now = datetime.now(); fecha_hoy_str = "(Error Locale)"; fecha_texto_str = "(Error Locale/num2words)"
            if locale_ok:
                 try: fecha_hoy_str = now.strftime("%d de %B de %Y").lower()
                 except Exception as e_fecha: print(f"Warn: Error formateando [FECHA]: {e_fecha}")
                 if NUM2WORDS_INSTALLED:
                     try:
                         dia_num = now.day; ano_num = now.year; dia_palabra = num2words(dia_num, lang='es'); ano_palabra = num2words(ano_num, lang='es')
                         medio_fecha = now.strftime("de %B de %Y").lower(); fecha_texto_str = f"{dia_num}-{dia_palabra} {medio_fecha}-{ano_palabra}"
                     except Exception as e_fecha_texto: print(f"Warn: Error formateando [FECHA_TEXTO]: {e_fecha_texto}")
                 else: fecha_texto_str = "(num2words no instalado)"
            ano_placeholder = str(min(anos_activos)) if anos_activos else str(now.year)
            if not anos_activos: print(f"Warn: No se encontraron años activos para Exp {expediente_actual}, usando año actual para [AÑO].")

            try:
                doc = Document(plantilla_path); print(f"  Plantilla '{plantilla_path}' cargada.")
            except Exception as e:
                print(f"  Error CRITICO al abrir plantilla '{plantilla_path}': {e}. Saltando expediente.")
                estado_final_bd_para_este_exp = "Error Carga Plantilla"; continue

            periodo_str = str(row.get("PERIODO", "")).strip(); periodo_año = 0; periodo_bimestre = 0
            if re.match(r"^\d{6}$", periodo_str):
                 try: periodo_año = int(periodo_str[:4]); periodo_bimestre = int(periodo_str[4:6]); bimestre_para_bd = f"{periodo_bimestre}-{periodo_año}"
                 except: print(f"  Advertencia: PERIODO inválido '{periodo_str}', no se pudo parsear."); bimestre_para_bd = ""
            else: print(f"  Advertencia: PERIODO ausente o mal formado '{periodo_str}'."); bimestre_para_bd = ""
            
            initial_replacements = {};
            placeholders_sumas_total = ["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", "[SUMA DE MONTOS ACTUALIZADOS]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL]", "[SUMA DE LA SANCION POR BIMESTRE]", "[TOTAL]", "[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]", "[SUMA DE MONTOS ACTUALIZADOS TEXTO]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]", "[SUMA DE LA SANCION POR BIMESTRE TEXTO]", "[TOTAL PESOS CORREGIDOS]", "[LEYENDA_FINAL_PAG17]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"]
            initial_replacements["[FECHA]"] = fecha_hoy_str; initial_replacements["[FECHA_TEXTO]"] = fecha_texto_str; initial_replacements["[AÑO]"] = ano_placeholder
            for col_name in nombres_columnas_reales:
                 placeholder = f"[{col_name}]"
                 if placeholder in placeholders_sumas_total or placeholder in initial_replacements or col_name.startswith("IMPUESTO PREDIAL DEL AÑO"): continue
                 valor_final = ""; col_year = None; match_year = re.search(r"(?:^|\b|_|\s)(\d{4})$", col_name.strip())
                 if match_year: year_str = match_year.group(1); col_year = int(year_str) if year_str in ["2022", "2023", "2024"] else None
                 if col_year is not None and col_year in anos_inactivos: valor_final = ""
                 # ...
                 elif col_name in row.index:
                     valor_original = row[col_name]
                     # Correctly call formatear_valor_v6 with all necessary arguments
                     valor_formateado = formatear_valor_v6(valor_original, col_name, tipos_columnas)

                     patrones_bimestre = [r"^BIMESTRE\s+(\d+)\s+(\d{4})$", r"^MONTO ACTUALIZADO PREDIAL POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^RECARGOS POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^SANCION POR BIMESTRE\s+(\d+)\s+(\d{4})$"]
                     for patron in patrones_bimestre:
                         match_bim = re.match(patron, col_name.upper())
                         if match_bim:
                             bim_num = int(match_bim.group(1))
                             bim_year = int(match_bim.group(2))
                             if (bim_year < periodo_año) or (bim_year == periodo_año and bim_num < periodo_bimestre):
                                 # If the bimestre is before the cutoff, format as "$0.00"
                                 valor_formateado = formatear_valor_v6("0", col_name, tipos_columnas)
                             else:
                                 # For current/future bimestres, check if original value is effectively zero
                                 valor_limpio = str(valor_original).replace(",", "").replace("$", "").strip()
                                 try:
                                     numero_local_bim = float(valor_limpio) # Use a distinct variable name
                                 except:
                                     numero_local_bim = 0

                                 if numero_local_bim <= 0: # If original number is zero or negative
                                     valor_formateado = " " # Represent as a space
                                 # else: valor_formateado (from the initial call) is already correct
                             break # Exit patron loop
                     valor_final = valor_formateado
                     if col_name == "EXPEDIENTE": valor_str_exp = str(valor_final).strip(); valor_final = '0' + valor_str_exp if len(valor_str_exp) == 7 else valor_str_exp
                 initial_replacements[placeholder] = valor_final

            print("  Aplicando reemplazos iniciales...")
            reemplazar_en_documento_v2(doc, initial_replacements, fase="inicial")

            print("  Eliminando secciones/filas de años inactivos...")
            tablas_protegidas = []; eliminar_elementos_inactivos_v_usuario(doc, anos_inactivos, periodo_año, periodo_bimestre, tablas_protegidas)
            
            print("  Procesando tablas de suelo/construcción...")
            procesar_tablas_suelo_construccion(doc, anos_activos)

            print("  Calculando sumas finales...")
            impuesto_anual_por_año = {2022: 0.0, 2023: 0.0, 2024: 0.0}; sum_impuesto_anual_total = 0.0; sum_monto_actualizado_total = 0.0; sum_recargos_total = 0.0; sum_sancion_total = 0.0
            for year in [2022, 2023, 2024]:
                 for bim in range(1, 7):
                     if (year < periodo_año) or (year == periodo_año and bim < periodo_bimestre): continue
                     campo_bim = f"BIMESTRE {bim} {year}"; campo_act = f"MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim} {year}"; campo_rec = f"RECARGOS POR BIMESTRE {bim} {year}"; campo_san = f"SANCION POR BIMESTRE {bim} {year}"
                     valor_bim_actual = safe_float(row.get(campo_bim, 0.0)); impuesto_anual_por_año[year] += valor_bim_actual; sum_impuesto_anual_total += valor_bim_actual
                     sum_monto_actualizado_total += safe_float(row.get(campo_act, 0.0)); sum_recargos_total += safe_float(row.get(campo_rec, 0.0)); sum_sancion_total += safe_float(row.get(campo_san, 0.0))
            monto_actualizado_restados = sum_monto_actualizado_total - sum_impuesto_anual_total
            total_general_recalculado = sum_impuesto_anual_total + monto_actualizado_restados + sum_recargos_total + sum_sancion_total
            total_redondeado_bd = round(total_general_recalculado) # Redondeo aquí
            monto_para_bd = f"$ {locale.format_string('%.2f', total_redondeado_bd, grouping=True)}"


            final_replacements = {}; bimestres_ordinales = {1: "PRIMER", 2: "SEGUNDO", 3: "TERCERO", 4: "CUARTO", 5: "QUINTO", 6: "SEXTO"}
            final_replacements["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]"] = f"$ {formatear_valor_v6(sum_impuesto_anual_total, '[SUMA...IMP]', {'[SUMA...IMP]': 'moneda'})}"
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS]"] = f"$ {formatear_valor_v6(sum_monto_actualizado_total, '[SUMA...ACT]', {'[SUMA...ACT]': 'moneda'})}"
            final_replacements["[SUMA DE LOS MONTOS DE RECARGO TOTAL]"] = f"$ {formatear_valor_v6(sum_recargos_total, '[SUMA...REC]', {'[SUMA...REC]': 'moneda'})}"
            final_replacements["[SUMA DE LA SANCION POR BIMESTRE]"] = f"$ {formatear_valor_v6(sum_sancion_total, '[SUMA...SAN]', {'[SUMA...SAN]': 'moneda'})}"
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS RESTADOS]"] = f"$ {formatear_valor_v6(monto_actualizado_restados, '[SUMA DE MONTOS ACTUALIZADOS RESTADOS]', {'[SUMA DE MONTOS ACTUALIZADOS RESTADOS]': 'moneda'})}"
            final_replacements["[TOTAL]"] = f"$ {locale.format_string('%.2f', total_redondeado_bd, grouping=True)}" # Usar el total redondeado
            for año_ip in [2022, 2023, 2024]:
                 valor_ip_anual = round(impuesto_anual_por_año[año_ip], 2); placeholder_ip = f"[IMPUESTO PREDIAL DEL AÑO {año_ip}]"; tipo_dato_ip = {f'IP{año_ip}': 'moneda'}
                 if valor_ip_anual > 0: final_replacements[placeholder_ip] = f"$ {formatear_valor_v6(valor_ip_anual, f'IP{año_ip}', tipo_dato_ip)}"
                 else: final_replacements[placeholder_ip] = " " # O "$ 0.00" si prefieres
            ordinal_texto = bimestres_ordinales.get(periodo_bimestre, ""); final_replacements["[BIMESTRE_ORDINAL]"] = ordinal_texto
            bm_por_anio = {}; anio_base = periodo_año; bimestre_base = periodo_bimestre
            for anio in [2022, 2023, 2024]:
                 if anio == anio_base: bm_por_anio[anio] = bimestre_base
                 elif anio > anio_base: bm_por_anio[anio] = 1
                 else: continue # No incluir años anteriores al base
            for anio_bm, bm_val in bm_por_anio.items(): final_replacements[f"[BM{anio_bm}]"] = f"{bm_val}°" # Ejemplo [BM2023] = 1°
            final_replacements["[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]"] = formatear_texto_moneda(sum_impuesto_anual_total)
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total)
            final_replacements["[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]"] = formatear_texto_moneda(sum_recargos_total)
            final_replacements["[SUMA DE LA SANCION POR BIMESTRE TEXTO]"] = formatear_texto_moneda(sum_sancion_total)
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"] = formatear_texto_moneda(monto_actualizado_restados)
            final_replacements["[TOTAL PESOS CORREGIDOS]"] = formatear_texto_moneda(total_redondeado_bd) # Usar el total redondeado

            es_pm = expediente_actual in pm_set
            final_replacements["[LEYENDA_FINAL_PAG17]"] = "" if es_pm else LEYENDA_PAG17_TEXTO

            print("  Aplicando reemplazos finales...")
            reemplazar_en_documento_v2(doc, final_replacements, fase="final")

             # --- INICIO DEL NUEVO BLOQUE PARA ELIMINAR "CONSTRUCCION" (VERSIÓN 2) ---
            if verificar_ausencia_info_construccion_v2(row, anos_activos):
                print("  Ausencia de información de construcción (columnas clave son cero). Procediendo a eliminar la palabra 'CONSTRUCCION'.")
                eliminar_palabra_especifica_del_documento(doc, "CONSTRUCCION")
            else:
                print("  Información de construcción presente (columnas clave tienen valor). No se eliminará la palabra 'CONSTRUCCION'.")
            # --- FIN DEL NUEVO BLOQUE ---

            try:
                doc.save(ruta_temp_docx)
                print(f"  DOCX temporal base guardado: {os.path.basename(ruta_temp_docx)}")

                # ----- INICIO: BLOQUE PARA COPIAR PDF AL SERVIDOR -----
                try:
                    # Reemplaza esto con el nombre de tu variable que contiene la ruta al PDF FINAL.
                    ruta_origen_del_pdf = Path(ruta_final_pdf)

                    # Usamos la variable 'row' de tu bucle para obtener la ruta de destino.
                    ruta_relativa_destino = row["Ruta PDF Generado"]

                    # Construimos la ruta de destino completa en el servidor.
                    ruta_destino_servidor = RUTA_BASE_PDF_SERVIDOR_PREDIAL / ruta_relativa_destino

                    # Aseguramos que la carpeta de destino en el servidor exista.
                    ruta_destino_servidor.parent.mkdir(parents=True, exist_ok=True)

                    # Copiamos el PDF final al servidor.
                    print(f"    -> Copiando PDF final a la red: {ruta_destino_servidor}")
                    shutil.copy2(ruta_origen_del_pdf, ruta_destino_servidor)
                    print("    -> ¡PDF copiado a la red con éxito!")

                except Exception as e:
                    print(f"    -> ERROR CRÍTICO AL COPIAR PDF A LA RED: {e}")
                    # Opcional: registrar el error en la base de datos si lo deseas.
                    # estado_final_bd = "Error Copia Servidor"
                # ----- FIN: BLOQUE PARA COPIAR PDF AL SERVIDOR -----
            except PermissionError:
                print(f"  Error PERMISO al guardar DOCX temporal '{os.path.basename(ruta_temp_docx)}'. ¿Archivo abierto? Saltando expediente.")
                estado_final_bd_para_este_exp = "Error Permiso DOCX"; continue
            except Exception as e_save_docx:
                print(f"  Error CRITICO al guardar DOCX temporal '{os.path.basename(ruta_temp_docx)}': {e_save_docx}")
                traceback.print_exc(); estado_final_bd_para_este_exp = "Error Guardar DOCX"; continue
            # --- FIN LÓGICA DOCX ---

            # --- GENERACIÓN PDF (siempre completo) ---
            pdf_creado_ok = False
            if DOCX2PDF_INSTALLED:
                try:
                    print(f"  Convirtiendo DOCX a PDF final...")
                    convert(ruta_temp_docx, ruta_final_pdf)
                    print(f"  -> PDF COMPLETO generado: {os.path.basename(ruta_final_pdf)}")
                    pdf_creado_ok = True
                except Exception as e_conv:
                    print(f"  (!) Error CRITICO al convertir DOCX a PDF completo: {e_conv}"); traceback.print_exc()
                    estado_final_bd_para_este_exp = "Error Conversion Completo"
                    if os.path.exists(ruta_final_pdf): # Limpiar si falló
                        try: os.remove(ruta_final_pdf)
                        except: pass
            else:
                print("  (!) WARN: 'docx2pdf' no instalado. No se puede generar PDF.");
                estado_final_bd_para_este_exp = "Error Falta docx2pdf"

            if pdf_creado_ok:
                pdf_generado_final = True

                if paginas_finales is not None and paginas_finales >= 0 :
                    hojas_calculadas = math.ceil(paginas_finales / 2.0) * 2 if paginas_finales > 0 else 0 # Asumiendo 2 copias si eso significa "hojas"
                    print(f"  Páginas contadas PDF final: {paginas_finales}, Hojas calculadas: {hojas_calculadas}")
                else:
                    print(f"  (!) Advertencia: PDF final '{os.path.basename(ruta_final_pdf)}' existe pero no se pudieron contar sus páginas.")
                    hojas_calculadas = pd.NA 
                
                estado_final_bd_para_este_exp = ESTADO_GEN_COMPLETO # Siempre es generación completa

        except Exception as main_loop_error:
            print(f"  (!) Error INESPERADO procesando expediente {expediente_actual}: {main_loop_error}")
            traceback.print_exc()
            estado_final_bd_para_este_exp = ESTADO_ERROR_GENERACION
        
        finally:
            if eliminar_docx_intermedio and os.path.exists(ruta_temp_docx):
                try: os.remove(ruta_temp_docx)
                except Exception as e_del_docx: print(f"  Warn: No se pudo eliminar DOCX temp '{os.path.basename(ruta_temp_docx)}': {e_del_docx}")

            if registrar_bd and df_bd_maestra is not None and expediente_actual:
                
                # --- INICIO DE LA CORRECCIÓN ---
                # Obtener el valor de TIPO DE SUELO desde la fila de datos.
                # Esta es la línea que faltaba agregar aquí.
                tipo_de_suelo_valor = str(row.get("TIPO DE SUELO", "")).strip()
                # --- FIN DE LA CORRECCIÓN ---

                registro_data_a_actualizar = {
                    COL_EXPEDIENTE: expediente_actual,
                    COL_NOMBRE_CONTRIBUYENTE: nombre_contribuyente,
                    COL_DIRECCION: direccion_contribuyente,
                    COL_COLONIA: colonia_para_bd,
                    COL_BIMESTRE: bimestre_para_bd,
                    COL_ANOS: anos_para_bd,
                    "TIPO": tipo_de_suelo_valor, # <-- Ahora esta línea funcionará
                    COL_MONTO: monto_para_bd,
                    COL_ESTADO: estado_final_bd_para_este_exp, # Estado de ESTA generación
                    COL_HOJAS_DOC: hojas_calculadas if pd.notna(hojas_calculadas) else pd.NA,
                }
                df_bd_maestra = actualizar_o_agregar_registro_bd(df_bd_maestra, registro_data_a_actualizar)
                registros_actualizados_en_bd += 1
            
            if pdf_generado_final:
                generados_count += 1

    print(f"\n--- Proceso de Generación para Expedientes Específicos Finalizado ---")
    print(f"Documentos PDF generados/intentados: {generados_count}")
    if registrar_bd:
        print(f"Registros actualizados/agregados en BD Maestra: {registros_actualizados_en_bd}")
        if df_bd_maestra is not None and registros_actualizados_en_bd > 0:
            print("Guardando estado final de la Base de Datos Maestra...")
            if not guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA):
                 print("(!) ERROR GRAVE: Falló el guardado FINAL de la BD Maestra.")
        elif df_bd_maestra is not None:
            print("No hubo cambios para guardar en la Base de Datos Maestra.")

# --- Función Principal REESTRUCTURADA y CORREGIDA ---
def generar_documentos(csv_data_path, config_path, pm_csv_path, plantilla_path, modo_generacion, max_archivos):
    """
    Función principal para generar documentos según el modo especificado.
    Lee estados de la BD Maestra para filtrar y escribe el estado resultante.
    Args:
        csv_data_path (str): Ruta al CSV de datos.
        config_path (str): Ruta al CSV de configuración de columnas.
        pm_csv_path (str): Ruta al CSV de expedientes PM.
        plantilla_path (str): Ruta a la plantilla DOCX.
        modo_generacion (str): MODO_COMPLETO, MODO_ULTIMA, o MODO_RESTO.
        max_archivos (int): Número máximo de archivos a generar (0 para todos).
    """
    print(f"\n--- Iniciando Generación en Modo: {modo_generacion} ---")
    print(f"--- Máximo a generar: {'Todos' if max_archivos == 0 else max_archivos} ---")

    locale_ok = configurar_locale()
    if not locale_ok:
        print("Error CRITICO: No se pudo configurar el locale español. Abortando.")
        return

    # Cargar configuración de columnas y lista PM
    print(f"Cargando config: {config_path}")
    nombres_columnas_reales, tipos_columnas = cargar_config_columnas(config_path)
    if nombres_columnas_reales is None:
        print("Error: No se pudo cargar config. Abortando.")
        return
    if "NOMBRE" not in nombres_columnas_reales:
         print("Error CRITICO: La columna 'NOMBRE' (para el nombre del contribuyente) no está definida en config_columnas.csv. Abortando.")
         return
    if "DIRECCION" not in nombres_columnas_reales:
         print("Advertencia: La columna 'DIRECCION' no está definida en config_columnas.csv. Se dejará vacía en la BD Maestra.")

    print(f"Cargando lista PM: {pm_csv_path}")
    pm_set = cargar_lista_pm(pm_csv_path)

    # --- Cargar Base de Datos Maestra ---
    df_bd_maestra = cargar_bd_maestra(RUTA_BD_MAESTRA)
    # Crear un mapeo Expediente -> Estado para consulta rápida
    mapa_estado_expediente = {}
    if COL_EXPEDIENTE in df_bd_maestra.columns and COL_ESTADO in df_bd_maestra.columns:
         df_bd_maestra[COL_EXPEDIENTE] = df_bd_maestra[COL_EXPEDIENTE].astype(str)
         df_bd_maestra[COL_ESTADO] = df_bd_maestra[COL_ESTADO].fillna(ESTADO_PENDIENTE)
         mapa_estado_expediente = df_bd_maestra.drop_duplicates(subset=[COL_EXPEDIENTE], keep='last').set_index(COL_EXPEDIENTE)[COL_ESTADO].to_dict()
    print(f"Se mapearon estados para {len(mapa_estado_expediente)} expedientes únicos desde la BD Maestra.")

    # --- Leer y Preparar Datos de Entrada ---
    print(f"Leyendo datos de entrada: {csv_data_path}")
    df_datos = None # Inicializar
    try:
        # Lógica existente para leer el CSV de datos (BASE_DE_DATOS.csv)
        read_encoding = 'utf-8'
        read_args = {'delimiter': '|', 'engine': 'python', 'dtype': str, 'on_bad_lines': 'warn', 'header': None, 'names': nombres_columnas_reales, 'skiprows': 1}
        try:
            df_datos = pd.read_csv(csv_data_path, **read_args, encoding=read_encoding)
        except UnicodeDecodeError:
            read_encoding = 'latin-1'; read_args['encoding'] = read_encoding
            print(f"Reintentando lectura con {read_encoding}...")
            df_datos = pd.read_csv(csv_data_path, **read_args)
        # Manejo de errores de lectura
        except ValueError as ve:
             try:
                 with open(csv_data_path, 'r', encoding=read_encoding) as f: first_line = f.readline(); num_cols_real = len(first_line.split('|'))
                 print(f"Error Columnas CSV: {ve}\nEsperaba {len(nombres_columnas_reales)}, archivo tiene {num_cols_real}.")
             except Exception as e_inner: print(f"Error Columnas CSV: {ve}\nEsperaba {len(nombres_columnas_reales)}. No se pudo verificar: {e_inner}")
             return
        except FileNotFoundError: print(f"Error CRITICO: Datos '{csv_data_path}' no encontrado."); return
        except Exception as read_err: print(f"Error CRITICO lectura CSV: {read_err}"); traceback.print_exc(); return

        df_datos = df_datos.fillna("")
        print(f"Lectura de datos OK. Registros iniciales: {len(df_datos)}")
    except Exception as e_setup:
        print(f"Error CRITICO setup lectura datos: {e_setup}")
        traceback.print_exc()
        return

    if df_datos is None or df_datos.empty:
        print("Error: No hay datos en el archivo CSV para procesar.")
        guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA) # Guardar por si se añadieron columnas
        return
    # Asegurar que EXPEDIENTE en df_datos sea string
    df_datos[COL_EXPEDIENTE] = df_datos[COL_EXPEDIENTE].astype(str)

    # --- Filtrar df_datos según el modo de generación y el estado en BD Maestra ---
    print(f"Filtrando {len(df_datos)} registros leídos según modo '{modo_generacion}' y estados en BD...")
    indices_a_mantener = []
    for index, row in df_datos.iterrows():
        expediente = row[COL_EXPEDIENTE]
        estado_actual = mapa_estado_expediente.get(expediente, ESTADO_PENDIENTE)

        mantener = False
        if modo_generacion == MODO_COMPLETO:
            # Generar completo si NO está ya Generado COMPLETO o Impreso COMPLETO.
            if estado_actual not in [ESTADO_GEN_COMPLETO, ESTADO_IMP_COMPLETO]:
                mantener = True
        elif modo_generacion == MODO_ULTIMA:
            # Generar última hoja SOLO si está pendiente o tuvo error de generación previo.
            if estado_actual == ESTADO_PENDIENTE or estado_actual == ESTADO_ERROR_GENERACION:
                mantener = True
        elif modo_generacion == MODO_RESTO:
            # Generar resto si está generado última o impreso última.
            if estado_actual in [ESTADO_GEN_ULTIMA, ESTADO_IMP_ULTIMA]:
                mantener = True

        if mantener:
            indices_a_mantener.append(index)

    df_datos_filtrados = df_datos.loc[indices_a_mantener].copy()
    num_filtrados = len(df_datos) - len(df_datos_filtrados)
    print(f"Registros después de filtrar por estado/modo: {len(df_datos_filtrados)} (Se omitieron {num_filtrados})")

    if df_datos_filtrados.empty:
         print(f"No hay registros pendientes para procesar en modo '{modo_generacion}' después de filtrar.")
         guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA) # Guardar por si se añadieron columnas
         return

# --- NUEVA LÓGICA DE ORDENAMIENTO DE 3 NIVELES ---
    print("Aplicando nuevo ordenamiento de 3 niveles...")

    # 1. Preparar las columnas para el ordenamiento
    df_a_ordenar = df_datos_filtrados.copy()
    df_a_ordenar['COLONIA_PROC'] = df_a_ordenar['COLONIA'].astype(str).fillna('VACIA').str.strip().str.upper()
    df_a_ordenar.loc[df_a_ordenar['COLONIA_PROC'] == '', 'COLONIA_PROC'] = 'VACIA'

    # Convertir la columna TOTAL a un número para poder ordenar por monto.
    # Se asume que la columna se llama 'TOTAL' como indicaste en el requisito.
    # Se crea una columna temporal numérica para evitar problemas con símbolos de moneda o comas.
    if 'TOTAL' in df_a_ordenar.columns:
        df_a_ordenar['__MONTO_SORT__'] = pd.to_numeric(
            df_a_ordenar['TOTAL'].astype(str).str.replace('$', '').str.replace(',', ''),
            errors='coerce'
        ).fillna(0.0)
        print("      - Columna 'TOTAL' encontrada y convertida a número para ordenar.")
    else:
        print("      - ADVERTENCIA: No se encontró la columna 'TOTAL'. El orden por monto no se aplicará.")
        df_a_ordenar['__MONTO_SORT__'] = 0.0

    # 2. Calcular el conteo de expedientes por colonia
    df_a_ordenar['__CONTEO_COLONIA__'] = df_a_ordenar.groupby('COLONIA_PROC')['EXPEDIENTE'].transform('nunique')
    
    # 3. Aplicar el ordenamiento de 3 niveles
    #    - Nivel 1: Conteo por colonia, de mayor a menor (descending)
    #    - Nivel 2: Nombre de la colonia, alfabéticamente (ascending)
    #    - Nivel 3: Monto total, de mayor a menor (descending)
    df_datos_ordenados = df_a_ordenar.sort_values(
        by=['__CONTEO_COLONIA__', 'COLONIA_PROC', '__MONTO_SORT__'],
        ascending=[False, True, False]
    ).reset_index(drop=True)

    # Limpiar las columnas temporales que usamos para ordenar
    df_datos_ordenados = df_datos_ordenados.drop(columns=['__MONTO_SORT__', '__CONTEO_COLONIA__'])
    
    print(f"Total de registros a procesar en esta ejecución (filtrados y ordenados): {len(df_datos_ordenados)}")
    print("Orden de procesamiento final establecido.")
    # --- FIN DE LA NUEVA LÓGICA DE ORDENAMIENTO ---

    # --- Bucle Principal de Generación (Itera sobre df_datos_ordenados) ---
    generados_count = 0
    registros_actualizados_en_bd = 0

    # Crear carpetas de salida
    carpeta_colonias_out = os.path.join(CARPETA_PRINCIPAL, "COLONIAS")
    carpeta_vacias_out = os.path.join(CARPETA_PRINCIPAL, "VACIAS")
    os.makedirs(carpeta_colonias_out, exist_ok=True)
    os.makedirs(carpeta_vacias_out, exist_ok=True)

    for index, row in df_datos_ordenados.iterrows():
        if max_archivos > 0 and generados_count >= max_archivos:
            print(f"\nAlcanzado límite de {max_archivos} archivos generados en esta ejecución.")
            break

        # Extraer datos clave de la fila actual
        expediente_actual = str(row.get(COL_EXPEDIENTE, "")).strip()
        nombre_contribuyente = str(row.get("NOMBRE", "")).strip() # Asumiendo que la columna se llama 'NOMBRE'
        direccion_contribuyente = str(row.get("DIRECCION", "")).strip() # Asumiendo columna 'DIRECCION'
        colonia_actual = str(row.get("COLONIA", "")).strip()

        if not expediente_actual:
            print(f"(!) Advertencia: Registro {index+1} sin EXPEDIENTE válido. Saltando.")
            continue
        if not nombre_contribuyente:
            print(f"(!) Advertencia: Registro {index+1} (Exp: {expediente_actual}) sin NOMBRE. Usando 'SIN_NOMBRE'.")
            nombre_contribuyente = "SIN_NOMBRE"

        print(f"\n--- ({generados_count + 1}/{len(df_datos_ordenados) if max_archivos == 0 else min(max_archivos, len(df_datos_ordenados))}) Procesando Exp: {expediente_actual} | Nombre: {nombre_contribuyente} | Modo: {modo_generacion} ---")

        # Definir Rutas y Nombres de Archivo
        nombre_base_limpio = limpiar_texto(f"{expediente_actual}_{nombre_contribuyente}")
        if not nombre_base_limpio or len(nombre_base_limpio) < len(expediente_actual):
            nombre_base_limpio = limpiar_texto(expediente_actual) if limpiar_texto(expediente_actual) else f"exp_{expediente_actual}"

        if not colonia_actual or colonia_actual == "0":
            ruta_salida_expediente = carpeta_vacias_out
            colonia_para_bd = "VACIAS"
        else:
            colonia_nombre_limpio = limpiar_texto(colonia_actual)
            ruta_salida_expediente = os.path.join(carpeta_colonias_out, colonia_nombre_limpio)
            os.makedirs(ruta_salida_expediente, exist_ok=True)
            colonia_para_bd = colonia_actual

        ruta_temp_docx = os.path.join(ruta_salida_expediente, f"~{nombre_base_limpio}.docx")
        ruta_final_pdf = os.path.join(ruta_salida_expediente, f"{nombre_base_limpio}.pdf")
        ruta_temp_pdf_extraccion = os.path.join(ruta_salida_expediente, f"~{nombre_base_limpio}_temp_extract.pdf")

        # Inicializar variables para este expediente
        pdf_generado_final = False
        estado_final_bd = ESTADO_ERROR_GENERACION # Asumir error por defecto
        paginas_finales = None
        hojas_calculadas = pd.NA
        monto_para_bd = ""
        bimestre_para_bd = ""
        anos_para_bd = ""

        # --- INICIO Bloque Try/Except/Finally ---
        try:
            # --- INICIO LÓGICA DE GENERACIÓN DEL DOCX (EXISTENTE) ---

            # --- Determinación Años Activos/Inactivos ---
            anos_activos = []
            anos_inactivos = []
            possible_year_cols = {2022: "2022", 2023: "2023", 2024: "2024"}
            impuesto_year_cols = {year: f"IMPUESTO PREDIAL DEL AÑO {year}" for year in [2022, 2023, 2024]}
            for year in [2022, 2023, 2024]:
                activo = False; valor_check = None; year_col_name = possible_year_cols.get(year)
                if year_col_name in row.index: valor_check = str(row.get(year_col_name, '0')).strip()
                elif impuesto_year_cols[year] in row.index: valor_check = str(row.get(impuesto_year_cols[year], '0')).strip()
                if valor_check and valor_check != '0' and valor_check.lower() not in ['nan', 'na', '', '-']:
                    if safe_float(valor_check, 0.0) != 0.0: activo = True
                if activo: anos_activos.append(year)
                else: anos_inactivos.append(year)
            print(f"  Anios activos detectados: {anos_activos}")
            if anos_activos: anos_para_bd = f"{min(anos_activos)}-{max(anos_activos)}"
            else: anos_para_bd = ""

            # --- Determinar Placeholders Fecha/Año ---
            now = datetime.now(); fecha_hoy_str = "(Error Locale)"; fecha_texto_str = "(Error Locale/num2words)"
            if locale_ok:
                 try: fecha_hoy_str = now.strftime("%d de %B de %Y").lower()
                 except Exception as e_fecha: print(f"Warn: Error formateando [FECHA]: {e_fecha}")
                 if NUM2WORDS_INSTALLED:
                     try:
                         dia_num = now.day; ano_num = now.year; dia_palabra = num2words(dia_num, lang='es'); ano_palabra = num2words(ano_num, lang='es')
                         medio_fecha = now.strftime("de %B de %Y").lower(); fecha_texto_str = f"{dia_num}-{dia_palabra} {medio_fecha}-{ano_palabra}"
                     except Exception as e_fecha_texto: print(f"Warn: Error formateando [FECHA_TEXTO]: {e_fecha_texto}")
                 else: fecha_texto_str = "(num2words no instalado)"
            ano_placeholder = str(min(anos_activos)) if anos_activos else str(now.year)
            if not anos_activos: print(f"Warn: No se encontraron años activos para Exp {expediente_actual}, usando año actual para [AÑO].")

            # --- Cargar Plantilla ---
            try:
                doc = Document(plantilla_path); print(f"  Plantilla '{plantilla_path}' cargada.")
            except Exception as e:
                print(f"  Error CRITICO al abrir plantilla '{plantilla_path}': {e}. Saltando expediente.")
                estado_final_bd = "Error Carga Plantilla"; continue

            # --- Extraer PERIODO ---
            periodo_str = str(row.get("PERIODO", "")).strip(); periodo_año = 0; periodo_bimestre = 0
            if re.match(r"^\d{6}$", periodo_str):
                 try: periodo_año = int(periodo_str[:4]); periodo_bimestre = int(periodo_str[4:6]); bimestre_para_bd = f"{periodo_bimestre}-{periodo_año}"
                 except: print(f"  Advertencia: PERIODO inválido '{periodo_str}', no se pudo parsear."); bimestre_para_bd = ""
            else: print(f"  Advertencia: PERIODO ausente o mal formado '{periodo_str}'."); bimestre_para_bd = ""

            # --- Crear Diccionario Reemplazos INICIAL ---
            initial_replacements = {};
            placeholders_sumas_total = ["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]", "[SUMA DE MONTOS ACTUALIZADOS]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL]", "[SUMA DE LA SANCION POR BIMESTRE]", "[TOTAL]", "[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]", "[SUMA DE MONTOS ACTUALIZADOS TEXTO]", "[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]", "[SUMA DE LA SANCION POR BIMESTRE TEXTO]", "[TOTAL PESOS CORREGIDOS]", "[LEYENDA_FINAL_PAG17]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS]", "[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"]
            initial_replacements["[FECHA]"] = fecha_hoy_str; initial_replacements["[FECHA_TEXTO]"] = fecha_texto_str; initial_replacements["[AÑO]"] = ano_placeholder
            for col_name in nombres_columnas_reales:
                 placeholder = f"[{col_name}]"
                 if placeholder in placeholders_sumas_total or placeholder in initial_replacements or col_name.startswith("IMPUESTO PREDIAL DEL AÑO"): continue
                 valor_final = ""; col_year = None; match_year = re.search(r"(?:^|\b|_|\s)(\d{4})$", col_name.strip())
                 if match_year: year_str = match_year.group(1); col_year = int(year_str) if year_str in ["2022", "2023", "2024"] else None
                 if col_year is not None and col_year in anos_inactivos: valor_final = ""
                 elif col_name in row.index:
                     valor_original = row[col_name]; 
                     valor_formateado = formatear_valor_v6(valor_original, col_name, tipos_columnas)
                     campos_con_pesos = ["SUELO_VALOR_UNITARIO", "VALOR CATASTRAL", "VALOR_CATASTRAL_2022", "VALOR_CATASTRAL_2023", "VALOR_CATASTRAL_2024", "VALOR UNITARIO DE CONSTRUCCION", "VALOR CATASTRAL CONSTRUCCION"]
                     es_campo_especial_pesos = any(texto in col_name.upper() for texto in campos_con_pesos)
                     if es_campo_especial_pesos:
                         try: numero = float(str(valor_formateado).replace(",", "").replace("$", "").strip());
                         except: numero = 0 # Fallback si no es número
                         if numero > 50: valor_formateado = f"$ {valor_formateado}"

                     es_bimestre_detectado = False
                     patrones_bimestre = [r"^BIMESTRE\s+(\d+)\s+(\d{4})$", r"^MONTO ACTUALIZADO PREDIAL POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^RECARGOS POR BIMESTRE\s+(\d+)\s+(\d{4})$", r"^SANCION POR BIMESTRE\s+(\d+)\s+(\d{4})$"]
                     for patron in patrones_bimestre:
                         match_bim = re.match(patron, col_name.upper())
                         if match_bim:
                             bim_num = int(match_bim.group(1))
                             bim_year = int(match_bim.group(2))
                             if (bim_year < periodo_año) or (bim_year == periodo_año and bim_num < periodo_bimestre):
                                 # If the bimestre is before the cutoff, format as "$0.00"
                                 valor_formateado = formatear_valor_v6("0", col_name, tipos_columnas)
                             else:
                                 # For current/future bimestres, check if original value is effectively zero
                                 valor_limpio = str(valor_original).replace(",", "").replace("$", "").strip()
                                 try:
                                     numero_local_bim = float(valor_limpio) # Use a distinct variable name
                                 except:
                                     numero_local_bim = 0

                                 if numero_local_bim <= 0: # If original number is zero or negative
                                     valor_formateado = " " # Represent as a space
                                 # else: valor_formateado (from the initial call) is already correct
                             break # Exit patron loop
                     valor_final = valor_formateado
                     if col_name == "EXPEDIENTE": valor_str_exp = str(valor_final).strip(); valor_final = '0' + valor_str_exp if len(valor_str_exp) == 7 else valor_str_exp
                 initial_replacements[placeholder] = valor_final

            # --- Aplicar Reemplazos Iniciales ---
            print("  Aplicando reemplazos iniciales...")
            reemplazar_en_documento_v2(doc, initial_replacements, fase="inicial")

            # --- ELIMINAR Elementos Inactivos/Inválidos ---
            print("  Eliminando secciones/filas de años inactivos...")
            tablas_protegidas = []; eliminar_elementos_inactivos_v_usuario(doc, anos_inactivos, periodo_año, periodo_bimestre, tablas_protegidas)

            # --- Procesar Tablas Suelo/Construcción ---
            print("  Procesando tablas de suelo/construcción...")
            procesar_tablas_suelo_construccion(doc, anos_activos)

            # --- Calcular Sumas ---
            print("  Calculando sumas finales...")
            impuesto_anual_por_año = {2022: 0.0, 2023: 0.0, 2024: 0.0}; sum_impuesto_anual_total = 0.0; sum_monto_actualizado_total = 0.0; sum_recargos_total = 0.0; sum_sancion_total = 0.0
            for year in [2022, 2023, 2024]:
                 for bim in range(1, 7):
                     if (year < periodo_año) or (year == periodo_año and bim < periodo_bimestre): continue
                     campo_bim = f"BIMESTRE {bim} {year}"; campo_act = f"MONTO ACTUALIZADO PREDIAL POR BIMESTRE {bim} {year}"; campo_rec = f"RECARGOS POR BIMESTRE {bim} {year}"; campo_san = f"SANCION POR BIMESTRE {bim} {year}"
                     valor_bim_actual = safe_float(row.get(campo_bim, 0.0)); impuesto_anual_por_año[year] += valor_bim_actual; sum_impuesto_anual_total += valor_bim_actual
                     sum_monto_actualizado_total += safe_float(row.get(campo_act, 0.0)); sum_recargos_total += safe_float(row.get(campo_rec, 0.0)); sum_sancion_total += safe_float(row.get(campo_san, 0.0))
            monto_actualizado_restados = sum_monto_actualizado_total - sum_impuesto_anual_total
            total_general_recalculado = sum_impuesto_anual_total + monto_actualizado_restados + sum_recargos_total + sum_sancion_total
            total_redondeado_bd = round(total_general_recalculado)
            monto_para_bd = f"$ {locale.format_string('%.2f', total_redondeado_bd, grouping=True)}"

            # --- Crear Diccionario Reemplazos FINAL ---
            final_replacements = {}; bimestres_ordinales = {1: "PRIMER", 2: "SEGUNDO", 3: "TERCERO", 4: "CUARTO", 5: "QUINTO", 6: "SEXTO"}
            final_replacements["[SUMA DEL IMPUESTO TOTAL DE TODOS LOS AÑOS]"] = f"$ {formatear_valor_v6(sum_impuesto_anual_total, '[SUMA...IMP]', {'[SUMA...IMP]': 'moneda'})}"
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS]"] = f"$ {formatear_valor_v6(sum_monto_actualizado_total, '[SUMA...ACT]', {'[SUMA...ACT]': 'moneda'})}"
            final_replacements["[SUMA DE LOS MONTOS DE RECARGO TOTAL]"] = f"$ {formatear_valor_v6(sum_recargos_total, '[SUMA...REC]', {'[SUMA...REC]': 'moneda'})}"
            final_replacements["[SUMA DE LA SANCION POR BIMESTRE]"] = f"$ {formatear_valor_v6(sum_sancion_total, '[SUMA...SAN]', {'[SUMA...SAN]': 'moneda'})}"
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS RESTADOS]"] = f"$ {formatear_valor_v6(monto_actualizado_restados, '[SUMA DE MONTOS ACTUALIZADOS RESTADOS]', {'[SUMA DE MONTOS ACTUALIZADOS RESTADOS]': 'moneda'})}"
            final_replacements["[TOTAL]"] = f"$ {locale.format_string('%.2f', total_redondeado_bd, grouping=True)}"
            for año_ip in [2022, 2023, 2024]:
                 valor_ip_anual = round(impuesto_anual_por_año[año_ip], 2); placeholder_ip = f"[IMPUESTO PREDIAL DEL AÑO {año_ip}]"; tipo_dato_ip = {f'IP{año_ip}': 'moneda'}
                 if valor_ip_anual > 0: final_replacements[placeholder_ip] = f"$ {formatear_valor_v6(valor_ip_anual, f'IP{año_ip}', tipo_dato_ip)}"
                 else: final_replacements[placeholder_ip] = " "
            ordinal_texto = bimestres_ordinales.get(periodo_bimestre, ""); final_replacements["[BIMESTRE_ORDINAL]"] = ordinal_texto
            bm_por_anio = {}; anio_base = periodo_año; bimestre_base = periodo_bimestre
            for anio in [2022, 2023, 2024]:
                 if anio == anio_base: bm_por_anio[anio] = bimestre_base
                 elif anio > anio_base: bm_por_anio[anio] = 1
                 else: continue
            for anio_bm, bm_val in bm_por_anio.items(): final_replacements[f"[BM{anio_bm}]"] = f"{bm_val}°"
            final_replacements["[SUMA DEL IMPUESTO PREDIAL ADEUDADO DE TODOS LOS AÑOS]"] = formatear_texto_moneda(sum_impuesto_anual_total)
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS TEXTO]"] = formatear_texto_moneda(sum_monto_actualizado_total)
            final_replacements["[SUMA DE LOS MONTOS DE RECARGO TOTAL TEXTO]"] = formatear_texto_moneda(sum_recargos_total)
            final_replacements["[SUMA DE LA SANCION POR BIMESTRE TEXTO]"] = formatear_texto_moneda(sum_sancion_total)
            final_replacements["[SUMA DE MONTOS ACTUALIZADOS RESTADOS TEXTO]"] = formatear_texto_moneda(monto_actualizado_restados)
            final_replacements["[TOTAL PESOS CORREGIDOS]"] = formatear_texto_moneda(total_redondeado_bd)

            # --- CORREGIDO: Leyenda final siempre en DOCX si no es PM ---
            es_pm = expediente_actual in pm_set
            final_replacements["[LEYENDA_FINAL_PAG17]"] = "" if es_pm else LEYENDA_PAG17_TEXTO

            # --- Aplicar Reemplazos FINALES ---
            print("  Aplicando reemplazos finales...")
            reemplazar_en_documento_v2(doc, final_replacements, fase="final")

             # --- INICIO DEL NUEVO BLOQUE PARA ELIMINAR "CONSTRUCCION" (VERSIÓN 2) ---
            if verificar_ausencia_info_construccion_v2(row, anos_activos):
                print("  Ausencia de información de construcción (columnas clave son cero). Procediendo a eliminar la palabra 'CONSTRUCCION'.")
                eliminar_palabra_especifica_del_documento(doc, "CONSTRUCCION")
            else:
                print("  Información de construcción presente (columnas clave tienen valor). No se eliminará la palabra 'CONSTRUCCION'.")
            # --- FIN DEL NUEVO BLOQUE ---

            # --- Guardar el DOCX Temporal ---
            try:
                doc.save(ruta_temp_docx)
                print(f"  DOCX temporal base guardado: {os.path.basename(ruta_temp_docx)}")
            except PermissionError:
                print(f"  Error PERMISO al guardar DOCX temporal '{os.path.basename(ruta_temp_docx)}'. ¿Archivo abierto? Saltando expediente.")
                estado_final_bd = "Error Permiso DOCX"; continue
            except Exception as e_save_docx:
                print(f"  Error CRITICO al guardar DOCX temporal '{os.path.basename(ruta_temp_docx)}': {e_save_docx}")
                traceback.print_exc(); estado_final_bd = "Error Guardar DOCX"; continue

            # --- FIN LÓGICA DE GENERACIÓN DEL DOCX ---

            # --- INICIO GENERACIÓN PDF FINAL SEGÚN MODO ---
            pdf_creado_ok = False
            if os.path.exists(ruta_temp_pdf_extraccion): os.remove(ruta_temp_pdf_extraccion) # Limpiar previo

            if modo_generacion == MODO_COMPLETO:
                if DOCX2PDF_INSTALLED:
                    try:
                        print(f"  Convirtiendo DOCX completo a PDF final...")
                        convert(ruta_temp_docx, ruta_final_pdf)
                        print(f"  -> PDF COMPLETO generado: {os.path.basename(ruta_final_pdf)}")
                        paginas_finales = contar_paginas_pdf(ruta_final_pdf)  # ← aquí se cuenta TODO el documento

                        pdf_creado_ok = True
                    except Exception as e_conv:
                        print(f"  (!) Error CRITICO al convertir DOCX a PDF completo: {e_conv}"); traceback.print_exc(); estado_final_bd = "Error Conversion Completo"
                        if os.path.exists(ruta_final_pdf): os.remove(ruta_final_pdf)
                        paginas_finales = pd.NA

                else: print("  (!) WARN: 'docx2pdf' no instalado."); estado_final_bd = "Error Falta docx2pdf"

            elif modo_generacion == MODO_ULTIMA: 
                if DOCX2PDF_INSTALLED:
                    pdf_conversion_temp_ok = False
                    try:
                        # 1. Convertir DOCX completo a PDF temporal
                        print(f"  Convirtiendo DOCX a PDF temporal..."); 
                        convert(ruta_temp_docx, ruta_temp_pdf_extraccion); 
                        pdf_conversion_temp_ok = True; 
                        print(f"  PDF temporal creado: {os.path.basename(ruta_temp_pdf_extraccion)}")
                    except Exception as e_conv_temp: 
                        print(f"  (!) Error CRITICO al convertir DOCX a PDF temporal: {e_conv_temp}"); 
                        traceback.print_exc(); 
                        estado_final_bd = "Error Conversion Temp (Ultima)"
                    
                    if pdf_conversion_temp_ok:
                        print(f"  Extrayendo ÚLTIMA página a PDF final...")
                        # 2. Intentar extraer la última página
                        extraccion_ok = extraer_paginas_pdf(ruta_temp_pdf_extraccion, ruta_final_pdf, "ULTIMA")
                        
                        if extraccion_ok:
                            print(f"  -> PDF ÚLTIMA PÁGINA supuestamente generado: {os.path.basename(ruta_final_pdf)}")
                            print(f"     DEBUG: Ruta completa PDF final: '{os.path.abspath(ruta_final_pdf)}'")
                            
                            # 3. Pausa y Verificación de Existencia (¡IMPORTANTE!)
                            print("     DEBUG: Pausando 0.5 segundos antes de verificar/contar...")
                            time.sleep(0.5) # Aumentar ligeramente la pausa por si acaso
                            
                            if not os.path.exists(ruta_final_pdf):
                                print(f"     DEBUG: ¡VERIFICACIÓN FALLIDA! El archivo '{ruta_final_pdf}' NO existe después de la pausa.")
                                estado_final_bd = "Error PDF Final No Encontrado Post-Extraccion"
                                pdf_creado_ok = False # Asegurar que no se cuente como éxito
                            else:
                                print(f"     DEBUG: ¡VERIFICACIÓN OK! El archivo '{ruta_final_pdf}' SÍ existe después de la pausa.")
                                pdf_creado_ok = True # Marcar como OK solo si existe
                        else: 
                            print(f"  (!) Error al extraer la última página (extraer_paginas_pdf devolvió False)."); 
                            estado_final_bd = "Error Extraccion Ultima";
                            pdf_creado_ok = False # No se creó
                            # Limpiar si la extracción falló pero el archivo se creó parcialmente o vacío
                            if os.path.exists(ruta_final_pdf): 
                                try:
                                    os.remove(ruta_final_pdf)
                                    print(f"     DEBUG: PDF final '{ruta_final_pdf}' eliminado debido a fallo en extracción.")
                                except Exception as e_del_fail:
                                    print(f"     DEBUG: No se pudo eliminar PDF final '{ruta_final_pdf}' tras fallo: {e_del_fail}")
                else: 
                    print("  (!) WARN: 'docx2pdf' no instalado."); 
                    estado_final_bd = "Error Falta docx2pdf"

            elif modo_generacion == MODO_RESTO:
                if DOCX2PDF_INSTALLED:
                    pdf_conversion_temp_ok = False
                    try:
                         print(f"  Convirtiendo DOCX a PDF temporal..."); convert(ruta_temp_docx, ruta_temp_pdf_extraccion); pdf_conversion_temp_ok = True; print(f"  PDF temporal creado: {os.path.basename(ruta_temp_pdf_extraccion)}")
                    except Exception as e_conv_temp: print(f"  (!) Error CRITICO al convertir DOCX a PDF temporal: {e_conv_temp}"); traceback.print_exc(); estado_final_bd = "Error Conversion Temp (Resto)"
                    if pdf_conversion_temp_ok:
                        print(f"  Extrayendo RESTO de páginas a PDF final...")
                        if extraer_paginas_pdf(ruta_temp_pdf_extraccion, ruta_final_pdf, "RESTO"):
                            paginas_resto = contar_paginas_pdf(ruta_final_pdf)
                            if paginas_resto is not None: print(f"  -> PDF RESTO ({paginas_resto} pág.) generado: {os.path.basename(ruta_final_pdf)}"); pdf_creado_ok = True
                            else: print(f"  (!) Error contando páginas del PDF RESTO generado."); estado_final_bd = "Error Conteo Resto";
                            if os.path.exists(ruta_final_pdf): os.remove(ruta_final_pdf)
                        else: print(f"  (!) Error al extraer el resto de páginas."); estado_final_bd = "Error Extraccion Resto";
                        if os.path.exists(ruta_final_pdf): os.remove(ruta_final_pdf)
                else: print("  (!) WARN: 'docx2pdf' no instalado."); estado_final_bd = "Error Falta docx2pdf"

            # --- Contar páginas y asignar estado SI el PDF se creó OK ---
            if pdf_creado_ok: # <--- Esta variable ahora se establece correctamente
                pdf_generado_final = True
                # 4. Intentar contar páginas SOLO si pdf_creado_ok es True

                if paginas_finales is not None and paginas_finales >= 0:
                    hojas_calculadas = math.ceil(paginas_finales / 2.0) * 2 if paginas_finales > 0 else 0
                    print(f"  Páginas contadas PDF final: {paginas_finales}, Hojas calculadas (2 copias): {hojas_calculadas}")
                else: 
                    # Si pdf_creado_ok era True pero contar_paginas_pdf falla ahora, es un problema diferente.
                    print(f"  (!) Advertencia: PDF final '{os.path.basename(ruta_final_pdf)}' existe pero no se pudieron contar sus páginas.")
                    hojas_calculadas = pd.NA 
                    # Podrías querer cambiar el estado aquí si el conteo es crucial, 
                    # pero por ahora mantenemos el estado de generación exitosa.
                    
                # Asignar estado de ÉXITO (ya que el PDF se creó)
                if modo_generacion == MODO_ULTIMA: estado_final_bd = ESTADO_GEN_ULTIMA
                elif modo_generacion in [MODO_COMPLETO, MODO_RESTO]: estado_final_bd = ESTADO_GEN_COMPLETO
                # else: estado_final_bd ya tendría un valor de error si pdf_creado_ok fuera False

            # else: Si pdf_creado_ok es False, estado_final_bd ya tiene un valor de error

        # --- Bloque Except General ---
        except Exception as main_loop_error:
            print(f"  (!) Error INESPERADO procesando expediente {expediente_actual}: {main_loop_error}")
            traceback.print_exc()
            estado_final_bd = ESTADO_ERROR_GENERACION

        # --- Bloque Finally ---
        finally:
            # Limpiar DOCX temporal
            if eliminar_docx_intermedio and os.path.exists(ruta_temp_docx):
                try: os.remove(ruta_temp_docx); #print(f"  DOCX temporal eliminado.") # Log reducido
                except Exception as e_del_docx: print(f"  Warn: No se pudo eliminar DOCX temp '{os.path.basename(ruta_temp_docx)}': {e_del_docx}")
            # Limpiar PDF de extracción temporal
            if os.path.exists(ruta_temp_pdf_extraccion):
                try: os.remove(ruta_temp_pdf_extraccion); #print(f"  PDF temporal de extracción eliminado.") # Log reducido
                except Exception as e_del_extract: print(f"  Warn: No se pudo eliminar PDF temp extract '{os.path.basename(ruta_temp_pdf_extraccion)}': {e_del_extract}")

            # --- Actualizar Base de Datos Maestra ---
        # --- Actualizar Base de Datos Maestra ---
        if expediente_actual: # Solo si tenemos expediente
            # Crear el diccionario SOLO con las columnas que Generador.py
            # establece o actualiza activamente para ESTE expediente.
            registro_data_a_actualizar = {
                COL_EXPEDIENTE: expediente_actual,
                COL_NOMBRE_CONTRIBUYENTE: nombre_contribuyente, # Valor del CSV o procesado
                COL_DIRECCION: direccion_contribuyente,       # Valor del CSV o procesado
                COL_COLONIA: colonia_para_bd,               # Valor del CSV o procesado
                COL_BIMESTRE: bimestre_para_bd,             # Valor del CSV o procesado
                COL_ANOS: anos_para_bd,                   # Calculado por Generador
                COL_MONTO: monto_para_bd,                   # Calculado por Generador
                COL_ESTADO: estado_final_bd,                # Establecido por Generador
                COL_HOJAS_DOC: paginas_finales,
                # NO INCLUIR explícitamente:
                # COL_ID_LOTE, 
                # COL_PRIORIDAD,
                # COL_IMPRESORA,
                # COL_FECHA_HORA_IMPRESION
                # La función `actualizar_o_agregar_registro_bd` se encargará de:
                #  - Si la fila es NUEVA: inicializar esas columnas con NA/"" (según su tipo).
                #  - Si la fila EXISTE: MANTENER los valores que ya tenían esas columnas en `df_bd_maestra`.
            }

            df_bd_maestra = actualizar_o_agregar_registro_bd(df_bd_maestra, registro_data_a_actualizar)
            registros_actualizados_en_bd += 1

            # Guardar BD periódicamente
            if registros_actualizados_en_bd > 0 and registros_actualizados_en_bd % 50 == 0: # Guardar cada 50
                print(f"\nGuardando progreso intermedio de BD Maestra ({registros_actualizados_en_bd} actualizados)...")
                if not guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA):
                     print("(!) ERROR GRAVE: Falló el guardado intermedio de la BD Maestra.")

            # Incrementar contador de generados solo si el PDF final se creó ok
            if pdf_generado_final:
                generados_count += 1

    # --- Fin del Bucle Principal ---
    print(f"\n--- Proceso de Generación Finalizado (Modo: {modo_generacion}) ---")
    print(f"Documentos PDF generados/intentados en esta ejecución: {generados_count}")
    print(f"Registros actualizados/agregados en BD Maestra en esta ejecución: {registros_actualizados_en_bd}")

    # --- Guardar Estado Final de la BD Maestra ---
    if registros_actualizados_en_bd > 0:
        print("Guardando estado final de la Base de Datos Maestra...")
        if not guardar_bd_maestra(df_bd_maestra, RUTA_BD_MAESTRA):
             print("(!) ERROR GRAVE: Falló el guardado FINAL de la BD Maestra.")
    else:
        print("No hubo cambios detectados en la Base de Datos Maestra para guardar.")

# --- Fin de la función generar_documentos ---

def reemplazar_texto_en_tablas(doc, reemplazos):
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    texto_original = parrafo.text
                    texto_modificado = texto_original
                    for marcador, valor in reemplazos.items():
                        if marcador in texto_modificado:
                            texto_modificado = texto_modificado.replace(marcador, str(valor))
                    if texto_modificado != texto_original:
                        for i in range(len(parrafo.runs) - 1, -1, -1):
                            parrafo._element.remove(parrafo.runs[i]._element)
                        parrafo.add_run(texto_modificado)


def reemplazar_texto_en_parrafos(doc, reemplazos):
    for parrafo in doc.paragraphs:
        texto_original = parrafo.text
        texto_modificado = texto_original
        for marcador, valor in reemplazos.items():
            if marcador in texto_modificado:
                texto_modificado = texto_modificado.replace(marcador, str(valor))
        if texto_modificado != texto_original:
            # Sobrescribir todo el párrafo
            for i in range(len(parrafo.runs) - 1, -1, -1):
                parrafo._element.remove(parrafo.runs[i]._element)
            parrafo.add_run(texto_modificado)

def verificar_ausencia_info_construccion_v2(row_data, anos_activos):
    """
    Verifica si no hay información significativa de construcción basándose en columnas específicas.
    La ausencia se determina si TODAS las columnas relevantes de construcción para los años activos
    tienen un valor numérico de cero.

    'row_data' es la fila (Serie de pandas) del expediente actual.
    'anos_activos' es una lista de los años que se consideran activos para este expediente.
    """
    # Nombres EXACTOS de las columnas en tu BASE_DE_DATOS.csv que corresponden a esos placeholders.
    # Es crucial que estos nombres de columna coincidan con los de tu archivo CSV.
    # Asumo que los placeholders que me diste se mapean a nombres de columna sin los corchetes.
    # ¡VERIFICA ESTOS NOMBRES DE COLUMNA!
    columnas_clave_construccion_por_ano = {
        2022: "VALOR CATASTRAL CONSTRUCCION 2022", # Ajusta si el nombre real de la columna es diferente
        2023: "VALOR CATASTRAL CONSTRUCCION 2023", # Ajusta si el nombre real de la columna es diferente
        2024: "VALOR CATASTRAL CONSTRUCCION 2024"  # Ajusta si el nombre real de la columna es diferente (ej. "VALOR CATASTRLA...")
                                                  # Nota: Mencionaste "VALOR CATASTRLA CONSTRUCCION 2024", asegúrate de usar el nombre correcto aquí.
                                                  # Voy a usar "VALOR CATASTRAL CONSTRUCCION 2024" asumiendo un typo, pero verifica.
    }

    # print(f"DEBUG: Verificando ausencia de construcción v2. Años activos: {anos_activos}")

    if not anos_activos:
        # Si no hay años activos para el expediente, se considera que no hay información de construcción relevante que mostrar.
        # Por lo tanto, se podría proceder a eliminar la palabra "CONSTRUCCION".
        # print("DEBUG v2: No hay años activos, considerando ausencia de info construcción.")
        return True

    hay_alguna_col_construccion_relevante_con_valor = False

    for ano in anos_activos:
        nombre_columna_especifica = columnas_clave_construccion_por_ano.get(ano)

        if nombre_columna_especifica:
            if nombre_columna_especifica in row_data:
                valor_str = str(row_data.get(nombre_columna_especifica, "0")).strip()
                
                # Intentar convertir a float para una comparación numérica robusta
                try:
                    valor_num_str = valor_str.replace(",", "").replace("$", "").strip()
                    if not valor_num_str: # Si después de limpiar queda vacío, es como cero
                        valor_num = 0.0
                    else:
                        valor_num = float(valor_num_str)
                    
                    if abs(valor_num) > 1e-9: # Usar una pequeña tolerancia para comparar con cero
                        # print(f"DEBUG v2: Info de construcción ENCONTRADA para año {ano} en columna '{nombre_columna_especifica}', valor numérico: {valor_num}")
                        hay_alguna_col_construccion_relevante_con_valor = True
                        break # Si encontramos un valor, ya no necesitamos seguir revisando
                except ValueError:
                    # Si no se puede convertir a float, y no es una cadena vacía obvia,
                    # podría ser un error en los datos o un texto.
                    # Para ser estrictos con "si detecta 0", un texto no es 0.
                    # Si la columna DEBE ser numérica, esto indica un problema de datos.
                    # Si puede ser texto (ej. "N/A"), y "N/A" no es "0", entonces hay "algo".
                    if valor_str and valor_str.lower() not in ['0', '0.0', '0.00', '', 'nan', 'na', '-', '0.000', '0.0000']:
                        # print(f"DEBUG v2: Info de construcción (texto no cero) ENCONTRADA para año {ano} en '{nombre_columna_especifica}', valor: '{valor_str}'")
                        hay_alguna_col_construccion_relevante_con_valor = True
                        break 
            # else:
                # print(f"DEBUG v2: Columna '{nombre_columna_especifica}' no encontrada en row_data para año activo {ano}.")
                # Si la columna esperada no existe en los datos para un año activo,
                # podríamos considerarlo como "sin valor" o "cero" para esa columna.
                # La lógica actual no la penaliza, simplemente no encuentra valor.

    if hay_alguna_col_construccion_relevante_con_valor:
        # print("DEBUG v2: Se encontró al menos un valor de construcción > 0 en años activos.")
        return False # Se encontró información de construcción
    else:
        # print("DEBUG v2: NO se encontró valor de construcción > 0 en columnas clave de años activos.")
        return True # Confirma ausencia de información de construcción (todo es cero o no relevante)
    
def eliminar_palabra_especifica_del_documento(document, palabra_a_eliminar):
    """
    Elimina todas las ocurrencias de una palabra específica (case-sensitive)
    en todos los párrafos y tablas del documento.
    Modifica el texto de los párrafos directamente.
    """
    print(f"      Iniciando búsqueda y eliminación de la palabra: '{palabra_a_eliminar}'")
    
    def procesar_lista_de_parrafos(paragraphs_list):
        for p_idx, paragraph in enumerate(list(paragraphs_list)):
            if paragraph is None or not hasattr(paragraph, 'text'):
                continue
            if palabra_a_eliminar in paragraph.text:
                texto_original_parrafo = paragraph.text
                texto_modificado_parrafo = texto_original_parrafo.replace(palabra_a_eliminar, "")
                if texto_original_parrafo != texto_modificado_parrafo:
                    for run in list(paragraph.runs):
                        p_element = paragraph._element
                        try:
                            p_element.remove(run._element)
                        except ValueError:
                            pass
                    if texto_modificado_parrafo.strip():
                        paragraph.add_run(texto_modificado_parrafo)
    
    procesar_lista_de_parrafos(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                procesar_lista_de_parrafos(cell.paragraphs)
    for section in document.sections:
        for header_footer_part in [section.header, section.footer, 
                                   section.first_page_header, section.first_page_footer, 
                                   section.even_page_header, section.even_page_footer]:
            if header_footer_part is not None:
                procesar_lista_de_parrafos(header_footer_part.paragraphs)
                if hasattr(header_footer_part, 'tables'):
                    for table_hf in header_footer_part.tables:
                        for row_hf in table_hf.rows:
                            for cell_hf in table_hf.cells:
                                procesar_lista_de_parrafos(cell_hf.paragraphs)
    print(f"      Búsqueda y eliminación de '{palabra_a_eliminar}' completada.")

def generar_vista_rapida_dos_paginas_predial(
    df_datos_para_procesar_ordenados, # DataFrame ya ordenado y limitado por generador.py
    pm_set_actual,
    config_predial_actual, # Diccionario del script maestro con rutas, etc.
    nombres_columnas_csv,
    tipos_columnas_csv
):
    """
    Genera un PDF de las primeras dos páginas para visualización rápida.
    NO interactúa con la Base de Datos Maestra.
    Guarda los PDFs en una subcarpeta "VISTAS_RAPIDAS".
    """
    print(f"\n--- [Logica Predial] Iniciando Generación de VISTA RÁPIDA (Primeras Dos Páginas) ---")
    if df_datos_para_procesar_ordenados.empty:
        print("  [Logica Predial] No hay registros proporcionados para la vista rápida.")
        return

    plantilla_path = config_predial_actual["template_file_path"]
    # Usar la carpeta de salida principal del modo, pero crear una subcarpeta para estas vistas
    carpeta_salida_base_modo = Path(config_predial_actual["output_docs_path"])
    carpeta_salida_vistas_rapidas = carpeta_salida_base_modo / "VISTAS_RAPIDAS"
    carpeta_salida_vistas_rapidas.mkdir(parents=True, exist_ok=True)

    # Determinar nombres de columnas clave del CSV (como en tu generar_documentos_predial_core)
    col_expediente_datos = ""
    if "EXPEDIENTE" in nombres_columnas_csv and "EXPEDIENTE" in df_datos_para_procesar_ordenados.columns:
        col_expediente_datos = "EXPEDIENTE"
    elif nombres_columnas_csv:
        col_expediente_datos = nombres_columnas_csv[0]

    if not col_expediente_datos or col_expediente_datos not in df_datos_para_procesar_ordenados.columns:
        print(f"    (!) Error [Vista Rápida]: No se pudo determinar la columna de expediente. Cols: {df_datos_para_procesar_ordenados.columns.tolist()}")
        return

    col_nombre_contrib_datos = config_predial_actual.get("col_nombre_csv_original", "NOMBRE")
    col_colonia_datos = config_predial_actual.get("col_colonia_csv_original", "COLONIA")
    periodo_col_datos = "PERIODO" # Asume que se llama así en tu CSV

    # Configurar locale si aún no se ha hecho globalmente por el script que llama
    if not configurar_locale(): # Llama a tu función configurar_locale
         print("(!) Error crítico [Vista Rápida]: No se pudo configurar el locale español. Abortando vista rápida.")
         return


    generados_count_vista = 0
    total_a_procesar_vista = len(df_datos_para_procesar_ordenados)

    for index, row in df_datos_para_procesar_ordenados.iterrows():
        expediente_actual = str(row.get(col_expediente_datos, "")).strip()
        nombre_contribuyente = str(row.get(col_nombre_contrib_datos, "SIN_NOMBRE")).strip()
        # ... (extraer otros datos necesarios de 'row' para los placeholders)

        print(f"\n  --- [Vista Rápida] ({generados_count_vista + 1}/{total_a_procesar_vista}) Exp: {expediente_actual} ---")

        nombre_base_limpio = limpiar_texto(f"{expediente_actual}_{nombre_contribuyente}")
        if not nombre_base_limpio: nombre_base_limpio = f"exp_{expediente_actual}"

        colonia_actual_datos = str(row.get(col_colonia_datos, "")).strip()
        if not colonia_actual_datos or colonia_actual_datos == "0" or colonia_actual_datos.upper() == "VACIA":
            subcarpeta_colonia_vista = carpeta_salida_vistas_rapidas / "VACIAS"
        else:
            subcarpeta_colonia_vista = carpeta_salida_vistas_rapidas / "COLONIAS" / limpiar_texto(colonia_actual_datos)
        subcarpeta_colonia_vista.mkdir(parents=True, exist_ok=True)

        ruta_temp_docx_vista = subcarpeta_colonia_vista / f"~{nombre_base_limpio}_vista.docx"
        ruta_final_pdf_vista = subcarpeta_colonia_vista / f"{nombre_base_limpio}_vista_2pg.pdf"
        ruta_temp_pdf_completo_vista = subcarpeta_colonia_vista / f"~{nombre_base_limpio}_vistafull_TEMP.pdf"

        try:
            # --- INICIO DE LA LÓGICA DE GENERACIÓN DEL DOCUMENTO WORD ---
            # Esta sección debe ser una copia casi idéntica de la lógica que tienes en
            # `generar_documentos_predial_core` (o `generar_documentos` si esa es tu función principal)
            # para calcular todos los placeholders, incluyendo los complejos como [TOTAL] y [TOTAL PESOS CORREGIDOS].
            # No omitas cálculos aquí pensando en optimizar, ya que los placeholders de las
            # primeras páginas podrían depender de ellos.

            # 1. Determinar años activos/inactivos
            anos_activos = []
            anos_inactivos = []
            # ... (tu lógica completa para llenar anos_activos y anos_inactivos usando 'row') ...
            # Ejemplo simplificado:
            for year_val_check in [2022, 2023, 2024]:
                # Supongamos que tienes una columna como "IMPUESTO PREDIAL DEL AÑO 202X" o similar
                # o una columna llamada "202X" con un valor.
                year_col_check_name = f"IMPUESTO PREDIAL DEL AÑO {year_val_check}" # Ajusta a tu nombre de columna real
                if year_col_check_name in row.index and pd.notna(row[year_col_check_name]) and safe_float(row[year_col_check_name], 0.0) != 0.0:
                    anos_activos.append(year_val_check)
                else:
                    anos_inactivos.append(year_val_check)

            # 2. Placeholders de fecha, año base
            now_vista = datetime.now()
            fecha_hoy_str_vista = now_vista.strftime("%d de %B de %Y").lower()
            fecha_texto_str_vista = "(num2words no disponible)"
            if NUM2WORDS_INSTALLED:
                try:
                    dia_num_vista = now_vista.day
                    ano_num_vista = now_vista.year
                    dia_palabra_vista = num2words(dia_num_vista, lang='es')
                    ano_palabra_vista = num2words(ano_num_vista, lang='es')
                    medio_fecha_vista = now_vista.strftime("de %B de %Y").lower()
                    fecha_texto_str_vista = f"{dia_num_vista}-{dia_palabra_vista} {medio_fecha_vista}-{ano_palabra_vista}"
                except Exception as e_n2w:
                    print(f"    Advertencia (Vista Rápida): Error generando fecha_texto: {e_n2w}")

            ano_placeholder_vista = str(min(anos_activos)) if anos_activos else str(now_vista.year)

            # 3. Cargar plantilla
            doc_vista = Document(plantilla_path)

            # 4. Extraer periodo y calcular bimestre/año para la BD (aunque no se use para BD aquí, puede ser para placeholders)
            periodo_str_vista = str(row.get(periodo_col_datos, "")).strip()
            periodo_año_vista = 0
            periodo_bim_vista = 0
            bimestre_ordinal_vista_texto = ""
            bimestres_ordinales_map = {1: "PRIMER", 2: "SEGUNDO", 3: "TERCERO", 4: "CUARTO", 5: "QUINTO", 6: "SEXTO"}

            if re.match(r"^\d{6}$", periodo_str_vista):
                try:
                    periodo_año_vista = int(periodo_str_vista[:4])
                    periodo_bim_vista = int(periodo_str_vista[4:6])
                    bimestre_ordinal_vista_texto = bimestres_ordinales_map.get(periodo_bim_vista, "")
                except ValueError:
                    pass # periodo_año_vista y periodo_bim_vista seguirán siendo 0


            # 5. Reemplazos Iniciales (todos los que puedan estar en las 2 primeras páginas)
            initial_replacements_vista = {
                "[FECHA]": fecha_hoy_str_vista,
                "[FECHA_TEXTO]": fecha_texto_str_vista,
                "[AÑO]": ano_placeholder_vista,
                "[EXPEDIENTE]": expediente_actual, # Este es clave
                # Añade aquí los placeholders de tu lista:
                "[NOMBRE]": str(row.get(col_nombre_contrib_datos, "")),
                "[DIRECCION]": str(row.get("DIRECCION", "")), # Asegúrate que "DIRECCION" es el nombre en tu CSV
                "[NUM_EXTERIOR]": str(row.get("NUM_EXTERIOR", "")),
                "[NUM_INTERIOR]": str(row.get("NUM_INTERIOR", "")),
                "[COLONIA]": str(row.get(col_colonia_datos, "")), # Usa la variable col_colonia_datos
                "[MUNICIPIO]": str(row.get("MUNICIPIO", "")), # Reemplaza "MUNICIPIO" con el nombre real de tu columna CSV
                "[CP]": str(row.get("CP", "")), # Reemplaza "CP" con el nombre real de tu columna CSV
                "[UBICACION]": str(row.get("UBICACION_DEL_PREDIO", "")), # Nombre ejemplo, usa tu columna CSV
                "[UNUM_EXTERIOR]": str(row.get("UNUM_EXTERIOR", "")),
                "[UNUM_INTERIOR]": str(row.get("UNUM_INTERIOR", "")),
                "[UCOLONIA]": str(row.get("UCOLONIA", "")),
                "[UMUNICIPIO]": str(row.get("UMUNICIPIO", "")),
            }
            # Añadir más reemplazos iniciales que provienen directamente del CSV y podrían estar en las primeras 2 pgs
            for col_csv_name_vista in nombres_columnas_csv:
                placeholder_vista = f"[{col_csv_name_vista}]"
                if placeholder_vista not in initial_replacements_vista: # Evitar sobrescribir los ya definidos
                    # Aplica la lógica de años inactivos y formateo como en generar_documentos_predial_core
                    # Esta es una simplificación, debes copiar la lógica completa de esa sección:
                    # if col_year_parsed is not None and col_year_parsed in anos_inactivos: valor_final_calc_vista = "" else ...
                    initial_replacements_vista[placeholder_vista] = formatear_valor_v6(row.get(col_csv_name_vista), col_csv_name_vista, tipos_columnas_csv)


            reemplazar_en_documento_v2(doc_vista, initial_replacements_vista, fase="inicial_vista")

            # 6. Eliminar elementos inactivos, procesar tablas (estas pueden afectar las primeras 2 pgs)
            # Es importante ejecutar esto si las tablas/secciones que se eliminan están al inicio.
            # Copia estas llamadas de tu función `generar_documentos_predial_core`
            # eliminar_elementos_inactivos_v_usuario(doc_vista, anos_inactivos, periodo_año_vista, periodo_bim_vista, [])
            # procesar_tablas_suelo_construccion(doc_vista, anos_activos)


            # 7. Calcular Sumas (NECESARIO PARA [TOTAL] y [TOTAL PESOS CORREGIDOS])
            # Copia EXACTAMENTE la lógica de cálculo de sumas de `generar_documentos_predial_core`
            # sum_impuesto_anual_total_calc_vista = 0.0
            # sum_monto_actualizado_total_calc_vista = 0.0
            # ... todos los bucles y sumas ...
            # total_general_recalculado_calc_vista = ...
            # total_redondeado_calc_vista = round(total_general_recalculado_calc_vista)
            # ¡Esta parte es crítica! Si no, [TOTAL] será incorrecto.
            # EJEMPLO (DEBES COPIAR TU LÓGICA COMPLETA DE SUMAS DE generar_documentos_predial_core):
            sum_impuesto_anual_total_calc_vista = safe_float(row.get("ADEUDO TOTAL", 0.0)) # Asumiendo que tienes una columna así para simplificar el ejemplo
            monto_actualizado_restados_calc_vista = safe_float(row.get("MONTO ACTUALIZADO", 0.0)) - sum_impuesto_anual_total_calc_vista
            sum_recargos_total_calc_vista = safe_float(row.get("RECARGOS",0.0))
            sum_sancion_total_calc_vista = safe_float(row.get("SANCIONES",0.0))
            total_redondeado_calc_vista = round(sum_impuesto_anual_total_calc_vista + monto_actualizado_restados_calc_vista + sum_recargos_total_calc_vista + sum_sancion_total_calc_vista)


            # 8. Reemplazos Finales (incluyendo [TOTAL], [TOTAL PESOS CORREGIDOS], etc.)
            final_replacements_vista = {
                "[BIMESTRE_ORDINAL]": bimestre_ordinal_vista_texto,
                "[TOTAL]": f"${locale.format_string('%.2f', total_redondeado_calc_vista, grouping=True)}",
                "[TOTAL PESOS CORREGIDOS]": formatear_texto_moneda(total_redondeado_calc_vista),
                # Incluye otros placeholders finales que puedan estar en las primeras 2 páginas
                # La leyenda de la página 17 usualmente no, pero el cálculo de es_pm es rápido
                "[LEYENDA_FINAL_PAG17]": "" if expediente_actual in pm_set_actual else LEYENDA_PAG17_TEXTO,
            }
            # Añade aquí los placeholders de [IMPUESTO PREDIAL DEL AÑO X] si están en las primeras págs
            # y las sumas en texto si también lo están.

            reemplazar_en_documento_v2(doc_vista, final_replacements_vista, fase="final_vista")

            # 9. Eliminar palabra "CONSTRUCCION" si aplica
            # if verificar_ausencia_info_construccion_v2(row, anos_activos):
            #    eliminar_palabra_especifica_del_documento(doc_vista, "CONSTRUCCION")

            # --- FIN DE LA LÓGICA DE GENERACIÓN DEL DOCUMENTO WORD ---

            doc_vista.save(ruta_temp_docx_vista)

            pdf_completo_ok = False
            if DOCX2PDF_INSTALLED:
                try:
                    convert(ruta_temp_docx_vista, ruta_temp_pdf_completo_vista)
                    if ruta_temp_pdf_completo_vista.exists() and ruta_temp_pdf_completo_vista.stat().st_size > 0:
                        pdf_completo_ok = True
                    else:
                         print(f"    (!) Advertencia [Vista Rápida]: PDF completo temporal para {expediente_actual} está vacío o no se creó.")
                except Exception as e_conv:
                    print(f"    (!) Error [Vista Rápida] convirtiendo a PDF completo para {expediente_actual}: {e_conv}")
            else:
                print("    ADVERTENCIA [Vista Rápida]: docx2pdf no instalado. No se puede generar PDF.")

            if pdf_completo_ok:
                            extraccion_exitosa = extraer_paginas_pdf(ruta_temp_pdf_completo_vista, ruta_final_pdf_vista, "PRIMERAS_DOS")
                            if extraccion_exitosa:
                                if ruta_final_pdf_vista.exists(): # Comprobación básica de existencia
                                    num_pags_final = contar_paginas_pdf(ruta_final_pdf_vista)
                                    if num_pags_final is not None and num_pags_final > 0:
                                        print(f"    -> Vista Rápida ({num_pags_final} pgs) generada: {ruta_final_pdf_vista.name}")
                                        generados_count_vista += 1
                                    elif num_pags_final == 0: # Puede ser 0 si el original tenía 0 pgs, o 1 pág y se extrajo bien.
                                        print(f"    -> Vista Rápida (0 o 1 pág resultante) generada: {ruta_final_pdf_vista.name}")
                                        generados_count_vista += 1
                                    else: # num_pags_final es None (error al contar) o algo inesperado
                                        print(f"    (*) Advertencia [Vista Rápida]: PDF final '{ruta_final_pdf_vista.name}' generado, pero hubo un problema al contar sus páginas (Resultado: {num_pags_final}). Se conserva.")
                                        generados_count_vista += 1 # Contar como generado si la extracción fue ok y el archivo existe
                                else:
                                    print(f"    (!) Error [Vista Rápida]: PDF final '{ruta_final_pdf_vista.name}' NO encontrado después de una supuesta extracción exitosa.")
                            else:
                                print(f"    (!) Error [Vista Rápida]: Falló la extracción de las primeras dos páginas para {expediente_actual}.")
            traceback.print_exc()
        finally:
            if ruta_temp_docx_vista.exists():
                try: ruta_temp_docx_vista.unlink()
                except Exception: pass
            if ruta_temp_pdf_completo_vista.exists():
                try: ruta_temp_pdf_completo_vista.unlink()
                except Exception: pass

    print(f"\n--- [Logica Predial] Generación de Vistas Rápidas (Primeras Dos Páginas) Finalizada ---")
    print(f"    Total de vistas rápidas generadas/intentadas: {generados_count_vista}")

# --- Bloque de Ejecución Principal (con Menú) ---
if __name__ == "__main__":
    while True:  # Bucle para mostrar el menú hasta que el usuario elija salir
        print("\n" + "=" * 25 + " MENÚ GENERADOR " + "=" * 25)
        print("1. Generar Documento COMPLETO (Proceso por lotes)")
        print("2. Generar SÓLO ÚLTIMA PÁGINA (Proceso por lotes)")
        print("3. Generar RESTO del Documento (para los de 'ultima hoja', Proceso por lotes)")
        print("4. Generar Expedientes ESPECÍFICOS")
        print("S. Salir")
        print("=" * 68)

        opcion = input("Seleccione una opción: ").strip().upper()

        csv_data_file = "BASE_DE_DATOS.csv"
        config_file = "config_columnas.csv"
        pm_csv_file = "PM.csv"
        plantilla_word = "PLANTILLA.docx"

        if opcion in ['1', '2', '3']:
            modo_seleccionado = None
            if opcion == '1':
                modo_seleccionado = "COMPLETO"
                print("\n--- Modo: Generar Documento COMPLETO ---")
            elif opcion == '2':
                modo_seleccionado = "ULTIMA"
                print("\n--- Modo: Generar SÓLO ÚLTIMA PÁGINA ---")
            elif opcion == '3':
                modo_seleccionado = "RESTO"
                print("\n--- Modo: Generar RESTO del Documento ---")

            max_a_generar = 0
            while True:
                try:
                    num_input = input("¿Cuántos documentos desea generar? (0 para todos los pendientes en este modo): ").strip()
                    max_a_generar = int(num_input)
                    if max_a_generar < 0:
                        print("Ingrese un número positivo o 0.")
                    else:
                        break
                except ValueError:
                    print("Entrada inválida. Ingrese un número.")
                except EOFError: # Manejo por si el usuario presiona Ctrl+D o similar
                    print("\nOperación cancelada.")
                    max_a_generar = -1 # Indicador para saltar la generación
                    break
            
            if max_a_generar < 0: # Si se canceló la entrada de número
                continue # Volver al menú principal

            generar_documentos(
                csv_data_path=csv_data_file,
                config_path=config_file,
                pm_csv_path=pm_csv_file,
                plantilla_path=plantilla_word,
                modo_generacion=modo_seleccionado,
                max_archivos=max_a_generar
            )
            input("\nPresiona Enter para volver al menú principal...")

        elif opcion == '4':
            print("\n--- Modo: Generar Expedientes ESPECÍFICOS ---")
            expedientes_str = input("Ingrese los expedientes separados por comas (ej. 1234567,0234567): ").strip()
            if not expedientes_str:
                print("No se ingresó ningún expediente. Regresando al menú.")
                input("\nPresiona Enter para volver al menú principal...")
                continue

            expedientes_lista = [exp.strip() for exp in expedientes_str.split(',') if exp.strip()]
            if not expedientes_lista:
                print("Lista de expedientes vacía después de procesar. Volviendo al menú.")
                input("\nPresiona Enter para volver al menú principal...")
                continue

            registrar_en_bd_str = input("¿Desea registrar los expedientes generados en la base de datos? (S/N): ").strip().upper()
            registrar_bd_bool = registrar_en_bd_str == 'S'

            # Asumiendo que la función generar_expedientes_especificos ya está definida en tu script
            generar_expedientes_especificos(
                csv_data_path=csv_data_file,
                config_path=config_file,
                pm_csv_path=pm_csv_file,
                plantilla_path=plantilla_word,
                lista_expedientes=expedientes_lista,
                registrar_bd=registrar_bd_bool
            )
            input("\nPresiona Enter para volver al menú principal...")

        elif opcion == 'S':
            print("Saliendo del programa.")
            break
        else:
            print("Opción inválida. Intente de nuevo.")
            input("\nPresiona Enter para volver al menú principal...")
            continue